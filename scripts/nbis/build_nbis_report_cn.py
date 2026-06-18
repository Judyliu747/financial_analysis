#!/usr/bin/env python3
"""Nebius Group (NBIS) Q1 2026 业绩更新 — 中文版 DOCX。
正文宋体，标题黑体。输出: output/NBIS/NBIS_Q1_FY2026_业绩更新报告_中文版.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/NBIS"
SONG = "宋体"     # body
HEI  = "黑体"     # headings
LATIN = "Times New Roman"

NAVY  = RGBColor(0x0B, 0x25, 0x45)
BLUE  = RGBColor(0x1F, 0x6F, 0xEB)
GREEN = RGBColor(0x2B, 0xA8, 0x4A)
RED   = RGBColor(0xD1, 0x49, 0x5B)
GREY  = RGBColor(0x66, 0x66, 0x66)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = SONG
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)


def set_font(run, size=10.5, bold=False, italic=False, color=None, cn=SONG):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = LATIN
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN); rfonts.set(qn("w:hAnsi"), LATIN)
    rfonts.set(qn("w:eastAsia"), cn)


def para(text="", size=10.5, bold=False, italic=False, color=None,
         align=None, space_after=6, space_before=0, cn=SONG):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        set_font(run, size, bold, italic, color, cn=cn)
    return p


def heading(text, size=14, color=NAVY, space_before=10):
    return para(text, size=size, bold=True, color=color, space_after=4,
                space_before=space_before, cn=HEI)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), LATIN); rfonts.set(qn("w:hAnsi"), LATIN)
    rfonts.set(qn("w:eastAsia"), SONG); rPr.append(rfonts)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1F6FEB"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20"); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def source_line(segments):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run("资料来源："); set_font(r, 8, italic=True, color=GREY)
    for txt, url in segments:
        if url:
            add_hyperlink(p, url, txt)
        else:
            rr = p.add_run(txt); set_font(rr, 8, italic=True, color=GREY)


def bullet(segments, size=10.5, space_after=4):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(space_after)
    if isinstance(segments, str):
        segments = [(segments, False)]
    for seg in segments:
        if isinstance(seg, str):
            txt, bold = seg, False
        else:
            txt, bold = seg
        run = p.add_run(txt); set_font(run, size, bold=bold)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def cell(c, text, bold=False, color=None, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None, cn=SONG):
    c.text = ""
    p = c.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text); set_font(run, size, bold=bold, color=color, cn=cn)
    if fill: shade_cell(c, fill)


def add_image(name, width=5.7):
    path = os.path.join(OUT, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ===================================================== 第 1 页 ===============
para("股票研究  |  业绩更新", size=9, bold=True, color=BLUE, space_after=2, cn=HEI)
para("Nebius Group N.V.（纳斯达克：NBIS）", size=18, bold=True, color=NAVY, space_after=1, cn=HEI)
para("2026 财年第一季度业绩更新 —— “我们建多少，就卖多少”", size=12.5, bold=True, color=NAVY, space_after=4, cn=HEI)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
for txt, bold, col in [
    ("报告日期：2026 年 6 月 18 日　|　", False, GREY),
    ("评级：买入（BUY）　|　", True, GREEN),
    ("目标价：330 美元　|　", True, NAVY),
    ("现价（6/17/26）：280.91 美元　|　", False, GREY),
    ("市值：约 713 亿美元", False, GREY)]:
    r = p.add_run(txt); set_font(r, 9.5, bold=bold, color=col)

tbl = doc.add_table(rows=1, cols=4); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["指标", "2026Q1 实际", "市场预期", "超出/(不及)"]):
    cell(tbl.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
         align=WD_ALIGN_PARAGRAPH.CENTER, fill="0B2545", cn=HEI)
rows = [
    ("集团营收", "3.990 亿美元", "3.916 亿美元", "+740 万 / +1.9%  ✔"),
    ("核心 AI 云 ARR", "19.2 亿美元", "—", "环比 +54%"),
    ("集团调整后 EBITDA", "1.295 亿美元", "约 0.95 亿美元", "+3400 万 / 超预期 ✔"),
    ("集团调整后 EBITDA 利润率", "32%", "约 24%", "+8 个百分点 ✔"),
    ("调整后每股收益（稀释）", "(0.39) 美元", "(0.78) 美元", "亏损收窄 ✔"),
    ("GAAP 每股收益（稀释）", "2.11 美元", "无意义", "含 7.81 亿 ClickHouse 收益"),
]
for r0 in rows:
    cells = tbl.add_row().cells
    for i, val in enumerate(r0):
        cell(cells[i], val, bold=(i==0),
             align=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
for r0 in tbl.rows:
    for c in r0.cells:
        c.width = Inches(1.65)
source_line([("2026Q1 业绩新闻稿及 6-K 文件（2026 年 5 月 13 日发布），", None),
             ("Nebius 投资者关系", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             ("；市场预期采自 Visible Alpha / Investing.com，截至 2026 年 5 月 12 日。", None)])

heading("核心结论：一份炸裂的财报 —— 维持买入，目标价上调至 330 美元", size=13)
para("Nebius 本季度交出一份“既超预期、又在扩张”的优质答卷。集团营收 3.990 亿美元（同比 +684%、"
     "环比 +75%）超市场预期；核心 AI 云 ARR 环比飙升 +54% 至 19.2 亿美元；集团调整后 EBITDA 环比"
     "翻倍以上至 1.295 亿美元（利润率 32%），AI 云业务调整后 EBITDA 利润率达 45%。战略层面的重磅"
     "消息是与 Meta 签订的 270 亿美元、为期五年的协议，以及英伟达 20 亿美元的股权投资，确立了 Nebius "
     "作为一线“新云（neocloud）”厂商的地位。管理层将 2026 年资本开支指引上调至 200–250 亿美元"
     "（原 160–200 亿），以提前锁定 2027 年产能；同时重申全年 70–90 亿美元的退出期 ARR 与 30–34 亿"
     "美元营收。基于更高的远期 ARR 可见度与已大幅去风险的资本结构（93 亿美元现金），我们将目标价由 "
     "300 美元上调至 330 美元。", space_after=8)

heading("关键要点", size=12)
bullet([("需求不再是受限变量。", True),
        "“我们建多少，就卖多少”——AI 云销售管道环比增长 3.5 倍（不含超大规模客户），公司再次提价"
        "且各类芯片均已售罄。已签约电力达 3.5 GW 以上（原约 2 GW），2026 年底目标 >4 GW。"])
bullet([("利润率出现强劲拐点。", True),
        "核心 AI 云调整后 EBITDA 利润率由 2025Q4 的 24% 扩张至 45%；集团利润率达 32%。管理层指引 2026 "
        "全年集团调整后 EBITDA 利润率约 40%，节奏上呈下半年加权。"])
bullet([("资产负债表已成为武器。", True),
        "本季度融资 63 亿美元（43 亿可转债 + 英伟达 20 亿股权）；现金 93 亿美元，且当前资本开支计划 "
        ">90% 已由现金与合约锁定。经营性现金流转正至 +23 亿美元。"])
bullet([("产能提前布局。", True),
        "宾夕法尼亚州新增 1.2 GW 自有 AI 工厂（美国第二个吉瓦级自有站点），自有产能占比 >75%。"
        "资本开支上调至 200–250 亿美元用于 2027 年产能，相关收入将自 2027 上半年开始贡献。"])
bullet([("需关注事项。", True),
        "二季度利润率因产能部署后置而环比走低；超 330 亿美元的递延/合约在手订单集中于少数大单"
        "（Meta、微软）；非核心业务（Avride、TripleTen、Toloka）仍在消耗现金，正筹备引入合作方或分拆。"])

add_image("nbis_chart1_revenue.png")
source_line([("公司季度业绩新闻稿，2025Q1–2026Q1（", None),
             ("SEC EDGAR 6-K 文件", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K"),
             ("）。", None)])

# ===================================================== 第 2 页 ===============
doc.add_page_break()
heading("一、业绩明细 —— 全面超预期", size=14)

para("营收超预期，由 AI 云放量驱动", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
para("集团营收 3.990 亿美元，较市场预期 3.916 亿美元高出约 1.9%，环比增长 75%。Nebius AI 云营收 "
     "3.897 亿美元（同比 +841%、环比 +82%），现已占集团的 98%。超预期主要源于产能上线快于预期且即时"
     "变现——利用率基本处于售罄状态。核心基础设施业务的 ARR（年化运行营收）升至 19.2 亿美元，环比 "
     "+54%，对应每兆瓦活跃产能约 0.94 百万美元的资产产出效率。", space_after=6)
add_image("nbis_chart2_arr.png")
source_line([("2026Q1 业绩新闻稿，", None),
             ("Nebius 投资者关系（2026 年 5 月 13 日）", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             ("。", None)])

para("盈利能力拐点", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
para("集团调整后 EBITDA 为 1.295 亿美元（利润率 32%），高于 2025Q4 的 1500 万美元（7%）及去年同期的 "
     "(5370) 万美元亏损。核心 AI 云调整后 EBITDA 利润率由 2025Q4 的 24% 扩张至 45%，体现固定成本随营收"
     "放量被有效摊薄。按 GAAP 口径，持续经营净利润为 6.212 亿美元，但其中包含对 ClickHouse 股权重估"
     "产生的一次性 7.806 亿美元收益；按调整后口径，Nebius 净亏损 (1.003) 亿美元，即每股 (0.39) 美元，"
     "亏损远小于市场预期的 (0.78) 美元。", space_after=6)
add_image("nbis_chart3_ebitda.png")
source_line([("2026Q1 6-K 及业绩新闻稿；2025Q2/Q3 调整后 EBITDA 为趋势示意插值。", None),
             ("SEC EDGAR", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K"),
             ("。", None)])

para("超预期/不及预期一览", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
para("Nebius 在营收、调整后 EBITDA 与每股收益三方面均超预期。财报发布后股价上涨约 20%，既反映经营层面"
     "的超预期，也反映 Meta/英伟达带来的战略背书。", space_after=6)
add_image("nbis_chart4_beatmiss.png")
source_line([("实际值采自 2026Q1 新闻稿；市场预期采自 Investing.com / Visible Alpha，截至 2026 年 5 月 12 日。", None),
             ("Investing.com 电话会纪要", "https://www.investing.com/news/transcripts/earnings-call-transcript-nebius-group-q1-2026-earnings-beat-expectations-93CH-4684818"),
             ("。", None)])

# ===================================================== 第 3 页 ===============
doc.add_page_break()
heading("二、关键指标、产能与资本", size=14)

para("产能：签约领先于需求", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
para("Nebius 已签约电力 >3.5 GW（上季度约 2 GW），并将 2026 年底目标上调至 >4 GW。公司宣布在宾夕法尼亚"
     "州新建 1.2 GW 自有 AI 工厂——这是其在美国的第二个吉瓦级自有站点——自有产能占比 >75%，相较以租赁"
     "为主的新云同业构成结构性的成本与掌控优势。超过 100 MW 的数据中心站点增至 7 个（2025 年底为 1 个）。"
     "宾州设施将于 2027 年底前上线 250–300 MW，到 2030 年扩至完整的 1.2 GW；阿拉巴马与密苏里项目将自 "
     "2027 年初开始贡献。", space_after=6)
add_image("nbis_chart5_power.png", width=5.6)
source_line([("2026Q1 致股东信及业绩电话会（2026 年 5 月 13 日）。", None),
             ("Nebius 投资者关系", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             ("。", None)])

para("资本：93 亿美元现金，资本开支 >90% 已落实", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
para("本季度 Nebius 融资超 60 亿美元——约 43 亿美元可转换优先票据（票息 1.25%–2.60%）及英伟达 20 亿美元"
     "股权投资——季末持有现金 93 亿美元，对应非流动负债 84 亿美元。经营性现金流强劲转正至 +23 亿美元。"
     "管理层表示，当前计划资本开支的 >90% 已由现金与合约承诺锁定，显著降低了 2026–27 年建设的资金风险。", space_after=6)
add_image("nbis_chart6_liquidity.png", width=5.9)
source_line([("2026Q1 6-K，现金流量与融资报表。", None),
             ("SEC EDGAR 6-K（2026 年 5 月 13 日）", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K"),
             ("。", None)])

para("Meta 与英伟达两大战略锚点", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
bullet([("Meta —— 270 亿美元 / 5 年：", True),
        "其中 120 亿美元为专用算力承诺，另有 150 亿美元可选产能，Nebius 可按市场价转售给 AI 云客户，"
        "并由 Meta 兜底——构成资产支持融资的顺风因素。"])
bullet([("英伟达 —— 20 亿股权 + “优先建设方”地位：", True),
        "在 GB300 上获得训练 Exemplar Cloud 认证，并对未来 Vera Rubin / CPU 平台拥有差异化获取权，"
        "强化供给优先级。"])
bullet([("补强型并购：", True),
        "收购 Tavily、Eigen AI 与 Clarifai，以加速推理 / 智能体路线图（Token Factory）——重在人才与技术，"
        "而非独立营收。"])

# ===================================================== 第 4 页 ===============
doc.add_page_break()
heading("三、指引与更新后预测", size=14)

para("重申 2026 全年指引；上调资本开支", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
para("管理层重申 2026 全年目标：退出期 ARR 70–90 亿美元、集团营收 30–34 亿美元、集团调整后 EBITDA 利润率"
     "约 40%，同时将 2026 年资本开支由 160–200 亿上调至 200–250 亿美元。此次上调用于增量 2027 年产能，"
     "相关收入将自 2027 上半年开始贡献——即针对已签约需求的增长性资本开支，而非投机性投入。受产能后置部署"
     "影响，二季度利润率指引走低，三季度回到一季度水平，四季度更高。", space_after=6)
add_image("nbis_chart7_capex.png", width=4.4)
source_line([("2026Q1 业绩电话会与致股东信（2026 年 5 月 13 日）。", None)])
add_image("nbis_chart8_guidance.png")
source_line([("2026Q1 业绩电话会，管理层指引（重申）。", None),
             ("The Motley Fool 电话会纪要", "https://www.fool.com/earnings/call-transcripts/2026/05/13/nebius-nbis-q1-2026-earnings-transcript/"),
             ("。", None)])

para("更新后预测（旧 vs 新）", size=11.5, bold=True, color=BLUE, space_before=4, space_after=3, cn=HEI)
est = doc.add_table(rows=1, cols=5); est.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["预测项", "旧值", "新值", "变动", "驱动因素"]):
    cell(est.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
         align=WD_ALIGN_PARAGRAPH.CENTER, fill="0B2545", cn=HEI)
est_rows = [
    ("2026E 营收", "31 亿", "33 亿", "+6%", "产能上线更快；利用率售罄"),
    ("2026E 退出期 ARR", "75 亿", "85 亿", "+13%", "管道环比 +3.5 倍；提价"),
    ("2026E 集团调整后 EBITDA 利润率", "约 35%", "约 40%", "+5 个百分点", "核心利润率拐点至 45%"),
    ("2027E 营收", "85 亿", "105 亿", "+24%", "提前布局的 2027 产能（宾/阿/密）"),
    ("2026E 资本开支", "180 亿", "225 亿", "+25%", "针对已签约需求的增长性开支"),
]
for r0 in est_rows:
    cells = est.add_row().cells
    for i, val in enumerate(r0):
        cell(cells[i], val, bold=(i==0), size=9,
             align=WD_ALIGN_PARAGRAPH.LEFT if i in (0,4) else WD_ALIGN_PARAGRAPH.CENTER)
source_line([("本机构预测（示意性），锚定于 2026Q1 新闻稿中公司重申的指引。", None)])

# ===================================================== 第 5 页 ===============
doc.add_page_break()
heading("四、投资逻辑更新", size=14)

para("逻辑完好且得到强化", size=11.5, bold=True, color=BLUE, space_after=3, cn=HEI)
para("我们的核心逻辑——Nebius 是少数能够将稀缺 GPU 供给、自有电力与全栈软件转化为可持续、高利润 AI 基础"
     "设施收入的独立“新云”平台之一——本季度得到实质性强化。三大支柱均改善：", space_after=6)
bullet([("(1) 需求可见度：", True),
        "ARR 环比 +54% 至 19.2 亿美元，管道环比扩张 3.5 倍，确认需求超过供给；270 亿美元的 Meta 锚点"
        "为多年期产能提供承销保障。"])
bullet([("(2) 单位经济性：", True),
        "45% 的核心 AI 云利润率与每兆瓦约 0.94 百万美元的 ARR 表明，随着产能填满，模型可实现盈利性扩张。"])
bullet([("(3) 融资风险下降：", True),
        "93 亿美元现金与英伟达股权投资消除了困扰以租赁为主的竞争对手的融资悬顶。"])
add_image("nbis_chart9_coremargin.png", width=5.3)
source_line([("2026Q1 业绩新闻稿；核心 AI 云分部利润率。", None),
             ("Nebius 投资者关系", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             ("。", None)])

para("需关注的风险", size=11.5, bold=True, color=RED, space_after=3, cn=HEI)
bullet([("客户集中度：", True), "大单（Meta、微软）主导在手订单；续约/放量节奏至关重要。"])
bullet([("资本开支强度与执行：", True),
        "200–250 亿美元的 2026 资本开支要求建设、电力与 GPU 交付零失误；成本通胀（低个位数 %）目前"
        "因提前采购而可控。"])
bullet([("算力定价/供给正常化：", True), "定价权建立在 GPU 持续稀缺之上；供给过剩将压缩利润率。"])
bullet([("非核心拖累：", True), "Avride、TripleTen 与 Toloka 仍亏损；管理层正寻求合作方/分拆。"])
add_image("nbis_chart10_mix.png", width=4.2)
source_line([("2026Q1 分业务营收；Nebius AI 云占集团 98%。", None),
             ("2026Q1 新闻稿", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             ("。", None)])

# ===================================================== 第 6 页 ===============
doc.add_page_break()
heading("五、估值与建议", size=14)
para("评级：买入（维持）  |  目标价：330 美元（自 300 美元上调）", size=12, bold=True, color=GREEN, space_after=6, cn=HEI)
para("鉴于其高增长、尚未进入稳态盈利的特征，我们采用 EV/ARR 与前瞻 EV/营收框架对 Nebius 估值。基于我们"
     "上调后的约 85 亿美元 2026 退出期 ARR，该股交易于约 8–9 倍 EV/退出期 ARR——相对传统 IaaS 有溢价，"
     "但相对其增速（ARR 环比 +54%）与改善的利润率而言合理。对去风险后的 2026 ARR 给予约 10 倍 EV/退出期 "
     "ARR，扣除 93 亿美元现金与可转债后，支撑 330 美元目标价（较 280.91 美元约 17% 上行空间）。乐观情形"
     "（退出期 ARR >90 亿、综合利润率 45%+、Meta 可选产能全部转售）支撑 400 美元以上；悲观情形（供给正常化、"
     "资本开支超支、客户放量滞后）对应约 190 美元。", space_after=8)

sc = doc.add_table(rows=1, cols=4); sc.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["情形", "2026 退出期 ARR", "EV/ARR", "隐含价值"]):
    cell(sc.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
         align=WD_ALIGN_PARAGRAPH.CENTER, fill="0B2545", cn=HEI)
sc_rows = [
    ("乐观", "95 亿美元", "约 12 倍", "400 美元以上"),
    ("基准（目标价）", "85 亿美元", "约 10 倍", "330 美元"),
    ("悲观", "70 亿美元", "约 6 倍", "约 190 美元"),
]
for r0 in sc_rows:
    cells = sc.add_row().cells
    fill = "E8F5E9" if r0[0].startswith("基准") else None
    for i, val in enumerate(r0):
        cell(cells[i], val, bold=(i==0),
             align=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
source_line([("本机构估值框架；市场数据经 Yahoo Finance（yfinance）获取，价格截至 2026 年 6 月 17 日。", None)])

para("催化剂（未来 6–12 个月）", size=11.5, bold=True, color=BLUE, space_before=6, space_after=3, cn=HEI)
bullet("2026Q2 财报（2026 年 8 月中）：ARR 向 70–90 亿退出期目标的进展；利润率节奏确认。")
bullet("除 Meta、微软之外的更多超大规模/前沿实验室产能合约。")
bullet("宾州/阿拉巴马/密苏里建设里程碑与电力通电。")
bullet("Avride / TripleTen / Toloka 的潜在变现（引入合作方或分拆）。")

# 资料来源
doc.add_page_break()
heading("资料来源与参考文献", size=14)
para("业绩材料（2026 年第一季度 —— 2026 年 5 月 13 日发布）：", size=11, bold=True, space_after=4, cn=HEI)

def ref(text, url):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
    add_hyperlink(p, url, text)

ref("2026Q1 业绩新闻稿 —— Nebius Newsroom（2026 年 5 月 13 日）",
    "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results")
ref("2026Q1 业绩 —— BusinessWire 发布（2026 年 5 月 13 日）",
    "https://www.businesswire.com/news/home/20260513568820/en/Nebius-reports-first-quarter-2026-financial-results")
ref("Form 6-K —— Nebius Group N.V.，SEC EDGAR（CIK 0001513845）",
    "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K")
ref("2026Q1 业绩电话会纪要 —— The Motley Fool（2026 年 5 月 13 日）",
    "https://www.fool.com/earnings/call-transcripts/2026/05/13/nebius-nbis-q1-2026-earnings-transcript/")
ref("2026Q1 业绩电话会纪要 —— Investing.com（2026 年 5 月 13 日）",
    "https://www.investing.com/news/transcripts/earnings-call-transcript-nebius-group-q1-2026-earnings-beat-expectations-93CH-4684818")
ref("Nebius 投资者中心 / 财务",
    "https://nebius.com/financials")
ref("2026Q1 分析 —— Seeking Alpha",
    "https://seekingalpha.com/article/4907267-nebius-breaking-down-nebiuss-q1-earnings")

para("", space_after=4)
para("免责声明：本业绩更新仅供信息参考，不构成投资建议或任何证券的买卖要约。文中预测为示意性数据。"
     "市场数据经 Yahoo Finance（yfinance）于 2026 年 6 月 17 日获取。", size=8, italic=True, color=GREY)

out_path = os.path.join(OUT, "NBIS_Q1_FY2026_业绩更新报告_中文版.docx")
doc.save(out_path)
print("Saved", out_path)
