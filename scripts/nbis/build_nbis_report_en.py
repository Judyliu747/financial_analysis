#!/usr/bin/env python3
"""Nebius Group (NBIS) Q1 2026 Earnings Update — English DOCX.
Output: output/NBIS/NBIS_Q1_FY2026_Earnings_Update.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/NBIS"
FONT = "Times New Roman"

NAVY  = RGBColor(0x0B, 0x25, 0x45)
BLUE  = RGBColor(0x1F, 0x6F, 0xEB)
GREEN = RGBColor(0x2B, 0xA8, 0x4A)
RED   = RGBColor(0xD1, 0x49, 0x5B)
GREY  = RGBColor(0x66, 0x66, 0x66)

doc = Document()
for s in doc.styles:
    try:
        s.font.name = FONT
    except Exception:
        pass
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)

sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)


def set_font(run, size=10.5, bold=False, italic=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font); rfonts.set(qn("w:hAnsi"), font)


def para(text="", size=10.5, bold=False, italic=False, color=None,
         align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        set_font(run, size, bold, italic, color)
    return p


def heading(text, size=14, color=NAVY, space_before=10):
    p = para(text, size=size, bold=True, color=color, space_after=4, space_before=space_before)
    return p


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT); rfonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rfonts)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1F6FEB"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20"); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def source_line(paragraph_runs):
    """paragraph_runs: list of ('text', url_or_None)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Source: "); set_font(r, 8, italic=True, color=GREY)
    for txt, url in paragraph_runs:
        if url:
            add_hyperlink(p, url, txt)
        else:
            rr = p.add_run(txt); set_font(rr, 8, italic=True, color=GREY)
    return p


def bullet(text_segments, size=10.5, space_after=4):
    """text_segments: list of (text, bold) tuples, or a plain string."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(space_after)
    if isinstance(text_segments, str):
        text_segments = [(text_segments, False)]
    for seg in text_segments:
        if isinstance(seg, str):
            txt, bold = seg, False
        else:
            txt, bold = seg
        run = p.add_run(txt)
        set_font(run, size, bold=bold)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def style_table_cell(cell, text, bold=False, color=None, size=9.5,
                     align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    set_font(run, size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)


def add_image(name, width=6.6, caption=None):
    path = os.path.join(OUT, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================ PAGE 1 ==========
p = para("EQUITY RESEARCH  |  EARNINGS UPDATE", size=9, bold=True, color=BLUE, space_after=2)
para("Nebius Group N.V. (NASDAQ: NBIS)", size=18, bold=True, color=NAVY, space_after=1)
para("Q1 2026 Earnings Update — “Everything We Build, We Sell”", size=12.5, bold=True, color=NAVY, space_after=4)

# meta line
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
for txt, bold, col in [
    ("Report date: June 18, 2026   |   ", False, GREY),
    ("Rating: BUY   |   ", True, GREEN),
    ("Price Target: $330   |   ", True, NAVY),
    ("Price (6/17/26): $280.91   |   ", False, GREY),
    ("Mkt Cap: ~$71B", False, GREY)]:
    r = p.add_run(txt); set_font(r, 9.5, bold=bold, color=col)

# Summary box table
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ["Metric", "Q1 2026A", "Consensus", "Beat / (Miss)"]
for i, h in enumerate(hdr):
    style_table_cell(tbl.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                     align=WD_ALIGN_PARAGRAPH.CENTER, fill="0B2545")
rows = [
    ("Group revenue", "$399.0M", "$391.6M", "+$7.4M / +1.9%  ✔"),
    ("Core AI Cloud ARR", "$1.92B", "—", "+54% QoQ"),
    ("Group adj. EBITDA", "$129.5M", "~$95M", "+$34M / Beat ✔"),
    ("Group adj. EBITDA margin", "32%", "~24%", "+8 pts ✔"),
    ("Adj. EPS (diluted)", "$(0.39)", "$(0.78)", "Narrower loss ✔"),
    ("GAAP EPS (diluted)", "$2.11", "n/m", "Incl. $781M ClickHouse gain"),
]
for r0 in rows:
    cells = tbl.add_row().cells
    for i, val in enumerate(r0):
        style_table_cell(cells[i], val, bold=(i==0),
                         align=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
for r0 in tbl.rows:
    for c in r0.cells:
        c.width = Inches(1.65)
source_line([("Q1 2026 results press release & 6-K (filed May 13, 2026), ", None),
             ("Nebius IR", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             ("; consensus per Visible Alpha / Investing.com as of May 12, 2026.", None)])

heading("Bottom Line: A Blowout Quarter — Reiterate BUY, Raise PT to $330", size=13)
para("Nebius delivered a clean beat-and-build quarter. Group revenue of $399M (+684% YoY, "
     "+75% QoQ) topped consensus, core AI Cloud ARR vaulted +54% QoQ to $1.92B, and group "
     "adjusted EBITDA more than doubled QoQ to $129.5M (32% margin) as the AI cloud unit hit a "
     "45% adjusted EBITDA margin. The strategic headline was a $27B, five-year Meta agreement plus "
     "a $2B NVIDIA equity investment, validating Nebius as a tier-one neocloud. Management raised "
     "2026 capex to $20–25B (from $16–20B) to pull forward 2027 capacity, while reiterating "
     "$7–9B exit-ARR and $3.0–3.4B revenue. We raise our PT to $330 (from $300) on higher "
     "out-year ARR visibility and a de-risked capital structure ($9.3B cash).", space_after=8)

heading("Key Takeaways", size=12)
bullet([("Demand is the bottleneck-free variable. ", True),
        ("“Everything we build, we sell” — AI cloud pipeline grew 3.5x QoQ (ex-hyperscalers), "
         "Nebius raised prices again and is sold out across all chip types. Contracted power reached "
         ">3.5 GW (from ~2 GW), with a YE-2026 target of >4 GW.")])
bullet([("Margins inflected hard. ", True),
        ("Core AI cloud adj. EBITDA margin expanded to 45% (from 24% in Q4'25); group margin hit 32%. "
         "Management guides ~40% group adj. EBITDA margin for FY2026, with a back-half-weighted cadence.")])
bullet([("Balance sheet is now a weapon. ", True),
        ("$6.3B raised in Q1 ($4.3B converts + $2B NVIDIA equity); $9.3B cash and >90% of current "
         "capex secured by cash/contracts. Operating cash flow swung to +$2.3B.")])
bullet([("Capacity pulled forward. ", True),
        ("New 1.2 GW owned AI factory in Pennsylvania (2nd US GW-scale site); >75% of capacity owned. "
         "Capex raised to $20–25B funds 2027 supply that begins contributing revenue in H1 2027.")])
bullet([("Watch items. ", True),
        ("Q2 margins step down on back-loaded deployment; ~$33B+ in deferred/contracted backlog is "
         "concentrated in a few mega-deals (Meta, Microsoft); non-core units (Avride, TripleTen, Toloka) "
         "remain cash-consumptive and are being prepped for partners/spin.")])

add_image("nbis_chart1_revenue.png", width=5.7)
source_line([("Company quarterly press releases, Q1'25–Q1'26 (", None),
             ("SEC EDGAR 6-K filings", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K"),
             (").", None)])

# ============================================================ PAGE 2 ==========
doc.add_page_break()
heading("1.  Results Detail — Beat Across the Board", size=14)

para("Revenue Beat Driven by AI Cloud Ramp", size=11.5, bold=True, color=BLUE, space_after=3)
para("Group revenue of $399.0M beat consensus of $391.6M by ~1.9% and grew 75% sequentially. "
     "Nebius AI Cloud revenue was $389.7M (+841% YoY, +82% QoQ), now 98% of the group. The "
     "outperformance was driven by faster-than-expected capacity coming online and immediate "
     "monetization — utilization remains effectively sold-out. ARR (annualized run-rate of the core "
     "infrastructure business) rose to $1.92B, +54% QoQ, implying asset productivity of ~$0.94M of "
     "ARR per active MW.", space_after=6)

add_image("nbis_chart2_arr.png", width=5.7)
source_line([("Q1 2026 results press release, ", None),
             ("Nebius IR (May 13, 2026)", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             (".", None)])

para("Profitability Inflection", size=11.5, bold=True, color=BLUE, space_after=3)
para("Group adjusted EBITDA was $129.5M (32% margin), up from $15.0M (7%) in Q4'25 and a "
     "$(53.7)M loss a year ago. The core AI cloud adjusted EBITDA margin expanded to 45% (from 24% "
     "in Q4'25) as fixed costs leveraged against the revenue ramp. Reported GAAP net income from "
     "continuing operations was $621.2M, but this is flattered by a one-time $780.6M gain on the "
     "revaluation of the ClickHouse equity stake; on an adjusted basis Nebius posted a net loss of "
     "$(100.3)M, or $(0.39) per share — a far narrower loss than the $(0.78) consensus.", space_after=6)

add_image("nbis_chart3_ebitda.png", width=5.7)
source_line([("Q1 2026 6-K and results release; Q2'25/Q3'25 adj. EBITDA interpolated for trend. ", None),
             ("SEC EDGAR", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K"),
             (".", None)])

para("Beat / Miss Summary", size=11.5, bold=True, color=BLUE, space_after=3)
para("Nebius beat on revenue, adjusted EBITDA, and EPS. The stock rallied ~20% on the print, "
     "reflecting both the operational beat and the strategic Meta/NVIDIA validation.", space_after=6)
add_image("nbis_chart4_beatmiss.png", width=5.7)
source_line([("Actuals per Q1 2026 release; consensus per Investing.com / Visible Alpha as of May 12, 2026. ", None),
             ("Investing.com transcript", "https://www.investing.com/news/transcripts/earnings-call-transcript-nebius-group-q1-2026-earnings-beat-expectations-93CH-4684818"),
             (".", None)])

# ============================================================ PAGE 3 ==========
doc.add_page_break()
heading("2.  Key Metrics, Capacity & Capital", size=14)

para("Capacity: Contracting Ahead of Demand", size=11.5, bold=True, color=BLUE, space_after=3)
para("Nebius contracted >3.5 GW of power (up from ~2 GW the prior quarter) and raised its YE-2026 "
     "target to >4 GW. It announced a new 1.2 GW owned AI factory in Pennsylvania — its second US "
     "gigawatt-scale owned site — and now owns >75% of its capacity stack, a structural cost and "
     "control advantage versus lease-heavy neocloud peers. Data-center sites above 100 MW grew to 7 "
     "(from 1 at YE-2025). The Pennsylvania facility brings 250–300 MW online by end-2027, scaling to "
     "the full 1.2 GW by 2030; Alabama and Missouri projects begin contributing in early 2027.", space_after=6)
add_image("nbis_chart5_power.png", width=5.6)
source_line([("Q1 2026 CEO letter to shareholders & earnings call (May 13, 2026). ", None),
             ("Nebius IR", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             (".", None)])

para("Capital: $9.3B Cash, >90% of Capex Funded", size=11.5, bold=True, color=BLUE, space_after=3)
para("Nebius raised over $6B in Q1 — ~$4.3B of convertible senior notes (1.25%–2.60% coupons) and a "
     "$2.0B equity investment from NVIDIA — closing the quarter with $9.3B cash against $8.4B "
     "non-current debt. Operating cash flow turned strongly positive at +$2.3B. Management states "
     ">90% of current-plan capex is already secured by cash and contractual commitments, materially "
     "de-risking the funding of the 2026–27 build-out.", space_after=6)
add_image("nbis_chart6_liquidity.png", width=5.9)
source_line([("Q1 2026 6-K, cash flow & financing statements. ", None),
             ("SEC EDGAR 6-K (May 13, 2026)", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K"),
             (".", None)])

para("The Meta & NVIDIA Strategic Anchors", size=11.5, bold=True, color=BLUE, space_after=3)
bullet([("Meta — $27B / 5 years: ", True),
        ("$12B of dedicated compute commitment plus $15B of optional capacity Nebius can resell to AI "
         "cloud customers at market rates, with Meta as backstop — an asset-backed financing tailwind.")])
bullet([("NVIDIA — $2B equity + preferred-builder status: ", True),
        ("Exemplar Cloud status on GB300 for training and differentiated access to future Vera Rubin / "
         "CPU platforms reinforces supply priority.")])
bullet([("Tuck-in M&A: ", True),
        ("Tavily, Eigen AI and Clarifai acquired to accelerate the inference / agent roadmap (Token "
         "Factory) — talent and tech, not standalone revenue.")])

# ============================================================ PAGE 4 ==========
doc.add_page_break()
heading("3.  Guidance & Updated Estimates", size=14)

para("FY2026 Guidance Reiterated; Capex Raised", size=11.5, bold=True, color=BLUE, space_after=3)
para("Management reiterated full-year 2026 targets of $7–9B exit-ARR, $3.0–3.4B group revenue, and "
     "~40% group adjusted EBITDA margin, while raising 2026 capex to $20–25B (from $16–20B). The "
     "capex raise funds additional 2027 capacity that should begin contributing revenue in H1 2027 — "
     "i.e., it is growth capex against contracted demand, not speculative. Q2 margins are guided "
     "lower on back-end-weighted capacity deployment, recovering to Q1 levels in Q3 and higher in Q4.", space_after=6)

add_image("nbis_chart7_capex.png", width=4.4)
source_line([("Q1 2026 earnings call & CEO letter (May 13, 2026).", None)])
add_image("nbis_chart8_guidance.png", width=5.7)
source_line([("Q1 2026 earnings call, management guidance (reiterated). ", None),
             ("The Motley Fool transcript", "https://www.fool.com/earnings/call-transcripts/2026/05/13/nebius-nbis-q1-2026-earnings-transcript/"),
             (".", None)])

para("Updated Estimates (Old vs. New)", size=11.5, bold=True, color=BLUE, space_before=4, space_after=3)
est = doc.add_table(rows=1, cols=5); est.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Estimate", "Old", "New", "Change", "Driver"]):
    style_table_cell(est.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                     align=WD_ALIGN_PARAGRAPH.CENTER, fill="0B2545")
est_rows = [
    ("FY26E revenue", "$3.1B", "$3.3B", "+6%", "Faster capacity online; sold-out utilization"),
    ("FY26E exit-ARR", "$7.5B", "$8.5B", "+13%", "Pipeline +3.5x QoQ; price increases"),
    ("FY26E grp adj. EBITDA mgn", "~35%", "~40%", "+5 pts", "Core margin inflection to 45%"),
    ("FY27E revenue", "$8.5B", "$10.5B", "+24%", "Pulled-forward 2027 capacity (PA/AL/MO)"),
    ("FY26E capex", "$18B", "$22.5B", "+25%", "Growth capex vs. contracted demand"),
]
for r0 in est_rows:
    cells = est.add_row().cells
    for i, val in enumerate(r0):
        style_table_cell(cells[i], val, bold=(i==0), size=9,
                         align=WD_ALIGN_PARAGRAPH.LEFT if i in (0,4) else WD_ALIGN_PARAGRAPH.CENTER)
source_line([("Firm estimates (illustrative), anchored to reiterated company guidance from the Q1 2026 release.", None)])

# ============================================================ PAGE 5 ==========
doc.add_page_break()
heading("4.  Updated Investment Thesis", size=14)

para("Thesis Intact and Strengthened", size=11.5, bold=True, color=BLUE, space_after=3)
para("Our thesis — that Nebius is among the few independent “neocloud” platforms able to convert "
     "scarce GPU supply, owned power, and full-stack software into durable, high-margin AI "
     "infrastructure revenue — was materially strengthened this quarter. Three pillars improved:", space_after=6)
bullet([("(1) Demand visibility: ", True),
        ("ARR +54% QoQ to $1.92B and a 3.5x QoQ pipeline expansion confirm demand outpaces supply; "
         "the $27B Meta anchor underwrites multi-year capacity.")])
bullet([("(2) Unit economics: ", True),
        ("45% core AI cloud margin and ~$0.94M ARR/MW demonstrate the model scales profitably as "
         "capacity fills.")])
bullet([("(3) Funding risk reduced: ", True),
        ("$9.3B cash and NVIDIA's equity stake remove the financing overhang that pressures lease-heavy "
         "competitors.")])
add_image("nbis_chart9_coremargin.png", width=5.3)
source_line([("Q1 2026 results release; core AI cloud segment margins. ", None),
             ("Nebius IR", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             (".", None)])

para("Risks to Monitor", size=11.5, bold=True, color=RED, space_after=3)
bullet([("Customer concentration: ", True),
        ("Mega-deals (Meta, Microsoft) dominate backlog; renewal/ramp timing matters.")])
bullet([("Capex intensity & execution: ", True),
        ("$20–25B 2026 capex requires flawless construction, power and GPU delivery; cost inflation "
         "(low-single-digit %) is so far contained by early procurement.")])
bullet([("Compute pricing / supply normalization: ", True),
        ("Pricing power assumes continued GPU scarcity; a supply glut would compress margins.")])
bullet([("Non-core drag: ", True),
        ("Avride, TripleTen and Toloka remain loss-making; management seeks partners/spin-offs.")])

add_image("nbis_chart10_mix.png", width=4.2)
source_line([("Q1 2026 revenue by business; Nebius AI Cloud = 98% of group. ", None),
             ("Q1 2026 release", "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results"),
             (".", None)])

# ============================================================ PAGE 6 ==========
doc.add_page_break()
heading("5.  Valuation & Recommendation", size=14)
para("Rating: BUY (maintained)  |  Price Target: $330 (raised from $300)", size=12, bold=True, color=GREEN, space_after=6)
para("We value Nebius on EV/ARR and a forward EV/revenue framework given the hyper-growth, "
     "pre-steady-state profile. On our raised ~$8.5B FY26 exit-ARR, the stock trades at roughly "
     "8–9x EV/exit-ARR — a premium to legacy IaaS but reasonable versus the growth rate (ARR +54% "
     "QoQ) and improving margins. Applying ~10x EV/exit-ARR on de-risked 2026 ARR, net of the $9.3B "
     "cash and convertible debt, supports a $330 target (~17% upside from $280.91). Bull case "
     "(>$9B exit-ARR, 45%+ blended margins, Meta optional capacity fully resold) supports $400+; bear "
     "case (supply normalization, capex overrun, customer ramp slippage) implies ~$190.", space_after=8)

# Scenario table
sc = doc.add_table(rows=1, cols=4); sc.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Scenario", "FY26 exit-ARR", "EV/ARR", "Implied value"]):
    style_table_cell(sc.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                     align=WD_ALIGN_PARAGRAPH.CENTER, fill="0B2545")
sc_rows = [
    ("Bull", "$9.5B", "~12x", "$400+"),
    ("Base (PT)", "$8.5B", "~10x", "$330"),
    ("Bear", "$7.0B", "~6x", "~$190"),
]
for r0 in sc_rows:
    cells = sc.add_row().cells
    fill = "E8F5E9" if r0[0]=="Base (PT)" else None
    for i, val in enumerate(r0):
        style_table_cell(cells[i], val, bold=(i==0),
                         align=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER,
                         fill=fill)
source_line([("Firm valuation framework; market data via Yahoo Finance (yfinance), price as of June 17, 2026.", None)])

para("Catalysts (next 6–12 months)", size=11.5, bold=True, color=BLUE, space_before=6, space_after=3)
bullet("Q2 2026 print (mid-Aug 2026): ARR trajectory toward $7–9B exit; margin cadence confirmation.")
bullet("Additional hyperscaler / frontier-lab capacity deals beyond Meta & Microsoft.")
bullet("Pennsylvania / Alabama / Missouri construction milestones and power energization.")
bullet("Potential monetization (partner or spin) of Avride / TripleTen / Toloka.")

# Sources section
doc.add_page_break()
heading("Sources & References", size=14)
para("Earnings Materials (Q1 2026 — reported May 13, 2026):", size=11, bold=True, space_after=4)

def ref(text, url):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
    add_hyperlink(p, url, text)

ref("Q1 2026 Results Press Release — Nebius Newsroom (May 13, 2026)",
    "https://nebius.com/newsroom/nebius-reports-first-quarter-2026-financial-results")
ref("Q1 2026 Results — BusinessWire distribution (May 13, 2026)",
    "https://www.businesswire.com/news/home/20260513568820/en/Nebius-reports-first-quarter-2026-financial-results")
ref("Form 6-K — Nebius Group N.V., SEC EDGAR (CIK 0001513845)",
    "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001513845&type=6-K")
ref("Q1 2026 Earnings Call Transcript — The Motley Fool (May 13, 2026)",
    "https://www.fool.com/earnings/call-transcripts/2026/05/13/nebius-nbis-q1-2026-earnings-transcript/")
ref("Q1 2026 Earnings Call Transcript — Investing.com (May 13, 2026)",
    "https://www.investing.com/news/transcripts/earnings-call-transcript-nebius-group-q1-2026-earnings-beat-expectations-93CH-4684818")
ref("Nebius Investor Hub / Financials",
    "https://nebius.com/financials")
ref("Q1 2026 Coverage — Seeking Alpha",
    "https://seekingalpha.com/article/4907267-nebius-breaking-down-nebiuss-q1-earnings")

para("", space_after=4)
para("Disclaimer: This earnings update is prepared for informational purposes only and does not "
     "constitute investment advice or an offer to buy or sell any security. Estimates are illustrative. "
     "Market data retrieved via Yahoo Finance (yfinance) as of June 17, 2026.", size=8, italic=True, color=GREY)

out_path = os.path.join(OUT, "NBIS_Q1_FY2026_Earnings_Update.docx")
doc.save(out_path)
print("Saved", out_path)
