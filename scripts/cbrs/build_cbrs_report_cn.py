#!/usr/bin/env python3
"""
Cerebras Systems (CBRS) IPO 暨 FY2025 财务分析报告 — 中文版 DOCX
机构研报格式：约10页，正文宋体，标题黑体
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, yfinance as yf

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/CBRS/"
CHARTS = OUT

# ── 市场数据 ─────────────────────────────────────────────────────────────────
def get_market_data(ticker):
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"警告：无法获取 {ticker} 市场数据：{e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("CBRS")
price = mkt["price"]
mcap = mkt["market_cap"]
mcap_b = round(mcap / 1e9, 1) if isinstance(mcap, (int, float)) else "N/A"

# ── 颜色 ─────────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
BLUE  = RGBColor(0x00, 0x66, 0xCC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x6B, 0x3A)
RED   = RGBColor(0xCC, 0x00, 0x00)

# ── 工具函数 ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink'); rPr.append(rStyle)
    color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '0066CC'); rPr.append(color_el)
    u_el = OxmlElement('w:u'); u_el.set(qn('w:val'), 'single'); rPr.append(u_el)
    sz_el = OxmlElement('w:sz'); sz_el.set(qn('w:val'), '20'); rPr.append(sz_el)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def heading(doc, text, level=1, color=NAVY, size=None, bold=True, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.bold = bold; run.font.color.rgb = color
    if size is None: size = {1:16, 2:13, 3:11}.get(level, 11)
    run.font.size = Pt(size)
    return p

def body(doc, text, size=10, bold=False, color=DGRAY, space_before=2, space_after=2, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic; run.font.color.rgb = color
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size); run.font.color.rgb = DGRAY
    return p

def add_chart(doc, fname, width=6.0, caption=None):
    path = CHARTS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2); cp.paragraph_format.space_after = Pt(6)
        for run in cp.runs:
            run.font.size = Pt(8); run.font.italic = True; run.font.color.rgb = DGRAY

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    run._r.append(br)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '003366')
    pBdr.append(bottom); pPr.append(pBdr)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""; p = cell.paragraphs[0]
        run = p.add_run(h); run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, '003366')
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = ""; p = cell.paragraphs[0]
            run = p.add_run(str(val)); run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9); run.font.color.rgb = DGRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i % 2 == 1: set_cell_bg(cell, 'F0F0F0')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# 生成报告
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
for section in doc.sections:
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# ── 第1页：封面与摘要 ────────────────────────────────────────────────────────
heading(doc, "CEREBRAS SYSTEMS, INC.（赛瑞布斯系统）", level=1, size=20, space_before=20, space_after=2)
heading(doc, "IPO 暨 FY2025 年度财务分析报告", level=1, size=14, color=BLUE, bold=False, space_before=2, space_after=2)
body(doc, f"纳斯达克代码：CBRS  |  当前股价：${price}  |  市值：${mcap_b}B  |  行业：AI基础设施/半导体",
     size=10, bold=True, color=NAVY, space_before=4, space_after=2)
body(doc, f"IPO日期：2026年5月14日  |  发行价：$185  |  首日收盘价：$311.07（+68.2%）  |  员工数：约708人",
     size=10, bold=False, color=DGRAY, space_before=2, space_after=6)

divider(doc)

heading(doc, "核心要点", level=2, space_before=8, space_after=4)
bullet(doc, "FY2025营收达5.10亿美元（同比+76%），由硬件系统（3.58亿美元）和云/推理服务（1.52亿美元）双轮驱动。")
bullet(doc, "GAAP净利润2.38亿美元，包含G42远期合约注销产生的一次性非现金收益3.63亿美元；调整后净亏损7,570万美元。")
bullet(doc, "FY2025毛利率39.0%，较FY2024（42.3%）下降330个基点，主因硬件产品占比上升；云服务业务具备更高毛利率潜力。")
bullet(doc, "签署重大合约：OpenAI 200亿美元以上多年期协议（750MW推理算力）；AWS多年期合作伙伴关系（含2.70亿美元股权投资）。")
bullet(doc, "客户集中度仍为核心风险：FY2025营收的86%来自阿联酋关联实体（MBZUAI占62%，G42占24%）。")
bullet(doc, "IPO以每股$185定价募资55.5亿美元——为Uber（2019年）以来最大美国科技IPO。首日以$350开盘，$311收盘。")

divider(doc)

heading(doc, "FY2025 业绩概览", level=2, space_before=8, space_after=4)
make_table(doc,
    ["指标", "FY2025", "FY2024", "同比变化"],
    [
        ["营收",               "$5.10亿",     "$2.90亿",    "+75.7%"],
        ["毛利润",             "$1.99亿",     "$1.23亿",    "+62.2%"],
        ["毛利率",             "39.0%",       "42.3%",      "-330 bps"],
        ["研发费用",           "$2.43亿",     "$1.58亿",    "+53.8%"],
        ["销售及管理费用",     "$1.02亿",     "$0.66亿",    "+54.1%"],
        ["运营亏损",           "($1.46亿)",   "($1.01亿)",  "扩大"],
        ["GAAP净利润",         "$2.38亿",     "($4.82亿)",  "扭亏为盈*"],
        ["调整后净亏损",       "($0.76亿)",   "—",          "—"],
        ["摊薄每股收益",       "$1.38",       "($9.90)",    "扭亏"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5]
)
body(doc, "*GAAP净利润包含G42远期合约清偿所产生的一次性非现金收益3.633亿美元。",
     size=8, italic=True, color=DGRAY, space_before=2, space_after=2)
body(doc, "数据来源：Cerebras Systems S-1招股说明书（SEC EDGAR，2026年4月17日）；StockAnalysis.com。",
     size=8, italic=True, color=DGRAY, space_before=1, space_after=4)

page_break(doc)

# ── 第2-3页：营收详析 ────────────────────────────────────────────────────────
heading(doc, "营收详细分析", level=1, size=16, space_before=6, space_after=6)

heading(doc, "营收增长轨迹：从初创到5.1亿美元", level=2, space_before=6, space_after=4)
body(doc, "Cerebras Systems展现了卓越的顶线增长能力，营收从FY2022年的2,460万美元增长至FY2025年的5.10亿美元——三年内实现20倍增长。FY2025年营收同比增长76%，虽较FY2024年的269%增速有所放缓，但系基数效应所致。增长轨迹反映出企业对GPU架构之外的AI专用算力基础设施需求正在加速释放。")

add_chart(doc, "cbrs_chart1_revenue.png", width=5.5,
          caption="图1：Cerebras年度营收趋势（FY2022-FY2025）。数据来源：S-1招股说明书，SEC EDGAR。")

heading(doc, "营收结构：硬件与云服务", level=2, space_before=8, space_after=4)
body(doc, "FY2025年营收由硬件系统（3.584亿美元，占比70.3%）和云/推理服务（1.516亿美元，占比29.7%）构成。云服务占比已从H1 2024年的约25%提升至近30%，标志着业务模式的关键转型。云服务板块基于Cerebras Inference推理平台，具备更高的毛利率潜力，并有望形成经常性收入来源。")
body(doc, "硬件营收主要来自基于WSE-3晶圆级芯片的CS-3系统销售。云服务营收则来自Cerebras即服务（CaaS）推理部署，客户按算力使用量付费而非直接购买硬件。")

add_chart(doc, "cbrs_chart2_segment.png", width=5.8,
          caption="图2：FY2025营收构成及分部趋势。数据来源：S-1招股说明书，SEC EDGAR。")

heading(doc, "客户集中度：最关键的风险因素", level=2, space_before=8, space_after=4)
body(doc, "最为突出的风险在于极端的客户集中度。FY2025年86%的营收来自两家阿联酋关联实体：穆罕默德·本·扎耶德人工智能大学（MBZUAI）贡献62%，G42贡献24%。截至2025年12月31日，仅MBZUAI就占应收账款的77.9%。两家实体被认定为关联方，实质上构成单一国家（阿布扎比）集中风险。")
body(doc, "公司招股说明书明确揭示了这一风险。但2026年初签署的OpenAI合约（200亿美元以上）和AWS合作伙伴关系有望大幅改善客户多元化程度。预计到FY2026年，OpenAI可能成为最大营收来源，将集中风险从中东主权实体转向美国超大规模云厂商。")

add_chart(doc, "cbrs_chart6_customers.png", width=5.8,
          caption="图3：客户集中度分析。数据来源：S-1招股说明书，SEC EDGAR。")

page_break(doc)

# ── 第4-5页：盈利能力与利润率 ────────────────────────────────────────────────
heading(doc, "盈利能力与利润率分析", level=1, size=16, space_before=6, space_after=6)

heading(doc, "毛利率：规模效益部分被产品组合变化抵消", level=2, space_before=6, space_after=4)
body(doc, "FY2025年毛利率为39.0%，较FY2024年的42.3%下降330个基点。下降主要反映硬件营收占比上升（硬件毛利率约36-37%），而云服务毛利率更高。但毛利润绝对值增长62%至1.991亿美元，表明营收规模正在推动可观的利润美元增长。")
body(doc, "展望未来，公司向云推理方向的战略转型（通过OpenAI和AWS合作伙伴关系推动）有望支撑毛利率回升，因为类SaaS的云服务模式将在营收组合中占据更大比例。")

add_chart(doc, "cbrs_chart3_margin.png", width=5.5,
          caption="图4：毛利率与毛利润趋势（FY2022-FY2025）。数据来源：S-1招股说明书，SEC EDGAR。")

heading(doc, "运营费用：持续大规模研发投入", level=2, space_before=8, space_after=4)
body(doc, "FY2025年总运营费用达3.449亿美元，同比增长54%。研发费用2.433亿美元（占营收47.7%）仍是最大支出项，反映出对下一代晶圆级技术和Cerebras软件平台的持续投资。销售及管理费用增至1.016亿美元，系公司在OpenAI和AWS业务放量前提前扩充商业团队所致。")
body(doc, "研发强度已从FY2023年占营收177.9%降至FY2025年的47.7%，但相对半导体同行（通常15-25%）仍属偏高水平，这与公司处于盈利前增长阶段及晶圆级技术资本密集的特性一致。")

add_chart(doc, "cbrs_chart4_opex.png", width=5.5,
          caption="图5：运营费用构成（FY2022-FY2025）。数据来源：S-1招股说明书，SEC EDGAR。")

heading(doc, "运营亏损与盈利路径", level=2, space_before=8, space_after=4)
body(doc, "GAAP运营亏损从FY2024年的1.014亿美元扩大至FY2025年的1.459亿美元，主因运营费用增长54%。运营利润率从-35.0%改善至-28.6%，反映出营收增速快于成本增速带来的经营杠杆效应。")
body(doc, "GAAP净利润2.378亿美元具有误导性——包含一次性非现金收益3.633亿美元（来自G42远期合约清偿）。剔除该一次性项目及股权激励费用后，调整后净亏损约为7,570万美元，同比恶化247%。")

add_chart(doc, "cbrs_chart5_oploss.png", width=5.5,
          caption="图6：运营亏损与运营利润率趋势。数据来源：S-1招股说明书，SEC EDGAR。")

page_break(doc)

# ── 第6-7页：技术与竞争格局 ──────────────────────────────────────────────────
heading(doc, "技术实力与竞争格局", level=1, size=16, space_before=6, space_after=6)

heading(doc, "晶圆级引擎：架构层面的差异化优势", level=2, space_before=6, space_after=4)
body(doc, "Cerebras的核心竞争优势在于其晶圆级引擎（WSE-3），这是一颗占据整片300mm硅晶圆（21.5cm×21.5cm）的单芯片。关键参数包括：4万亿晶体管、90万AI优化核心、44GB片上SRAM（内存带宽21 PB/s）以及CS-3系统125 PetaFLOPs FP16算力。")
body(doc, "通过保持台积电制造的晶圆完整而非切割为独立芯片，Cerebras消除了芯片间通信瓶颈——这是制约GPU集群扩展的主要瓶颈。此架构在AI推理负载（尤其是延迟敏感型应用）上具有显著优势。")

heading(doc, "性能基准对比：Cerebras vs. NVIDIA", level=3, space_before=6, space_after=4)
make_table(doc,
    ["规格参数", "Cerebras WSE-3 / CS-3", "NVIDIA B200"],
    [
        ["晶体管数量",   "4万亿",          "2,080亿"],
        ["AI核心数",     "900,000",        "约17,000"],
        ["片上存储",     "44 GB SRAM",     "192 GB HBM3e"],
        ["内存带宽",     "21 PB/s",        "8 TB/s"],
        ["推理速度",     "2,100 tok/s/用户*", "约260 tok/s/用户"],
        ["系统功耗",     "约23 kW",        "约10 kW（每GPU）"],
        ["外形尺寸",     "15U机架",        "8U（每台DGX）"],
    ],
    col_widths=[2.0, 2.2, 2.2]
)
body(doc, "*基于Llama 3.1 70B基准测试。数据来源：Cerebras公司资料；Futurum Group分析报告。",
     size=8, italic=True, color=DGRAY, space_before=2, space_after=4)

heading(doc, "战略合约：OpenAI与AWS", level=2, space_before=8, space_after=4)
body(doc, "2026年1月公布的OpenAI协议是一份具有变革性意义的合约，总价值超200亿美元。OpenAI承诺购买750兆瓦推理算力（可扩展至2030年的2吉瓦）。OpenAI同时以6%年利率向Cerebras提供10亿美元贷款用于建设专用数据中心，并获得最多3,340万股无投票权股票的认股权证——若OpenAI总支出达300亿美元，认股权证可转换为约10%的Cerebras股权。")
body(doc, "AWS合作伙伴关系（2026年3月）将Cerebras芯片引入亚马逊数据中心生态系统，与亚马逊自研的Trainium芯片协同工作。亚马逊投资2.70亿美元购入Cerebras股权作为合作的一部分，验证了Cerebras技术与现有云基础设施的互补性（而非仅限于竞争关系）。")

heading(doc, "台积电依赖：关键供应链风险", level=2, space_before=8, space_after=4)
body(doc, "Cerebras所有晶圆均由台积电代工，且无正式的长期供应或产能分配承诺。目前没有其他代工厂能够量产如此大尺寸的晶圆级芯片。台积电的产能分配优先倾向于更大体量客户（苹果、英伟达、AMD），可能对Cerebras产生供应约束。任何涉及台湾的地缘政治冲突都将直接影响Cerebras的生产能力。这种无近期替代方案的单一供应商依赖构成重大风险。")

add_chart(doc, "cbrs_chart8_rd.png", width=5.5,
          caption="图7：研发投入趋势。数据来源：S-1招股说明书，SEC EDGAR。")

page_break(doc)

# ── 第8-9页：估值与财务总览 ──────────────────────────────────────────────────
heading(doc, "估值与财务总览", level=1, size=16, space_before=6, space_after=6)

heading(doc, "IPO定价与首日表现", level=2, space_before=6, space_after=4)
body(doc, "Cerebras于2026年5月13日以每股$185定价IPO，高于此前已上调的$150-$160区间（最初为$115-$125）。发行获约20倍超额认购。公司出售3,000万股A类普通股，募集资金55.5亿美元——为2019年Uber以来最大的美国科技IPO。股票以$350开盘，首日收于$311.07，较发行价溢价68.2%。")

add_chart(doc, "cbrs_chart10_ipo.png", width=5.8,
          caption="图8：IPO定价与资本结构。数据来源：CNBC；公司招股说明书。")

heading(doc, "GAAP与Non-GAAP盈利能力差异", level=2, space_before=8, space_after=4)
body(doc, "GAAP与Non-GAAP结果之间的差异对投资者至关重要。GAAP净利润2.378亿美元包含一次性非现金收益3.633亿美元（来自G42远期合约清偿）。剔除该一次性项目和股权激励费用后，公司调整后净亏损约为7,570万美元。经营业务本身仍未实现盈利，FY2025年运营亏损为1.459亿美元。")

add_chart(doc, "cbrs_chart7_netincome.png", width=5.5,
          caption="图9：GAAP与Non-GAAP净利润对比。数据来源：S-1招股说明书，SEC EDGAR。")

heading(doc, "估值分析：高溢价反映增长预期", level=2, space_before=8, space_after=4)
body(doc, f"以当前股价${price}计算，Cerebras的滞后市销率约为131倍——显著高于AI半导体同行（英伟达约29倍，AMD约9倍，博通约18倍）。高溢价反映了：(1)具有变革性的OpenAI和AWS合约（可能推动FY2027年营收达20-30亿美元以上）；(2)晶圆级技术的独特性，缺乏直接竞争对手；(3)更广泛的AI基础设施资本开支超级周期。")
body(doc, "然而，高溢价也隐含着巨大的执行风险。公司必须成功推进OpenAI合约落地、实现客户多元化、管理台积电供应约束并实现运营盈利——同时还要与英伟达根深蒂固的GPU生态系统（CUDA）竞争。鉴于公司刚完成IPO，目前尚无卖方分析师覆盖。")

add_chart(doc, "cbrs_chart9_valuation.png", width=5.5,
          caption="图10：市销率对比——Cerebras vs. AI半导体同行。数据来源：StockAnalysis.com；Yahoo Finance。")

page_break(doc)

# ── 第10页：投资论点与风险 ───────────────────────────────────────────────────
heading(doc, "投资论点与风险因素", level=1, size=16, space_before=6, space_after=6)

heading(doc, "看多逻辑", level=2, space_before=6, space_after=4)
bullet(doc, "OpenAI合约（200亿美元以上）重塑营收增长轨迹——有望推动FY2027-2028年年营收达20-30亿美元以上。")
bullet(doc, "晶圆级架构在推理负载方面相对GPU集群具备真正的性能差异化优势。")
bullet(doc, "云/推理服务模式推动利润率扩张——随着营收组合从硬件转向经常性收入。")
bullet(doc, "AWS合作伙伴关系验证技术实力，并打开全球最大云基础设施生态系统的入口。")
bullet(doc, "AI推理需求增速超过训练需求，Cerebras正处于增长拐点的有利位置。")

heading(doc, "看空逻辑 / 主要风险", level=2, space_before=8, space_after=4)
bullet(doc, "极端客户集中度：FY2025年86%营收来自两家阿联酋实体；OpenAI合约转移风险但未消除集中度问题。")
bullet(doc, "台积电单一供应商依赖，无正式产能分配承诺，且无其他代工厂具备晶圆级生产能力。")
bullet(doc, "运营亏损持续：FY2025年运营亏损1.459亿美元；实现盈利需要显著的营收规模提升和利润率扩张。")
bullet(doc, "估值溢价（约131倍市销率）几乎不容许执行失误或市场情绪转变。")
bullet(doc, "英伟达竞争反应：Blackwell架构和软件生态系统（CUDA）构成强大的竞争护城河。")
bullet(doc, "地缘政治/出口管制风险：阿联酋客户群及潜在受限市场扩张可能触发监管审查。")
bullet(doc, "OpenAI关系复杂性：OpenAI持有可转换为约10%股权的认股权证；贷款与股权结构形成绑定风险。")

divider(doc)

heading(doc, "核心财务指标汇总", level=2, space_before=8, space_after=4)
make_table(doc,
    ["指标", "FY2023", "FY2024", "FY2025"],
    [
        ["营收",                 "$0.79亿",    "$2.90亿",    "$5.10亿"],
        ["营收增速",             "+220%",      "+269%",      "+76%"],
        ["毛利率",               "33.5%",      "42.3%",      "39.0%"],
        ["运营亏损",             "($1.34亿)",  "($1.01亿)",  "($1.46亿)"],
        ["运营利润率",           "-170.1%",    "-35.0%",     "-28.6%"],
        ["研发费用占营收比",     "177.9%",     "54.5%",      "47.7%"],
        ["GAAP净利润/(亏损)",    "($1.27亿)",  "($4.82亿)",  "$2.38亿*"],
        ["摊薄每股收益",         "($2.92)",    "($9.90)",    "$1.38*"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5]
)
body(doc, "*包含一次性非现金收益3.633亿美元。调整后净亏损：7,570万美元。",
     size=8, italic=True, color=DGRAY, space_before=2, space_after=6)

divider(doc)

# ── 资料来源 ─────────────────────────────────────────────────────────────────
heading(doc, "资料来源与参考文献", level=2, space_before=8, space_after=4)

p = body(doc, "", size=9, space_before=2, space_after=2)
p.clear()
run = p.add_run("IPO/财报资料：\n")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• S-1注册声明（2026年4月17日提交）：")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "SEC EDGAR文件", "https://www.sec.gov/Archives/edgar/data/2021728/000162828026025762/cerebras-sx1april2026.htm")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• 修订S-1（2026年5月4日提交）：")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "SEC EDGAR", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0002021728&type=S-1&dateb=&owner=include&count=40")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• IPO定价报道（2026年5月14日）：")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "CNBC报道", "https://www.cnbc.com/2026/05/14/cerebras-cbrs-stock-trade-nasdaq-ipo.html")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• 财务数据：")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "StockAnalysis.com — CBRS财务数据", "https://stockanalysis.com/stocks/cbrs/financials/")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• S-1深度解析报告：")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "Futurum Group", "https://futurumgroup.com/insights/cerebras-s-1-teardown-is-the-23b-wafer-scale-ipo-the-end-of-gpu-homogeneity/")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• 市场数据：Yahoo Finance，yfinance API（截至2026年5月16日）")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY

body(doc, "", size=6, space_before=10, space_after=0)
body(doc, "免责声明：本报告仅供参考，不构成投资建议。所有财务数据均来源于公开文件和第三方数据提供商。"
     "过往业绩不代表未来表现。投资者在做出投资决策前应自行进行尽职调查。",
     size=7, italic=True, color=DGRAY, space_before=6, space_after=2)

# ── 保存 ─────────────────────────────────────────────────────────────────────
outfile = OUT + "CBRS_FY2025_IPO暨年度财务分析_中文版.docx"
doc.save(outfile)
print(f"报告已保存：{outfile}")
print(f"页数：约10页 | 图表：10张 | 表格：3张")
