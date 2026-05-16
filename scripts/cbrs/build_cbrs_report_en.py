#!/usr/bin/env python3
"""
Cerebras Systems (CBRS) IPO & FY2025 Financial Review — English DOCX Report
Institutional format: 10 pages, Times New Roman
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, yfinance as yf

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/CBRS/"
CHARTS = OUT

# ── Market Data ──────────────────────────────────────────────────────────────
def get_market_data(ticker):
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: Could not fetch market data for {ticker}: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("CBRS")
price = mkt["price"]
mcap = mkt["market_cap"]
mcap_b = round(mcap / 1e9, 1) if isinstance(mcap, (int, float)) else "N/A"

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
BLUE  = RGBColor(0x00, 0x66, 0xCC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x6B, 0x3A)
RED   = RGBColor(0xCC, 0x00, 0x00)
LGRAY = RGBColor(0xF0, 0xF0, 0xF0)
GOLD  = RGBColor(0xCC, 0x99, 0x00)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink'); rPr.append(rStyle)
    color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '0066CC'); rPr.append(color_el)
    u_el = OxmlElement('w:u'); u_el.set(qn('w:val'), 'single'); rPr.append(u_el)
    sz_el = OxmlElement('w:sz'); sz_el.set(qn('w:val'), '20'); rPr.append(sz_el)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def heading(doc, text, level=1, color=NAVY, size=None, bold=True, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.bold = bold; run.font.color.rgb = color
    if size is None: size = {1:16, 2:13, 3:11}.get(level, 11)
    run.font.size = Pt(size)
    return p

def body(doc, text, size=10, bold=False, color=DGRAY, space_before=2, space_after=2, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(size); run.font.bold = bold
    run.font.italic = italic; run.font.color.rgb = color
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(size); run.font.color.rgb = DGRAY
    return p

def add_chart(doc, fname, width=6.0, caption=None):
    path = CHARTS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2); cp.paragraph_format.space_after = Pt(6)
        for run in cp.runs:
            run.font.size = Pt(8); run.font.italic = True; run.font.color.rgb = DGRAY

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    run._r.append(br)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '003366')
    pBdr.append(bottom); pPr.append(pBdr)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""; p = cell.paragraphs[0]
        run = p.add_run(h); run.font.name = 'Times New Roman'
        run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, '003366')
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = ""; p = cell.paragraphs[0]
            run = p.add_run(str(val)); run.font.name = 'Times New Roman'
            run.font.size = Pt(9); run.font.color.rgb = DGRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i % 2 == 1: set_cell_bg(cell, 'F0F0F0')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# BUILD REPORT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(10)
for section in doc.sections:
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# ── PAGE 1: TITLE & SUMMARY ─────────────────────────────────────────────────
heading(doc, "CEREBRAS SYSTEMS, INC.", level=1, size=20, space_before=20, space_after=2)
heading(doc, "IPO & FY2025 Financial Review", level=1, size=14, color=BLUE, bold=False, space_before=2, space_after=2)
body(doc, f"NASDAQ: CBRS  |  Current Price: ${price}  |  Market Cap: ${mcap_b}B  |  Sector: AI Infrastructure / Semiconductors",
     size=10, bold=True, color=NAVY, space_before=4, space_after=2)
body(doc, f"IPO Date: May 14, 2026  |  IPO Price: $185  |  Day-1 Close: $311.07 (+68.2%)  |  Employees: ~708",
     size=10, bold=False, color=DGRAY, space_before=2, space_after=6)

divider(doc)

heading(doc, "KEY TAKEAWAYS", level=2, space_before=8, space_after=4)
bullet(doc, "FY2025 revenue reached $510M (+76% YoY), driven by hardware systems ($358M) and cloud/inference services ($152M).")
bullet(doc, "GAAP net income of $238M includes a one-time $363M non-cash gain from G42 forward-contract extinguishment; adjusted net loss of $(76)M.")
bullet(doc, "Gross margin of 39.0% in FY2025, down 330bps from FY2024 (42.3%) as hardware mix shifted; cloud segment carries higher margin potential.")
bullet(doc, "Mega-contracts secured: OpenAI $20B+ multi-year deal (750MW inference capacity); AWS multi-year partnership with $270M equity investment.")
bullet(doc, "Customer concentration remains the key risk: 86% of FY2025 revenue from UAE-linked entities (MBZUAI 62%, G42 24%).")
bullet(doc, "IPO raised $5.55B at $185/share — the largest US tech IPO since Uber (2019). Stock opened at $350, closed Day 1 at $311.")

divider(doc)

heading(doc, "FY2025 RESULTS SNAPSHOT", level=2, space_before=8, space_after=4)
make_table(doc,
    ["Metric", "FY2025", "FY2024", "YoY Change"],
    [
        ["Revenue",           "$510.0M",   "$290.3M",  "+75.7%"],
        ["Gross Profit",      "$199.1M",   "$122.7M",  "+62.2%"],
        ["Gross Margin",      "39.0%",     "42.3%",    "-330 bps"],
        ["R&D Expense",       "$243.3M",   "$158.2M",  "+53.8%"],
        ["SG&A Expense",      "$101.6M",   "$65.9M",   "+54.1%"],
        ["Operating Loss",    "($145.9M)", "($101.4M)", "Widened"],
        ["GAAP Net Income",   "$237.8M",   "($481.6M)","Swing to profit*"],
        ["Adj. Net Loss",     "($75.7M)",  "—",        "—"],
        ["EPS (Diluted)",     "$1.38",     "($9.90)",  "Swing"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5]
)
body(doc, "*GAAP net income includes $363.3M one-time non-cash gain from extinguishment of G42 forward-contract liability.",
     size=8, italic=True, color=DGRAY, space_before=2, space_after=2)
body(doc, "Source: Cerebras Systems S-1 Filing (SEC EDGAR, April 17, 2026); StockAnalysis.com.",
     size=8, italic=True, color=DGRAY, space_before=1, space_after=4)

page_break(doc)

# ── PAGE 2-3: DETAILED RESULTS ──────────────────────────────────────────────
heading(doc, "REVENUE ANALYSIS", level=1, size=16, space_before=6, space_after=6)

heading(doc, "Revenue Trajectory: From Startup to $510M", level=2, space_before=6, space_after=4)
body(doc, "Cerebras Systems has demonstrated exceptional top-line growth, scaling from $24.6M in FY2022 to $510.0M in FY2025 — a 20x increase in just three years. The FY2025 revenue of $510.0M represented 76% year-over-year growth, decelerating from the 269% growth achieved in FY2024, but reflecting a much larger base effect. The growth trajectory underscores accelerating enterprise demand for AI-specific compute infrastructure beyond conventional GPU architectures.")

add_chart(doc, "cbrs_chart1_revenue.png", width=5.5,
          caption="Figure 1: Cerebras Annual Revenue (FY2022–FY2025). Source: S-1 Filing, SEC EDGAR.")

heading(doc, "Revenue Composition: Hardware vs. Cloud Services", level=2, space_before=8, space_after=4)
body(doc, "FY2025 revenue breaks down into hardware systems ($358.4M, 70.3% of total) and cloud/inference services ($151.6M, 29.7%). This represents a meaningful shift: cloud services grew to nearly 30% of the mix, up from approximately 25% in H1 2024. The cloud segment, powered by the Cerebras Inference platform, carries higher gross margin potential and offers a path to recurring revenue.")
body(doc, "Hardware revenue reflects sales of CS-3 systems built around the WSE-3 wafer-scale chip. Cloud revenue comes from Cerebras-as-a-Service inference deployments where customers pay for compute capacity rather than purchasing hardware outright.")

add_chart(doc, "cbrs_chart2_segment.png", width=5.8,
          caption="Figure 2: FY2025 Revenue by Segment & Segment Trends. Source: S-1 Filing, SEC EDGAR.")

heading(doc, "Customer Concentration: The Defining Risk", level=2, space_before=8, space_after=4)
body(doc, "The most significant risk factor is extreme customer concentration. In FY2025, 86% of revenue came from two UAE-linked entities: Mohamed bin Zayed University of Artificial Intelligence (MBZUAI) at 62% and G42 at 24%. MBZUAI alone accounted for 77.9% of accounts receivable as of December 31, 2025. These entities are considered related parties with respect to each other, effectively representing single-country (Abu Dhabi) concentration.")
body(doc, "The company's own prospectus acknowledges this risk explicitly. However, the signing of the OpenAI contract ($20B+) and the AWS partnership in early 2026 is expected to dramatically diversify the customer base going forward. By FY2026E, OpenAI could represent the majority of revenue, shifting concentration from Middle Eastern sovereign entities to US-based hyperscalers.")

add_chart(doc, "cbrs_chart6_customers.png", width=5.8,
          caption="Figure 3: Customer Concentration Analysis. Source: S-1 Filing, SEC EDGAR.")

page_break(doc)

# ── PAGE 4-5: MARGIN & PROFITABILITY ─────────────────────────────────────────
heading(doc, "PROFITABILITY & MARGINS", level=1, size=16, space_before=6, space_after=6)

heading(doc, "Gross Margin: Scale Benefits Partially Offset by Mix Shift", level=2, space_before=6, space_after=4)
body(doc, "Gross margin declined 330 basis points to 39.0% in FY2025 from 42.3% in FY2024. The decline primarily reflects a higher mix of hardware revenue (which carries ~36-37% gross margins) relative to cloud services (which carry higher margins). However, the absolute gross profit expanded 62% to $199.1M, demonstrating that revenue scale is driving meaningful profit dollar growth even as margin percentages compress slightly.")
body(doc, "Looking ahead, the company's strategic pivot toward cloud inference (via the OpenAI and AWS partnerships) should support gross margin expansion as the recurring-revenue, software-like cloud model becomes a larger portion of the mix.")

add_chart(doc, "cbrs_chart3_margin.png", width=5.5,
          caption="Figure 4: Gross Margin & Gross Profit Trends (FY2022–FY2025). Source: S-1 Filing, SEC EDGAR.")

heading(doc, "Operating Expenses: Heavy R&D Investment", level=2, space_before=8, space_after=4)
body(doc, "Total operating expenses reached $344.9M in FY2025, up 54% from $224.2M in FY2024. R&D spending of $243.3M (47.7% of revenue) remains the dominant expense category, reflecting continued investment in next-generation wafer-scale technology and the Cerebras Software Platform. SG&A grew to $101.6M as the company scaled its commercial organization ahead of the OpenAI and AWS ramp.")
body(doc, "While R&D intensity has declined from 177.9% of revenue in FY2023 to 47.7% in FY2025, it remains elevated relative to semiconductor peers (typically 15-25%), consistent with the company's pre-profitability growth stage and the capital-intensive nature of developing wafer-scale technology.")

add_chart(doc, "cbrs_chart4_opex.png", width=5.5,
          caption="Figure 5: Operating Expense Breakdown (FY2022–FY2025). Source: S-1 Filing, SEC EDGAR.")

heading(doc, "Operating Loss & Path to Profitability", level=2, space_before=8, space_after=4)
body(doc, "The GAAP operating loss widened to $(145.9M) in FY2025 from $(101.4M) in FY2024, primarily driven by the 54% increase in operating expenses. Operating margin improved to -28.6% from -35.0%, reflecting operating leverage as revenue scaled faster than costs.")
body(doc, "The GAAP net income of $237.8M is misleading — it includes a one-time, non-cash gain of $363.3M from extinguishing a forward-contract liability tied to G42's original investment. Stripping out this item plus stock-based compensation, the non-GAAP adjusted net loss was approximately $(75.7M), representing a 247% deterioration year-over-year on an adjusted basis.")

add_chart(doc, "cbrs_chart5_oploss.png", width=5.5,
          caption="Figure 6: Operating Loss & Operating Margin Trends. Source: S-1 Filing, SEC EDGAR.")

page_break(doc)

# ── PAGE 6-7: TECHNOLOGY & COMPETITIVE POSITION ──────────────────────────────
heading(doc, "TECHNOLOGY & COMPETITIVE POSITION", level=1, size=16, space_before=6, space_after=6)

heading(doc, "Wafer-Scale Engine: Architectural Differentiation", level=2, space_before=6, space_after=4)
body(doc, "Cerebras' core competitive advantage lies in its Wafer-Scale Engine (WSE-3), a single chip that occupies an entire 300mm silicon wafer (21.5cm x 21.5cm). Key specifications include 4 trillion transistors, 900,000 AI-optimized cores, 44GB of on-chip SRAM with 21 PB/s memory bandwidth, and 125 PetaFLOPs FP16 performance in the CS-3 system form factor.")
body(doc, "By keeping the TSMC-manufactured wafer whole rather than dicing it into discrete chips, Cerebras eliminates inter-chip communication bottlenecks — the primary constraint limiting GPU cluster scaling. This architecture delivers significant advantages for AI inference workloads, particularly in latency-sensitive applications.")

heading(doc, "Performance Benchmarks vs. NVIDIA", level=3, space_before=6, space_after=4)
make_table(doc,
    ["Specification", "Cerebras WSE-3 / CS-3", "NVIDIA B200"],
    [
        ["Transistors",     "4 trillion",         "208 billion"],
        ["AI Cores",        "900,000",            "~17,000"],
        ["On-Chip Memory",  "44 GB SRAM",         "192 GB HBM3e"],
        ["Memory Bandwidth","21 PB/s",            "8 TB/s"],
        ["Inference Speed", "2,100 tok/s/user*",  "~260 tok/s/user"],
        ["Power (System)",  "~23 kW",             "~10 kW (per GPU)"],
        ["Form Factor",     "15U rack",           "8U per DGX"],
    ],
    col_widths=[2.0, 2.2, 2.2]
)
body(doc, "*Llama 3.1 70B benchmark. Source: Cerebras corporate materials; Futurum Group analysis.",
     size=8, italic=True, color=DGRAY, space_before=2, space_after=4)

heading(doc, "Strategic Contracts: OpenAI & AWS", level=2, space_before=8, space_after=4)
body(doc, "The OpenAI agreement, announced January 2026, represents a transformational contract valued at over $20 billion. OpenAI has committed to purchasing 750 megawatts of inference capacity (expandable to 2 GW by 2030). OpenAI also provided a $1 billion loan at 6% annual interest to help Cerebras build dedicated data centers, and received warrants for up to 33.4 million shares of non-voting stock that convert to ~10% ownership if total spending reaches $30 billion.")
body(doc, "The AWS partnership (March 2026) brings Cerebras chips into Amazon's data center ecosystem, working alongside Amazon's proprietary Trainium chips. Amazon invested $270 million in Cerebras equity as part of the collaboration, validating Cerebras' technology as complementary to (not just competitive with) existing cloud infrastructure.")

heading(doc, "TSMC Dependency: Key Supply-Chain Risk", level=2, space_before=8, space_after=4)
body(doc, "All Cerebras wafers are manufactured by TSMC, with no formalized long-term supply or allocation commitment. No other foundry can viably produce wafer-scale chips of this size. TSMC's allocation priorities favor higher-volume customers (Apple, NVIDIA, AMD), creating potential supply constraints. Any geopolitical disruption to Taiwan-based manufacturing would directly impair Cerebras' production capacity. This single-supplier dependency with no near-term alternative represents a material risk.")

add_chart(doc, "cbrs_chart8_rd.png", width=5.5,
          caption="Figure 7: R&D Investment Trend. Source: S-1 Filing, SEC EDGAR.")

page_break(doc)

# ── PAGE 8-9: VALUATION & FINANCIAL OVERVIEW ─────────────────────────────────
heading(doc, "VALUATION & FINANCIAL OVERVIEW", level=1, size=16, space_before=6, space_after=6)

heading(doc, "IPO Pricing & First-Day Performance", level=2, space_before=6, space_after=4)
body(doc, "Cerebras priced its IPO at $185 per share on May 13, 2026, above the already-raised range of $150-$160 (originally $115-$125). The offering was approximately 20x oversubscribed. The company sold 30 million Class A shares, raising $5.55 billion in gross proceeds — the largest US tech IPO since Uber in 2019. Shares opened at $350 and closed Day 1 at $311.07, representing a 68.2% premium to the IPO price.")

add_chart(doc, "cbrs_chart10_ipo.png", width=5.8,
          caption="Figure 8: IPO Pricing & Capital Structure. Source: CNBC; Company Prospectus.")

heading(doc, "GAAP vs. Non-GAAP Profitability", level=2, space_before=8, space_after=4)
body(doc, "The divergence between GAAP and non-GAAP results is critical for investors to understand. GAAP net income of $237.8M includes a $363.3M non-cash gain from extinguishing a forward-contract liability tied to G42's original investment. Excluding this one-time item and stock-based compensation, the company reported an adjusted net loss of approximately $(75.7M). The operating business remains unprofitable, with an operating loss of $(145.9M) in FY2025.")

add_chart(doc, "cbrs_chart7_netincome.png", width=5.5,
          caption="Figure 9: GAAP vs. Non-GAAP Net Income. Source: S-1 Filing, SEC EDGAR.")

heading(doc, "Valuation Context: Premium Reflects Growth Potential", level=2, space_before=8, space_after=4)
body(doc, f"At the current price of ${price}, Cerebras trades at a trailing P/S ratio of approximately 131x — a significant premium to AI semiconductor peers (NVIDIA at ~29x, AMD at ~9x, Broadcom at ~18x). The premium reflects: (1) the transformational OpenAI and AWS contracts that could drive revenue to $2-3B+ by FY2027; (2) the unique wafer-scale technology with limited direct competition; and (3) the broader AI infrastructure capital expenditure super-cycle.")
body(doc, "However, the premium also embeds substantial execution risk. The company must successfully ramp the OpenAI contract, diversify beyond UAE customers, manage TSMC supply constraints, and achieve operating profitability — all while competing against NVIDIA's established GPU ecosystem. No sell-side analyst coverage exists yet given the recent IPO.")

add_chart(doc, "cbrs_chart9_valuation.png", width=5.5,
          caption="Figure 10: P/S Ratio vs. AI Semiconductor Peers. Source: StockAnalysis.com; Yahoo Finance.")

page_break(doc)

# ── PAGE 10: INVESTMENT THESIS & RISKS ───────────────────────────────────────
heading(doc, "INVESTMENT THESIS & RISK FACTORS", level=1, size=16, space_before=6, space_after=6)

heading(doc, "Bull Case", level=2, space_before=6, space_after=4)
bullet(doc, "OpenAI contract ($20B+) transforms revenue trajectory — could drive $2-3B+ annual revenue by FY2027-2028.")
bullet(doc, "Wafer-scale architecture offers genuine performance differentiation vs. GPU clusters for inference workloads.")
bullet(doc, "Cloud/inference services model drives margin expansion as mix shifts from hardware to recurring revenue.")
bullet(doc, "AWS partnership validates technology and opens access to the world's largest cloud infrastructure ecosystem.")
bullet(doc, "AI inference demand is scaling faster than training demand, positioning Cerebras at the growth inflection point.")

heading(doc, "Bear Case / Key Risks", level=2, space_before=8, space_after=4)
bullet(doc, "Extreme customer concentration: 86% of FY2025 revenue from two UAE entities; OpenAI contract shifts risk but doesn't eliminate concentration.")
bullet(doc, "TSMC single-supplier dependency with no formalized allocation commitment and no alternative foundry capable of wafer-scale production.")
bullet(doc, "Operating losses persist: $(145.9M) operating loss in FY2025; path to profitability requires significant revenue scale and margin expansion.")
bullet(doc, "Valuation premium (~131x P/S) leaves little margin for execution missteps or market sentiment shifts.")
bullet(doc, "NVIDIA competitive response: NVIDIA's Blackwell architecture and software ecosystem (CUDA) represent a formidable competitive moat.")
bullet(doc, "Geopolitical/export control risks: UAE customer base and potential expansion into restricted markets could trigger regulatory scrutiny.")
bullet(doc, "OpenAI relationship complexity: OpenAI holds warrants convertible to ~10% ownership; loan and equity structure creates entanglement risk.")

divider(doc)

heading(doc, "Key Financial Summary", level=2, space_before=8, space_after=4)
make_table(doc,
    ["Metric", "FY2023", "FY2024", "FY2025"],
    [
        ["Revenue",                "$78.7M",    "$290.3M",   "$510.0M"],
        ["Revenue Growth",         "+220%",     "+269%",     "+76%"],
        ["Gross Margin",           "33.5%",     "42.3%",     "39.0%"],
        ["Operating Loss",         "($133.9M)", "($101.4M)", "($145.9M)"],
        ["Operating Margin",       "-170.1%",   "-35.0%",    "-28.6%"],
        ["R&D (% of Revenue)",     "177.9%",    "54.5%",     "47.7%"],
        ["GAAP Net Income/(Loss)", "($127.2M)", "($481.6M)", "$237.8M*"],
        ["EPS (Diluted)",          "($2.92)",   "($9.90)",   "$1.38*"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5]
)
body(doc, "*Includes $363.3M one-time non-cash gain. Adjusted net loss: $(75.7M).",
     size=8, italic=True, color=DGRAY, space_before=2, space_after=6)

divider(doc)

# ── SOURCES ──────────────────────────────────────────────────────────────────
heading(doc, "SOURCES & REFERENCES", level=2, space_before=8, space_after=4)

p = body(doc, "", size=9, space_before=2, space_after=2)
p.clear()
run = p.add_run("Earnings / IPO Materials:\n")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• S-1 Registration Statement (Filed April 17, 2026): ")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "SEC EDGAR Filing", "https://www.sec.gov/Archives/edgar/data/2021728/000162828026025762/cerebras-sx1april2026.htm")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• Amended S-1 (Filed May 4, 2026): ")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "SEC EDGAR", "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0002021728&type=S-1&dateb=&owner=include&count=40")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• IPO Pricing Report (May 14, 2026): ")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "CNBC Coverage", "https://www.cnbc.com/2026/05/14/cerebras-cbrs-stock-trade-nasdaq-ipo.html")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• Financial Data: ")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "StockAnalysis.com — CBRS Financials", "https://stockanalysis.com/stocks/cbrs/financials/")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• S-1 Teardown Analysis: ")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
add_hyperlink(p, "Futurum Group", "https://futurumgroup.com/insights/cerebras-s-1-teardown-is-the-23b-wafer-scale-ipo-the-end-of-gpu-homogeneity/")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
run = p.add_run("• Market Data: Yahoo Finance, yfinance API (as of May 16, 2026)")
run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY

body(doc, "", size=6, space_before=10, space_after=0)
body(doc, "DISCLAIMER: This report is for informational purposes only and does not constitute investment advice. "
     "All financial data is sourced from public filings and third-party data providers. Past performance is not indicative of future results. "
     "Investors should conduct their own due diligence before making investment decisions.",
     size=7, italic=True, color=DGRAY, space_before=6, space_after=2)

# ── SAVE ─────────────────────────────────────────────────────────────────────
outfile = OUT + "CBRS_FY2025_IPO_Financial_Review.docx"
doc.save(outfile)
print(f"Report saved: {outfile}")
print(f"Pages: ~10 | Charts: 10 | Tables: 3")
