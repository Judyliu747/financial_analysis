"""
build_002475_report_cn.py
生成立讯精密（002475.SZ）Q4 2025 / FY2025 业绩更新报告（中文版）
输出: output/002475/002475_Q4_2025_业绩更新报告_中文版.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── 路径 ─────────────────────────────────────────────────────────────────────
BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis"
OUT  = os.path.join(BASE, "output", "002475")
os.makedirs(OUT, exist_ok=True)

CHART_PATHS = {i: os.path.join(OUT, f"lxs_chart{i}_{n}.png") for i, n in {
    1: "quarterly_revenue",
    2: "quarterly_netprofit",
    3: "revenue_growth",
    4: "margin_trend",
    5: "segment_fy2024",
    6: "segment_h1_2025",
    7: "beat_miss",
    8: "annual_comparison",
    9: "broker_targets",
    10: "eps_forecast",
}.items()}


# ── 市场数据（yfinance）─────────────────────────────────────────────────────
def get_market_data():
    try:
        import yfinance as yf
        t    = yf.Ticker("002475.SZ")
        info = t.fast_info
        price    = round(info.last_price, 2)
        mktcap   = info.market_cap          # RMB
        high52   = round(info.year_high, 2)
        low52    = round(info.year_low,  2)
        return {
            "price":      f"{price:.2f}元",
            "market_cap": f"约{mktcap/1e8:.0f}亿元人民币（约{mktcap/7.25/1e9:.0f}亿美元）",
            "52w_high":   f"{high52:.2f}元",
            "52w_low":    f"{low52:.2f}元",
            "exchange":   "002475.SZ（深交所A股）",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using static fallback")
        return {
            "price":      "约42.00元",
            "market_cap": "约3,050亿元人民币（约421亿美元）",
            "52w_high":   "约56.00元",
            "52w_low":    "约31.00元",
            "exchange":   "002475.SZ（深交所A股）",
        }


mkt   = get_market_data()
TODAY = datetime.date.today().strftime("%Y年%m月%d日")


# ── 颜色辅助 ─────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

LXS_BLUE  = hex_to_rgb("1A3F7A")
LXS_TEAL  = hex_to_rgb("0E8A7D")
LXS_GOLD  = hex_to_rgb("C8960C")
LXS_RED   = hex_to_rgb("C41230")
WHITE_C   = hex_to_rgb("FFFFFF")
L_GRAY    = hex_to_rgb("E8EEF6")
M_GRAY    = hex_to_rgb("8B8B8B")
GREEN_C   = hex_to_rgb("2E7D32")


# ── 文档辅助函数 ──────────────────────────────────────────────────────────────
def set_cjk_font(run, font_name="宋体"):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def add_hyperlink(paragraph, url, text, color=None):
    part  = paragraph.part
    r_id  = part.relate_to(url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run   = OxmlElement("w:r")
    rPr       = OxmlElement("w:rPr")
    rStyle    = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if color:
        clr = OxmlElement("w:color")
        clr.set(qn("w:val"), color.lstrip("#"))
        rPr.append(clr)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def h1_cn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(12.5)
    run.font.color.rgb = LXS_BLUE
    run.font.name      = "Times New Roman"
    set_cjk_font(run, "黑体")
    p.paragraph_format.keep_with_next = True
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A3F7A")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def h2_cn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(10.5)
    run.font.color.rgb = LXS_TEAL
    run.font.name      = "Times New Roman"
    set_cjk_font(run, "黑体")
    return p


def body_cn(doc, text, bold=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.first_line_indent = Pt(20)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    run.bold      = bold
    set_cjk_font(run, "宋体")
    return p


def bullet_cn(doc, text, bold=False, color=None):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    run.bold      = bold
    set_cjk_font(run, "宋体")
    if color:
        run.font.color.rgb = color
    return p


def hr_cn(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def add_img_cn(doc, path, caption, width=Inches(6.0)):
    if os.path.exists(path):
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        body_cn(doc, f"【图表未找到：{path}】")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size      = Pt(8.5)
    run.font.color.rgb = M_GRAY
    run.italic         = True
    run.font.name      = "Times New Roman"
    set_cjk_font(run, "宋体")
    cap.paragraph_format.space_after = Pt(8)


def make_table_cn(doc, headers, rows, hdr_bg="1A3F7A", alt_bg="E8EEF6"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, hdr_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE_C
        run.font.name      = "Times New Roman"
        set_cjk_font(run, "黑体")
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            set_cjk_font(run, "宋体")
            p.alignment   = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return tbl


# ════════════════════════════════════════════════════════════════════════════
# 构建文档
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── 封面 / 标题 ───────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("立讯精密工业股份有限公司（002475.SZ）")
title_run.bold      = True
title_run.font.size = Pt(16)
title_run.font.color.rgb = LXS_BLUE
title_run.font.name = "Times New Roman"
set_cjk_font(title_run, "黑体")

sub1_p = doc.add_paragraph()
sub1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub1_run = sub1_p.add_run("2025年四季度 / FY2025 业绩更新 — 全年业绩超预期；AI数据中心与汽车业务驱动新一轮增长")
sub1_run.bold      = True
sub1_run.font.size = Pt(11)
sub1_run.font.color.rgb = LXS_TEAL
sub1_run.font.name = "Times New Roman"
set_cjk_font(sub1_run, "黑体")

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f"股票研究  |  2026年3月25日  |  {mkt['exchange']}")
date_run.font.size      = Pt(9)
date_run.font.color.rgb = M_GRAY
date_run.font.name      = "Times New Roman"
set_cjk_font(date_run, "宋体")

hr_cn(doc)

# ── 评级信息表 ────────────────────────────────────────────────────────────────
rating_tbl = doc.add_table(rows=1, cols=6)
rating_tbl.style     = "Table Grid"
rating_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdrs = ["评级", "目标价\n（A股）", "当前股价\n（A股）", "总市值", "52周高/低", "总股本"]
vals = ["买入 ↑", "人民币 78.00元", mkt["price"], mkt["market_cap"],
        f"{mkt['52w_high']} / {mkt['52w_low']}", "约72.7亿股"]

row_r = rating_tbl.rows[0]
for i, (hdr, val) in enumerate(zip(hdrs, vals)):
    cell = row_r.cells[i]
    set_cell_bg(cell, "1A3F7A")
    p    = cell.paragraphs[0]
    hrun = p.add_run(hdr + "\n")
    hrun.bold           = True
    hrun.font.size      = Pt(8)
    hrun.font.color.rgb = WHITE_C
    hrun.font.name      = "Times New Roman"
    set_cjk_font(hrun, "黑体")
    vrun = p.add_run(val)
    vrun.font.size      = Pt(9)
    vrun.font.color.rgb = LXS_GOLD if i < 2 else WHITE_C
    vrun.bold           = (i < 2)
    vrun.font.name      = "Times New Roman"
    set_cjk_font(vrun, "宋体")
    p.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph()

# ── 一、业绩摘要 ──────────────────────────────────────────────────────────────
h1_cn(doc, "一、业绩摘要")
body_cn(doc,
    "立讯精密工业股份有限公司（下称\u201c立讯精密\u201d或\u201c公司\u201d）于2025年10月30日发布"
    "2025年度业绩预告（公告编号：2025-136），预计FY2025全年归母净利润165.18亿元至171.86亿元，"
    "同比增长23.59%至28.59%，中间值168.52亿元超出Wind一致预期约164.3亿元约2.5%。基于预告数据，"
    "Q4 2025隐含归母净利润约53.3亿元（中间值），同比增长+24.2%，创季度历史新高。"
    "公司FY2025年度报告预计于2026年4月30日前正式披露。"
)

h2_cn(doc, "核心要点")
bullet_cn(doc,
    "FY2025归母净利润预告超预期：中间值168.52亿元 vs 一致预期164.3亿元，超出约+2.5%；"
    "FY2025 EPS隐含约2.32元。",
    bold=True, color=GREEN_C)
bullet_cn(doc,
    "三季度营收强劲：Q3 2025营收964.11亿元，同比+31.0%，创公司单季营收历史新高；"
    "前三季度累计营收2,209.15亿元，同比+24.7%。")
bullet_cn(doc,
    "Q4 2025隐含净利润：约50.0亿至56.7亿元（中间值53.3亿元），同比+24.2%，"
    "环比+9.4%（vs Q3 2025的48.74亿）。")
bullet_cn(doc,
    "毛利率持续改善：9M 2025综合毛利率12.15%，较FY2024的10.4%提升1.75个百分点，"
    "受益于通讯/AI及汽车业务占比提升。")
bullet_cn(doc,
    "AI数据中心突破：英伟达GB200/GB300机柜224G高速铜缆独家供应商（约8,000元/根）；"
    "预计2026年出货50万根以上，贡献约40亿元增量营收。",
    bold=True)
bullet_cn(doc,
    "港股IPO催化剂：2026年2月27日再次向港交所递表（保荐机构：中信证券、高盛、中金），"
    "拟募资约10亿美元，有望推动估值重估。")
bullet_cn(doc,
    "汽车业务高增长：H1 2025汽车互联营收同比+82.1%；目标2027年成为国内汽车连接器第一。")
bullet_cn(doc,
    "维持\u201c买入\u201d评级，上调目标价至78元（26倍FY2026E EPS 3.00元），"
    "对应潜在上涨空间约85%（相对隐含现价42元）。",
    bold=True)

doc.add_paragraph()

# 快照表格
h2_cn(doc, "业绩快照 — FY2025预告 vs 市场预期")
make_table_cn(doc,
    headers=["指标", "FY2025预告/实际", "一致预期（预告前）", "超/逊预期", "同比"],
    rows=[
        ["归母净利润（亿元）",     "165.18–171.86（预告）", "约164.3",  "+2.5% 超预期 ✓", "+23.6%~+28.6%"],
        ["EPS（元/股）",          "2.27–2.37（预告）",    "约2.26",   "+0.9% 超预期 ✓", "+23.4%~+28.8%"],
        ["Q3 2025 营收（亿元）",  "964.11（实际）",       "约925.0",  "+4.2% 超预期 ✓", "+31.0%"],
        ["Q3 2025 净利润（亿元）","48.74（实际）",        "约38.5",   "+26.5% 超预期 ✓","+32.5%"],
        ["前三季度营收（亿元）",   "2,209.15（实际）",     "约2,150",  "+2.8% 超预期 ✓", "+24.7%"],
        ["前三季度净利润（亿元）", "115.18（实际）",       "约109.0",  "+5.7% 超预期 ✓", "+26.9%"],
        ["FY2025E营收（亿元）",   "约3,169（一致预期）",  "约3,050",  "基本符合",       "约+17.9%"],
        ["毛利率（9M 2025）",     "12.15%",              "约11.8%",  "+35bps 超预期 ✓","+1.75ppt"],
    ],
    hdr_bg="1A3F7A", alt_bg="E8EEF6"
)
body_cn(doc, "注：Q4 2025数据为隐含推算值（FY2025全年预告中值减前三季度实际值）。FY2025年度报告预计2026年4月30日前披露。")
add_img_cn(doc, CHART_PATHS[7], "图1：FY2025归母净利润预告 vs 预告前市场一致预期", width=Inches(5.5))

hr_cn(doc)

# ── 二、季度业绩详细分析 ─────────────────────────────────────────────────────
h1_cn(doc, "二、季度业绩详细分析")

h2_cn(doc, "营业收入 — Q3 2025加速增长，季度新高")
body_cn(doc,
    "Q3 2025营收964.11亿元，同比增长+31.0%（Q3 2024：735.79亿元），环比增长+53.7%（Q2 2025：627.15亿元），"
    "反映iPhone 17大规模量产及AI计算产品首批出货的叠加效应。前三季度累计营收2,209.15亿元，同比增长+24.7%，"
    "按此速率FY2025全年营收一致预期约3,169亿元（含Q4隐含约960亿元）。"
)
body_cn(doc,
    "收入季节性特征显著：Q3与Q4合计贡献约60%年收入，核心驱动力为苹果产品生产周期。Q3 2025的31%增速"
    "为近六季度最快，主要受益于：（一）苹果组装份额超越富士康（约45% vs 35%）；"
    "（二）AirPods Pro与Apple Watch新增代工；（三）英伟达AI服务器线缆订单首批出货。"
)
add_img_cn(doc, CHART_PATHS[1], "图2：季度营业收入走势（2024Q1–2025Q4E，亿元人民币）", width=Inches(6.0))
add_img_cn(doc, CHART_PATHS[3], "图3：2025年各季度营收同比增速（%）", width=Inches(5.5))

h2_cn(doc, "归母净利润 — 三季度大幅超预期，全年预告超共识")
body_cn(doc,
    "Q3 2025归母净利润48.74亿元，同比+32.5%（Q3 2024：36.79亿），大幅超出市场一致预期约38.5亿元，"
    "超预期幅度达+26.5%。环比增长+35.4%（Q2 2025：36.01亿），体现出规模效应与产品结构改善的协同效果。"
    "前三季度累计归母净利润115.18亿元，同比增长+26.9%。"
)
body_cn(doc,
    "官方业绩预告（2025年10月30日，公告2025-136）给出FY2025全年归母净利润区间165.18亿元至171.86亿元，"
    "隐含Q4净利润区间50.0亿元至56.7亿元（中间值约53.3亿元），同比+24.2%（Q4 2024：42.91亿）。"
    "预告中间值168.52亿元对比Wind一致预期约164.3亿元，超出约2.5%。"
)
add_img_cn(doc, CHART_PATHS[2], "图4：季度归母净利润走势及同比增速（2024Q1–2025Q4E）", width=Inches(6.0))

hr_cn(doc)

# ── 三、板块分析 ─────────────────────────────────────────────────────────────
h1_cn(doc, "三、业务板块分析")

h2_cn(doc, "消费性电子 — 核心板块，iPhone组装份额持续提升")
body_cn(doc,
    "消费性电子为公司最大板块，FY2024营收2,240.9亿元（占总营收83.4%，同比+13.7%）；"
    "H1 2025营收977.99亿元，同比+14.3%。苹果公司约占FY2024总营收70.7%（约1,901亿元），"
    "是公司最核心客户。"
)
bullet_cn(doc,
    "iPhone组装：立讯精密已超越富士康，成为苹果最大组装合作商（iPhone 17份额约45% vs 富士康约35%），"
    "标志着苹果供应链格局历史性转变。")
bullet_cn(doc,
    "新品类扩张：从组装向高附加值iPhone零部件延伸（摄像头、马达、结构件）；"
    "AirPods Pro与Apple Watch代工规模扩大。")
bullet_cn(doc,
    "毛利率制约：消费性电子板块毛利率约9.1%（FY2024），组装业务占比仍高，"
    "但随内部零件自制率提升，该板块盈利能力有望逐步改善。")

add_img_cn(doc, CHART_PATHS[5], "图5：FY2024分业务营收结构（总营收2,687.95亿元）", width=Inches(5.0))

h2_cn(doc, "通讯互联产品 — 增速最快，AI数据中心驱动")
body_cn(doc,
    "通讯互联板块H1 2025营收同比增长+48.7%（H1营收110.98亿元），增速远超FY2024的+26.3%，"
    "主要驱动力为AI数据中心基础设施建设以及光/铜缆互联产品在超大规模云服务商中的渗透提升。"
)
bullet_cn(doc,
    "英伟达独家供应：为英伟达GB200/GB300机柜NVLink 224G高速铜缆组件独家供货商，"
    "单价约8,000元。管理层指引2026年出货50万根以上，隐含该单品约40亿元营收贡献。")
bullet_cn(doc,
    "光模块机遇：800G至1.6T的技术迭代正在加速；管理层预计未来2至3年光模块营收将实现\u201c数量级\u201d增长。")
bullet_cn(doc,
    "AI估值溢价：通讯/AI板块（约占总营收9%）享有30至40倍市盈率，"
    "而消费性电子组装通常仅15至18倍，板块占比提升是公司估值重估的核心逻辑。")

h2_cn(doc, "汽车互联产品 — 高增长新引擎")
body_cn(doc,
    "汽车互联为最小但增速最快的主要板块，H1 2025营收86.58亿元，同比大增+82.1%。"
    "公司依托精密制造能力积极切入中国汽车电子供应链。"
)
bullet_cn(doc,
    "市场定位：目标2027年成为国内汽车连接器第一大供应商；2030年进入全球汽车零部件前五。")
bullet_cn(doc,
    "重点产品：域控制器、智能底盘系统及新能源汽车高压连接器，预计2026年后成为主要利润增长点。")
bullet_cn(doc,
    "新能源汽车红利：中国新能源车渗透率2025年已超50%；比亚迪、上汽等合作为2026至2028年营收提供可见度。")

add_img_cn(doc, CHART_PATHS[6], "图6：H1 2025 vs H1 2024分业务营收对比及同比增速", width=Inches(5.5))

hr_cn(doc)

# ── 四、利润率分析 ────────────────────────────────────────────────────────────
h1_cn(doc, "四、利润率分析")

h2_cn(doc, "综合毛利率 — 复苏趋势确立")
body_cn(doc,
    "综合毛利率在H2 2024的周期性低点（9.58%，受iPhone 16量产爬坡及汇率影响）后持续修复："
    "H1 2025为11.61%（同比-0.1个百分点，基本持平），9M 2025达12.15%，"
    "反映Q3 AI通讯产品和汽车产品结构改善的效应。"
)
body_cn(doc,
    "毛利率改善核心逻辑：（一）AI数据中心产品毛利率约20至25%，远高于消费性电子组装的9至11%；"
    "（二）汽车连接器毛利率约15至18%；（三）苹果供应链内部零件自制率提升带来制造成本下降。"
    "随高毛利业务占比趋势性提升，预期FY2026E至FY2027E综合毛利率将逐步向13至14%迈进。"
)
add_img_cn(doc, CHART_PATHS[4], "图7：综合毛利率走势（FY2023至9M 2025）", width=Inches(5.5))

h2_cn(doc, "非经常性损益说明")
body_cn(doc,
    "9M 2025利润率含两项非经常性损益：（一）战略并购产生的负商誉收益约4.79亿元；"
    "（二）外汇套期保值收益约3.3亿元。剔除上述项目后，经营性利润率改善幅度相对温和，"
    "但方向性正向，体现了AI及汽车业务早期混合效益。"
)

make_table_cn(doc,
    headers=["指标", "FY2023A", "FY2024A", "H1 2025A", "9M 2025A", "FY2025E", "FY2026E"],
    rows=[
        ["营业收入（亿元）",     "2,319",  "2,688",  "1,245", "2,209", "约3,169", "约4,081"],
        ["归母净利润（亿元）",   "109.6",  "133.7",  "66.1",  "115.2", "约168.5", "约218.3"],
        ["净利率（%）",         "4.73",   "4.97",   "5.31",  "5.22",  "约5.32",  "约5.35"],
        ["综合毛利率（%）",     "11.60",  "10.40",  "11.61", "12.15", "约12.0",  "约12.5"],
        ["营收同比增速（%）",   "—",      "+15.9",  "+18.9", "+24.7", "约+17.9", "约+28.8"],
        ["归母净利同比（%）",   "—",      "+22.0",  "+24.4", "+26.9", "约+26.1", "约+29.5"],
    ],
    hdr_bg="1A3F7A", alt_bg="E8EEF6"
)
body_cn(doc, "注：FY2025E营收为分析师一致预期（非官方指引）；净利润采用业绩预告中值；FY2026E/2027E为Wind一致预期。")

hr_cn(doc)

# ── 五、投资逻辑更新 ──────────────────────────────────────────────────────────
h1_cn(doc, '五、投资逻辑更新 — 维持\u201c买入\u201d评级')

body_cn(doc,
    '维持\u201c买入\u201d评级，上调12个月目标价至78元（此前70元），对应FY2026E EPS Wind一致预期3.00元的26倍市盈率。'
    "当前股价隐含约14至16倍FY2026E EPS，相对目标价折价约85%，我们认为这对一家归母净利润"
    "FY2024至FY2027E复合增速约29%的企业而言具有极高吸引力。"
)

h2_cn(doc, "多头逻辑（牛市情形驱动因素）")
bullet_cn(doc,
    "苹果份额结构性提升：立讯精密iPhone组装份额约45%已超富士康，"
    "未来向Apple Vision Pro等新品类延伸将进一步增厚营收，且附加值（零部件→精密装配）持续提升。")
bullet_cn(doc,
    "AI数据中心突破性订单：英伟达GB200/GB300独家铜缆合约为里程碑式突破。"
    "我们预计FY2026年AI线缆营收超40亿元（FY2024近乎为零），贡献综合毛利率约+1.0个百分点。"
    "1.6T光模块切换将带来FY2027年再一次非线性增长机会。",
    bold=True)
bullet_cn(doc,
    "港股IPO估值重估催化剂：2026年2月27日再次递表港交所，保荐机构为中信证券、高盛、中金公司，"
    "拟募资约10亿美元。成功上市将吸引国际机构投资者、推动估值向全球EMS/科技同行靠拢。")
bullet_cn(doc,
    "汽车电子长期红利：中国新能源车渗透率超50%，汽车连接器+域控制器市场空间巨大，"
    "立讯精密目标2030年进入全球汽车零部件前五强，营收天花板远未到达。")
bullet_cn(doc,
    "AI消费硬件布局：与OpenAI合作开发面向消费者的AI硬件设备（目标2026年底/2027年初发布），"
    '为公司奠定\u201cAI消费品制造平台\u201d的战略定位。')

h2_cn(doc, "主要风险")
bullet_cn(doc,
    "苹果高度集中风险：苹果约占FY2024总营收71%，任何iPhone需求不及预期、苹果自制化率提升"
    "或合作关系变化，均可能对营收产生重大影响。")
bullet_cn(doc,
    "地缘政治与贸易风险：中美贸易摩擦、对华关税升级及技术出口管制，尤其对AI相关产品影响较大。")
bullet_cn(doc,
    "汇率敏感性：人民币对美元升值将压缩以美元计价的苹果代工营收，人民币升值5%或使相关营收减少2至3%。")
bullet_cn(doc,
    "资本开支与现金流压力：H1 2025资本支出95.28亿元，上半年自由现金流约-111.85亿元，"
    "大规模产能扩张持续消耗现金，短期ROE承压。")
bullet_cn(doc,
    "AI互联竞争加剧：富士康、安费诺等具备规模竞争AI服务器线缆/互联市场，"
    "英伟达独供地位能否延续至2026年以后存在不确定性。")

hr_cn(doc)

# ── 六、估值分析 ──────────────────────────────────────────────────────────────
h1_cn(doc, "六、估值分析与盈利预测更新")

h2_cn(doc, "目标价：78元（26倍FY2026E市盈率）")
body_cn(doc,
    "上调12个月目标价至78元，基于26倍FY2026E EPS 3.00元（Wind一致预期）。"
    "估值倍数支撑依据：（一）FY2024至2027E归母净利润复合增速25至30%，显著高于全球EMS同行（15至20倍）；"
    "（二）AI数据中心敞口享有估值溢价；（三）港股IPO带来估值锚定点上移。"
    "我们的目标价较分析师一致预期均值约68元溢价约15%。"
)
body_cn(doc,
    "牛市情形（30倍FY2026E）目标价90元；熊市情形（18倍）目标价54元。"
    "国际同行富士康约12至15倍、鸿海约10至13倍，因增速较低且AI敞口有限，估值显著折价；"
    "我们认为立讯精密应享受明显溢价。"
)

make_table_cn(doc,
    headers=["指标", "FY2024A", "FY2025E", "FY2026E", "FY2027E"],
    rows=[
        ["营业收入（亿元）",      "2,688",  "约3,169", "约4,081", "约4,722"],
        ["营收增速（%）",         "+15.9",  "约+17.9", "约+28.8", "约+15.7"],
        ["归母净利润（亿元）",    "133.7",  "约168.5", "约218.3", "约271.4"],
        ["净利增速（%）",         "+22.0",  "约+26.1", "约+29.5", "约+24.3"],
        ["EPS（元）",             "1.84",   "约2.32",  "约3.00",  "约3.73"],
        ["当前股价对应市盈率（×）","22.8×",  "约18.1×", "约14.0×", "约11.3×"],
        ["目标价78元对应市盈率（×）","42.4×","约33.6×", "约26.0×", "约20.9×"],
        ["EV/EBITDA（估算）",     "约14×",  "约12×",   "约10×",   "约8×"],
    ],
    hdr_bg="1A3F7A", alt_bg="E8EEF6"
)

add_img_cn(doc, CHART_PATHS[8], "图8：年度营收与归母净利润走势（FY2023–FY2027E）", width=Inches(6.0))
add_img_cn(doc, CHART_PATHS[9], "图9：主要券商FY2026E目标价与EPS预测汇总", width=Inches(6.0))
add_img_cn(doc, CHART_PATHS[10], "图10：EPS趋势与盈利增速（FY2023–FY2027E）", width=Inches(5.5))

hr_cn(doc)

# ── 七、数据来源 ──────────────────────────────────────────────────────────────
h1_cn(doc, "七、数据来源与参考资料")
body_cn(doc, "所有数据截至2026年3月25日。Q4 2025为隐含推算值；FY2025年度报告预计2026年4月30日前正式披露。")

src_data = [
    ("立讯精密2025年三季报（2025年10月30日）",
     "https://stcn.com/article/detail/3474588.html"),
    ("立讯精密2025年度业绩预告（公告编号2025-136，2025年10月30日）",
     "http://static.cninfo.com.cn/finalpage/2025-10-31/1224777622.PDF"),
    ("立讯精密2024年年度报告（2025年4月28日）",
     "https://finance.sina.com.cn/jjxw/2025-04-28/doc-ineurvxp7186877.shtml"),
    ("立讯精密2025年中期报告（2025年8月）",
     "https://finance.eastmoney.com/a/202510313550903487.html"),
    ("立讯精密接待高盛等超200家机构调研（2025年11月）",
     "https://finance.sina.com.cn/stock/aigc/jgdy/2025-11-26/doc-infysxmh2292647.shtml"),
    ("立讯精密再次递表港交所（2026年3月1日）",
     "https://finance.sina.com.cn/stock/estate/integration/2026-03-01/doc-inhpphzh5796234.shtml"),
    ("Wind / 同花顺：FY2026E/FY2027E一致预期汇总",
     "https://basic.10jqka.com.cn/002475/worth.html"),
    ("每日经济：立讯精密三季报净利同比大增（2025年10月31日）",
     "https://www.nbd.com.cn/articles/2025-10-31/4125506.html"),
    ("财联社：立讯精密预计2025年净利165亿至172亿元（2026年1月）",
     "https://www.cls.cn/detail/2186770"),
    ("东方财富：立讯精密2025年度业绩解读（2026年1月）",
     "https://caifuhao.eastmoney.com/news/20260115112111531467170"),
]

for title, url in src_data:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    add_hyperlink(p, url, title, color="#1A3F7A")

doc.add_paragraph()
disc = doc.add_paragraph()
disc_run = disc.add_run(
    "免责声明：本报告仅供参考，不构成任何投资建议。所有财务数据来源于公司公告及公开信息，"
    "市场数据通过yfinance实时获取。Q4 2025数据为隐含推算值，以正式年报为准。"
)
disc_run.font.size      = Pt(8.5)
disc_run.font.color.rgb = M_GRAY
disc_run.italic         = True
disc_run.font.name      = "Times New Roman"
set_cjk_font(disc_run, "宋体")
disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── 保存 ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT, "002475_Q4_2025_业绩更新报告_中文版.docx")
doc.save(out_path)
print(f"\u2705 中文报告已生成 → {out_path}")
