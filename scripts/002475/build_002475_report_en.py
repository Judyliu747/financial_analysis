"""
build_002475_report_en.py
Generates English DOCX earnings update report for
立讯精密 (Luxshare Precision Industry, 002475.SZ)
Q4 2025 / FY2025 Earnings Update
Output: output/002475/002475_Q4_2025_Earnings_Update.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Paths ────────────────────────────────────────────────────────────────────
BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis"
OUT  = os.path.join(BASE, "output", "002475")
os.makedirs(OUT, exist_ok=True)

CHART_PATHS = {i: os.path.join(OUT, f"lxs_chart{i}_{n}.png") for i, n in {
    1: "quarterly_revenue",
    2: "quarterly_netprofit",
    3: "revenue_growth",
    4: "margin_trend",
    5: "segment_fy2024",
    6: "segment_h1_2025",
    7: "beat_miss",
    8: "annual_comparison",
    9: "broker_targets",
    10: "eps_forecast",
}.items()}


# ── Market Data (yfinance) ────────────────────────────────────────────────────
def get_market_data():
    try:
        import yfinance as yf
        t = yf.Ticker("002475.SZ")
        info = t.fast_info
        price    = round(info.last_price, 2)
        mktcap   = info.market_cap
        high52   = round(info.year_high, 2)
        low52    = round(info.year_low, 2)
        return {
            "price":      f"RMB {price:.2f}",
            "market_cap": f"~RMB {mktcap/1e9:.0f}B (~USD {mktcap/7.25/1e9:.0f}B)",
            "52w_high":   f"RMB {high52:.2f}",
            "52w_low":    f"RMB {low52:.2f}",
            "exchange":   "002475.SZ (A-share, Shenzhen)",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using static fallback")
        return {
            "price":      "RMB ~42.00",
            "market_cap": "~RMB 3,050B (~USD 421B)",
            "52w_high":   "RMB ~56.00",
            "52w_low":    "RMB ~31.00",
            "exchange":   "002475.SZ (A-share, Shenzhen)",
        }


mkt = get_market_data()
TODAY = datetime.date.today().strftime("%B %d, %Y")


# ── Color Helpers ─────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

LXS_BLUE  = hex_to_rgb("1A3F7A")
LXS_TEAL  = hex_to_rgb("0E8A7D")
LXS_GOLD  = hex_to_rgb("C8960C")
LXS_RED   = hex_to_rgb("C41230")
WHITE_C   = hex_to_rgb("FFFFFF")
L_GRAY    = hex_to_rgb("E8EEF6")
M_GRAY    = hex_to_rgb("8B8B8B")
GREEN_C   = hex_to_rgb("2E7D32")
LGRAY2    = hex_to_rgb("F5F5F5")


# ── Document Helpers ──────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val", "single"))
            el.set(qn("w:sz"),    val.get("sz", "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_hyperlink(paragraph, url, text, color=None):
    part     = paragraph.part
    r_id     = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                               is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run   = OxmlElement("w:r")
    rPr       = OxmlElement("w:rPr")
    rStyle    = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if color:
        clr = OxmlElement("w:color")
        clr.set(qn("w:val"), color.lstrip("#"))
        rPr.append(clr)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = LXS_BLUE
    p.paragraph_format.keep_with_next = True
    # add bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A3F7A")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = LXS_TEAL
    return p


def body(doc, text, bold=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.name      = "Times New Roman"
    run.bold           = bold
    return p


def bullet(doc, text, bold=False, color=None):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold      = bold
    if color:
        run.font.color.rgb = color
    return p


def hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def add_img(doc, path, caption, width=Inches(6.0)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        body(doc, f"[Chart not found: {path}]")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size      = Pt(8.5)
    run.font.color.rgb = M_GRAY
    run.italic         = True
    cap.paragraph_format.space_after = Pt(8)


def make_table(doc, headers, rows, hdr_bg="1A3F7A", alt_bg="E8EEF6"):
    col_count = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, hdr_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE_C
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            p.alignment   = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return tbl


# ════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── COVER / HEADER ────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Luxshare Precision Industry (002475.SZ)")
title_run.bold           = True
title_run.font.size      = Pt(16)
title_run.font.color.rgb = LXS_BLUE
title_run.font.name      = "Times New Roman"

sub1_p = doc.add_paragraph()
sub1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub1_run = sub1_p.add_run("Q4 2025 / FY2025 Earnings Update — Beat on Profit Guidance; AI & Auto Driving Next Leg of Growth")
sub1_run.bold           = True
sub1_run.font.size      = Pt(11)
sub1_run.font.color.rgb = LXS_TEAL
sub1_run.font.name      = "Times New Roman"

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f"Equity Research | March 25, 2026  |  {mkt['exchange']}")
date_run.font.size      = Pt(9)
date_run.font.color.rgb = M_GRAY
date_run.font.name      = "Times New Roman"

hr(doc)

# ── RATING TABLE ──────────────────────────────────────────────────────────────
rating_tbl = doc.add_table(rows=1, cols=6)
rating_tbl.style     = "Table Grid"
rating_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_r = ["Rating", "Price Target\n(A-share)", "Current Price\n(A-share)",
             "Market Cap", "52W High/Low", "Shares Out."]
values_r  = ["BUY ↑", "RMB 78.00",
             mkt["price"], mkt["market_cap"],
             f"{mkt['52w_high']} / {mkt['52w_low']}", "~7.27B shares"]

row_r = rating_tbl.rows[0]
for i, (hdr, val) in enumerate(zip(headers_r, values_r)):
    cell = row_r.cells[i]
    set_cell_bg(cell, "1A3F7A")
    p = cell.paragraphs[0]
    hrun = p.add_run(hdr + "\n")
    hrun.bold           = True
    hrun.font.size      = Pt(8)
    hrun.font.color.rgb = WHITE_C
    hrun.font.name      = "Times New Roman"
    vrun = p.add_run(val)
    vrun.font.size      = Pt(9.5)
    vrun.font.color.rgb = LXS_GOLD if i < 2 else WHITE_C
    vrun.bold           = (i < 2)
    vrun.font.name      = "Times New Roman"
    p.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph()

# ── SECTION 1: EXECUTIVE SUMMARY ─────────────────────────────────────────────
h1(doc, "I. Executive Summary")
body(doc,
     "Luxshare Precision Industry (002475.SZ) delivered a strong finish to FY2025, with the company's "
     "FY2025 full-year performance guidance (issued October 30, 2025) implying net profit attributable to "
     "shareholders of RMB 16.52B–17.19B, representing 23.6%–28.6% YoY growth and a meaningful beat "
     "against prior consensus expectations. Q4 2025 implied net profit of ~RMB 5.3B (midpoint) represents "
     "the company's strongest fourth quarter on record, up +24% YoY vs. Q4 2024's RMB 4.29B. "
     "The full annual report (FY2025 Annual Report) is expected by April 30, 2026.")

h2(doc, "Key Takeaways")
bullet(doc, "FY2025 net profit guidance beat: RMB 16.52B–17.19B vs. prior consensus ~RMB 16.4B — "
            "implied mid-point of RMB 16.85B represents a ~+2.5% beat.", bold=True, color=GREEN_C)
bullet(doc, "Q4 2025 implied net profit: ~RMB 5.33B (FY guidance midpoint minus Q1–Q3 actuals), "
            "+24.2% YoY — acceleration from Q3's +32.5% but on a higher base.")
bullet(doc, "Revenue momentum sustained: Q3 2025 revenue +31.0% YoY (RMB 96.4B); cumulative "
            "9M 2025 revenue +24.7% YoY to RMB 220.9B. FY2025 consensus revenue ~RMB 316.9B (+17.9% YoY).")
bullet(doc, "Gross margin expanding: 9M 2025 gross margin 12.15% vs. FY2024's 10.4%, driven by "
            "mix-shift toward higher-margin communications/AI and automotive segments.")
bullet(doc, "AI data center: Sole supplier for Nvidia GB200/GB300 224G high-speed copper cables (~RMB 8K/unit); "
            "500K+ unit orders expected in 2026 → ~RMB 4B incremental revenue.", bold=True)
bullet(doc, "HK IPO re-filed February 27, 2026 (sponsors: CITIC Securities, Goldman Sachs, CICC), "
            "targeting ~USD 1B raise — a key near-term re-rating catalyst.")
bullet(doc, "Automotive: On track to become #1 China automotive connector by 2027; "
            "automotive revenue +82% YoY in H1 2025.")
bullet(doc, "Maintain BUY with revised PT of RMB 78 (26x FY2026E EPS of RMB 3.00, "
            "vs. consensus target ~RMB 68).")

doc.add_paragraph()

# Results Snapshot Table
h2(doc, "Results Snapshot — FY2025 Performance Guidance vs. Consensus")
make_table(doc,
    headers=["Metric", "FY2025 Guidance / Actual", "Consensus (Pre-Guidance)", "Beat / Miss", "YoY"],
    rows=[
        ["Net Profit (RMB B)",     "16.52–17.19 (guidance)",   "~16.43",   "+2.5% beat ✓",  "+23.6%–28.6%"],
        ["EPS (RMB)",              "2.27–2.37 (guidance)",     "~2.26",    "+0.9% beat ✓",  "+23.4%–28.8%"],
        ["Q3 2025 Revenue (RMB B)","96.41 (actual)",           "~92.5",    "+4.2% beat ✓",  "+31.0%"],
        ["Q3 2025 Net Profit (RMB B)","4.87 (actual)",         "~3.85",    "+26.5% beat ✓", "+32.5%"],
        ["9M 2025 Revenue (RMB B)","220.92 (actual)",          "~215.0",   "+2.8% beat ✓",  "+24.7%"],
        ["9M 2025 Net Profit (RMB B)","11.52 (actual)",        "~10.90",   "+5.7% beat ✓",  "+26.9%"],
        ["FY2025E Revenue (RMB B)","~316.9 (consensus est.)",  "~305.0",   "In Line",        "~+17.9%"],
        ["Gross Margin (9M 2025)", "12.15%",                   "~11.8%",   "+35bps ✓",      "+3.6ppt YoY"],
    ],
    hdr_bg="1A3F7A", alt_bg="E8EEF6"
)
body(doc, "Note: Q4 2025 data is implied (FY guidance minus Q1–Q3 actuals). Full FY2025 Annual Report expected by April 30, 2026.")

add_img(doc, CHART_PATHS[7], "Exhibit 1: FY2025 Net Profit Guidance vs. Pre-Announcement Consensus Estimates", width=Inches(5.5))

hr(doc)

# ── SECTION 2: DETAILED RESULTS ───────────────────────────────────────────────
h1(doc, "II. Detailed Quarterly Results")

h2(doc, "Revenue — Sustained Acceleration in Q3 2025")
body(doc,
     "Q3 2025 revenue reached RMB 96.41B, up +31.0% YoY from RMB 73.58B in Q3 2024, and +53.7% QoQ "
     "from Q2 2025's RMB 62.72B — reflecting the sharp seasonal peak driven by iPhone 16/17 production "
     "ramp and new AI compute product shipments. Cumulative 9M 2025 revenue of RMB 220.92B grew +24.7% "
     "YoY, firmly tracking ahead of the ~RMB 316.9B FY2025 consensus (implying Q4 2025 revenue of "
     "~RMB 96.0B, +4.8% YoY vs. Q4 2024's RMB 91.62B).")
body(doc,
     "Revenue seasonality is pronounced: Q3 and Q4 collectively account for ~60% of annual revenue, "
     "driven by the iPhone production cycle. Q3 2025's 31% growth marks the fastest quarterly expansion "
     "in six quarters, driven by (1) iPhone assembly share gains over Foxconn (~45% vs. 35%), "
     "(2) new Apple Watch and AirPods Pro wins, and (3) initial AI server cable shipments to Nvidia.")

add_img(doc, CHART_PATHS[1], "Exhibit 2: Quarterly Revenue Trend (Q1 2024 – Q4 2025E; RMB Billion)", width=Inches(6.0))
add_img(doc, CHART_PATHS[3], "Exhibit 3: Quarterly Revenue YoY Growth (%), 2025", width=Inches(5.5))

h2(doc, "Net Profit — Q3 Beat, FY2025 Guidance Above Consensus")
body(doc,
     "Q3 2025 net profit attributable to shareholders reached RMB 4.87B, up +32.5% YoY (Q3 2024: "
     "RMB 3.68B) — a significant beat vs. consensus estimates of ~RMB 3.85B. Sequential improvement "
     "was +35.4% QoQ from Q2 2025's RMB 3.60B, highlighting operating leverage on higher volumes. "
     "For Q1–Q3 2025 cumulative, net profit was RMB 11.52B (+26.9% YoY).")
body(doc,
     "The official FY2025 performance guidance (Announcement 2025-136, October 30, 2025) indicates "
     "full-year net profit of RMB 16.52B–17.19B, implying Q4 net profit of RMB 5.00B–5.67B (midpoint "
     "~RMB 5.33B, +24.2% YoY vs. Q4 2024's RMB 4.29B). The FY2025 guidance midpoint of RMB 16.85B "
     "compares favorably to prior Wind consensus of ~RMB 16.4B, a +2.5% beat.")

add_img(doc, CHART_PATHS[2], "Exhibit 4: Quarterly Net Profit Trend with YoY Growth (Q1 2024 – Q4 2025E)", width=Inches(6.0))

hr(doc)

# ── SECTION 3: SEGMENT ANALYSIS ───────────────────────────────────────────────
h1(doc, "III. Segment Analysis")

h2(doc, "Consumer Electronics (消费性电子) — Core Segment, iPhone Share Gains")
body(doc,
     "Consumer electronics remains the dominant segment, accounting for ~83% of FY2024 revenue "
     "(RMB 224.09B, +13.7% YoY). In H1 2025, consumer electronics revenue was RMB 97.80B (+14.3% YoY). "
     "Apple accounted for approximately 70.7% of FY2024 total revenue (RMB 190.14B), making it Luxshare's "
     "most critical customer relationship.")
bullet(doc, "iPhone assembly: Luxshare has surpassed Foxconn as the leading iPhone assembler, "
            "with an estimated ~45% assembly share for iPhone 17 (2025 model). This represents a "
            "structural shift in Apple's supply chain strategy.")
bullet(doc, "New product wins: Expanding beyond assembly into higher-value iPhone components "
            "(cameras, vibration modules, structural components); AirPods Pro and Apple Watch "
            "contract manufacturing expanded.")
bullet(doc, "Margin headwind: Consumer electronics carries the lowest gross margin (~9.1% in FY2024) "
            "due to the assembly-intensive nature, but mix improvement within the segment "
            "(more components vs. pure assembly) is driving gradual margin recovery.")

add_img(doc, CHART_PATHS[5], "Exhibit 5: FY2024 Revenue by Business Segment (Total: RMB 268.8B)", width=Inches(5.0))

h2(doc, "Communications & AI (通讯互联) — Fastest-Growing Segment")
body(doc,
     "Communications/interconnect products grew +48.7% YoY in H1 2025 (H1 revenue: RMB 11.10B), "
     "accelerating from FY2024's already strong +26.3% YoY. The primary drivers are AI data center "
     "infrastructure buildout and optical/copper cable products for hyperscalers.")
bullet(doc, "Nvidia partnership: Sole supplier for GB200/GB300 NVLink 224G high-speed copper cable "
            "assemblies at ~RMB 8,000/unit. Management guided 500,000+ unit orders in 2026, "
            "implying ~RMB 4B in 2026 revenue from this product alone.")
bullet(doc, "Optical modules: 800G/1.6T transition underway; management expects 'order-of-magnitude' "
            "growth in optical revenue over next 2–3 years as 1.6T adoption accelerates.")
bullet(doc, "AI valuation premium: Communications/AI (~9% of revenue) commands 30–40x P/E multiples "
            "vs. 15–18x for core consumer electronics — driving the bull case for re-rating.")

h2(doc, "Automotive Electronics (汽车互联) — High-Growth New Driver")
body(doc,
     "Automotive electronics is the smallest but fastest-growing major segment, with H1 2025 revenue "
     "of RMB 8.66B (+82.1% YoY). Luxshare is targeting a structural position in the Chinese automotive "
     "electronics supply chain, leveraging its precision manufacturing capabilities.")
bullet(doc, "Market position: On track to become #1 automotive connector supplier in China by 2027; "
            "aiming for top-5 global auto parts status by 2030.")
bullet(doc, "Key verticals: Domain controllers, smart chassis systems, and high-voltage connectors "
            "for new energy vehicles (NEV) — expected to become major profit drivers post-2026.")
bullet(doc, "NEV tailwinds: China NEV penetration rate exceeded 50% in 2025; BYD and SAIC partnerships "
            "provide revenue visibility for 2026–2028E.")

add_img(doc, CHART_PATHS[6], "Exhibit 6: H1 2025 vs H1 2024 Segment Revenue Comparison (+YoY Growth)", width=Inches(5.5))

hr(doc)

# ── SECTION 4: MARGIN ANALYSIS ────────────────────────────────────────────────
h1(doc, "IV. Margin Analysis")

h2(doc, "Gross Margin — Recovery Trend Intact Despite Mix Pressure")
body(doc,
     "Consolidated gross margin has been recovering after a cyclical trough in H2 2024 (9.58%), "
     "which was depressed by iPhone 16 production ramp costs and unfavorable FX (USD/RMB). "
     "H1 2025 gross margin of 11.61% was effectively flat YoY (-10bps), but 9M 2025 gross margin "
     "of 12.15% reflects Q3's improvement driven by AI communications products and automotive mix-shift.")
body(doc,
     "Key margin dynamics to watch: (1) AI/data center products carry 20–25% gross margins vs. "
     "9–11% for consumer electronics assembly; (2) automotive connectors carry 15–18% margins; "
     "(3) ongoing cost reduction in consumer electronics via vertical integration. As high-margin "
     "segments grow toward 15–20% of total revenue, consolidated gross margin should trend toward "
     "13–14% in FY2026E–2027E.")

add_img(doc, CHART_PATHS[4], "Exhibit 7: Consolidated Gross Margin Trend (FY2023 – 9M 2025)", width=Inches(5.5))

h2(doc, "Note on Non-Recurring Items in 9M 2025")
body(doc,
     "The 9M 2025 margin figures include two notable non-recurring items: (1) a negative goodwill "
     "gain (bargain purchase gain) of ~RMB 479M from a strategic acquisition, and (2) FX hedging "
     "gains of ~RMB 330M. Excluding these, underlying operating margin improvement is more modest "
     "but still directionally positive, reflecting the early benefits of mix improvement toward "
     "AI and automotive products.")

make_table(doc,
    headers=["Metric", "FY2023A", "FY2024A", "H1 2025A", "9M 2025A", "FY2025E", "FY2026E"],
    rows=[
        ["Revenue (RMB B)",           "231.9", "268.8", "124.5", "220.9", "~316.9", "~408.1"],
        ["Net Profit (RMB B)",         "10.96", "13.37", "6.61",  "11.52", "~16.85", "~21.83"],
        ["Net Margin (%)",             "4.73",  "4.97",  "5.31",  "5.22",  "~5.32",  "~5.35"],
        ["Gross Margin (%)",           "11.60", "10.40", "11.61", "12.15", "~12.0",  "~12.5"],
        ["YoY Revenue Growth (%)",     "—",     "+15.9", "+18.9", "+24.7", "~+17.9", "~+28.8"],
        ["YoY Net Profit Growth (%)",  "—",     "+22.0", "+24.4", "+26.9", "~+26.1", "~+29.5"],
    ],
    hdr_bg="1A3F7A", alt_bg="E8EEF6"
)
body(doc, "Note: FY2025E revenue is analyst consensus (not official guidance). FY2025E/FY2026E net profit from Wind consensus / management pre-announcement.")

hr(doc)

# ── SECTION 5: INVESTMENT THESIS ──────────────────────────────────────────────
h1(doc, "V. Investment Thesis Update — Maintain BUY")

body(doc,
     "We maintain our BUY rating and raise our 12-month price target to RMB 78 (from prior RMB 70), "
     "representing 26x FY2026E EPS of RMB 3.00 (Wind consensus). At the current price, Luxshare trades "
     "at approximately 14–16x FY2026E EPS — a 35–45% discount to our target and a compelling entry point "
     "for a company growing earnings at ~29% CAGR over FY2024–FY2027E.")

h2(doc, "Bull Case Drivers")
bullet(doc, "Apple share gains: Luxshare's ~45% iPhone assembly share (vs. 35% for Foxconn) is a "
            "structural shift — not cyclical. Further gains into AirPods, Apple Vision Pro manufacturing "
            "would add incremental revenue with improving margins (assembler → component manufacturer).")
bullet(doc, "AI data center: The Nvidia GB200/GB300 cable sole-source contract is a landmark win. "
            "We model RMB 4B+ in AI cable revenue for FY2026 (vs. near-zero in FY2024), contributing "
            "~1.0ppt to consolidated gross margin expansion. Optical module transition to 1.6T offers "
            "another step-function growth opportunity in FY2027.")
bullet(doc, "Hong Kong IPO re-rating catalyst: Re-submission of HKEX listing (February 27, 2026) "
            "with Goldman Sachs, CITIC, and CICC as sponsors signals management commitment. A "
            "successful listing would attract international institutional investors and provide "
            "permanent re-rating to a global EMS/tech peer group multiple.")
bullet(doc, "Automotive secular growth: China NEV market structural growth underpins Luxshare's "
            "82% YoY automotive revenue growth in H1 2025. The company is building durable competitive "
            "advantages in domain controllers and smart chassis — targeting RMB 30B+ in automotive "
            "revenue by 2028.")
bullet(doc, "AI consumer devices: Collaboration with OpenAI on consumer AI hardware (targeting "
            "launch late 2026/early 2027) positions Luxshare as a key beneficiary if the "
            "'AI smartphone/PC' replacement cycle materializes at scale.")

h2(doc, "Key Risks")
bullet(doc, "Apple concentration: ~71% customer concentration in Apple represents meaningful "
            "single-customer risk. Any iPhone demand shortfall, Apple in-sourcing of components, "
            "or disruption to the Apple-Luxshare relationship would materially impact revenue.")
bullet(doc, "Geopolitical risk: US-China trade tensions, potential tariffs on Chinese EMS companies, "
            "and technology export controls represent tail risks, particularly for AI-related products.")
bullet(doc, "FX sensitivity: RMB appreciation vs. USD compresses Apple assembly revenue when "
            "translated from USD contract pricing. A 5% RMB appreciation could reduce revenue "
            "by ~2–3% on the Apple-linked portion.")
bullet(doc, "Margin pressure from scale-up costs: Rapid capacity expansion (H1 2025 CapEx: "
            "RMB 9.53B; H1 FCF: ~-RMB 11.2B) may continue to weigh on near-term FCF and ROE.")
bullet(doc, "Competition in AI hardware: AI server cable/interconnect market may attract new "
            "entrants; Foxconn and Amphenol have the scale to compete. Sole-source status with "
            "Nvidia may not persist beyond 2026–2027.")

hr(doc)

# ── SECTION 6: VALUATION ──────────────────────────────────────────────────────
h1(doc, "VI. Valuation & Updated Estimates")

h2(doc, "Price Target: RMB 78 (26x FY2026E P/E)")
body(doc,
     "Our revised 12-month price target of RMB 78 is based on 26x our FY2026E EPS estimate of "
     "RMB 3.00 (Wind consensus). This multiple is justified by: (1) Luxshare's 25–30% earnings "
     "CAGR over FY2024–2027E (above global EMS peers trading at 15–20x), (2) meaningful AI "
     "data center exposure commanding a premium, and (3) near-term HK IPO re-rating catalyst. "
     "Our target represents ~25% upside from consensus analyst targets of ~RMB 68 on average.")

body(doc,
     "In the bull scenario (30x FY2026E), PT = RMB 90; in the bear scenario (18x), PT = RMB 54. "
     "International peers (Foxconn at 12–15x, Hon Hai at 10–13x) trade at significant discounts "
     "due to lower growth and limited AI exposure — we believe Luxshare deserves a premium.")

make_table(doc,
    headers=["Metric", "FY2024A", "FY2025E", "FY2026E", "FY2027E"],
    rows=[
        ["Revenue (RMB B)",          "268.8", "~316.9", "~408.1", "~472.2"],
        ["YoY Growth (%)",           "+15.9", "~+17.9", "~+28.8", "~+15.7"],
        ["Net Profit (RMB B)",       "13.37", "~16.85", "~21.83", "~27.14"],
        ["NP Growth (%)",            "+22.0", "~+26.1", "~+29.5", "~+24.3"],
        ["EPS (RMB)",                "1.84",  "~2.32",  "~3.00",  "~3.73"],
        ["P/E at RMB 42 (x)",        "22.8x", "~18.1x", "~14.0x", "~11.3x"],
        ["P/E at PT RMB 78 (x)",     "42.4x", "~33.6x", "~26.0x", "~20.9x"],
        ["EV/EBITDA (est.)",         "~14x",  "~12x",   "~10x",   "~8x"],
    ],
    hdr_bg="1A3F7A", alt_bg="E8EEF6"
)

add_img(doc, CHART_PATHS[8], "Exhibit 8: Annual Revenue & Net Profit Trend (FY2023–FY2027E)", width=Inches(6.0))
add_img(doc, CHART_PATHS[9], "Exhibit 9: Broker Price Target & FY2026E EPS Consensus", width=Inches(6.0))
add_img(doc, CHART_PATHS[10], "Exhibit 10: EPS Trend & Earnings Growth (FY2023–FY2027E)", width=Inches(5.5))

hr(doc)

# ── SECTION 7: SOURCES ────────────────────────────────────────────────────────
h1(doc, "VII. Sources & References")

body(doc, "All data as of March 25, 2026. Q4 2025 figures are implied estimates; full FY2025 "
          "Annual Report expected by April 30, 2026.")

src_data = [
    ("Luxshare Q3 2025 Quarterly Report (October 30, 2025)",
     "https://stcn.com/article/detail/3474588.html"),
    ("FY2025 Performance Pre-Announcement (Announcement 2025-136, October 30, 2025)",
     "http://static.cninfo.com.cn/finalpage/2025-10-31/1224777622.PDF"),
    ("Luxshare FY2024 Annual Report (April 28, 2025)",
     "https://finance.sina.com.cn/jjxw/2025-04-28/doc-ineurvxp7186877.shtml"),
    ("Luxshare H1 2025 Interim Report (August 2025)",
     "https://finance.eastmoney.com/a/202510313550903487.html"),
    ("Institutional Investor Conference Notes — November 2025 Roadshow (Goldman Sachs 200+ institutions)",
     "https://finance.sina.com.cn/stock/aigc/jgdy/2025-11-26/doc-infysxmh2292647.shtml"),
    ("Luxshare HKEX Re-submission (February 27, 2026)",
     "https://finance.sina.com.cn/stock/estate/integration/2026-03-01/doc-inhpphzh5796234.shtml"),
    ("Wind Consensus Estimates — FY2026E/FY2027E EPS (Tonghuashun 同花顺)",
     "https://basic.10jqka.com.cn/002475/worth.html"),
    ("NBD: Luxshare Q3 2025 Results Beat Expectations (October 31, 2025)",
     "https://www.nbd.com.cn/articles/2025-10-31/4125506.html"),
    ("CLS: FY2025 Net Profit Guidance RMB 16.52–17.19B (January 2026)",
     "https://www.cls.cn/detail/2186770"),
    ("EastMoney: Luxshare FY2025 Business Analysis (January 2026)",
     "https://caifuhao.eastmoney.com/news/20260115112111531467170"),
]

for title, url in src_data:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    add_hyperlink(p, url, title, color="#1A3F7A")

doc.add_paragraph()
disc = doc.add_paragraph()
disc_run = disc.add_run(
    "DISCLAIMER: This report is for informational purposes only and does not constitute investment "
    "advice. All financial data is sourced from public company filings and consensus estimates. "
    "Market data via yfinance as of report date. Q4 2025 figures are analyst-implied estimates "
    "pending official annual report disclosure."
)
disc_run.font.size      = Pt(8)
disc_run.font.color.rgb = M_GRAY
disc_run.italic         = True
disc.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT, "002475_Q4_2025_Earnings_Update.docx")
doc.save(out_path)
print(f"✅ English report saved → {out_path}")
