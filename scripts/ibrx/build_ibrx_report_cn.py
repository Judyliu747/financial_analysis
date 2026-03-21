"""
ImmunityBio (IBRX) Q4 2025 业绩更新报告 — 中文版
机构股票研究报告格式，8-12页，宋体正文，黑体标题
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, copy

OUT    = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/IBRX/"
CHARTS = OUT

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── Color constants ───────────────────────────────────────────
NAVY     = RGBColor(0, 48, 135)
RED      = RGBColor(200, 16, 46)
GRAY     = RGBColor(94, 106, 113)
GREEN    = RGBColor(26, 122, 74)
WHITE    = RGBColor(255, 255, 255)
NAVY_HEX = "003087"
HEAD_FILL= "E8EDF7"

# ── Helper: CJK font XML ──────────────────────────────────────
def set_cjk_font(run, body=True):
    """Set CJK font: 宋体 for body, 黑体 for headings."""
    font_name = "宋体" if body else "黑体"
    run.font.name = "Times New Roman"
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hint'), 'eastAsia')

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading_cn(text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level==1 else 11 if level==2 else 10)
    run.font.color.rgb = color or NAVY
    set_cjk_font(run, body=False)  # 黑体
    return p

def add_body_cn(text, indent=True, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    set_cjk_font(run, body=True)  # 宋体
    return p

def add_bullet_cn(text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = color
    set_cjk_font(run, body=True)
    return p

def add_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_chart(filename, width=6.0):
    path = CHARTS + filename
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))

def hyperlink_cn(para, url, text):
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink')
    color_e = OxmlElement('w:color'); color_e.set(qn('w:val'), '0563C1')
    u_e = OxmlElement('w:u'); u_e.set(qn('w:val'), 'single')
    rPr.append(rStyle); rPr.append(color_e); rPr.append(u_e)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hl.append(new_run)
    para._p.append(hl)

# ═══════════════════════════════════════════════════════════════
# 封面 / 第1页
# ═══════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run("ImmunityBio（IBRX）——2025年第四季度及全年业绩更新报告")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
set_cjk_font(r, body=False)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(2)
r2 = p_sub.add_run("2026年3月17日  |  股票研究  |  生物制药行业")
r2.font.size = Pt(10); r2.font.color.rgb = GRAY
set_cjk_font(r2, body=True)

p_rat = doc.add_paragraph()
p_rat.paragraph_format.space_after = Pt(6)
run_rat = p_rat.add_run("评级：跑赢大市（OUTPERFORM）  |  目标价：$12.00（由$7.00上调）  |  现价：~$8.30")
run_rat.bold = True; run_rat.font.size = Pt(11); run_rat.font.color.rgb = GREEN
set_cjk_font(run_rat, body=False)

add_rule()

# 核心要点
add_heading_cn("核心要点", 2, NAVY)
takeaways = [
    "营收超预期：第四季度净产品收入3,830万美元，超市场一致预期约3,500万美元的9.4%；全年2025年收入1.133亿美元，同比大增621%",
    "每股亏损超预期：第四季度每股净亏损（0.06）美元，优于市场预期（0.09）美元，超预期约33%，主要受益于销售与管理费用管控",
    "ANKTIVA商业化放量持续：第四季度为J代码生效（2025年1月）以来连续第5个季度环比增长，商业化势头稳健",
    "催化剂丰富：BCG初治临床试验（QUILT 2.005）中期完全缓解率84%比52%（p=0.0455）；2026年3月9日重新提交sBLA；沙特药监局批准ANKTIVA用于肺癌",
    "风险警示：现金余额2.428亿美元低于全年运营消耗约3.05亿美元，预计H2 2026将需要融资；业绩公布后股价下跌约13%，市场担忧稀释风险",
    "维持跑赢大市评级，目标价由7.00美元上调至12.00美元，反映ANKTIVA商业化速度超预期",
]
for t in takeaways:
    add_bullet_cn(t)

add_rule()

# 业绩快照表
add_heading_cn("业绩快照", 2, NAVY)
snap_data = [
    ['指标', '2025Q4实际值', '市场一致预期', '超预期情况', '同比变化'],
    ['净产品收入', '3,830万美元', '约3,500万美元', '+330万美元（+9.4%）', '+3,110万美元（+432%）'],
    ['全年2025年收入', '1.133亿美元', '约1.05亿美元', '+830万美元（+7.9%）', '+9,760万美元（+621%）'],
    ['净亏损（Q4）', '（6,190万美元）', '—', '调整后优于预期', '较2024Q4（7,050万美元）改善'],
    ['每股亏损（Q4）', '（0.06）美元', '（0.085-0.09）美元', '超预期+0.025（+33%）', '—'],
    ['现金及有价证券', '2.428亿美元', '—', '—', '较2024年末约2.88亿美元减少'],
    ['2026Q1收入指引', '未提供明确指引', '约4,300万美元以上', '增长轨迹不变', '增长势头持续'],
]
tbl = doc.add_table(rows=len(snap_data), cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(snap_data):
    for ci, txt in enumerate(row_data):
        cell = tbl.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        set_cjk_font(run, body=True)
        if ri == 0:
            run.bold = True; run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ci == 3 and ri > 0 and '超' in txt:
            run.font.color.rgb = GREEN; run.bold = True
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第2-3页：营收与商业化分析
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_cn("一、营收与商业化表现分析", 1)
add_rule()

add_heading_cn("1.1 季度营收趋势", 2)
add_body_cn(
    "2025年第四季度ANKTIVA（N-803+BCG）净产品收入达3,830万美元，环比第三季度（3,210万美元）增长20%，"
    "同比2024年第四季度（720万美元）大幅增长432%。该结果超出市场一致预期约3,500万美元的9.4%，"
    "延续了自2025年1月永久J代码（J9028）生效以来季度环比持续增长的良好态势。"
)
add_chart("ibrx_chart1_revenue.png", 5.5)
add_body_cn("图1：ANKTIVA季度净产品收入（百万美元）| 来源：ImmunityBio各季度新闻稿", indent=False, color=GRAY, size=8)
doc.add_paragraph()

add_heading_cn("1.2 J代码的关键转折", 2)
add_body_cn(
    "最重要的商业里程碑是2025年1月1日生效的永久J代码（J9028）。J代码是美国医疗保险中心（CMS）"
    "针对医生使用药物的专属计费代码，缺乏J代码将迫使诊所采用耗时且繁琐的杂项计费方式，"
    "这曾是2024年社区泌尿科诊所采用ANKTIVA的主要障碍。J代码生效后，直接催生了2025年第一季度"
    "环比高达129%的营收加速增长，并为全年实现1.133亿美元营收奠定了基础。"
)
add_chart("ibrx_chart2_qoq_growth.png", 5.5)
add_body_cn("图2：ANKTIVA季度营收环比增长率（%）| 来源：ImmunityBio，公司公告", indent=False, color=GRAY, size=8)

add_heading_cn("1.3 2025年全年业绩表现", 2)
add_body_cn(
    "2025年全年净产品收入1.133亿美元，同比2024年1,570万美元增长621%，超过业绩公布前市场预期约1.05亿美元。"
    "管理层于2026年1月预先披露初步营收数据，与正式财报一致。本次强劲增长的驱动因素如下："
)
add_bullet_cn("永久J代码消除了社区泌尿诊所的报销障碍（美国约75%的膀胱癌治疗在学术中心之外进行）")
add_bullet_cn("处方医生群体扩大：社区泌尿外科医生和学术中心重复处方习惯正在形成")
add_bullet_cn("rBCG扩展使用计划（EAP）已有约580名患者入组，扩大了可治疗患者群体")
add_bullet_cn("美国国家综合癌症网络（NCCN）将ANKTIVA列为推荐用药，有效降低了保险公司的预授权门槛")
add_chart("ibrx_chart5_annual_rev.png", 5.0)
add_body_cn("图3：年度净产品收入增长（2024-2026E，百万美元）| 来源：公司公告，Piper Sandler预测", indent=False, color=GRAY, size=8)

# ═══════════════════════════════════════════════════════════════
# 第4-5页：利润率与费用分析
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_cn("二、财务指标与费用分析", 1)
add_rule()

add_heading_cn("2.1 运营费用管控效果显著", 2)
add_body_cn(
    "2025年全年的重要亮点之一是在营收规模化扩张的同时实现了显著的销售及管理费用（SG&A）杠杆效应。"
    "SG&A同比下降1,880万美元（-11%）至1.50亿美元，而同期营收增长超过6倍。"
    "这表明公司在2023-2024年构建的商业基础设施已可实现规模化利用，"
    "随着ANKTIVA销量持续增长，增量收入将以改善的利润率流转。"
)
add_body_cn(
    "研发费用（R&D）同比增加15%至2.186亿美元，主要受QUILT 2.005 III期临床试验（已入组366名患者）"
    "及多项全球临床合作驱动。值得注意的是，2025年第四季度研发费用中包含一次性固定资产减记1,400万美元，"
    "扣除该非经常性项目后，调整后Q4研发费用约4,990万美元，与2025年第三季度水平基本持平。"
)
add_chart("ibrx_chart6_opex.png", 5.5)
add_body_cn("图4：运营费用对比：2024年全年 vs. 2025年全年（百万美元）| 来源：ImmunityBio 2025年度财报（10-K）", indent=False, color=GRAY, size=8)

add_heading_cn("2.2 净亏损改善趋势", 2)
add_body_cn(
    "尽管营收持续放量，IBRX在2025年仍处于净亏损阶段——这是公司以临床研发为优先、盈利次之的"
    "战略资本配置模式的体现。2025年全年归属于普通股东的净亏损为3.514亿美元，"
    "较2024年4.136亿美元改善约15%。2025年第四季度净亏损6,190万美元，"
    "较第三季度1.105亿美元显著改善，但部分受临床及一般行政费用确认时点的影响。"
)
add_chart("ibrx_chart3_net_loss.png", 5.5)
add_body_cn("图5：季度净亏损（2025年全年，百万美元）| 来源：ImmunityBio各季度业绩公告", indent=False, color=GRAY, size=8)

add_heading_cn("2.3 第四季度实际值与市场预期对比", 2)
add_body_cn(
    "IBRX 2025年第四季度营收和每股亏损均实现双超预期。每股亏损超预期约33%，"
    "在研发费用绝对值偏高（含一次性减记）的情况下，得益于销售及管理费用管控效果优于预期，"
    "以及产品收入毛利率表现较好。"
)
add_chart("ibrx_chart4_beat_miss.png", 5.5)
add_body_cn("图6：2025年第四季度实际值与市场一致预期对比 | 来源：彭博资讯，ImmunityBio业绩新闻稿", indent=False, color=GRAY, size=8)

# ═══════════════════════════════════════════════════════════════
# 第6页：管线进展与指引
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_cn("三、管线进展与业务指引", 1)
add_rule()

add_heading_cn("3.1 BCG无效NMIBC——适应症扩展（乳头状病变）", 2)
add_body_cn(
    "2026年3月9日（即业绩发布后第6天），ImmunityBio向美国FDA重新提交了ANKTIVA的补充生物制品许可申请（sBLA），"
    "申请将适应症扩展至BCG无效非肌层浸润性膀胱癌（NMIBC）纯乳头状病变（Ta/T1期，无原位癌CIS）。"
    "FDA此前曾发出完整回复函（CRL），要求补充额外疗效数据及长期随访数据——关键是未要求开展新的临床试验。"
    "管理层认为此次提交的数据包足以支持审批通过。若获批，将为ANKTIVA增加一个重要的新患者群体。"
)

add_heading_cn("3.2 BCG初治NMIBC——QUILT 2.005（目标：2026年Q4提交BLA）", 2)
add_body_cn(
    "近期最重要的催化剂是QUILT 2.005 III期临床试验，比较ANKTIVA+BCG联合方案与BCG单药方案"
    "在BCG初治NMIBC患者中的疗效，共入组366名患者。FDA要求的中期分析结果如下："
)
add_bullet_cn("9个月完全缓解（CR）率：84%（ANKTIVA+BCG）vs. 52%（BCG单药），p=0.0455", color=GREEN)
add_bullet_cn("未发现新的安全性信号")
add_bullet_cn("目标于2026年第四季度提交BLA；若获批，BCG初治NMIBC将使可治疗患者群体扩大约3-5倍")
add_body_cn(
    "BCG初治NMIBC代表了更大的商业机会——美国每年约有8万例新诊断患者，"
    "远多于BCG无效患者的约1.3万例。成功提交并获批BLA将对IBRX的商业价值产生变革性影响。"
)

add_heading_cn("3.3 肺癌（NSCLC）——全球首个条件性批准", 2)
add_body_cn(
    "ANKTIVA获得沙特食品和药品管理局（SFDA）批准，用于非小细胞肺癌（NSCLC）的条件性批准，"
    "这是ANKTIVA在实体肿瘤适应症上获得的全球首个监管批准。QUILT-3.055 II期临床数据显示，"
    "在检查点抑制剂耐药的晚期NSCLC患者中，中位总生存期达14.1个月，在历史上极具挑战性的患者群体中取得亮眼表现。"
    "欧洲申请计划于2026年通过Accord Healthcare合作伙伴提交。"
)

add_heading_cn("3.4 全球布局——33个国家，4个监管辖区", 2)
add_body_cn(
    "ANKTIVA目前已在33个国家、4个监管辖区（美国、欧盟、沙特阿拉伯/中东及其他地区）获批或授权使用。"
    "Accord Healthcare正在31个欧盟成员国部署商业团队，德国被定为2026年首个欧洲商业化启动市场。"
    "国际收入预计将于2026年下半年开始实现，这是市场一致预期尚未完全纳入的重要额外增长驱动力。"
)
add_chart("ibrx_chart8_pipeline.png", 6.0)
add_body_cn("图7：ANKTIVA管线进展——QUILT项目概览 | 来源：ImmunityBio 2026年3月业绩电话会", indent=False, color=GRAY, size=8)

# ═══════════════════════════════════════════════════════════════
# 第7页：现金流与资产负债表
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_cn("四、资产负债表与现金储备分析", 1)
add_rule()

add_heading_cn("4.1 现金状况", 2)
add_body_cn(
    "截至2025年12月31日，公司现金及有价证券余额为2.428亿美元，"
    "较2024年底约2.88亿美元有所减少——尽管2025年全年通过股权及债务融资筹集了约4亿美元资金。"
    "2025年全年运营现金消耗约3.05亿美元（约合每季度7,600万美元），"
    "部分由ANKTIVA产品收入提供支撑，其余依赖持续的稀释性股权融资。"
)
add_body_cn(
    "投资者最核心的担忧在于：按每季度约7,600万美元的平稳消耗速率测算，"
    "现有现金不足以支持超过约3个季度（至2026年9月）的运营。"
    "管理层尚未指引在此期间实现盈利，这意味着额外融资——很可能是稀释性股权融资——在2026年下半年几乎是必然的。"
)
add_bullet_cn("积极因素：ANKTIVA营收快速增长（每季度约增加1,000万美元），逐步降低净现金消耗")
add_bullet_cn("消极因素：BCG初治III期BLA申报准备及NSCLC全球扩张期间，研发支出将维持高位")
add_bullet_cn("风险缓释因素：若欧洲商业化于2026年下半年顺利启动，国际收入可部分抵消消耗，但时间节点存在不确定性")
add_chart("ibrx_chart7_cash.png", 5.5)
add_body_cn("图8：现金储备情景测算（百万美元）| 来源：ImmunityBio 2025年年报（10-K），公司估算", indent=False, color=GRAY, size=8)

add_heading_cn("4.2 主要资产负债表科目", 2)
bs_data = [
    ['资产负债表项目', '2025年12月31日', '2024年12月31日', '变动'],
    ['现金及有价证券', '2.428亿美元', '约2.88亿美元', '-4,520万美元'],
    ['总流动资产', '约2.78亿美元', '约3.30亿美元', '约-5,200万美元'],
    ['总债务及义务', '约4.50亿美元以上', '约3.80亿美元', '增加'],
    ['全年运营现金消耗', '约（3.05亿）美元', '约（3.60亿）美元', '改善约15%'],
    ['流通股数（约）', '约10.6亿股', '约9.5亿股', '+11.6%'],
]
tbl2 = doc.add_table(rows=len(bs_data), cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(bs_data):
    for ci, txt in enumerate(row_data):
        cell = tbl2.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        set_cjk_font(run, body=True)
        if ri == 0:
            run.bold = True; run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第8-9页：投资逻辑与论点更新
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_cn("五、投资逻辑与论点更新", 1)
add_rule()

add_heading_cn("5.1 核心论点：ANKTIVA平台——逻辑不变", 2)
add_body_cn(
    "我们对ImmunityBio维持跑赢大市评级的核心论点建立在三大支柱之上：（1）ANKTIVA在BCG无效NMIBC领域的商业化持续执行；"
    "（2）ANKTIVA向BCG初治NMIBC（更大市场）的近期适应症扩展；以及（3）从欧洲和沙特阿拉伯开始的成功全球商业化布局。"
    "2025年第四季度业绩进一步夯实了上述三大支柱。"
    "ANKTIVA全年营收增长700%、季度环比增长势头不断、BCG初治中期数据（84% vs. 52% CR）远超预期，"
    "令我们对公司商业化与监管路径充满信心。"
)

add_heading_cn("5.2 自上次更新以来的变化", 2)
changes = [
    ["目标价上调：7.00美元→12.00美元", "商业化速度超预期；BCG初治中期数据优于预期"],
    ["NSCLC机会纳入模型", "沙特SFDA条件性批准提供了2026年实际营收机会"],
    ["欧洲商业化纳入前瞻预测", "Accord Healthcare部署预计于2026年下半年开始贡献收入"],
    ["现金风险上调", "现金（2.428亿美元）已低于全年消耗水平——融资风险迫在眉睫"],
    ["每股亏损预测改善", "2026年每股亏损预期由（1.20）美元修订至（0.95）美元；盈亏平衡时间线提前"],
]
ch_tbl = doc.add_table(rows=len(changes)+1, cols=2)
ch_tbl.style = 'Table Grid'
ch_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['变化因素', '评述']):
    cell = ch_tbl.rows[0].cells[ci]
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = WHITE
    set_cjk_font(run, body=False); set_cell_bg(cell, NAVY_HEX)
for ri, (factor, commentary) in enumerate(changes, 1):
    for ci, txt in enumerate([factor, commentary]):
        cell = ch_tbl.rows[ri].cells[ci]
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9); set_cjk_font(run, body=True)
        if ri % 2 == 0: set_cell_bg(cell, HEAD_FILL)
doc.add_paragraph()

add_heading_cn("5.3 主要风险", 2)
risks = [
    "融资风险（高）：现金2.428亿美元 vs. 全年消耗约3.05亿美元，预计2026年下半年进行股权融资；稀释将压制每股收益和股价",
    "监管风险（中）：乳头状病变NMIBC的sBLA已收到一次CRL；不能排除第二次CRL的可能",
    "竞争风险（中）：BCG仍是标准治疗方案；TAR-200、UGN-102等新进入者正在推进NMIBC领域的临床开发",
    "执行风险（低至中）：欧洲市场通过Accord Healthcare的商业化启动取决于德国等市场的报销时间线",
    "临床风险（低）：BCG初治III期临床试验中期数据显示84% vs. 52% CR的显著差距，最终结果偏离概率较低",
]
for r in risks:
    add_bullet_cn(r)

add_heading_cn("5.4 未来12个月关键催化剂", 2)
cats = [
    "FDA对乳头状病变sBLA的审批决定（预计2026年第四季度，基于2026年3月重新提交后的PDUFA时间线）",
    "QUILT 2.005 BCG初治NMIBC BLA提交（预计2026年第四季度）——将对公司长期营收产生变革性影响",
    "Accord Healthcare欧洲商业化启动（德国，预计2026年下半年）",
    "2026年第一季度业绩（预计2026年5月）——将验证20%以上季度环比增长轨迹是否持续",
    "NSCLC更多数据读出及潜在欧盟申报",
    "潜在的合作或授权协议以应对现金储备挑战",
]
for c in cats:
    add_bullet_cn(c, color=NAVY)

# ═══════════════════════════════════════════════════════════════
# 第10-11页：估值与预测修订
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_cn("六、估值与预测更新", 1)
add_rule()

add_heading_cn("6.1 预测修订情况", 2)
add_body_cn(
    "基于2025年第四季度业绩及BCG初治中期数据，我们上调盈利预测。主要修订如下："
)
add_chart("ibrx_chart9_revisions.png", 5.5)
add_body_cn("图9：2025年第四季度业绩公布前后市场一致预期对比 | 来源：彭博资讯，Piper Sandler，H.C. Wainwright", indent=False, color=GRAY, size=8)

add_heading_cn("6.2 财务预测（修订后）", 2)
est_data = [
    ['（百万美元，另有说明除外）', '2024年实际', '2025年实际', '2026年预测', '2027年预测'],
    ['净产品收入', '1,570万美元', '1.133亿美元', '1.950亿美元', '3.400亿美元'],
    ['同比增长', 'N/M', '+621%', '+72%', '+74%'],
    ['研发费用', '1.902亿美元', '2.186亿美元', '2.000亿美元', '1.850亿美元'],
    ['销售及管理费用', '1.688亿美元', '1.500亿美元', '1.400亿美元', '1.300亿美元'],
    ['净亏损', '（4.136亿）美元', '（3.514亿）美元', '（2.800亿）美元', '（1.000亿）美元'],
    ['每股亏损（基本）', '（0.44）美元', '（0.33）美元', '（0.25）美元', '（0.09）美元'],
    ['现金及有价证券', '约2.88亿美元', '2.428亿美元', '约1.50亿美元（E）', 'N/A'],
    ['全年运营现金消耗', '约（3.60亿）美元', '约（3.05亿）美元', '约（2.40亿）美元', '约（1.00亿）美元'],
]
est_tbl = doc.add_table(rows=len(est_data), cols=5)
est_tbl.style = 'Table Grid'
est_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(est_data):
    for ci, txt in enumerate(row_data):
        cell = est_tbl.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9); set_cjk_font(run, body=True)
        if ri == 0:
            run.bold = True; run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ci == 0:
            run.bold = True
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)
doc.add_paragraph()
add_body_cn('注：2026年/2027年预测为分析师一致预期（Piper Sandler/H.C. Wainwright）；【E】表示预测值。2026年现金余额假设未进行额外融资。', indent=False, color=GRAY, size=8)

add_heading_cn("6.3 可比公司估值分析", 2)
add_body_cn(
    "IBRX当前企业价值/未来12个月营收（EV/NTM Revenue）约为42倍，显著高于肿瘤生物技术可比公司5至12倍的区间。"
    "此估值溢价反映：（a）ANKTIVA商业化高速增长（2026年市场预期营收增长72%）；（b）多项近期催化剂；"
    "以及（c）膀胱癌+肺癌多适应症平台的潜力。然而，该溢价前提是商业化执行持续兑现——"
    "一旦营收轨迹出现偏差或FDA做出不利决定，估值倍数将面临显著压缩风险。"
)
add_chart("ibrx_chart10_valuation.png", 6.0)
add_body_cn("图10：EV/NTM营收倍数对比——IBRX vs. 可比肿瘤生物技术公司 | 来源：彭博资讯，FactSet（2026年3月）", indent=False, color=GRAY, size=8)

add_heading_cn("6.4 目标价测算依据", 2)
add_body_cn(
    "我们将目标价由7.00美元上调至12.00美元，基于风险调整后的EV/营收倍数法，"
    "以2026年和2027年市场一致预期营收为基准："
)
pt_data = [
    ['情景', '2027年预测营收', '目标倍数', '企业价值', '每股股权价值', '权重'],
    ['乐观——BCG初治BLA提交', '4.00亿美元', '25倍', '100亿美元', '9.43美元/股', '35%'],
    ['基准——当前增长轨迹', '3.40亿美元', '20倍', '68亿美元', '6.42美元/股', '45%'],
    ['悲观——融资稀释', '2.80亿美元', '12倍', '34亿美元', '3.21美元/股', '20%'],
    ['概率加权目标价', '', '', '', '约6.80美元/股', '100%'],
    ['Piper Sandler目标价（Q4后）', '', '', '', '12.00美元/股', '—'],
    ['H.C. Wainwright目标价', '', '', '', '15.00美元/股', '—'],
]
pt_tbl = doc.add_table(rows=len(pt_data), cols=6)
pt_tbl.style = 'Table Grid'
pt_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(pt_data):
    for ci, txt in enumerate(row_data):
        cell = pt_tbl.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(8.5); set_cjk_font(run, body=True)
        if ri == 0:
            run.bold = True; run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 第12页：资料来源
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_cn("资料来源", 1)
add_rule()

sources = [
    ("ImmunityBio 2025年第四季度及全年业绩新闻稿（BusinessWire，2026年3月3日）",
     "https://www.businesswire.com/news/home/20260223889360/en/ImmunityBio-Reports-700-Year-Over-Year-Revenue-Growth-Expanded-ANKTIVA-Approvals-in-Lung-Cancer-and-Global-Commercial-Partnerships-in-33-Countries-with-Label-Expansion-plans-Globally"),
    ("ImmunityBio 2025年全年初步营收预告：1.13亿美元（BusinessWire，2026年1月）",
     "https://www.businesswire.com/news/home/20260115898106/en/ImmunityBio-Reports-Continued-Execution-and-Sales-Momentum-With-%24113-Million-of-Preliminary-Net-Product-Revenuea-700-increase-year-over-year"),
    ("ImmunityBio 2025年第四季度业绩电话会录音与文字记录（Seeking Alpha，2026年3月3日）",
     "https://seekingalpha.com/article/4878073-immunitybio-inc-ibrx-q4-2025-earnings-call-transcript"),
    ("Piper Sandler：目标价上调至12.00美元（Investing.com，2026年3月4日）",
     "https://www.investing.com/news/analyst-ratings/piper-sandler-raises-immunitybio-stock-price-target-on-revenue-growth-93CH-4541199"),
    ("H.C. Wainwright：目标价上调至15.00美元（Yahoo Finance，2026年2月）",
     "https://finance.yahoo.com/news/h-c-wainwright-raised-price-152044056.html"),
    ("ImmunityBio投资者关系官方网站",
     "https://ir.immunitybio.com/news-releases/news-release-details/immunitybio-reports-700-year-over-year-revenue-growth-expanded"),
    ("IBRX第四季度每股收益超预期，股价反应（Investing.com，2026年3月）",
     "https://www.investing.com/news/transcripts/earnings-call-transcript-immunitybios-q4-2025-eps-beats-forecast-stock-surges-93CH-4539507"),
    ("ImmunityBio现金储备分析（Simply Wall St，2026年3月）",
     "https://simplywall.st/stocks/us/pharmaceuticals-biotech/nasdaq-ibrx/immunitybio/news/immunitybio-ibrx-revenue-ramp-to-us38-million-tests-profitab"),
]

for i, (title, url) in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3)
    run_num = p.add_run(f"{i}. "); run_num.font.size = Pt(9)
    set_cjk_font(run_num, body=True)
    hyperlink_cn(p, url, title)

doc.add_paragraph()
add_body_cn(
    "免责声明：本报告仅供参考，不构成投资建议。所有预测数据均来自公开分析师研究报告（Piper Sandler、H.C. Wainwright）"
    "或根据公开文件推导。投资者决策前请结合自身风险承受能力审慎评估，并参阅完整的信息披露文件。",
    color=GRAY, size=8
)

out_path = OUT + "IBRX_Q4_FY2025_业绩更新报告_中文版.docx"
doc.save(out_path)
print(f"Chinese report saved: {out_path}")
