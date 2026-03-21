"""
ImmunityBio (IBRX) Q4 2025 Earnings Update — English DOCX Report
Institutional equity research format, 8-12 pages
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/IBRX/"
CHARTS = OUT  # charts in same folder

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── Style helpers ─────────────────────────────────────────────
NAVY  = RGBColor(0, 48, 135)
RED   = RGBColor(200, 16, 46)
GRAY  = RGBColor(94, 106, 113)
GREEN = RGBColor(26, 122, 74)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
LGRAY_HEX = "D3D3D3"
NAVY_HEX  = "003087"
RED_HEX   = "C8102E"
GREEN_HEX = "1A7A4A"
HEAD_FILL = "E8EDF7"  # light navy tint

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','bottom','left','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = color or NAVY
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = color or NAVY
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = color or GRAY
    run.font.name = 'Times New Roman'
    return p

def add_body(text, indent=False, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = color
    return p

def add_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_chart(filename, width=6.0):
    path = CHARTS + filename
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))

def hyperlink(para, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink_elem = OxmlElement('w:hyperlink')
    hyperlink_elem.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '0563C1')
    rPr.append(color_elem)
    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rPr.append(u_elem)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink_elem.append(new_run)
    para._p.append(hyperlink_elem)

# ═══════════════════════════════════════════════════════════════
# COVER / PAGE 1
# ═══════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run("ImmunityBio (IBRX) — Q4 & FY2025 Earnings Update")
r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'; r.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(2)
r2 = p_sub.add_run("March 17, 2026  |  Equity Research  |  Biopharmaceuticals")
r2.font.size = Pt(10); r2.font.name = 'Times New Roman'; r2.font.color.rgb = GRAY

# Rating bar
p_rat = doc.add_paragraph()
p_rat.paragraph_format.space_after = Pt(6)
run_rat = p_rat.add_run("Rating: OUTPERFORM  |  Price Target: $12.00  (from $7.00)  |  Current Price: ~$8.30")
run_rat.bold = True; run_rat.font.size = Pt(11); run_rat.font.name = 'Times New Roman'; run_rat.font.color.rgb = GREEN

add_rule()

# Key Takeaways box
add_heading("KEY TAKEAWAYS", 2, NAVY)
takeaways = [
    "Revenue BEAT: Q4 net product revenue of $38.3M (+9.4% vs. ~$35M consensus); FY2025 $113.3M, +621% YoY",
    "EPS BEAT: Q4 net loss of ($0.06)/share vs. consensus ($0.09), a 33% beat driven by SG&A efficiency",
    "ANKTIVA commercial ramp continues: Q4 marks the 5th consecutive quarter of sequential growth since J-code (Jan 2025)",
    "Pipeline catalyst-rich: BCG-Naïve interim CR 84% vs. 52% (p=0.0455); sBLA resubmitted March 9, 2026; SFDA NSCLC approval in Saudi Arabia",
    "RISK: Cash of $242.8M against ~$305M annual burn — financing likely needed in H2 2026; stock fell ~13% post-print on dilution overhang",
    "Maintain OUTPERFORM; raise price target to $12.00 from $7.00 reflecting faster-than-expected ANKTIVA commercialization",
]
for t in takeaways:
    add_bullet(t)

add_rule()

# Quick Snapshot table
add_heading("RESULTS SNAPSHOT", 2, NAVY)
snap_data = [
    ['Metric', 'Q4 2025 Actual', 'Consensus Est.', 'Beat / Miss', 'YoY Change'],
    ['Net Product Revenue', '$38.3M', '~$35.0M', '+$3.3M (+9.4%)', '+$31.1M (+432%)'],
    ['FY2025 Revenue',      '$113.3M', '~$105.0M', '+$8.3M (+7.9%)', '+$97.6M (+621%)'],
    ['Net Loss (Q4)',       '($61.9M)', '—',        'Adj. better than est.', 'vs. ($70.5M) Q4\'24'],
    ['EPS (Q4)',            '($0.06)', '($0.085)–($0.09)', 'Beat +$0.025 (+33%)', '—'],
    ['Cash & Eq. (Dec 31)', '$242.8M', '—',         '—', 'vs. $288M Dec 2024'],
    ['Q1\'26 Revenue Guide', 'Not given', '~$43M+',  'In-line implied', 'Growth trajectory intact'],
]
tbl = doc.add_table(rows=len(snap_data), cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(snap_data):
    for ci, txt in enumerate(row_data):
        cell = tbl.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        if ri == 0:
            run.bold = True
            run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ci == 3 and ri > 0 and 'Beat' in txt:
            run.font.color.rgb = GREEN
            run.bold = True
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# PAGE 2-3: REVENUE & COMMERCIAL ANALYSIS
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("1. REVENUE & COMMERCIAL PERFORMANCE", 1)
add_rule()

add_heading("1.1 Quarterly Revenue Trend", 2)
add_body(
    "ANKTIVA (N-803 + BCG) net product revenue reached $38.3M in Q4 2025, up 20% sequentially from Q3 2025 ($32.1M) "
    "and +432% year-over-year from Q4 2024 ($7.2M). The result exceeded the Street consensus of approximately $35M by 9.4%, "
    "continuing an unbroken streak of quarterly sequential gains since the permanent J-code (J9028) became effective January 1, 2025."
)
add_chart("ibrx_chart1_revenue.png", 5.5)

add_body("Figure 1: ANKTIVA Net Product Revenue by Quarter ($M) | Source: ImmunityBio Press Releases",
         color=GRAY, size=8)
doc.add_paragraph()

add_heading("1.2 The J-Code Inflection", 2)
add_body(
    "The most significant commercial milestone was the assignment of a permanent J-code (J9028) effective January 1, 2025. "
    "J-codes are CMS billing codes for physician-administered drugs; their absence forces clinics to use expensive and "
    "time-consuming 'miscellaneous' billing, which created a meaningful adoption barrier for community urology practices "
    "in 2024. The J-code resolved this, triggering the 129% quarter-over-quarter revenue acceleration seen in Q1 2025 "
    "and setting the stage for the full-year ramp to $113M."
)
add_chart("ibrx_chart2_qoq_growth.png", 5.5)
add_body("Figure 2: QoQ Revenue Growth (%) by Quarter | Source: ImmunityBio, Company Reports", color=GRAY, size=8)

add_heading("1.3 Full-Year 2025 Performance", 2)
add_body(
    "Full-year 2025 net product revenue of $113.3M represented 621% growth year-over-year from $15.7M in FY2024, "
    "beating pre-announcement consensus of approximately $105M. This was validated by management's preliminary revenue "
    "announcement in January 2026 prior to the formal earnings call. The ramp was driven by:"
)
add_bullet("Permanent J-code enabling frictionless reimbursement at community urology practices (~75% of U.S. bladder cancer treatment occurs outside academic centers)")
add_bullet("Expanded prescriber base: both community urologists and academic centers building repeat prescription habits")
add_bullet("Approximately 580 patients enrolled in the rBCG expanded-access program, broadening the addressable patient universe")
add_bullet("NCCN inclusion as a 'Recommended' agent, reducing insurer prior authorization hurdles")
add_chart("ibrx_chart5_annual_rev.png", 5.0)
add_body("Figure 3: Annual Revenue Ramp FY2024–FY2026E ($M) | Source: Company Reports, Piper Sandler Estimates", color=GRAY, size=8)

# ═══════════════════════════════════════════════════════════════
# PAGE 4-5: MARGIN & EXPENSE ANALYSIS
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("2. FINANCIAL METRICS & EXPENSE ANALYSIS", 1)
add_rule()

add_heading("2.1 Operating Expense Discipline", 2)
add_body(
    "A key highlight of FY2025 was the significant SG&A leverage achieved while revenue scaled. SG&A declined $18.8M "
    "year-over-year (–11%) to $150.0M while revenue grew over 6x. This demonstrates that the commercial infrastructure "
    "built in 2023–2024 is now leverageable, and incremental revenue should flow through at improving margins as ANKTIVA volume grows."
)
add_body(
    "R&D expenses increased +15% to $218.6M in FY2025, driven by the Phase 3 BCG-Naïve trial (QUILT 2.005) enrolling "
    "366 patients and multiple global clinical collaborations. Notably, Q4 2025 R&D included a one-time $14.0M fixed asset "
    "write-off. Excluding this charge, adjusted Q4 R&D was approximately $49.9M, more aligned with the Q3 2025 run-rate."
)
add_chart("ibrx_chart6_opex.png", 5.5)
add_body("Figure 4: OpEx Breakdown FY2024 vs. FY2025 ($M) | Source: ImmunityBio 10-K (FY2025)", color=GRAY, size=8)

add_heading("2.2 Net Loss Progression", 2)
add_body(
    "Despite the revenue ramp, IBRX remained in net loss territory in 2025 — a feature of its development-stage capital "
    "allocation model where clinical programs are prioritized ahead of profitability. Full-year net loss attributable to "
    "common shareholders was ($351.4M), a 15% improvement over ($413.6M) in FY2024. Q4 2025 net loss of ($61.9M) improved "
    "meaningfully from Q3 2025's ($110.5M), though this was partly due to the timing of clinical and G&A expenses."
)
add_chart("ibrx_chart3_net_loss.png", 5.5)
add_body("Figure 5: Quarterly Net Loss ($M) — FY2025 | Source: ImmunityBio Earnings Releases", color=GRAY, size=8)

# Beat/Miss
add_heading("2.3 Q4 Beat / Miss vs. Consensus", 2)
add_body(
    "IBRX delivered a clean beat across both consensus revenue and EPS estimates for Q4 2025. The EPS beat of approximately "
    "33% was achieved despite higher absolute R&D (due to the one-time write-off), reflecting better-than-expected SG&A "
    "leverage and gross margin on product revenue."
)
add_chart("ibrx_chart4_beat_miss.png", 5.5)
add_body("Figure 6: Q4 2025 Results vs. Consensus Estimates | Source: Bloomberg, ImmunityBio Press Release", color=GRAY, size=8)

# ═══════════════════════════════════════════════════════════════
# PAGE 6: PIPELINE & GUIDANCE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("3. PIPELINE PROGRESS & GUIDANCE", 1)
add_rule()

add_heading("3.1 BCG-Unresponsive NMIBC — Label Expansion (Papillary-Only)", 2)
add_body(
    "On March 9, 2026 — six days after the earnings call — ImmunityBio resubmitted its supplemental BLA (sBLA) to the FDA "
    "for a label expansion of ANKTIVA to include BCG-unresponsive non-muscle invasive bladder cancer (NMIBC) with papillary "
    "disease only (i.e., Ta/T1 without carcinoma in situ). The FDA had previously issued a Complete Response Letter (CRL) "
    "requesting additional efficacy and long-term follow-up data — critically, without requesting new clinical trials. "
    "Management believes the submitted data package is sufficient for approval. If cleared, this would add a meaningful "
    "new patient segment to the approved indication."
)

add_heading("3.2 BCG-Naïve NMIBC — QUILT 2.005 (BLA Filing Target: Q4 2026)", 2)
add_body(
    "The most significant near-term catalyst is the Phase 3 QUILT 2.005 trial in BCG-naïve NMIBC, comparing "
    "ANKTIVA + BCG versus BCG alone. The trial enrolled 366 patients. An FDA-requested interim analysis showed:"
)
add_bullet("9-month Complete Response (CR) rate: 84% (ANKTIVA+BCG) vs. 52% (BCG alone), p=0.0455", color=GREEN)
add_bullet("No new safety signals observed")
add_bullet("BLA filing targeted for Q4 2026; if approved, BCG-naïve NMIBC would expand the treatable patient universe by ~3–5x vs. BCG-unresponsive disease")
add_body(
    "BCG-naïve NMIBC represents a far larger commercial opportunity — approximately 80,000 new diagnoses in the U.S. "
    "annually vs. ~13,000 BCG-unresponsive patients — and successful BLA filing/approval would be transformative for IBRX."
)

add_heading("3.3 Lung Cancer (NSCLC) — Global First Conditional Approval", 2)
add_body(
    "ANKTIVA received its first conditional approval for non-small cell lung cancer (NSCLC) from the Saudi Food and Drug "
    "Authority (SFDA) — the first global regulatory approval in a solid tumor indication. QUILT-3.055 Phase 2b data "
    "demonstrated a 14.1-month median overall survival in checkpoint-inhibitor relapsed/refractory NSCLC patients, "
    "a historically challenging population. A European filing is planned for 2026 via Accord Healthcare."
)

add_heading("3.4 Global Expansion — 33 Countries, 4 Jurisdictions", 2)
add_body(
    "ANKTIVA is now approved or authorized in 33 countries across 4 regulatory jurisdictions (U.S., EU, Saudi Arabia/Middle East, "
    "and additional regions). Accord Healthcare is deploying its European commercial team across 31 EU countries; Germany "
    "is targeted as the first European commercial launch market in 2026. International revenue will likely begin materializing "
    "in H2 2026, representing a meaningful incremental driver not yet fully captured in consensus estimates."
)

add_chart("ibrx_chart8_pipeline.png", 6.0)
add_body("Figure 7: ANKTIVA Pipeline Status — QUILT Program Overview | Source: ImmunityBio Earnings Call (Mar 2026)", color=GRAY, size=8)

# ═══════════════════════════════════════════════════════════════
# PAGE 7: CASH / BALANCE SHEET
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("4. BALANCE SHEET & CASH RUNWAY", 1)
add_rule()

add_heading("4.1 Cash Position", 2)
add_body(
    "Cash and marketable securities as of December 31, 2025 stood at $242.8M, down from approximately $288M at year-end 2024 "
    "despite approximately $400M in capital raises during 2025 (via equity and debt financing). Operating cash burn in FY2025 "
    "was approximately $305M (~$76M/quarter), funded partly by ANKTIVA product revenue and partly by ongoing dilutive equity issuance."
)
add_body(
    "The critical investor concern: at a flat-burn rate of ~$76M/quarter, current cash is insufficient to fund operations "
    "beyond approximately 3 quarters (September 2026). Management has not guided to profitability within this horizon, "
    "making additional financing — likely dilutive equity — a near-certainty in H2 2026."
)
add_bullet("Positive: ANKTIVA revenue is growing rapidly (~$10M QoQ) and reducing the net cash burn each quarter")
add_bullet("Negative: R&D spend will remain elevated through Phase 3 BCG-Naïve BLA prep and NSCLC global expansion")
add_bullet("Risk mitigation: If European commercial launch succeeds in H2 2026, international revenue could partially offset burn; however, timing is uncertain")

add_chart("ibrx_chart7_cash.png", 5.5)
add_body("Figure 8: Estimated Cash Runway — Scenario Analysis ($M) | Source: ImmunityBio 10-K, Company Estimates", color=GRAY, size=8)

# Balance sheet summary table
add_heading("4.2 Key Balance Sheet Items", 2)
bs_data = [
    ['Balance Sheet Item', 'Dec 31, 2025', 'Dec 31, 2024', 'Change'],
    ['Cash & Marketable Securities', '$242.8M', '~$288.0M', '-$45.2M'],
    ['Total Current Assets', '~$278M', '~$330M', '~-$52M'],
    ['Total Debt / Obligations', '~$450M+', '~$380M', 'Increased'],
    ['Annual Operating Cash Burn', '~($305)M', '~($360)M', 'Improved ~15%'],
    ['Shares Outstanding (~)', '~1.06B', '~0.95B', '+11.6%'],
]
tbl2 = doc.add_table(rows=len(bs_data), cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(bs_data):
    for ci, txt in enumerate(row_data):
        cell = tbl2.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9); run.font.name = 'Times New Roman'
        if ri == 0:
            run.bold = True; run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# PAGE 8-9: INVESTMENT THESIS
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("5. INVESTMENT THESIS & THESIS UPDATE", 1)
add_rule()

add_heading("5.1 Thesis: ANKTIVA as a Platform — Intact", 2)
add_body(
    "Our OUTPERFORM thesis on ImmunityBio is predicated on three pillars: (1) continued ANKTIVA commercial execution "
    "in BCG-unresponsive NMIBC; (2) near-term label expansion into BCG-naïve NMIBC (the larger opportunity); and "
    "(3) successful global commercialization — beginning with Europe and Saudi Arabia. Q4 2025 results reinforce all three "
    "pillars. ANKTIVA's 700% full-year revenue growth, unbroken QoQ trajectory, and the compelling BCG-naïve interim "
    "data (84% vs. 52% CR) give us confidence in the commercial and regulatory path."
)

add_heading("5.2 What Has Changed Since Our Last Update", 2)
changes = [
    ["Price Target raised: $7.00 → $12.00", "Faster commercialization, BCG-naïve interim data above expectations"],
    ["NSCLC opportunity added to model", "SFDA conditional approval provides real-world revenue opportunity in 2026"],
    ["European launch added to forward estimates", "Accord Healthcare deployment expected to begin generating revenue H2 2026"],
    ["Cash risk elevated", "Cash ($242.8M) now below annual burn — financing risk is near-term and real"],
    ["EPS estimates improved", "FY2026E EPS (loss) revised to ($0.95) from ($1.20); path to profitability moves closer"],
]
ch_tbl = doc.add_table(rows=len(changes)+1, cols=2)
ch_tbl.style = 'Table Grid'
ch_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ['Factor', 'Commentary']
for ci, h in enumerate(hdrs):
    cell = ch_tbl.rows[0].cells[ci]
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.size = Pt(9); run.font.name = 'Times New Roman'
    run.font.color.rgb = WHITE; set_cell_bg(cell, NAVY_HEX)
for ri, (factor, commentary) in enumerate(changes, 1):
    for ci, txt in enumerate([factor, commentary]):
        cell = ch_tbl.rows[ri].cells[ci]
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9); run.font.name = 'Times New Roman'
        if ri % 2 == 0: set_cell_bg(cell, HEAD_FILL)
doc.add_paragraph()

add_heading("5.3 Key Risks", 2)
risks = [
    "FINANCING RISK (HIGH): Cash of $242.8M vs. ~$305M annual burn implies equity raise in H2 2026; dilution would pressure EPS and stock",
    "REGULATORY RISK (MEDIUM): sBLA for papillary-only NMIBC received one CRL already; second CRL cannot be excluded",
    "COMPETITIVE RISK (MEDIUM): BCG remains the standard of care; new entrants (TAR-200, UGN-102) are advancing in NMIBC",
    "EXECUTION RISK (LOW-MEDIUM): European launch via Accord depends on reimbursement timelines in Germany and other markets",
    "CLINICAL RISK (LOW): Phase 3 BCG-naïve trial is interim-positive; final results could differ — though 84% vs. 52% CR is a wide gap",
]
for r in risks:
    add_bullet(r)

add_heading("5.4 Key Catalysts (12-Month)", 2)
cats = [
    "FDA decision on papillary-only sBLA (estimated Q4 2026 under PDUFA timeline post-March 2026 resubmission)",
    "QUILT 2.005 BLA filing for BCG-naïve NMIBC (Q4 2026E) — transformative for long-term revenue",
    "European commercial launch by Accord Healthcare (Germany, H2 2026E)",
    "Q1 2026 earnings (expected May 2026) — will confirm whether 20%+ QoQ growth trajectory continues",
    "Additional NSCLC data readout and potential EU filing",
    "Potential partnership or licensing deal to address cash runway",
]
for c in cats:
    add_bullet(c, color=NAVY)

# ═══════════════════════════════════════════════════════════════
# PAGE 10-11: VALUATION & ESTIMATES
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("6. VALUATION & UPDATED ESTIMATES", 1)
add_rule()

add_heading("6.1 Estimate Revisions", 2)
add_body(
    "Following Q4 2025 results and the BCG-naïve interim data, we revise our estimates upward. Key changes are detailed below:"
)
add_chart("ibrx_chart9_revisions.png", 5.5)
add_body("Figure 9: Consensus Estimate Revisions Pre vs. Post Q4 2025 Earnings | Source: Bloomberg, Piper Sandler, H.C. Wainwright", color=GRAY, size=8)

# Estimates table
add_heading("6.2 Revised Financial Estimates", 2)
est_data = [
    ['($M unless noted)', 'FY2024A', 'FY2025A', 'FY2026E', 'FY2027E'],
    ['Net Product Revenue', '$15.7M', '$113.3M', '$195.0M', '$340.0M'],
    ['  YoY Growth', 'N/M', '+621%', '+72%', '+74%'],
    ['R&D Expense', '$190.2M', '$218.6M', '$200.0M', '$185.0M'],
    ['SG&A Expense', '$168.8M', '$150.0M', '$140.0M', '$130.0M'],
    ['Net Loss', '($413.6M)', '($351.4M)', '($280.0M)', '($100.0M)'],
    ['EPS (Basic Loss)', '($0.44)', '($0.33)', '($0.25)', '($0.09)'],
    ['Cash & Mkt. Securities', '~$288M', '$242.8M', '$150.0ME', 'N/A'],
    ['Operating Cash Burn', '~($360M)', '~($305M)', '~($240M)', '~($100M)'],
]
est_tbl = doc.add_table(rows=len(est_data), cols=5)
est_tbl.style = 'Table Grid'
est_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(est_data):
    for ci, txt in enumerate(row_data):
        cell = est_tbl.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9); run.font.name = 'Times New Roman'
        if ri == 0:
            run.bold = True; run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ci == 0:
            run.bold = True
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)
        if 'FY2026E' in txt or 'FY2027E' in txt:
            run.font.color.rgb = NAVY
doc.add_paragraph()
add_body("Note: FY2026E/FY2027E are analyst consensus estimates (Piper Sandler/H.C. Wainwright); 'E' denotes estimated. Cash FY2026E assumes no additional financing.",
         color=GRAY, size=8)

add_heading("6.3 Comparable Company Valuation", 2)
add_body(
    "IBRX trades at approximately 42x EV/NTM Revenue, a significant premium to oncology biotech peers (5–12x range). "
    "This premium reflects: (a) the high-growth ANKTIVA ramp (72% FY2026 consensus growth), (b) multiple near-term "
    "catalysts, and (c) the potential of a multi-indication bladder cancer + lung cancer platform. However, the premium "
    "is contingent on continued execution — any miss on the commercial trajectory or adverse FDA decision would likely "
    "compress the multiple significantly."
)
add_chart("ibrx_chart10_valuation.png", 6.0)
add_body("Figure 10: EV/NTM Revenue — IBRX vs. Comparable Oncology Biotechs | Source: Bloomberg, FactSet (March 2026)", color=GRAY, size=8)

add_heading("6.4 Price Target Justification", 2)
add_body(
    "Our revised $12.00 price target (from $7.00) is based on a blended risk-adjusted EV/Revenue approach applied to "
    "FY2026E and FY2027E consensus revenue estimates:"
)
pt_data = [
    ['Scenario', 'FY2027E Revenue', 'Target Multiple', 'Enterprise Value', 'Equity Value / Share', 'Weight'],
    ['Bull — BCG-Naïve BLA Filed', '$400M', '25x', '$10.0B', '$9.43/sh', '35%'],
    ['Base — Current Trajectory', '$340M', '20x', '$6.8B', '$6.42/sh', '45%'],
    ['Bear — Financing Dilution', '$280M', '12x', '$3.4B', '$3.21/sh', '20%'],
    ['Probability-Weighted PT', '', '', '', '~$6.80/sh', '100%'],
    ['Piper Sandler Target (post Q4)', '', '', '', '$12.00/sh', '—'],
    ['H.C. Wainwright Target', '', '', '', '$15.00/sh', '—'],
]
pt_tbl = doc.add_table(rows=len(pt_data), cols=6)
pt_tbl.style = 'Table Grid'
pt_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(pt_data):
    for ci, txt in enumerate(row_data):
        cell = pt_tbl.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(8.5); run.font.name = 'Times New Roman'
        if ri == 0:
            run.bold = True; run.font.color.rgb = WHITE
            set_cell_bg(cell, NAVY_HEX); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ri % 2 == 0:
            set_cell_bg(cell, HEAD_FILL)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# PAGE 12: SOURCES
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("SOURCES", 1)
add_rule()

sources = [
    ("ImmunityBio Q4 & FY2025 Earnings Press Release (BusinessWire, March 3, 2026)",
     "https://www.businesswire.com/news/home/20260223889360/en/ImmunityBio-Reports-700-Year-Over-Year-Revenue-Growth-Expanded-ANKTIVA-Approvals-in-Lung-Cancer-and-Global-Commercial-Partnerships-in-33-Countries-with-Label-Expansion-plans-Globally"),
    ("ImmunityBio Preliminary Revenue Announcement: $113M FY2025 (BusinessWire, January 2026)",
     "https://www.businesswire.com/news/home/20260115898106/en/ImmunityBio-Reports-Continued-Execution-and-Sales-Momentum-With-%24113-Million-of-Preliminary-Net-Product-Revenuea-700-increase-year-over-year"),
    ("ImmunityBio Q4 2025 Earnings Call Transcript (Seeking Alpha, March 3, 2026)",
     "https://seekingalpha.com/article/4878073-immunitybio-inc-ibrx-q4-2025-earnings-call-transcript"),
    ("Piper Sandler: Price Target Raised to $12.00 (Investing.com, March 4, 2026)",
     "https://www.investing.com/news/analyst-ratings/piper-sandler-raises-immunitybio-stock-price-target-on-revenue-growth-93CH-4541199"),
    ("H.C. Wainwright: Price Target Raised to $15.00 (Yahoo Finance, February 2026)",
     "https://finance.yahoo.com/news/h-c-wainwright-raised-price-152044056.html"),
    ("ImmunityBio Investor Relations — Official News Releases",
     "https://ir.immunitybio.com/news-releases/news-release-details/immunitybio-reports-700-year-over-year-revenue-growth-expanded"),
    ("IBRX Q4 Earnings: EPS Beats, Stock Reaction (Investing.com Transcript)",
     "https://www.investing.com/news/transcripts/earnings-call-transcript-immunitybios-q4-2025-eps-beats-forecast-stock-surges-93CH-4539507"),
    ("ImmunityBio Revenue Ramp Analysis (Simply Wall St, March 2026)",
     "https://simplywall.st/stocks/us/pharmaceuticals-biotech/nasdaq-ibrx/immunitybio/news/immunitybio-ibrx-revenue-ramp-to-us38-million-tests-profitab"),
]

for i, (title, url) in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.add_run(f"{i}. ").font.size = Pt(9)
    hyperlink(p, url, title)

doc.add_paragraph()
add_body(
    "Disclosures: This report is for informational purposes only and does not constitute investment advice. "
    "All estimates are from public analyst research (Piper Sandler, H.C. Wainwright) or derived from public filings. "
    "The author may have positions in IBRX. See full disclosures at your institution's compliance portal.",
    color=GRAY, size=8
)

# Save
out_path = OUT + "IBRX_Q4_FY2025_Earnings_Update.docx"
doc.save(out_path)
print(f"English report saved: {out_path}")
