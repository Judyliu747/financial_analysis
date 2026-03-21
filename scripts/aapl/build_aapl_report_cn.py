"""
苹果公司 (AAPL) — 2025年日历第四季度 / FQ1 2026 业绩更新
股票研究报告（中文版）DOCX 生成器
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/AAPL"

# ── 颜色定义 ──────────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1C, 0x1C, 0x4D)
C_BLUE   = RGBColor(0x00, 0x71, 0xE3)
C_GREEN  = RGBColor(0x1E, 0x8A, 0x44)
C_RED    = RGBColor(0xCC, 0x00, 0x00)
C_GREY   = RGBColor(0x8E, 0x8E, 0x93)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── 工具函数 ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)

def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "0071E3")
    rPr.append(color_elem)
    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rPr.append(u_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)

def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.name = "SimHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1C1C4D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Pt(20) if not bold else None
    return p

def bullet(doc, text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"[图表：{filename} 未找到]", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
            run.font.name = "SimSun"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        cp.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "1C1C4D")
        set_cell_border(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "SimHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
                run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F5F5F7" if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="E5E5EA")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    if "+" in str(val) and ci > 1:
                        run.font.color.rgb = C_GREEN
                    elif str(val).startswith("−") and ci > 0:
                        run.font.color.rgb = C_RED
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()
    return table

def page_break(doc):
    doc.add_page_break()

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────────────────────────────────────
# 构建报告
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)

doc.styles["Normal"].font.name = "SimSun"
doc.styles["Normal"].font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# 第一页：封面 / 业绩摘要
# ══════════════════════════════════════════════════════════════════════════════
p_ticker = doc.add_paragraph()
p_ticker.clear()
run = p_ticker.add_run("苹果公司（AAPL：纳斯达克）")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_NAVY
run.font.name = "SimHei"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
p_ticker.paragraph_format.space_after = Pt(2)

p_sub = doc.add_paragraph()
p_sub.clear()
run2 = p_sub.add_run("股票研究  |  科技硬件与半导体板块")
run2.font.size = Pt(10)
run2.font.color.rgb = C_GREY
run2.font.name = "SimSun"
run2._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
p_sub.paragraph_format.space_after = Pt(6)

# 评级信息栏
rating_table = doc.add_table(rows=1, cols=5)
rating_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["评级", "目标价", "当前价格", "52周区间", "总市值"]
values = ["跑赢大市 ↑", "$288.30", "$250.12", "$169.21–$288.62", "$3.68万亿"]
colors_v = [C_GREEN, C_BLUE, C_NAVY, C_GREY, C_NAVY]

for i, (lbl, val, col) in enumerate(zip(labels, values, colors_v)):
    cell = rating_table.rows[0].cells[i]
    set_cell_bg(cell, "F0F4FF")
    set_cell_border(cell, color="D0D8F0")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lbl + "\n")
    r1.font.size = Pt(8)
    r1.font.color.rgb = C_GREY
    r1.font.name = "SimSun"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r2 = p1.add_run(val)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = col
    r2.font.name = "SimSun"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

doc.add_paragraph().paragraph_format.space_after = Pt(4)
hr(doc)

p_title = doc.add_paragraph()
p_title.clear()
rt = p_title.add_run(
    "2025年四季报（会计年度Q1 FY2026）业绩更新："
    "全线刷新历史记录，大陆市场强势反弹成最大亮点"
)
rt.font.size = Pt(14)
rt.font.bold = True
rt.font.color.rgb = C_NAVY
rt.font.name = "SimHei"
rt._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
p_title.paragraph_format.space_after = Pt(3)

p_date = doc.add_paragraph()
p_date.clear()
rd = p_date.add_run(
    "报告日期：2026年3月16日  |  财报发布：2026年1月29日  |  财季截止日：2025年12月27日"
)
rd.font.size = Pt(8.5)
rd.font.color.rgb = C_GREY
rd.font.name = "SimSun"
rd._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
p_date.paragraph_format.space_after = Pt(8)
hr(doc)

# 核心业绩汇总表
heading(doc, "核心业绩数据汇总", level=2)
hl_headers = ["指标", "Q1 FY2026实际值", "Q1 FY2025", "同比变动", "市场一致预期", "超预期幅度"]
hl_rows = [
    ["总营收",        "$1,437.7亿", "$1,243.0亿", "+15.7%", "$1,384亿",  "+$53.7亿 / +3.9%"],
    ["摊薄后EPS（美国通用会计准则）", "$2.84", "$2.40", "+18.3%", "$2.68",  "+$0.16 / +6.0%"],
    ["综合毛利率",    "48.2%",       "46.9%",       "+130bps","~47.0%",   "+120bps"],
    ["服务营收",      "$300.1亿",    "$263.4亿",    "+13.9%", "$269亿",   "+$31.1亿 / +11.6%"],
    ["净利润",        "$421.0亿",    "$363.3亿",    "+15.9%", "~$400亿",  "+$21亿"],
    ["经营现金流",    "$539亿",      "~$399亿",     "+35.1%", "N/A",      "史上单季最高"],
]
add_table(doc, hl_headers, hl_rows,
          col_widths=[Inches(1.55), Inches(1.3), Inches(1.1), Inches(0.9), Inches(1.15), Inches(1.45)])

# 核心要点
heading(doc, "核心要点", level=2)
key_bullets = [
    "【全线刷新历史纪录】营收、EPS、毛利率及全部五大产品线均超一致预期；"
    "Q1 FY2026创下总营收、iPhone营收、服务营收、净利润及经营现金流全部历史最高纪录。",

    "【大陆强势反弹是最大亮点】大中华区同比+37.9%（vs. Q1 FY2025同比-13%），"
    "驱动因素为iPhone 17需求爆发及中国消费电子以旧换新补贴政策。管理层明确表示实现市场份额提升。",

    "【服务业务单季突破$300亿门槛】服务营收$300.1亿（同比+14%），延续两位数增长轨迹。"
    "FY2025全年服务营收首次突破$1,000亿。高毛利服务业务（毛利率76.5%）"
    "现已贡献总毛利润的约43%，确认了毛利率结构性上移的投资逻辑。",

    "【毛利率明显扩张（同比+130bps至48.2%）】iPhone 17 Pro/Pro Max机型组合优化"
    "推动产品毛利率升至约40.7%，显著高于近期35-36%区间；服务毛利率同步提升至76.5%。",

    "【苹果智能/AI Siri进展顺利】管理层确认大多数符合条件的iPhone用户已开始使用"
    "Apple Intelligence功能。搭载Google Gemini的全新Siri计划随iOS 26.4于2026年春季发布，"
    "是服务货币化的下一个重要催化剂。",

    "【近期关注点：供应链约束+内存成本上涨】库克表示因iPhone季度同比增长23%致"
    "先进制程产能吃紧，Q2灵活性受限。CFO帕雷克指出内存（DRAM/NAND）价格"
    "因AI数据中心需求显著攀升，相关影响已纳入Q2毛利率指引48%-49%。",

    "【市场反应偏负但基本面强劲】苹果股价在财报发布后次日（1月30日）下跌1.9%，"
    "市场聚焦于供应链约束及内存成本压力。当前股价$250.12，较52周高位折让约13%，"
    "对应分析师一致目标价$288.30，隐含上涨空间+15.3%。",
]
for bk in key_bullets:
    bullet(doc, bk)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第二~三页：详细业绩分析
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "详细业绩分析", level=1)

heading(doc, "各产品线营收", level=2)
add_image(doc, "chart3_segments.png", width=Inches(6.4),
          caption="图1 — 各产品线营收对比：Q1 FY2026 vs. Q1 FY2025\n"
                  "资料来源：苹果公司Q1 FY2026业绩公告，2026年1月29日（apple.com/newsroom）")

seg_headers = ["产品线", "Q1 FY2026", "Q1 FY2025", "同比", "营收占比", "核心看点"]
seg_rows = [
    ["iPhone",      "$852.7亿", "$691.4亿", "+23.3%", "59.3%", "历史最高；iPhone 17结构优化+大陆反弹"],
    ["服务",        "$300.1亿", "$263.4亿", "+13.9%", "20.9%", "历史最高；首次突破$300亿季度门槛"],
    ["可穿戴/家居/配件","$115.0亿","$117.6亿","−2.2%","8.0%",  "AirPods Pro 3供应受限"],
    ["iPad",        "$86.0亿",  "$80.9亿",  "+6.3%",  "6.0%",  "M5芯片iPad Pro新品驱动"],
    ["Mac",         "$84.0亿",  "$90.1亿",  "−6.7%",  "5.8%",  "高基数效应（去年同期M4 Mac密集发布）"],
    ["合计",        "$1,437.7亿","$1,243.0亿","+15.7%","100%",  "公司历史单季最高营收"],
]
add_table(doc, seg_headers, seg_rows,
          col_widths=[Inches(1.2), Inches(0.95), Inches(0.95), Inches(0.75), Inches(0.85), Inches(1.95)])

body(doc,
     "iPhone成为本季度的绝对主角。23.3%的同比增速背后有三重驱动：其一，iPhone 17"
     "全季度上市（去年同期iPhone 16在财季内仅有部分时间贡献）；其二，大陆以旧换新补贴"
     "政策叠加iPhone 17 Pro系列需求旺盛；其三，管理层明确确认实现市场份额提升。"
     "Mac同比下滑6.7%，主要系去年同期M4芯片MacBook Pro、Mac Mini与iMac集中发布"
     "形成的高基数所致，属于结构性原因，而非需求下滑。")

heading(doc, "季度营收走势", level=3)
add_image(doc, "chart1_revenue.png", width=Inches(6.4),
          caption="图2 — 季度营收走势（Q2 FY2024 – Q1 FY2026）\n"
                  "资料来源：苹果公司历次业绩公告（apple.com/newsroom）")

heading(doc, "地区营收——大陆市场成最大正面催化剂", level=2)
add_image(doc, "chart6_geography.png", width=Inches(6.4),
          caption="图3 — 地区营收对比及同比增速（Q1 FY2026 vs. Q1 FY2025）\n"
                  "资料来源：苹果公司Q1 FY2026业绩公告；苹果公司Q1 FY2025业绩公告")

geo_headers = ["地区", "Q1 FY2026", "Q1 FY2025", "同比", "备注"]
geo_rows = [
    ["美洲",        "$585.0亿", "$506.5亿", "+15.5%", "历史最高"],
    ["欧洲",        "$381.0亿", "$338.9亿", "+12.4%", "历史最高"],
    ["大中华区",    "$255.0亿", "$185.1亿", "+37.9%", "最大惊喜；政策补贴+iPhone 17爆发"],
    ["日本",        "~$75亿",   "~$71亿",   "+5.6%",  "历史最高（估算）"],
    ["其他亚太",    "~$142亿",  "~$142亿",  "~0%",    "历史最高（管理层表述）"],
]
add_table(doc, geo_headers, geo_rows,
          col_widths=[Inches(1.3), Inches(1.0), Inches(1.0), Inches(0.85), Inches(2.2)])

body(doc,
     "大中华区同比增长37.9%，实现了对Q1 FY2025（-12.9%）的强势反转，是本季度"
     "最超市场预期的增量贡献因素。库克特别提到在华市场份额提升，与iPhone 17"
     "在高端价位段压制国产旗舰的市场反馈一致。中国政府消费电子以旧换新补贴政策"
     "延续至2026年，形成持续性结构利好。大陆市场的持续超预期表现，"
     "是我们FY2026业绩预测的核心上行风险来源。")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第四~五页：盈利能力与关键指标
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "盈利能力与关键指标", level=1)

heading(doc, "毛利率：产品毛利率重估 + 服务飞轮效应持续", level=2)
add_image(doc, "chart4_margins.png", width=Inches(6.4),
          caption="图4 — 毛利率走势：综合、产品及服务（Q2 FY2024 – Q1 FY2026）\n"
                  "资料来源：苹果公司历次业绩公告；Q1 FY2026 10-Q季报，2026年2月提交（SEC EDGAR）")

body(doc,
     "综合毛利率同比扩张130个基点至48.2%，为近八个季度最高水平。核心驱动因素有二：")
bullet(doc,
       "产品毛利率升至约40.7%（环比+540bps，Q4 FY2025约35.2%）："
       "iPhone 17 Pro/Pro Max机型占比提升、平均售价（ASP）上移是主要驱动力。"
       "Q1为产品毛利率的季节性强季，高端iPhone机型集中出货叠加运营杠杆，"
       "毛利率结构性改善信号明显。", color=C_NAVY)
bullet(doc,
       "服务毛利率升至76.5%（环比+120bps）：服务业务固定成本摊薄效应持续发挥，"
       "服务毛利润约$230亿已占公司总毛利润约43%，"
       "首次超过iPhone毛利润贡献，验证毛利率结构性上移的核心投资逻辑。", color=C_NAVY)
bullet(doc,
       "近期隐患：CFO帕雷克明确提示内存价格（DRAM/NAND同比上涨20-30%，"
       "系AI数据中心需求挤压消费级NAND供给所致）将对Q2 FY2026产品毛利率形成"
       "一定拖累，已计入Q2毛利率指引范围（48%-49%）。", color=C_RED)

heading(doc, "超预期幅度", level=2)
add_image(doc, "chart7_beat_miss.png", width=Inches(6.0),
          caption="图5 — Q1 FY2026实际值 vs. 市场一致预期\n"
                  "资料来源：苹果公司Q1 FY2026业绩公告；彭博一致预期（截至2026年1月28日）")

heading(doc, "EPS走势", level=2)
add_image(doc, "chart2_eps.png", width=Inches(6.4),
          caption="图6 — 摊薄后EPS走势（Q2 FY2024 – Q1 FY2026）\n"
                  "资料来源：苹果公司历次业绩公告；彭博一致预期")

body(doc,
     "摊薄后EPS为$2.84，超一致预期$0.16（+6.0%），同比增速18.3%，净利润达$421亿。"
     "超预期的主要来源：其一，营收超预期部分以约65-70%增量利润率流转至利润端；"
     "其二，毛利率扩张；其三，持续回购稀释摊薄效应——"
     "在外流通股数同比减少约3.5%，贡献EPS约$0.09的同比增量。")

heading(doc, "Q2 FY2026业绩指引分析", level=2)
guide_headers = ["指标", "Q2 FY2026指引", "Q2 FY2025实际", "隐含同比", "分析师解读"]
guide_rows = [
    ["营收增速",    "同比+13%至+16%",    "$953.6亿", "→$1,078亿–$1,107亿", "超此前市场预期（约$1,030亿中值）"],
    ["综合毛利率",  "48.0%–49.0%",        "~47.0%",   "+100至+200bps",      "内存成本压力已纳入区间"],
    ["服务增速",    "约同比+14%",         "约+12%",   "加速",               "Gemini Siri发布带来增量"],
    ["运营费用",    "$153亿–$155亿",      "~$143亿",  "同比约+8%",          "AI研发持续投入，低于营收增速"],
    ["税率",        "约16.5%",            "约16.4%",  "基本稳定",           "正常区间"],
]
add_table(doc, guide_headers, guide_rows,
          col_widths=[Inches(1.15), Inches(1.35), Inches(1.1), Inches(1.1), Inches(1.85)])

body(doc,
     "指引中值约$1,093亿，高于此前华尔街约$1,030亿的市场预期。库克所提"
     "「追供应」表述实为高质量问题——在需求旺盛背景下的产能约束，"
     "而非需求萎缩。随台积电及苹果供应链伙伴先进制程产能于未来2-4个季度"
     "持续扩充，供应瓶颈有望逐步缓解。")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第六~七页：服务业务、AI战略与投资逻辑更新
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "服务业务与苹果智能——结构性投资逻辑", level=1)

heading(doc, "服务业务：季度营收突破$300亿——利润引擎全面激活", level=2)
add_image(doc, "chart5_services.png", width=Inches(6.4),
          caption="图7 — 服务营收增长轨迹（Q1 FY2024 – Q1 FY2026）\n"
                  "资料来源：苹果公司历次业绩公告；苹果公司FY2025年度报告")

body(doc,
     "服务营收已从Q1 FY2024的$231.2亿增长至Q1 FY2026的$300.1亿，"
     "八个季度累计增长29.9%，始终维持约14%的年化增速。$300亿季度里程碑"
     "意义不仅在于数字本身——其背后的毛利润（76.5%毛利率下约$230亿）"
     "已超越iPhone毛利润贡献，连续第二季度成为公司最大毛利来源，"
     "正式确认我们自覆盖启动以来持续强调的毛利率结构性重估逻辑。")
body(doc,
     "FY2025全年服务营收$1,091.6亿——首次突破$1,000亿——证明飞轮模型的可持续性："
     "装机用户规模扩大（现约22亿台活跃设备）× 货币化率提升（订阅提价、"
     "广告收入、App Store抽成、Apple Pay、谷歌搜索授权费）。"
     "单台活跃设备服务营收仍在持续提升，反驳了「规模天花板效应」的市场担忧。")

heading(doc, "Apple Intelligence与谷歌Gemini——Siri进化的下一程", level=2)
body(doc, "Q1 FY2026标志着苹果AI叙事的关键拐点：", bold=True, space_after=2)
bullet(doc,
       "用户渗透率：库克确认大多数符合条件的iPhone用户已开始使用Apple Intelligence功能。"
       "考虑到iPhone 15 Pro、全系iPhone 16及iPhone 17均符合条件，"
       "可寻址设备规模约4-5亿台——意味着苹果AI在发布后约12个月内实现广泛覆盖。")
bullet(doc,
       "谷歌Gemini合作正式确认：苹果将整合Google Gemini，驱动具备「世界知识」"
       "能力的全新Siri，计划随iOS 26.4于2026年春季推出。"
       "据报道该多年合作每年费用约$10亿——对比$1,091亿年度服务营收不过九牛一毛，"
       "但战略意义重大，可有效阻断消费者向竞争性AI原生生态系统的迁移。")
bullet(doc,
       "服务货币化循环：增强版Siri预计将提升App Store互动频次（拉升开发者流量）、"
       "提高订阅转化率（Apple One、Apple Intelligence付费层级开发中），"
       "并加深锁定效应以支撑硬件换机周期。我们测算，这将为苹果"
       "带来FY2028年前约30-50亿美元的服务营收增量机会。")

heading(doc, "投资逻辑更新", level=2)
body(doc, "我们的苹果投资逻辑建立在三大支柱之上，本季度各项均得到强化：", space_after=2)
body(doc, "支柱一：高端硬件×装机用户基础形成可持续服务平台", bold=True, size=10, space_after=2)
body(doc,
     "Q1 FY2026再次验证苹果服务货币化引擎正常运转。iPhone同比+23.3%持续扩大装机基数，"
     "服务同比+14%叠加毛利率扩张（76.5%），印证飞轮模型的复利特性。"
     "至今没有竞争对手在全球范围内复制出苹果「高端硬件+差异化操作系统+深度整合服务」"
     "的三位一体生态体系。")
body(doc, "支柱二：大陆市场——从风险因子转为机遇", bold=True, size=10, space_after=2)
body(doc,
     "大陆+37.9%的强势反转，是本报告中对长期投资逻辑验证价值最高的单一数据点。"
     "经历连续三年的市场担忧（份额流失、监管风险、华为国产竞争），"
     "苹果在全球最大智能手机市场的品牌韧性得到有力证明。"
     "尽管大陆仍存在地缘政治结构性风险，Q1 FY2026数据表明苹果的竞争壁垒"
     "比空头预期的更为深厚。")
body(doc, "支柱三：资本配置纪律驱动EPS持续复利增长", bold=True, size=10, space_after=2)
body(doc,
     "本季度$250亿回购+$39亿分红（合计$290亿），折年化资本回报率约相当于"
     "当前市值的7.7%。$540亿净现金储备叠加约$488亿季度自由现金流，"
     "资产负债表完全自给自足。我们预计苹果将在2026年5月股东大会上"
     "宣布新一期$1,000-$1,200亿回购授权。")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第八~十页：估值与预测更新
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "估值与预测更新", level=1)

heading(doc, "当前估值水平", level=2)
val_headers = ["估值指标", "当前值", "52周低点", "52周高点", "5年均值", "解读"]
val_rows = [
    ["股价",           "$250.12",  "$169.21",  "$288.62",  "~$185",   "较52周高点折让约13%"],
    ["总市值",         "$3.68万亿", "—",       "—",        "—",       "全球第二大市值公司"],
    ["NTM市盈率（P/E）", "32.1x",  "~22x",     "~38x",     "~29.5x",  "较5年均值溢价约8%"],
    ["EV/EBITDA",      "23.7x",    "—",        "—",        "—",       "反映服务业务组合优化"],
    ["EV/营收",        "8.8x",     "—",        "—",        "—",       "高于硬件可比公司均值"],
    ["自由现金流收益率","~5.3%",   "—",        "—",        "—",       "基于约$1,950亿年化FCF"],
    ["股息率",         "~0.4%",    "—",        "—",        "—",       "$0.26/季；规模小但持续增长"],
]
add_table(doc, val_headers, val_rows,
          col_widths=[Inches(1.5), Inches(0.95), Inches(0.9), Inches(0.9), Inches(0.9), Inches(1.55)])

add_image(doc, "chart10_valuation.png", width=Inches(6.4),
          caption="图8 — AAPL NTM市盈率走势（Q2 FY2023 – 2026年3月）\n"
                  "资料来源：彭博NTM一致预期；AAPL每日价格数据。当前价格$250.12（截至2026年3月13日）")

body(doc,
     "当前NTM P/E为32.1x，较5年均值约29.5x溢价约8%。支撑该溢价的逻辑在于："
     "（1）服务占比提升驱动毛利率结构性扩张；"
     "（2）Apple Intelligence提供多年期产品催化剂；"
     "（3）大陆市场复苏降低了核心下行风险情景的概率。"
     "当前股价较2025年12月历史高点（$288.62）低约13.3%，"
     "相对于分析师一致目标价$288.30形成较具吸引力的技术性入场机会。")

heading(doc, "预测更新", level=2)
est_headers = ["财季/财年", "旧营收预测", "新营收预测", "调整幅度", "旧EPS预测", "新EPS预测", "调整幅度"]
est_rows = [
    ["Q2 FY2026E", "$1,030亿", "$1,093亿", "+$63亿 / +6.1%",  "$1.63", "$1.76", "+$0.13 / +8.0%"],
    ["FY2026E",    "$4,035亿", "$4,180亿", "+$145亿 / +3.6%", "$7.05", "$7.45", "+$0.40 / +5.7%"],
    ["FY2027E",    "$4,280亿", "$4,450亿", "+$170亿 / +4.0%", "$7.80", "$8.30", "+$0.50 / +6.4%"],
]
add_table(doc, est_headers, est_rows,
          col_widths=[Inches(1.1), Inches(1.0), Inches(1.0), Inches(1.35), Inches(0.9), Inches(0.9), Inches(1.15)])

body(doc,
     "我们上调FY2026E营收$145亿、EPS $0.40，主要驱动因素：（1）Q1超预期贡献的"
     "基础效应传导；（2）Q2指引高于预期；（3）大中华区假设从全年持平上调至+10%同比；"
     "（4）服务营收增速小幅上修。我们将12个月目标价$288.30建立在"
     "FY2027E EPS $8.30的35倍市盈率基础上——处于苹果历史估值区间上端，"
     "以反映服务驱动的毛利率重估及Apple Intelligence的期权价值。")

heading(doc, "现金流与资本回报", level=2)
add_image(doc, "chart8_cashflow.png", width=Inches(6.2),
          caption="图9 — 经营现金流与估算自由现金流（Q2 FY2024 – Q1 FY2026）\n"
                  "资料来源：苹果公司历次业绩公告；FCF = 经营现金流 - 资本支出（估算）")

add_image(doc, "chart9_buybacks.png", width=Inches(6.2),
          caption="图10 — 资本回报计划：股票回购与分红（FY2021 – FY2025 + Q1 FY2026）\n"
                  "资料来源：苹果公司FY2021-FY2025年度报告；Q1 FY2026业绩公告（2026年1月29日）")

body(doc,
     "Q1 FY2026经营现金流$539亿，创苹果公司单季历史最高纪录。"
     "本季度向股东返还约$290亿（$250亿回购 + $39亿分红），"
     "延续高度自律的资本配置策略。2025年5月宣布的$1,000亿回购授权"
     "尚余约$748亿，我们预计苹果将在2026年5月年度股东大会上"
     "宣布新一轮$1,000-$1,200亿回购计划。"
     "有机自由现金流产生与高效资本回报的结合，"
     "每年从回购机制本身贡献约3-4%的EPS增量复利。")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第十一页：风险因素
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "主要风险因素", level=1)

body(doc, "上行风险：", bold=True, size=10, space_after=2)
for risk in [
    "大中华区持续复苏——若iPhone 17换机周期延续为长期份额提升故事，"
    "FY2026中国营收或超$950亿（vs. 我们基准预测约$850亿），"
    "带来$20-30亿额外营收。",
    "Apple Intelligence/Siri货币化——若推出付费订阅层（据传$4.99-9.99/月），"
    "到FY2028年可为服务营收额外贡献$50-80亿。",
    "服务估值重估——若服务占比升至25%（vs. 当前20.9%），"
    "叠加约77%毛利率，增量组合效益支持NTM P/E向37-40x扩张。",
    "台积电先进制程产能改善——供应约束缓解将驱动Q3 FY2026 iPhone 18"
    "发布季节效应更为顺畅，周期超预期。",
]:
    bullet(doc, risk, color=C_GREEN)

body(doc, "下行风险：", bold=True, size=10, space_after=2)
for risk in [
    "App Store监管风险——美国司法部/欧盟对App Store手续费及侧载要求的持续审查"
    "仍是结构性利空；不利裁决可能使服务毛利率承压100-200bps，"
    "削弱货币化飞轮。",
    "中美地缘政治恶化——贸易摩擦升温或中国政府针对苹果供应商/渠道的行动"
    "可能逆转Q1 FY2026涨幅；若回归Q1 FY2025水平，"
    "年度营收将损失约$70亿。",
    "内存成本持续攀升——若DRAM/NAND价格超出当前预期继续上涨"
    "（AI数据中心需求持续挤压消费级供应），"
    "产品毛利率可能在整个FY2026维持在37-38%以下，限制总毛利率扩张空间。",
    "谷歌Gemini合作受阻——若因监管或竞争因素中断合作，"
    "苹果Siri升级时间表将推迟12-18个月，"
    "与安卓AI能力之间出现竞争缺口。",
    "汇率逆风——苹果约60%营收来自美国以外。"
    "美元升值5%将在不变汇率下削减FY2026E营收约$80-100亿。",
]:
    bullet(doc, risk, color=C_RED)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第十二页：资料来源
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "资料来源", level=1)

heading(doc, "Q1 FY2026业绩相关文件", level=2)
sources = [
    ("苹果公司Q1 FY2026业绩公告（2026年1月29日）",
     "https://www.apple.com/newsroom/2026/01/apple-reports-first-quarter-results/"),
    ("苹果公司Form 10-Q季报（2026年2月提交，SEC EDGAR）",
     "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=AAPL&type=10-Q&dateb=&owner=include&count=40"),
    ("苹果公司Q1 FY2026电话会议记录（2026年1月29日）——The Motley Fool",
     "https://www.fool.com/earnings/call-transcripts/2026/01/29/apple-aapl-q1-2026-earnings-call-transcript/"),
    ("苹果公司Q1 FY2026电话会议记录——Investing.com",
     "https://www.investing.com/news/transcripts/earnings-call-transcript-apple-q1-2026-earnings-beat-expectations-93CH-4474928"),
    ("苹果公司Q1 FY2026业绩报道——CNBC",
     "https://www.cnbc.com/2026/01/29/apple-aapl-earnings-report-q1-2026.html"),
    ("苹果公司Q1 FY2026业绩报道——MacRumors",
     "https://www.macrumors.com/2026/01/29/apple-1q-2026-earnings/"),
]

heading(doc, "FQ4 2025参考数据（2025年9月季度）", level=2)
ref_sources = [
    ("苹果公司Q4 FY2025业绩公告（2025年10月30日）",
     "https://www.apple.com/newsroom/2025/10/apple-reports-fourth-quarter-results/"),
    ("苹果公司Q4 FY2025分析——The Acquirer's Multiple",
     "https://acquirersmultiple.com/2025/11/apple-q4-2025-record-102-5b-revenue-surging-services-and-ai-powered-future/"),
]

heading(doc, "市场数据与估值", level=2)
mkt_sources = [
    ("AAPL市盈率历史数据——FinanceCharts（2026年3月）",
     "https://www.financecharts.com/stocks/AAPL/value/pe-ratio"),
    ("AAPL分析师预测与目标价——MarketBeat（2026年3月）",
     "https://www.marketbeat.com/stocks/NASDAQ/AAPL/forecast/"),
]

heading(doc, "AI / Apple Intelligence相关报道", level=2)
ai_sources = [
    ("苹果Gemini Siri运行机制详解——MacRumors（2026年1月30日）",
     "https://www.macrumors.com/2026/01/30/apple-explains-how-gemini-powered-siri-will-work/"),
    ("苹果AI战略重组与iOS 26.4 Siri计划内幕——彭博社",
     "https://www.bloomberg.com/news/newsletters/2026-01-25/inside-apple-s-ai-shake-up-ai-safari-and-plans-for-new-siri-in-ios-26-4-ios-27-mktqy7xb"),
]

for src_group in [sources, ref_sources, mkt_sources, ai_sources]:
    for display, url in src_group:
        p_src = doc.add_paragraph()
        p_src.paragraph_format.left_indent = Inches(0.25)
        p_src.paragraph_format.space_after = Pt(3)
        p_src.add_run("• ")
        add_hyperlink(p_src, display, url)

doc.add_paragraph()
hr(doc)

disc = doc.add_paragraph()
disc.paragraph_format.space_before = Pt(6)
dr = disc.add_run(
    "免责声明：本报告仅供信息参考，不构成投资建议。财务数据来源于公司公开申报文件及"
    "新闻稿。市场一致预期数据来自彭博（截至2026年1月28日）。市场数据截至2026年3月13日。"
    "所有金额以美元计。预测数据为分析师估算，实际结果可能存在重大差异。"
)
dr.font.size = Pt(8)
dr.font.color.rgb = C_GREY
dr.font.name = "SimSun"
dr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
dr.font.italic = True

# ── 保存 ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(BASE, "AAPL_Q4_CY2025_业绩更新报告_中文版.docx")
doc.save(out_path)
print(f"中文报告已保存 → {out_path}")
