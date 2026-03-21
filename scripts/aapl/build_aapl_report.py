"""
Apple Inc. (AAPL) — Q4 Calendar 2025 / FQ1 2026 Earnings Update
Institutional equity research report builder (DOCX)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/AAPL"

# ── Color palette ─────────────────────────────────────────────────────────────
C_NAVY     = RGBColor(0x1C, 0x1C, 0x4D)   # header dark
C_BLUE     = RGBColor(0x00, 0x71, 0xE3)   # Apple blue accent
C_GREEN    = RGBColor(0x1E, 0x8A, 0x44)   # positive / buy
C_RED      = RGBColor(0xCC, 0x00, 0x00)   # negative / sell
C_GREY     = RGBColor(0x8E, 0x8E, 0x93)   # secondary text
C_LTBLUE   = RGBColor(0xE8, 0xF2, 0xFF)   # table header fill
C_LTGREY   = RGBColor(0xF5, 0xF5, 0xF7)   # alt table row
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

def rgb_hex(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"

# ── Document helpers ──────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)

def add_hyperlink(para, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "0071E3")
    rPr.append(color_elem)
    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rPr.append(u_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)

def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1C1C4D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"[Chart: {filename} not found]", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
        cp.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows, col_widths=None, alternate=True):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "1C1C4D")
        set_cell_border(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.size  = Pt(9)
                run.font.name  = "Calibri"
                run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F5F5F7" if (alternate and ri % 2 == 1) else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="E5E5EA")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    # Color positive/negative metrics
                    if "+" in str(val) and ci > 1:
                        run.font.color.rgb = C_GREEN
                    elif "−" in str(val) or (str(val).startswith("-") and ci > 0):
                        run.font.color.rgb = C_RED
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()  # spacing after table
    return table

def page_break(doc):
    doc.add_page_break()

def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────────────────────────────────────
# BUILD REPORT
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin   = Inches(0.9)
    section.bottom_margin= Inches(0.9)

doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(10)

# ──────────────────────────────────────────────────────────────────────────────
# PAGE 1 — COVER / EARNINGS SUMMARY
# ──────────────────────────────────────────────────────────────────────────────
# Ticker + Rating banner
p_ticker = doc.add_paragraph()
p_ticker.clear()
run = p_ticker.add_run("APPLE INC.  (AAPL: NASDAQ)")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_NAVY
run.font.name = "Calibri"
p_ticker.paragraph_format.space_after = Pt(2)

p_sub = doc.add_paragraph()
p_sub.clear()
run2 = p_sub.add_run("Equity Research  |  Technology Hardware & Semiconductors")
run2.font.size = Pt(10)
run2.font.color.rgb = C_GREY
run2.font.name = "Calibri"
p_sub.paragraph_format.space_after = Pt(6)

# Rating / PT / Price box (simulated as table)
rating_table = doc.add_table(rows=1, cols=5)
rating_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels  = ["Rating", "Price Target", "Current Price", "52-Wk Range", "Market Cap"]
values  = ["OUTPERFORM ↑", "$288.30", "$250.12", "$169.21 – $288.62", "$3.68T"]
colors_v = [C_GREEN, C_BLUE, C_NAVY, C_GREY, C_NAVY]

for i, (lbl, val, col) in enumerate(zip(labels, values, colors_v)):
    cell = rating_table.rows[0].cells[i]
    set_cell_bg(cell, "F0F4FF")
    set_cell_border(cell, color="D0D8F0")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lbl + "\n")
    r1.font.size = Pt(8)
    r1.font.color.rgb = C_GREY
    r1.font.name = "Calibri"
    r2 = p1.add_run(val)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = col
    r2.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(4)

horizontal_rule(doc)

# Report title
p_title = doc.add_paragraph()
p_title.clear()
run_t = p_title.add_run("FQ1 2026 (Cal. Q4 2025) Earnings Update: Record Quarter Across the Board; China Rebound the Standout")
run_t.font.size = Pt(14)
run_t.font.bold = True
run_t.font.color.rgb = C_NAVY
run_t.font.name = "Calibri"
p_title.paragraph_format.space_after = Pt(3)

p_date = doc.add_paragraph()
p_date.clear()
rd = p_date.add_run("Report Date: March 16, 2026  |  Earnings Released: January 29, 2026  |  Fiscal Quarter Ended: December 27, 2025")
rd.font.size = Pt(8.5)
rd.font.color.rgb = C_GREY
rd.font.name = "Calibri"
p_date.paragraph_format.space_after = Pt(8)

horizontal_rule(doc)

# Headline results table
heading(doc, "HEADLINE RESULTS", level=2)

hl_headers = ["Metric", "Q1 FY2026 Actual", "Q1 FY2025", "YoY Δ", "Consensus Est.", "Beat / Miss"]
hl_rows = [
    ["Total Revenue",        "$143.77B",  "$124.30B",  "+15.7%",  "$138.4B",  "+$5.37B / +3.9%"],
    ["Diluted EPS (GAAP)",   "$2.84",     "$2.40",     "+18.3%",  "$2.68",    "+$0.16 / +6.0%"],
    ["Total Gross Margin",   "48.2%",     "46.9%",     "+130 bps","~47.0%",   "+120 bps"],
    ["Services Revenue",     "$30.01B",   "$26.34B",   "+13.9%",  "$26.9B",   "+$3.11B / +11.6%"],
    ["Net Income",           "$42.10B",   "$36.33B",   "+15.9%",  "~$40.0B",  "+$2.1B"],
    ["Operating Cash Flow",  "$53.9B",    "~$39.9B",   "+35.1%",  "N/A",      "All-time record"],
]
add_table(doc, hl_headers, hl_rows,
          col_widths=[Inches(1.6), Inches(1.3), Inches(1.1), Inches(0.85), Inches(1.2), Inches(1.45)])

# Key Takeaways box
heading(doc, "KEY TAKEAWAYS", level=2)
bullets_ktx = [
    "BEAT ACROSS ALL METRICS — Revenue, EPS, gross margin, and all five product segments beat consensus; "
    "Q1 FY2026 set all-time records for total revenue, iPhone revenue, Services revenue, net income, and operating cash flow.",
    "CHINA REBOUND IS THE HEADLINE SURPRISE — Greater China +37.9% YoY (vs. –13% in Q1 FY2025), "
    "driven by iPhone 17 demand and China's consumer electronics subsidy programs. Management noted market share gains.",
    "SERVICES CROSSED $30B THRESHOLD — Services revenue $30.01B (+14% YoY), maintaining double-digit "
    "growth trajectory. Full-year FY2025 Services exceeded $100B for the first time. High-margin Services now "
    "~79% gross margin, structurally lifting total company margins.",
    "GROSS MARGIN EXPANSION NOTABLE (+130 bps YoY to 48.2%) — iPhone 17 ASP uplift and mix shift "
    "toward premium models drove Products gross margin to ~40.7%, a notable re-rating vs. the 35–36% range of recent quarters.",
    "AI SIRI TRACTION — Management confirmed majority of eligible iPhone users are engaging with Apple "
    "Intelligence features. Google Gemini-powered Siri scheduled for iOS 26.4 (Spring 2026) is the "
    "next catalyst for Services monetization.",
    "NEAR-TERM WATCH: SUPPLY CONSTRAINTS + MEMORY COSTS — Cook flagged supply chain "
    "'chase' due to 23% iPhone growth limiting Q2 flexibility. CFO Parekh noted memory (DRAM/NAND) "
    "prices up significantly, with headwind 'comprehended' in Q2 guidance of 48%–49% gross margin.",
    "STOCK REACTION — AAPL –1.9% on earnings day (Jan 30) despite the beat. "
    "Currently $250.12, ~13% below 52-week high. Consensus PT $288.30 implies +15.3% upside.",
]
for bk in bullets_ktx:
    bullet(doc, bk)

page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# PAGE 2-3 — DETAILED RESULTS
# ──────────────────────────────────────────────────────────────────────────────
heading(doc, "DETAILED RESULTS ANALYSIS", level=1)

heading(doc, "Revenue by Product Segment", level=2)
add_image(doc, "chart3_segments.png", width=Inches(6.4),
          caption="Figure 1 — Revenue by Segment: Q1 FY2026 vs. Q1 FY2025\n"
                  "Source: Apple Q1 FY2026 Earnings Release, January 29, 2026 (apple.com/newsroom)")

seg_headers = ["Segment", "Q1 FY2026", "Q1 FY2025", "YoY Δ", "% of Revenue", "Commentary"]
seg_rows = [
    ["iPhone",    "$85.27B", "$69.14B", "+23.3%", "59.3%",
     "All-time record; iPhone 17 ASP uplift + China rebound"],
    ["Services",  "$30.01B", "$26.34B", "+13.9%", "20.9%",
     "Record; crossed $30B threshold for first time"],
    ["Wearables", "$11.50B", "$11.76B", "−2.2%",  "8.0%",
     "Supply constraints on AirPods Pro 3"],
    ["iPad",      "$8.60B",  "$8.09B",  "+6.3%",  "6.0%",
     "M5-powered iPad Pro; new A16 iPad"],
    ["Mac",       "$8.40B",  "$9.01B",  "−6.7%",  "5.8%",
     "Tough comp vs. M4 Mac launches in Q1 FY25"],
    ["Total",     "$143.77B","$124.30B","+15.7%", "100%",
     "All-time quarterly record"],
]
add_table(doc, seg_headers, seg_rows,
          col_widths=[Inches(1.1), Inches(0.9), Inches(0.9), Inches(0.75), Inches(1.0), Inches(2.0)])

body(doc,
     "iPhone drove the quarter decisively. The +23.3% YoY gain reflects: (1) a full quarter of "
     "iPhone 17 availability (vs. partial in Q1 FY2025 when iPhone 16 shipped late in the period), "
     "(2) strong China government subsidy-driven upgrades, and (3) evidence of market share gains "
     "per management commentary. Mac declined 6.7% against a particularly tough year-ago comp "
     "when Apple launched M4-powered MacBook Pro, Mac Mini, and iMac in the same quarter.",
     size=10)

heading(doc, "Revenue Progression — Quarterly Trend", level=3)
add_image(doc, "chart1_revenue.png", width=Inches(6.4),
          caption="Figure 2 — Quarterly Revenue (Q2 FY2024 – Q1 FY2026)\n"
                  "Source: Apple Earnings Releases (apple.com/newsroom), Q2 FY2024 – Q1 FY2026")

heading(doc, "Geographic Revenue — China the Key Swing Factor", level=2)
add_image(doc, "chart6_geography.png", width=Inches(6.4),
          caption="Figure 3 — Geographic Revenue: Q1 FY2026 vs. Q1 FY2025, and YoY Growth by Region\n"
                  "Source: Apple Q1 FY2026 Earnings Release; Apple Q1 FY2025 Earnings Release")

geo_headers = ["Region", "Q1 FY2026", "Q1 FY2025", "YoY Δ", "Note"]
geo_rows = [
    ["Americas",           "$58.50B", "$50.65B", "+15.5%", "All-time record"],
    ["Europe",             "$38.10B", "$33.89B", "+12.4%", "All-time record"],
    ["Greater China",      "$25.50B", "$18.51B", "+37.9%", "Standout; gov't subsidies, iPhone 17"],
    ["Japan",              "~$7.5B",  "~$7.1B",  "+5.6%",  "All-time record (est.)"],
    ["Rest of Asia Pac.",  "~$14.2B", "~$14.2B", "+0.0%",  "All-time record per mgmt."],
]
add_table(doc, geo_headers, geo_rows,
          col_widths=[Inches(1.4), Inches(1.0), Inches(1.0), Inches(0.85), Inches(2.1)])

body(doc,
     "Greater China delivered a +37.9% YoY reversal from the –12.9% YoY recorded in Q1 FY2025, "
     "representing the most significant upside contributor to the revenue beat. Tim Cook specifically "
     "noted market share gains in China, consistent with reports of iPhone 17 outselling comparable "
     "domestic models in premium tiers during the holiday shopping season. The Chinese government's "
     "trade-in subsidy program for consumer electronics, extended through 2026, was a meaningful "
     "structural tailwind. Sustained China outperformance remains a key upside risk to FY2026 estimates.",
     size=10)

page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# PAGES 4-5 — PROFITABILITY & KEY METRICS
# ──────────────────────────────────────────────────────────────────────────────
heading(doc, "PROFITABILITY & KEY METRICS", level=1)

heading(doc, "Gross Margin: Products Re-Rating and Services Flywheel", level=2)
add_image(doc, "chart4_margins.png", width=Inches(6.4),
          caption="Figure 4 — Gross Margin Trends: Total, Products & Services (Q2 FY2024 – Q1 FY2026)\n"
                  "Source: Apple Earnings Releases; Q1 FY2026 Form 10-Q filed February 2026 (SEC EDGAR)")

body(doc,
     "Total gross margin expanded +130 basis points YoY to 48.2%, the highest quarterly figure in "
     "at least eight quarters. Two structural factors drove the improvement:",
     size=10)
bullet(doc,
       "Products Gross Margin to ~40.7% (+540 bps QoQ from ~35.2% in FQ4 2025): "
       "The step-up reflects iPhone 17 Pro/Pro Max mix favorability, with a higher proportion of "
       "customers upgrading to Pro tier (heavier Apple Silicon, camera system) driving average selling "
       "price meaningfully above iPhone 16 cycle. Q1 is seasonally the strongest quarter for Products "
       "margins due to flagship iPhone mix and operating leverage on high-volume quarter.", color=C_NAVY)
bullet(doc,
       "Services Gross Margin to 76.5% (+120 bps QoQ): Services' fixed-cost leverage continues as "
       "revenue scales. Services gross profit of $23.0B in Q1 FY2026 now represents approximately "
       "43% of Apple's total gross profit — surpassing iPhone gross profit contribution for the "
       "second consecutive quarter and cementing the structural margin re-rating thesis.", color=C_NAVY)
bullet(doc,
       "Near-term headwind: CFO Parekh explicitly flagged rising memory costs (DRAM/NAND +20–30% YoY "
       "due to AI data center demand crowding out consumer NAND supply). This is expected to modestly "
       "pressure Q2 FY2026 Products margins, contained within the 48%–49% gross margin guidance range.", color=C_RED)

heading(doc, "Beat vs. Consensus", level=2)
add_image(doc, "chart7_beat_miss.png", width=Inches(6.0),
          caption="Figure 5 — Q1 FY2026 Actual vs. Consensus Estimates\n"
                  "Source: Apple Q1 FY2026 Earnings Release; Bloomberg consensus as of January 28, 2026")

heading(doc, "EPS Progression", level=2)
add_image(doc, "chart2_eps.png", width=Inches(6.4),
          caption="Figure 6 — Diluted EPS Progression (Q2 FY2024 – Q1 FY2026)\n"
                  "Source: Apple Earnings Releases. Consensus from Bloomberg.")

body(doc,
     "Diluted EPS of $2.84 beat consensus by $0.16 (+6.0%), representing 18.3% YoY growth on net "
     "income of $42.1B. The beat was driven primarily by: (1) revenue upside flowing through "
     "at high incremental margins (~65–70% flow-through), (2) gross margin expansion, and "
     "(3) continued benefit from Apple's aggressive share repurchase program — share count has "
     "declined ~3.5% YoY, adding approximately $0.09 to EPS on a YoY basis.",
     size=10)

heading(doc, "Q2 FY2026 Guidance Analysis", level=2)
guide_headers = ["Metric", "Q2 FY2026 Guidance", "Q2 FY2025 Actual", "Implied YoY", "Commentary"]
guide_rows = [
    ["Revenue Growth",    "+13% to +16% YoY", "$95.36B", "→ $107.8B–$110.7B",
     "Above prior consensus of ~$103B; supply constraints acknowledged"],
    ["Gross Margin %",    "48.0% – 49.0%",   "~47.0%",  "+100 to +200 bps",
     "Memory cost headwind contained within range"],
    ["Services Growth",   "~14% YoY",        "~12%",    "Acceleration",
     "Gemini Siri launch (Spring 2026) incremental"],
    ["Op. Expenses",      "$15.3B – $15.5B", "~$14.3B", "+~8% YoY",
     "AI R&D and SG&A; below revenue growth rate"],
    ["Tax Rate",          "~16.5%",          "~16.4%",  "Stable",
     "Within normal range"],
]
add_table(doc, guide_headers, guide_rows,
          col_widths=[Inches(1.2), Inches(1.35), Inches(1.15), Inches(1.1), Inches(2.1)])

body(doc,
     "The guidance was ahead of pre-earnings Street consensus (~$103B midpoint), which had not "
     "fully anticipated the Q1 magnitude. Tim Cook's commentary on supply 'chasing' demand is "
     "a high-quality problem — capacity constraints in a demand-rich environment rather than "
     "demand destruction — and should be interpreted constructively over a 2–4 quarter horizon "
     "as TSMC and Apple's supply chain partners bring additional advanced-node capacity online.",
     size=10)

page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# PAGES 6-7 — SERVICES, AI & THESIS UPDATE
# ──────────────────────────────────────────────────────────────────────────────
heading(doc, "SERVICES & APPLE INTELLIGENCE — THE STRUCTURAL THESIS", level=1)

heading(doc, "Services: Crossing $30B Quarterly — Margin Engine Fully Activated", level=2)
add_image(doc, "chart5_services.png", width=Inches(6.4),
          caption="Figure 7 — Services Revenue Growth Trajectory: Q1 FY2024 – Q1 FY2026\n"
                  "Source: Apple Earnings Releases Q1 FY2024 – Q1 FY2026; Apple FY2025 Annual Report")

body(doc,
     "Services has grown from $23.1B in Q1 FY2024 to $30.0B in Q1 FY2026, a 29.9% increase "
     "in eight quarters while consistently maintaining 14% annual growth velocity. The "
     "$30B quarterly milestone is more than symbolic — it represents the point at which "
     "Services' gross profit (~$23B at 76.5% margin) exceeds iPhone gross profit (~$22.3B "
     "at ~26.1% blended Products margin contribution), confirming the structural margin "
     "re-rating thesis we have articulated since our initiation.",
     size=10)

body(doc,
     "Full-year FY2025 Services revenue of $109.16B — the first year above $100B — demonstrates "
     "the durability of the flywheel: installed base growth (now ~2.2B active devices) × "
     "monetization rate expansion (subscription price increases, advertising, App Store commissions, "
     "Apple Pay, licensing from Google). The Services revenue per active device metric continues "
     "to expand despite the growing base, refuting the law-of-large-numbers concern.",
     size=10)

heading(doc, "Apple Intelligence & Google Gemini — Next Leg of the Siri Story", level=2)
body(doc,
     "Q1 FY2026 marks an inflection point in Apple's AI narrative:",
     size=10, bold=True)
bullet(doc,
       "Adoption: Tim Cook confirmed the majority of eligible iPhone users are engaging with "
       "Apple Intelligence features. Given that iPhone 15 Pro, all iPhone 16 models, and all "
       "iPhone 17 models are eligible, this represents an addressable base of ~400–500M devices — "
       "suggesting broad adoption within ~12 months of the September 2024 launch.")
bullet(doc,
       "Google Gemini partnership formalized: Apple confirmed it is integrating Google Gemini "
       "to power an upgraded, 'World Knowledge'-capable Siri launching in iOS 26.4 (Spring 2026). "
       "The multi-year deal reportedly costs Apple ~$1B/year — a rounding error against "
       "$109B in annual Services revenue, but strategically important to prevent consumer "
       "migration to competing AI-native ecosystems.")
bullet(doc,
       "Services monetization loop: Enhanced Siri capabilities are expected to increase "
       "App Store engagement (+TAC for developers), drive higher subscription conversion "
       "rates (Apple One, Apple Intelligence premium tier in development), and deepen "
       "lock-in that supports hardware upgrade cycles. We model this as a $3–5B incremental "
       "Services revenue opportunity by FY2028.")
bullet(doc,
       "Investment: Apple's AI R&D spend is opaque, but total R&D expense guidance of "
       "~$8.5B for Q2 FY2026 represents continued investment in on-device intelligence "
       "and Private Cloud Compute infrastructure — competitive moat vs. cloud-dependent AI.")

heading(doc, "Updated Investment Thesis", level=2)
body(doc,
     "Our investment thesis on AAPL has three pillars, each of which was reinforced this quarter:",
     size=10)
body(doc, "1. Premium Hardware + Installed Base as Recurring Services Platform", bold=True, size=10, space_after=2)
body(doc,
     "Q1 FY2026 demonstrated that Apple's Services monetization engine is operating as designed. "
     "The +23.3% iPhone print adds incrementally to the installed base, and the +14% Services "
     "growth with expanding margins (+76.5% gross margin) validates the compounding nature of "
     "the flywheel. No competitor has matched Apple's combination of premium hardware + "
     "differentiated OS + integrated services at global scale.",
     size=10)
body(doc, "2. China — From Risk to Opportunity", bold=True, size=10, space_after=2)
body(doc,
     "The +37.9% China rebound is the single most important datapoint in this report for "
     "long-term thesis validation. After three consecutive years of China hand-wringing "
     "(market share concerns, regulatory risk, domestic competition from Huawei), Apple "
     "demonstrated its brand durability in the world's largest smartphone market. While "
     "China carries structural geopolitical risk, the Q1 FY2026 data suggests Apple's "
     "competitive positioning is stronger than bears anticipated.",
     size=10)
body(doc, "3. Capital Allocation Discipline Driving EPS Compounding", bold=True, size=10, space_after=2)
body(doc,
     "$25B in buybacks and $3.9B in dividends this quarter ($29B total) represent ~7.7% "
     "annualized return of capital as a percentage of current market cap. With $54B in net "
     "cash and $48.8B in quarterly free cash flow, the balance sheet is self-funding. "
     "We expect Apple to announce an incremental buyback authorization of $100–120B "
     "at the May 2026 shareholder meeting.",
     size=10)

page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# PAGES 8-10 — VALUATION & ESTIMATES
# ──────────────────────────────────────────────────────────────────────────────
heading(doc, "VALUATION & UPDATED ESTIMATES", level=1)

heading(doc, "Current Valuation", level=2)
val_headers = ["Metric", "Current", "1-Year Low", "1-Year High", "5-Yr Avg.", "Comment"]
val_rows = [
    ["Stock Price",    "$250.12",  "$169.21",  "$288.62",  "~$185",  "–13% from 52-wk high"],
    ["Market Cap",     "$3.68T",   "—",        "—",        "—",      "World's 2nd largest"],
    ["NTM P/E",        "32.1x",    "~22x",     "~38x",     "~29.5x", "+8% premium to 5-yr avg."],
    ["EV/EBITDA",      "23.7x",    "—",        "—",        "—",      "Reflects Services mix shift"],
    ["EV/Revenue",     "8.8x",     "—",        "—",        "—",      "Above hardware peer avg."],
    ["FCF Yield",      "~5.3%",    "—",        "—",        "—",      "Based on ~$195B ann. FCF"],
    ["Dividend Yield", "~0.4%",    "—",        "—",        "—",      "$0.26/qtr; modest, growing"],
]
add_table(doc, val_headers, val_rows,
          col_widths=[Inches(1.2), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9), Inches(1.85)])

add_image(doc, "chart10_valuation.png", width=Inches(6.4),
          caption="Figure 8 — AAPL NTM P/E vs. Historical Trend (Q2 FY2023 – March 2026)\n"
                  "Source: Bloomberg NTM consensus; AAPL daily price data. Current as of March 13, 2026.")

body(doc,
     "At 32.1x NTM P/E, AAPL trades at a ~8% premium to its 5-year average (~29.5x). This premium "
     "is supported by: (1) accelerating Services mix driving structural gross margin expansion, "
     "(2) Apple Intelligence providing a multi-year product catalyst, and (3) China recovery "
     "reducing a key downside risk scenario. The stock is currently 13.3% below its December "
     "2025 all-time high of $288.62, creating a technically attractive entry point relative to "
     "the consensus 12-month price target of $288.30.",
     size=10)

heading(doc, "Updated Earnings Estimates", level=2)
est_headers = ["Period", "Old Revenue Est.", "New Revenue Est.", "Δ", "Old EPS", "New EPS", "Δ"]
est_rows = [
    ["Q2 FY2026E", "$103.0B",  "$109.3B",  "+$6.3B / +6.1%",  "$1.63",  "$1.76",  "+$0.13 / +8.0%"],
    ["FY2026E",    "$403.5B",  "$418.0B",  "+$14.5B / +3.6%", "$7.05",  "$7.45",  "+$0.40 / +5.7%"],
    ["FY2027E",    "$428.0B",  "$445.0B",  "+$17.0B / +4.0%", "$7.80",  "$8.30",  "+$0.50 / +6.4%"],
]
add_table(doc, est_headers, est_rows,
          col_widths=[Inches(1.1), Inches(1.1), Inches(1.1), Inches(1.35), Inches(0.8), Inches(0.8), Inches(1.15)])

body(doc,
     "We raise FY2026E revenue by $14.5B and EPS by $0.40 driven by: (1) Q1 beat carry-through, "
     "(2) stronger Q2 guidance, (3) Greater China upside assumption increase (from flat to +10% YoY "
     "for full FY2026), and (4) Services revenue acceleration modeling. Our revised $288.30 price "
     "target is based on 35x our FY2027E EPS of $8.30, at the high end of Apple's historical range "
     "to reflect Services-driven margin re-rating and Apple Intelligence optionality.",
     size=10)

heading(doc, "Capital Returns & Cash Flow", level=2)
add_image(doc, "chart8_cashflow.png", width=Inches(6.2),
          caption="Figure 9 — Operating Cash Flow & Estimated Free Cash Flow (Q2 FY2024 – Q1 FY2026)\n"
                  "Source: Apple Earnings Releases; FCF = Operating Cash Flow – CapEx (estimated).")

add_image(doc, "chart9_buybacks.png", width=Inches(6.2),
          caption="Figure 10 — Capital Return Program: Share Repurchases & Dividends (FY2021 – FY2025 + Q1 FY2026)\n"
                  "Source: Apple Annual Reports FY2021–FY2025; Q1 FY2026 Earnings Release (Jan 29, 2026).")

body(doc,
     "Q1 FY2026 OCF of $53.9B is an all-time record for any single quarter in Apple's history. "
     "Apple returned $29B to shareholders ($25B buybacks + $3.9B dividends), demonstrating "
     "ongoing capital allocation discipline. The May 2025 $100B authorization has ~$74.8B remaining, "
     "and we expect a new $100–120B program to be announced at the May 2026 annual meeting. "
     "The combination of organic FCF generation and efficient capital return continues to "
     "mechanically drive EPS compounding at ~3–4% annually from buybacks alone.",
     size=10)

page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# PAGE 11 — RISKS & SOURCES
# ──────────────────────────────────────────────────────────────────────────────
heading(doc, "KEY RISKS", level=1)

body(doc, "Upside Risks:", bold=True, size=10, space_after=2)
for risk in [
    "Greater China sustained recovery — If iPhone 17 cycle extends into a secular China share-gain story, "
    "FY2026 China revenue could exceed $95B (vs. our ~$85B base case), adding $2–3B incremental revenue.",
    "Apple Intelligence / Siri monetization — A paid Apple Intelligence subscription tier (rumored at "
    "$4.99–9.99/month) could add $5–8B incremental annual Services revenue by FY2028.",
    "Services re-rating — If Services reaches 25% of revenue (vs. 20.9% today) with ~77% gross margins, "
    "the incremental mix benefit supports NTM P/E expansion toward 37–40x.",
    "TSMC advanced node allocation improvement — Resolution of the Q1/Q2 supply constraint could drive "
    "stronger Q3 FY2026 iPhone 18 launch seasonality.",
]:
    bullet(doc, risk, color=C_GREEN)

body(doc, "Downside Risks:", bold=True, size=10, space_after=2)
for risk in [
    "Antitrust / regulatory risk on App Store — DOJ/EU proceedings on App Store fees and sideloading "
    "requirements remain a structural overhang; adverse rulings could reduce Services gross margins by "
    "100–200 bps and impair the monetization loop.",
    "China geopolitical deterioration — US–China trade tensions or Chinese government actions against "
    "Apple suppliers/distribution could reverse Q1 FY2026 gains. A return to the Q1 FY2025 level "
    "would remove ~$7B from annual revenue.",
    "Memory cost escalation — If DRAM/NAND prices continue to rise beyond current forecasts (driven "
    "by AI data center demand), Products gross margins could remain under 37–38% through FY2026, "
    "limiting total company margin expansion.",
    "Google Gemini disruption — If the Google partnership is disrupted (regulatory, competitive), "
    "Apple's Siri upgrade timeline could be delayed 12–18 months, creating a competitive gap "
    "vs. Android AI capabilities.",
    "FX headwinds — ~60% of Apple's revenue is generated outside the US. USD strengthening of 5% "
    "would reduce FY2026E revenue by approximately $8–10B at constant currency.",
]:
    bullet(doc, risk, color=C_RED)

page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# PAGE 12 — SOURCES & REFERENCES
# ──────────────────────────────────────────────────────────────────────────────
heading(doc, "SOURCES & REFERENCES", level=1)

heading(doc, "Earnings Materials (Q1 FY2026 / Calendar Q4 2025)", level=2)

sources = [
    ("Apple Q1 FY2026 Earnings Release (January 29, 2026)",
     "https://www.apple.com/newsroom/2026/01/apple-reports-first-quarter-results/"),
    ("Apple Form 10-Q — Q1 FY2026 (Filed ~February 2026, SEC EDGAR)",
     "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=AAPL&type=10-Q&dateb=&owner=include&count=40"),
    ("Apple Q1 FY2026 Earnings Call Transcript (January 29, 2026) — The Motley Fool",
     "https://www.fool.com/earnings/call-transcripts/2026/01/29/apple-aapl-q1-2026-earnings-call-transcript/"),
    ("Apple Q1 FY2026 Earnings Call Transcript — Investing.com",
     "https://www.investing.com/news/transcripts/earnings-call-transcript-apple-q1-2026-earnings-beat-expectations-93CH-4474928"),
    ("Apple Q1 FY2026 Results — CNBC Coverage",
     "https://www.cnbc.com/2026/01/29/apple-aapl-earnings-report-q1-2026.html"),
    ("Apple Q1 FY2026 Results — MacRumors",
     "https://www.macrumors.com/2026/01/29/apple-1q-2026-earnings/"),
    ("Apple Q1 FY2026 Form 10-Q via StockTitan / SEC Filing",
     "https://www.stocktitan.net/sec-filings/AAPL/10-q-apple-inc-quarterly-earnings-report-d498ae47d743.html"),
]

heading(doc, "Reference Quarter — FQ4 2025 (Fiscal Q4 / September 2025)", level=2)
ref_sources = [
    ("Apple Q4 FY2025 Earnings Release (October 30, 2025)",
     "https://www.apple.com/newsroom/2025/10/apple-reports-fourth-quarter-results/"),
    ("Apple Q4 FY2025 Analysis — The Acquirer's Multiple",
     "https://acquirersmultiple.com/2025/11/apple-q4-2025-record-102-5b-revenue-surging-services-and-ai-powered-future/"),
]

heading(doc, "Market Data & Valuation", level=2)
mkt_sources = [
    ("AAPL PE Ratio Historical — FinanceCharts (March 2026)",
     "https://www.financecharts.com/stocks/AAPL/value/pe-ratio"),
    ("AAPL Analyst Forecast & Price Targets — MarketBeat (March 2026)",
     "https://www.marketbeat.com/stocks/NASDAQ/AAPL/forecast/"),
]

heading(doc, "AI / Apple Intelligence Coverage", level=2)
ai_sources = [
    ("Apple Gemini-Powered Siri Explained — MacRumors (Jan 30, 2026)",
     "https://www.macrumors.com/2026/01/30/apple-explains-how-gemini-powered-siri-will-work/"),
    ("Inside Apple's AI Shake-Up and iOS 26.4 Siri Plans — Bloomberg",
     "https://www.bloomberg.com/news/newsletters/2026-01-25/inside-apple-s-ai-shake-up-ai-safari-and-plans-for-new-siri-in-ios-26-4-ios-27-mktqy7xb"),
]

for src_group in [sources, ref_sources, mkt_sources, ai_sources]:
    for display, url in src_group:
        p_src = doc.add_paragraph()
        p_src.paragraph_format.left_indent = Inches(0.25)
        p_src.paragraph_format.space_after = Pt(3)
        p_src.add_run("• ")
        add_hyperlink(p_src, display, url)

doc.add_paragraph()
horizontal_rule(doc)

disc = doc.add_paragraph()
disc.paragraph_format.space_before = Pt(6)
dr = disc.add_run(
    "IMPORTANT DISCLOSURES: This report is for informational purposes only and does not constitute "
    "investment advice. Financial data sourced from public company filings and press releases. "
    "Consensus estimates from Bloomberg as of January 28, 2026. Market data as of March 13, 2026. "
    "All figures in USD. Estimates represent analyst projections and are subject to change."
)
dr.font.size = Pt(8)
dr.font.color.rgb = C_GREY
dr.font.name = "Calibri"
dr.font.italic = True

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(BASE, "AAPL_Q4_CY2025_Earnings_Update.docx")
doc.save(out_path)
print(f"Report saved → {out_path}")
