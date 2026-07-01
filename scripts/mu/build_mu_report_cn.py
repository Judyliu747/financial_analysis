"""
Micron Technology (MU) Q3 FY2026 业绩更新报告 -- 中文 DOCX report.
会计季度截止：2026年5月28日 | 财报发布日期：2026年6月24日
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from mu_q3_fy2026_common import (
    BUSINESS_UNITS,
    CONSENSUS,
    ESTIMATES,
    FILING_DATE,
    FILING_DATE,
    GUIDANCE_Q4,
    IBD_URL,
    INVESTOPEDIA_URL,
    LONG_BRIDGE_QUOTE,
    OUT,
    PERIOD_END_CN,
    PREPARED_REMARKS_URL,
    PRESENTATION_URL,
    PRESS_RELEASE_URL,
    QUARTER_CN,
    QUARTERLY_RESULTS_URL,
    RELEASE_DATE,
    REPORT_DATE_CN,
    RESULTS,
    SCAS,
    SOURCES,
    TEN_Q_URL,
)


DOCX = OUT / "MU_Q3_FY2026_业绩更新报告_中文版.docx"
CHARTS = OUT
BLUE = "1A6B8A"
DARK = "111827"
GRAY = "6B7280"
LIGHT_BLUE = "EAF4FA"
LIGHT_GRAY = "F3F4F6"
GREEN = "2E7D52"
RED = "C0392B"


def add_hyperlink(paragraph, url: str, text: str, color: str = BLUE):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run(run, size=10.5, bold=False, italic=False, color: str | None = None, font_name="宋体"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))


def para(doc, text="", size=10.5, bold=False, italic=False, color: str | None = None,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=5, heading_font=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color, font_name="黑体" if heading_font else "宋体")
    return p


def heading(doc, text: str, level=1):
    if level == 1:
        p = para(doc, text, size=13, bold=True, color=BLUE, before=10, after=4, heading_font=True)
        ppr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), BLUE)
        border.append(bottom)
        ppr.append(border)
        return p
    return para(doc, text, size=11, bold=True, color=DARK, before=7, after=3, heading_font=True)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT, size=9):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=color, font_name="黑体" if bold else "宋体")


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(tbl.rows[0].cells[i], BLUE)
        set_cell_text(tbl.rows[0].cells[i], h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
    for r_i, row in enumerate(rows):
        for c_i, value in enumerate(row):
            cell = tbl.rows[r_i + 1].cells[c_i]
            shade(cell, "FFFFFF" if r_i % 2 == 0 else LIGHT_GRAY)
            align = WD_ALIGN_PARAGRAPH.LEFT if c_i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            color = GREEN if any(x in str(value) for x in ["超预期", "+", "买入"]) else RED if "低于" in str(value) else None
            set_cell_text(cell, value, color=color, align=align, size=8.5)
    if widths:
        for i, width in enumerate(widths):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(width)
    return tbl


def source_line(doc, items):
    p = para(doc, "资料来源：", size=7.5, italic=True, color=GRAY, before=2, after=4)
    for idx, (label, url) in enumerate(items):
        if idx:
            r = p.add_run("；")
            set_run(r, size=7.5, italic=True, color=GRAY)
        add_hyperlink(p, url, label)
    return p


def image(doc, filename: str, caption: str, source_items):
    path = CHARTS / filename
    if not path.exists():
        para(doc, f"[缺少图表：{filename}]", color=RED)
        return
    p = para(doc, caption, size=8.5, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2, heading_font=True)
    p.paragraph_format.keep_with_next = True
    doc.add_picture(str(path), width=Inches(5.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    source_line(doc, source_items)


def bullet(doc, title: str, body: str, color=BLUE):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{title} ")
    set_run(r1, bold=True, color=color, font_name="黑体")
    r2 = p.add_run(body)
    set_run(r2)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    header = section.header.paragraphs[0]
    header.text = "美光科技（MU）| FY2026Q3 业绩更新报告"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run(run, size=8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = "仅供参考，不构成投资建议。"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run(run, size=8, color=GRAY)
    return doc


def build():
    doc = setup_doc()

    para(doc, "股票研究 - 业绩更新报告", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "美光科技（Micron Technology, Inc.）| NASDAQ: MU", size=19, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, heading_font=True)
    para(doc, f"{QUARTER_CN}：创纪录季度、长期合同能见度与再次超预期上调指引", size=12,
         italic=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    rating = table(
        doc,
        ["评级", "目标价", "当前股价", "市值", "报告日期"],
        [["买入 / 增持", "$1,600", f"${LONG_BRIDGE_QUOTE['last_close']:,.2f}", f"约${LONG_BRIDGE_QUOTE['market_cap_b']/1000:.2f}万亿", REPORT_DATE_CN]],
        [1.05, 1.2, 1.25, 1.25, 1.25],
    )
    for cell in rating.rows[1].cells:
        shade(cell, LIGHT_BLUE)
    para(doc, f"市场数据：{LONG_BRIDGE_QUOTE['source']}。本地运行环境未安装 yfinance，股价采用 Longbridge CLI 动态报价。",
         size=7.5, italic=True, color=GRAY, after=4)

    heading(doc, "业绩摘要")
    revenue_beat = RESULTS["revenue_b"] - CONSENSUS["revenue_b"]
    eps_beat = RESULTS["non_gaap_eps"] - CONSENSUS["eps"]
    rows = [
        ["营收", f"${RESULTS['revenue_b']:.2f}B", f"${CONSENSUS['revenue_b']:.2f}B", f"超预期 +${revenue_beat:.2f}B（+{revenue_beat / CONSENSUS['revenue_b']:.1%}）", "+346%"],
        ["Non-GAAP EPS", f"${RESULTS['non_gaap_eps']:.2f}", f"${CONSENSUS['eps']:.2f}", f"超预期 +${eps_beat:.2f}（+{eps_beat / CONSENSUS['eps']:.1%}）", "+1,215%"],
        ["Non-GAAP毛利率", "84.9%", "N/A", "公司历史新高", "+45.9个百分点"],
        ["经营现金流", "$25.39B", "N/A", "现金创造创纪录", "+451%"],
        ["FQ4营收指引", "$50.0B +/- $1.0B", "$43.58B", "高于一致预期约15%", "隐含同比+342%"],
    ]
    table(doc, ["指标", "实际 / 指引", "一致预期", "差异", "同比"], rows, [1.3, 1.4, 1.25, 1.75, 1.2])
    source_line(doc, [("美光业绩公告", PRESS_RELEASE_URL), ("IBD / FactSet一致预期", IBD_URL)])

    bullet(doc, "营收超预期幅度巨大且质量高。", "FY2026Q3营收414.6亿美元，高于FactSet一致预期55.5亿美元，环比增长74%，DRAM和NAND均创历史新高。")
    bullet(doc, "利润率改变估值框架。", "Non-GAAP毛利率84.9%，营业利润率81.2%，显示当前存储供需格局已不再是普通周期性复苏，而是短缺经济。")
    bullet(doc, "战略客户协议降低历史周期波动。", "公司披露16份SCA，签署后RPO约1,000亿美元，预计现金存款及相关承诺约220亿美元，其中现金存款约180亿美元。")
    bullet(doc, "维持买入，目标价上调至1,600美元。", "目标价反映FY2026/FY2027 EPS上调、收入能见度提升，以及更高的合同化存储业务估值倍数。")
    doc.add_page_break()

    heading(doc, "详细业绩分析")
    heading(doc, "营收：规模创纪录，环比加速", 2)
    para(doc,
         f"美光在截至{PERIOD_END_CN}的FY2026Q3实现营收414.6亿美元，同比增长346%、环比增长74%。"
         "业绩超预期的原因在于DRAM和NAND同时供需紧张、AI相关产品结构提升，以及价格上涨速度显著快于市场模型。"
         "55.5亿美元的营收超预期幅度本身就相当于去年同期营收的一半以上。")
    image(doc, "mu_chart1_revenue.png", "图1 - 季度营收走势", [("美光业绩公告", PRESS_RELEASE_URL), ("Form 10-Q", TEN_Q_URL), ("FactSet via IBD", IBD_URL)])

    heading(doc, "DRAM与NAND：两个业务均进入产能约束状态", 2)
    para(doc,
         "管理层准备发言显示，DRAM收入313亿美元，同比增长343%、环比增长67%，占总收入76%；NAND收入99亿美元，同比增长361%、环比增长99%，占总收入24%。"
         "DRAM ASP环比低60%区间上涨，NAND ASP环比中80%区间上涨，价格是本季利润弹性的核心。")
    image(doc, "mu_chart4_dram_nand.png", "图2 - DRAM与NAND分产品营收", [("管理层准备发言", PREPARED_REMARKS_URL), ("投资者演示材料", PRESENTATION_URL)])

    heading(doc, "业务单元：四大业务均创纪录", 2)
    bu_rows = [[name, f"${rev:.1f}B", f"{gm}%", f"+{qoq}%", f"+{yoy}%"] for name, rev, gm, qoq, yoy in BUSINESS_UNITS]
    table(doc, ["业务单元", "收入", "毛利率", "环比", "同比"], bu_rows, [1.8, 1.0, 1.0, 0.9, 0.9])
    source_line(doc, [("美光FY2026Q3业绩公告", PRESS_RELEASE_URL), ("投资者演示材料第22-24页", PRESENTATION_URL)])
    image(doc, "mu_chart5_business_units.png", "图3 - 各业务单元收入", [("美光投资者演示材料", PRESENTATION_URL)])
    doc.add_page_break()

    heading(doc, "盈利能力与现金流")
    heading(doc, "毛利率环比提升约10个百分点", 2)
    para(doc,
         "Non-GAAP毛利率84.9%，环比提升约10个百分点，较去年同期翻倍以上。核心驱动是价格，产品结构改善和制造执行也有贡献。"
         "这是美光从商品存储供给过剩转向稀缺经济的最清晰信号。")
    image(doc, "mu_chart3_margins.png", "图4 - 毛利率与营业利润率走势", [("业绩公告", PRESS_RELEASE_URL), ("投资者演示材料", PRESENTATION_URL)])

    heading(doc, "EPS与现金创造均显著超预期", 2)
    para(doc,
         "Non-GAAP EPS为25.11美元，高于FactSet一致预期4.25美元，超预期幅度20%。Non-GAAP营业利润336.8亿美元，对应81.2%营业利润率。"
         "经营现金流253.9亿美元，调整后自由现金流183亿美元，即便资本开支达70.8亿美元，仍创公司纪录。")
    image(doc, "mu_chart2_eps.png", "图5 - Non-GAAP EPS走势", [("业绩公告", PRESS_RELEASE_URL), ("FactSet via IBD", IBD_URL)])
    image(doc, "mu_chart7_cash_flow_capex.png", "图6 - 经营现金流与资本开支", [("管理层准备发言", PREPARED_REMARKS_URL), ("投资者演示材料", PRESENTATION_URL)])
    doc.add_page_break()

    heading(doc, "指引与展望")
    heading(doc, "FQ4指引再次抬高市场预期", 2)
    para(doc,
         "美光给出FQ4指引：营收500亿美元（上下10亿美元）、Non-GAAP毛利率约86%、Non-GAAP营业费用约16.5亿美元、Non-GAAP EPS 31.00美元（上下1.00美元）。"
         "该指引明显高于财报前市场预期的435.8亿美元营收和25.72美元EPS。")
    guide_rows = [
        ["营收", f"${GUIDANCE_Q4['revenue_b_mid']:.1f}B +/- $1.0B", f"${CONSENSUS['q4_revenue_b']:.2f}B", "高于市场约15%"],
        ["Non-GAAP毛利率", "约86%", "N/A", "环比+约110bps"],
        ["Non-GAAP EPS", f"${GUIDANCE_Q4['non_gaap_eps_mid']:.2f} +/- $1.00", f"${CONSENSUS['q4_eps']:.2f}", "高于市场约21%"],
        ["净资本开支", "约$10B", "N/A", "FY2026约$27B"],
    ]
    table(doc, ["指标", "美光FQ4指引", "一致预期", "含义"], guide_rows, [1.4, 1.7, 1.3, 1.7])
    source_line(doc, [("美光业绩公告", PRESS_RELEASE_URL), ("管理层准备发言", PREPARED_REMARKS_URL), ("IBD / FactSet", IBD_URL)])
    image(doc, "mu_chart9_guidance.png", "图7 - Q4指引 vs. 市场预期", [("美光业绩公告", PRESS_RELEASE_URL), ("IBD / FactSet", IBD_URL)])

    heading(doc, "SCA能见度成为新核心变量", 2)
    para(doc,
         "本季最重要的定性变化是战略客户协议（SCA）框架。公司已签署16份协议；FQ3末RPO超过50亿美元，计入季后签约后，基于最低采购量和最低价格，RPO约1,000亿美元。"
         "公司预计相关现金存款和承诺约220亿美元，其中现金存款约180亿美元。")
    image(doc, "mu_chart8_sca_visibility.png", "图8 - 战略客户协议带来的收入能见度", [("管理层准备发言", PREPARED_REMARKS_URL)])
    doc.add_page_break()

    heading(doc, "投资逻辑更新")
    bullet(doc, "逻辑一：AI存储正在成为战略基础设施。", "HBM、先进DRAM和数据中心NAND已不只是周期性零部件，而是AI基础设施的瓶颈资源。")
    bullet(doc, "逻辑二：合同化需求支持更高估值倍数。", "SCA无法完全消除周期性，但最低采购量与最低价格安排能降低未来下行周期深度。")
    bullet(doc, "逻辑三：供给追赶需要几年，而不是几个季度。", "管理层预计DRAM和NAND供需紧张将延续至2027年以后；美国、台湾、新加坡、日本产能建设不会立即释放。")
    bullet(doc, "逻辑四：资本强度上升，但现金流提升更快。", "FY2026资本开支预计约270亿美元，但Q3自由现金流已达183亿美元，且管理层预计FQ4自由现金流继续显著增加。")

    heading(doc, "主要风险", 2)
    bullet(doc, "客户行为风险：", "高存储价格可能促使云厂商、OEM和AI加速器厂商优化内存用量或推迟部署。", RED)
    bullet(doc, "周期风险：", "存储仍是周期行业。如果供给增量快于需求，当前利润率结构可能快速回落。", RED)
    bullet(doc, "地缘政治风险：", "美光全球产能布局和中国相关敞口仍受出口管制、贸易政策和客户限制影响。", RED)
    bullet(doc, "估值风险：", "股价大涨后，市场已经反映了一段较长时间的超常盈利。", RED)
    doc.add_page_break()

    heading(doc, "估值与盈利预测更新")
    heading(doc, "更新后的模型", 2)
    para(doc,
         "我们显著上调FY2026和FY2027预测。FY2026前三季度Non-GAAP EPS已累计约42.09美元，FQ4指引意味着全年Non-GAAP EPS约73美元。"
         "此前模型假设价格更快正常化；更新模型则假设供给紧张延续至2027年，符合管理层表述。")
    est_rows = []
    for metric, old, new, fy27 in ESTIMATES:
        metric_cn = {
            "Revenue ($B)": "营收（十亿美元）",
            "Gross Margin": "毛利率",
            "Non-GAAP EPS": "Non-GAAP EPS",
            "Free Cash Flow ($B)": "自由现金流（十亿美元）",
        }[metric]
        chg = (new / old - 1) if old else 0
        est_rows.append([metric_cn, f"{old:.1f}", f"{new:.1f}", f"+{chg:.0%}", f"{fy27:.1f}"])
    table(doc, ["指标", "旧FY2026E", "新FY2026E", "变化", "FY2027E"], est_rows, [1.65, 1.2, 1.2, 1.0, 1.1])
    source_line(doc, [("公司文件", TEN_Q_URL), ("美光管理层准备发言", PREPARED_REMARKS_URL), ("Codex分析", QUARTERLY_RESULTS_URL)])
    image(doc, "mu_chart10_estimates_valuation.png", "图9 - 盈利预测修订与目标价", [("公司文件", TEN_Q_URL), ("Codex分析", QUARTERLY_RESULTS_URL)])

    heading(doc, "目标价方法", 2)
    para(doc,
         "将示意性目标价从1,250美元上调至1,600美元。该目标价约对应FY2027E Non-GAAP EPS 115美元的13.9倍。"
         "相较普通商品存储周期倍数更高，但SCA能见度、净现金能力改善和AI存储含量结构性提升提供支持。"
         "相对6月30日收盘价1,154.29美元，隐含上行空间约39%。")
    para(doc,
         "若出现以下情形，我们将重新评估评级：（1）SCA现金存款与RPO转化不及预期；（2）Q4毛利率明显低于86%指引；（3）行业供需从“2027年以后仍紧张”转为“2027年内正常化”。")

    image(doc, "mu_chart6_beat_miss.png", "图10 - 核心指标超预期摘要", [("美光业绩公告", PRESS_RELEASE_URL), ("IBD / FactSet", IBD_URL)])
    doc.add_page_break()

    heading(doc, "资料来源与参考文献")
    para(doc, f"最新一期验证：今天为{REPORT_DATE_CN}；最新财报发布日期为{RELEASE_DATE}；10-Q备案日为{FILING_DATE}。该财报处于最近三个月内，所有材料均指向{QUARTER_CN}，季度截止日为{PERIOD_END_CN}。",
         size=9.5, bold=True, color=DARK, heading_font=True)
    cn_labels = {
        "Micron Q3 FY2026 earnings release": "美光FY2026Q3业绩公告",
        "Micron Q3 FY2026 investor presentation": "美光FY2026Q3投资者演示材料",
        "Micron Q3 FY2026 prepared remarks/webcast materials": "美光FY2026Q3管理层准备发言/电话会材料",
        "Micron Q3 FY2026 Form 10-Q": "美光FY2026Q3 Form 10-Q",
        "Micron quarterly results page": "美光季度业绩页面",
        "Investor's Business Daily / FactSet consensus": "Investor's Business Daily / FactSet一致预期",
        "Investopedia / Visible Alpha consensus": "Investopedia / Visible Alpha一致预期",
    }
    for label, url in SOURCES:
        p = para(doc, "", after=3)
        r = p.add_run("- ")
        set_run(r, size=9)
        add_hyperlink(p, url, cn_labels.get(label, label))
    para(doc, f"一致预期：{CONSENSUS['source']}；另以Investopedia / Visible Alpha进行交叉校验。", size=9, italic=True, color=GRAY)
    p = para(doc, "补充一致预期交叉校验：", size=9, italic=True, color=GRAY)
    add_hyperlink(p, INVESTOPEDIA_URL, "Investopedia / Visible Alpha")
    para(doc, "免责声明：本报告仅供参考，不构成投资建议或买卖证券建议。", size=8.5, italic=True, color=GRAY)

    doc.save(DOCX)
    print(f"中文报告已保存：{DOCX}")


if __name__ == "__main__":
    build()
