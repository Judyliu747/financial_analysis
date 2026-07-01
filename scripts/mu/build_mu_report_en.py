"""
Micron Technology (MU) Q3 FY2026 Earnings Update -- English DOCX report.
Quarter ended: May 28, 2026 | Reported: June 24, 2026
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from mu_q3_fy2026_common import (
    BUSINESS_UNITS,
    CONSENSUS,
    ESTIMATES,
    FILING_DATE,
    GUIDANCE_Q4,
    IBD_URL,
    INVESTOPEDIA_URL,
    LONG_BRIDGE_QUOTE,
    OUT,
    PERIOD_END,
    PREPARED_REMARKS_URL,
    PRESENTATION_URL,
    PRESS_RELEASE_URL,
    QUARTER,
    QUARTERLY_RESULTS_URL,
    RELEASE_DATE,
    REPORT_DATE,
    RESULTS,
    SCAS,
    SOURCES,
    TEN_Q_URL,
)


DOCX = OUT / "MU_Q3_FY2026_Earnings_Update.docx"
CHARTS = OUT
BLUE = "1A6B8A"
DARK = "111827"
GRAY = "6B7280"
LIGHT_BLUE = "EAF4FA"
LIGHT_GRAY = "F3F4F6"
GREEN = "2E7D52"
RED = "C0392B"


def add_hyperlink(paragraph, url: str, text: str, color: str = BLUE):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run(run, size=10.5, bold=False, italic=False, color: str | None = None, font_name="Times New Roman"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))


def para(doc, text="", size=10.5, bold=False, italic=False, color: str | None = None,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text: str, level=1):
    if level == 1:
        p = para(doc, text.upper(), size=13, bold=True, color=BLUE, before=10, after=4)
        ppr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), BLUE)
        border.append(bottom)
        ppr.append(border)
        return p
    return para(doc, text, size=11, bold=True, color=DARK, before=7, after=3)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT, size=9):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=color)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(tbl.rows[0].cells[i], BLUE)
        set_cell_text(tbl.rows[0].cells[i], h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
    for r_i, row in enumerate(rows):
        for c_i, value in enumerate(row):
            cell = tbl.rows[r_i + 1].cells[c_i]
            shade(cell, "FFFFFF" if r_i % 2 == 0 else LIGHT_GRAY)
            align = WD_ALIGN_PARAGRAPH.LEFT if c_i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            color = GREEN if any(x in str(value) for x in ["Beat", "+", "BUY"]) else RED if "Miss" in str(value) else None
            set_cell_text(cell, value, color=color, align=align, size=8.5)
    if widths:
        for i, width in enumerate(widths):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(width)
    return tbl


def source_line(doc, items):
    p = para(doc, "Source: ", size=7.5, italic=True, color=GRAY, before=2, after=4)
    for idx, (label, url) in enumerate(items):
        if idx:
            r = p.add_run("; ")
            set_run(r, size=7.5, italic=True, color=GRAY)
        add_hyperlink(p, url, label)
    return p


def image(doc, filename: str, caption: str, source_items):
    path = CHARTS / filename
    if not path.exists():
        para(doc, f"[Missing chart: {filename}]", color=RED)
        return
    p = para(doc, caption, size=8.5, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2)
    p.paragraph_format.keep_with_next = True
    doc.add_picture(str(path), width=Inches(5.85))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    source_line(doc, source_items)


def bullet(doc, title: str, body: str, color=BLUE):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{title} ")
    set_run(r1, bold=True, color=color)
    r2 = p.add_run(body)
    set_run(r2)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"

    header = section.header.paragraphs[0]
    header.text = "Micron Technology (MU) | Q3 FY2026 Earnings Update"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run(run, size=8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = "For informational purposes only. Not investment advice."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run(run, size=8, color=GRAY)
    return doc


def build():
    doc = setup_doc()

    # Page 1
    para(doc, "EQUITY RESEARCH - EARNINGS UPDATE", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "MICRON TECHNOLOGY, INC. (NASDAQ: MU)", size=19, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"{QUARTER}: Record Quarter, SCA Visibility, and Another Beat-and-Raise", size=12,
         italic=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    rating = table(
        doc,
        ["Rating", "Price Target", "Current Price", "Market Cap", "Report Date"],
        [["BUY / OW", "$1,600", f"${LONG_BRIDGE_QUOTE['last_close']:,.2f}", f"~${LONG_BRIDGE_QUOTE['market_cap_b']/1000:.2f}T", REPORT_DATE]],
        [1.05, 1.2, 1.25, 1.25, 1.25],
    )
    for cell in rating.rows[1].cells:
        shade(cell, LIGHT_BLUE)
    para(doc, f"Market data: {LONG_BRIDGE_QUOTE['source']}. yfinance was unavailable in the local runtime; price data was pulled dynamically via Longbridge.",
         size=7.5, italic=True, color=GRAY, after=4)

    heading(doc, "Earnings Summary")
    revenue_beat = RESULTS["revenue_b"] - CONSENSUS["revenue_b"]
    eps_beat = RESULTS["non_gaap_eps"] - CONSENSUS["eps"]
    rows = [
        ["Revenue", f"${RESULTS['revenue_b']:.2f}B", f"${CONSENSUS['revenue_b']:.2f}B", f"Beat +${revenue_beat:.2f}B (+{revenue_beat / CONSENSUS['revenue_b']:.1%})", "+346%"],
        ["Non-GAAP EPS", f"${RESULTS['non_gaap_eps']:.2f}", f"${CONSENSUS['eps']:.2f}", f"Beat +${eps_beat:.2f} (+{eps_beat / CONSENSUS['eps']:.1%})", "+1,215%"],
        ["Non-GAAP Gross Margin", "84.9%", "N/A", "Company record", "+45.9pp"],
        ["Operating Cash Flow", "$25.39B", "N/A", "Record cash generation", "+451%"],
        ["FQ4 Revenue Guide", "$50.0B +/- $1.0B", "$43.58B", "Guide +15% vs. Street", "+342% implied YoY"],
    ]
    table(doc, ["Metric", "Reported / Guide", "Consensus", "Variance", "YoY"], rows, [1.3, 1.4, 1.25, 1.75, 1.2])
    source_line(doc, [("Micron earnings release", PRESS_RELEASE_URL), ("IBD / FactSet consensus", IBD_URL)])

    bullet(doc, "Top-line beat was large and clean.", "Revenue of $41.46B was $5.55B above FactSet consensus and rose 74% sequentially, with DRAM and NAND both setting records.")
    bullet(doc, "Margins changed the valuation conversation.", "Non-GAAP gross margin reached 84.9% and operating margin reached 81.2%, levels that suggest the current memory shortage is not a normal cycle.")
    bullet(doc, "Strategic Customer Agreements reduce historical cyclicality.", "Micron has signed 16 SCAs; signed-agreement RPO is approximately $100B, with $22B of expected cash deposits and related commitments.")
    bullet(doc, "Maintaining BUY; raising price target to $1,600.", "Our target reflects higher FY2026/FY2027 EPS, improved revenue visibility, and a higher warranted multiple for a more contracted memory model.")
    doc.add_page_break()

    # Pages 2-3
    heading(doc, "Detailed Results Analysis")
    heading(doc, "Revenue: record scale and accelerating sequential growth", 2)
    para(doc,
         "Micron reported fiscal Q3 revenue of $41.46B for the quarter ended May 28, 2026, up 346% YoY and 74% QoQ. "
         "The beat was driven by a severe demand/supply imbalance across both DRAM and NAND, a richer AI-driven mix, and pricing that moved far faster than investors had modeled. "
         "The $5.55B revenue beat vs. FactSet was not a rounding issue: it was equivalent to more than half of the revenue Micron generated in the same quarter a year ago.")
    image(doc, "mu_chart1_revenue.png", "Figure 1 - Quarterly revenue progression", [("Micron earnings release", PRESS_RELEASE_URL), ("Form 10-Q", TEN_Q_URL), ("FactSet via IBD", IBD_URL)])

    heading(doc, "DRAM and NAND: both businesses are now capacity-constrained", 2)
    para(doc,
         "Prepared remarks show DRAM revenue of $31.3B, up 343% YoY and 67% QoQ, representing 76% of company revenue. "
         "NAND revenue was $9.9B, up 361% YoY and 99% QoQ, representing 24% of revenue. Management attributed the growth mostly to pricing: DRAM ASPs increased in the low-60% range sequentially, while NAND ASPs increased in the mid-80% range.")
    image(doc, "mu_chart4_dram_nand.png", "Figure 2 - DRAM and NAND revenue by quarter", [("Prepared remarks", PREPARED_REMARKS_URL), ("Investor presentation", PRESENTATION_URL)])

    heading(doc, "Business units: all four reached records", 2)
    bu_rows = [[name, f"${rev:.1f}B", f"{gm}%", f"+{qoq}%", f"+{yoy}%"] for name, rev, gm, qoq, yoy in BUSINESS_UNITS]
    table(doc, ["Business Unit", "Revenue", "Gross Margin", "QoQ", "YoY"], bu_rows, [1.8, 1.0, 1.0, 0.9, 0.9])
    source_line(doc, [("Micron Q3 FY2026 press release", PRESS_RELEASE_URL), ("Investor presentation slides 22-24", PRESENTATION_URL)])
    image(doc, "mu_chart5_business_units.png", "Figure 3 - Revenue by business unit", [("Micron investor presentation", PRESENTATION_URL)])
    doc.add_page_break()

    heading(doc, "Profitability and Cash Flow")
    heading(doc, "Gross margin expanded 10 points sequentially", 2)
    para(doc,
         "Non-GAAP gross margin reached 84.9%, up roughly 10 percentage points sequentially and more than double the year-ago level. The primary driver was pricing, with favorable mix and manufacturing execution also helping. This is the clearest signal that memory has temporarily shifted from commodity oversupply to scarcity economics.")
    image(doc, "mu_chart3_margins.png", "Figure 4 - Gross and operating margin trend", [("Earnings release", PRESS_RELEASE_URL), ("Investor presentation", PRESENTATION_URL)])

    heading(doc, "EPS and cash generation were both far ahead", 2)
    para(doc,
         "Non-GAAP EPS of $25.11 beat FactSet consensus by $4.25, or 20%. Non-GAAP operating income was $33.68B, equivalent to an 81.2% margin. Operating cash flow was $25.39B and adjusted free cash flow was $18.3B, a company record even with net capex of $7.1B.")
    image(doc, "mu_chart2_eps.png", "Figure 5 - Non-GAAP EPS progression", [("Earnings release", PRESS_RELEASE_URL), ("FactSet via IBD", IBD_URL)])
    image(doc, "mu_chart7_cash_flow_capex.png", "Figure 6 - Operating cash flow and net capex", [("Prepared remarks", PREPARED_REMARKS_URL), ("Investor presentation", PRESENTATION_URL)])
    doc.add_page_break()

    # Pages 4-5
    heading(doc, "Guidance and Outlook")
    heading(doc, "FQ4 guide again resets expectations", 2)
    para(doc,
         "For FQ4, Micron guided to revenue of $50.0B +/- $1.0B, non-GAAP gross margin of approximately 86%, non-GAAP operating expenses of approximately $1.65B, and non-GAAP EPS of $31.00 +/- $1.00. This was well above pre-earnings Street expectations of $43.58B revenue and $25.72 EPS.")
    guide_rows = [
        ["Revenue", f"${GUIDANCE_Q4['revenue_b_mid']:.1f}B +/- $1.0B", f"${CONSENSUS['q4_revenue_b']:.2f}B", "+15% vs. Street"],
        ["Non-GAAP Gross Margin", "Approx. 86%", "N/A", "+110bps QoQ"],
        ["Non-GAAP EPS", f"${GUIDANCE_Q4['non_gaap_eps_mid']:.2f} +/- $1.00", f"${CONSENSUS['q4_eps']:.2f}", "+21% vs. Street"],
        ["Net Capex", "~$10B", "N/A", "FY2026 capex ~$27B"],
    ]
    table(doc, ["Metric", "Micron FQ4 Guide", "Consensus", "Implication"], guide_rows, [1.4, 1.7, 1.3, 1.7])
    source_line(doc, [("Micron press release", PRESS_RELEASE_URL), ("Prepared remarks", PREPARED_REMARKS_URL), ("IBD / FactSet", IBD_URL)])
    image(doc, "mu_chart9_guidance.png", "Figure 7 - Q4 guidance vs. Street", [("Micron press release", PRESS_RELEASE_URL), ("IBD / FactSet", IBD_URL)])

    heading(doc, "SCA visibility is the new thesis variable", 2)
    para(doc,
         "The most important qualitative update was the Strategic Customer Agreement framework. Micron disclosed 16 signed agreements. RPO at FQ3 was over $5B; including agreements executed after quarter-end, RPO is approximately $100B based on minimum committed volumes and minimum pricing. Management expects $22B of cash deposits and related commitments, with approximately $18B in cash deposits.")
    image(doc, "mu_chart8_sca_visibility.png", "Figure 8 - Strategic customer agreement visibility", [("Prepared remarks", PREPARED_REMARKS_URL)])
    doc.add_page_break()

    # Pages 6-7
    heading(doc, "Updated Investment Thesis")
    bullet(doc, "Thesis pillar 1 - AI memory has become strategic infrastructure.", "The quarter confirms that HBM, advanced DRAM, and data-center NAND are no longer simply cyclical components. They are becoming bottleneck resources in AI infrastructure.")
    bullet(doc, "Thesis pillar 2 - Contracted demand supports a higher multiple.", "SCAs do not remove cyclicality, but they add minimum-volume, minimum-price visibility that should reduce the depth of future downcycles.")
    bullet(doc, "Thesis pillar 3 - Supply will take years, not quarters, to catch up.", "Prepared remarks state that DRAM and NAND supply-demand conditions are expected to remain tight beyond calendar 2027. New U.S., Taiwan, Singapore, and Japan capacity helps but is not immediate.")
    bullet(doc, "Thesis pillar 4 - Capital intensity rises, but cash flow rises faster.", "FY2026 capex is now expected at roughly $27B, yet Q3 FCF was $18.3B and management expects FQ4 FCF to increase substantially again.")

    heading(doc, "Risks", 2)
    bullet(doc, "Customer behavior risk:", "High memory prices could push hyperscalers, OEMs, and AI accelerator vendors to optimize memory usage or delay deployments.", RED)
    bullet(doc, "Cycle risk:", "Memory remains a cyclical market. If supply additions arrive faster than demand, today’s margin structure could compress quickly.", RED)
    bullet(doc, "Geopolitical risk:", "Micron’s global footprint and China exposure remain vulnerable to export controls, trade policy changes, and customer restrictions.", RED)
    bullet(doc, "Valuation risk:", "After a sharp stock move, the market is already pricing in a significant duration of supernormal earnings.", RED)
    doc.add_page_break()

    # Pages 8-10
    heading(doc, "Valuation and Estimate Revisions")
    heading(doc, "Updated model", 2)
    para(doc,
         "We raise our illustrative FY2026 and FY2027 model materially. Q1-Q3 FY2026 non-GAAP EPS already totals approximately $42.09, and FQ4 guidance implies full-year non-GAAP EPS of roughly $73. The old model was anchored to a faster normalization of pricing; the new model assumes tight supply extends through 2027, in line with management commentary.")
    est_rows = []
    for metric, old, new, fy27 in ESTIMATES:
        chg = (new / old - 1) if old else 0
        est_rows.append([metric, f"{old:.1f}", f"{new:.1f}", f"+{chg:.0%}", f"{fy27:.1f}"])
    table(doc, ["Metric", "Old FY2026E", "New FY2026E", "Change", "FY2027E"], est_rows, [1.65, 1.2, 1.2, 1.0, 1.1])
    source_line(doc, [("Company filings", TEN_Q_URL), ("Micron prepared remarks", PREPARED_REMARKS_URL), ("Codex analysis", QUARTERLY_RESULTS_URL)])
    image(doc, "mu_chart10_estimates_valuation.png", "Figure 9 - Estimate revisions and price target", [("Company filings", TEN_Q_URL), ("Codex analysis", QUARTERLY_RESULTS_URL)])

    heading(doc, "Price target methodology", 2)
    para(doc,
         "We raise the illustrative price target to $1,600 from $1,250. The target applies roughly 13.9x our FY2027E non-GAAP EPS estimate of $115, a premium to a normal commodity-memory multiple but justified by signed SCA visibility, a stronger net-cash balance, and structurally higher AI-memory content. The target implies approximately 39% upside from the June 30 close of $1,154.29.")
    para(doc,
         "We would revisit the rating if either: (1) SCA cash deposits and RPO conversion fail to appear on schedule, (2) Q4 gross margin falls meaningfully below the 86% guide, or (3) industry supply commentary changes from tight beyond 2027 to normalization inside 2027.")

    image(doc, "mu_chart6_beat_miss.png", "Figure 10 - Headline beat/miss summary", [("Micron press release", PRESS_RELEASE_URL), ("IBD / FactSet", IBD_URL)])
    doc.add_page_break()

    # Sources
    heading(doc, "Sources and References")
    para(doc, f"Latest-quarter verification: today is {REPORT_DATE}; the latest earnings release was dated {RELEASE_DATE}; the 10-Q was filed {FILING_DATE}. The release is within the last three months and all materials refer to {QUARTER}, quarter ended {PERIOD_END}.",
         size=9.5, bold=True, color=DARK)
    for label, url in SOURCES:
        p = para(doc, "", after=3)
        r = p.add_run("- ")
        set_run(r, size=9)
        add_hyperlink(p, url, label)
    para(doc, f"Consensus: {CONSENSUS['source']}; additional consensus cross-check from Investopedia / Visible Alpha.", size=9, italic=True, color=GRAY)
    p = para(doc, "Additional consensus cross-check: ", size=9, italic=True, color=GRAY)
    add_hyperlink(p, INVESTOPEDIA_URL, "Investopedia / Visible Alpha")
    para(doc,
         "Disclaimer: This report is for informational purposes only and does not constitute investment advice or a recommendation to buy or sell securities.",
         size=8.5, italic=True, color=GRAY)

    doc.save(DOCX)
    print(f"Report saved: {DOCX}")


if __name__ == "__main__":
    build()
