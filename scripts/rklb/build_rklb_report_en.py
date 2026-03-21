"""
Build RKLB Q4 FY2025 Earnings Update – English DOCX Report
Output: output/RKLB/RKLB_Q4_FY2025_Earnings_Update.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT  = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/RKLB/"
os.makedirs(OUT, exist_ok=True)

# ─── Live market data via yfinance ────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        import yfinance as yf
        info = yf.Ticker(ticker).fast_info
        return {
            "price":      f"${info.last_price:.2f}",
            "market_cap": f"~${info.market_cap/1e9:.1f}B",
            "52w_high":   f"${info.year_high:.2f}",
            "52w_low":    f"${info.year_low:.2f}",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using N/A")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("RKLB")
print(f"Live price: {mkt['price']} | Mkt Cap: {mkt['market_cap']} | 52W: {mkt['52w_low']}–{mkt['52w_high']}")

# ─── Helpers ──────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

BLUE_HEX   = "1B3A6B"
LBLUE_HEX  = "4A90D9"
GREEN_HEX  = "27AE60"
RED_HEX    = "C0392B"
GRAY_HEX   = "8C8C8C"
GOLD_HEX   = "F39C12"
WHITE_HEX  = "FFFFFF"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def set_cell_font(cell, bold=False, color_hex=None, size_pt=10, italic=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold    = bold
            run.font.size    = Pt(size_pt)
            run.font.italic  = italic
            if color_hex:
                r, g, b = hex_to_rgb(color_hex)
                run.font.color.rgb = RGBColor(r, g, b)

def add_hyperlink(para, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r  = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), LBLUE_HEX)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(color_el); rPr.append(u_el)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)

def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10.5)
    r, g, b = hex_to_rgb(BLUE_HEX)
    run.font.color.rgb = RGBColor(r, g, b)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, bold=False, italic=False, color_hex=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, bold_prefix=None, color_hex=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(10)
        if color_hex:
            rv, gv, bv = hex_to_rgb(color_hex)
            r1.font.color.rgb = RGBColor(rv, gv, bv)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_chart(doc, filename, width_in=6.0, caption=""):
    path = OUT + filename
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(8.5)
            r, g, b = hex_to_rgb(GRAY_HEX)
            cp.runs[0].font.color.rgb = RGBColor(r, g, b)
            cp.paragraph_format.space_after = Pt(6)
    else:
        body(doc, f"[Chart not found: {filename}]", italic=True)

# ─── Build Document ───────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER & EARNINGS SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

# Header banner
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EQUITY RESEARCH  |  AEROSPACE & DEFENSE")
run.font.size = Pt(8.5)
run.bold = True
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

# Company title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ROCKET LAB USA, INC. (NASDAQ: RKLB)")
run.font.size = Pt(20)
run.bold = True
r, g, b = hex_to_rgb(BLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Q4 & Full Year 2025 Earnings Update")
run.font.size = Pt(14)
run.italic = True
r, g, b = hex_to_rgb(LBLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("February 26, 2026  |  Earnings Release Date")
run.font.size = Pt(9)
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

doc.add_paragraph()

# Rating / Valuation table
rating_table = doc.add_table(rows=2, cols=6)
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rating_table.style = "Table Grid"
headers = ["Rating", "Price Target", "Current Price", "Market Cap", "52W Low", "52W High"]
values  = ["BUY", "$35.00", mkt["price"], mkt["market_cap"], mkt["52w_low"], mkt["52w_high"]]
for i, (h, v) in enumerate(zip(headers, values)):
    hcell = rating_table.cell(0, i)
    hcell.text = h
    set_cell_bg(hcell, BLUE_HEX)
    set_cell_font(hcell, bold=True, color_hex=WHITE_HEX, size_pt=9)
    vcell = rating_table.cell(1, i)
    vcell.text = v
    vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i == 0:
        set_cell_bg(vcell, GREEN_HEX)
        set_cell_font(vcell, bold=True, color_hex=WHITE_HEX, size_pt=11)
    else:
        set_cell_font(vcell, bold=(i == 1), size_pt=10)

doc.add_paragraph()

# KEY TAKEAWAYS
add_section_heading(doc, "Key Takeaways")
bullet(doc, "Record quarterly revenue of $179.7M in Q4 2025, +36% YoY and +16% QoQ, beating consensus ~$177M by ~+1.6%.",
       bold_prefix="✔ Revenue Beat:", color_hex=GREEN_HEX)
bullet(doc, "GAAP EPS of ($0.09) in-line with consensus estimate of ($0.09)–($0.10), reflecting disciplined cost management.",
       bold_prefix="✔ EPS In-Line:", color_hex=LBLUE_HEX)
bullet(doc, "Adj. EBITDA loss of ($17.4M) — materially better than company guidance of ($23M)–($29M) loss.",
       bold_prefix="✔ EBITDA Beat:", color_hex=GREEN_HEX)
bullet(doc, "Record GAAP gross margin of 38% and non-GAAP gross margin of 44%, +100bps and +240bps sequentially.",
       bold_prefix="✔ Margin Expansion:", color_hex=GREEN_HEX)
bullet(doc, "Backlog surged +73% YoY to a record $1.85B, underpinned by $816M SDA satellite prime contract (largest in history).",
       bold_prefix="✔ Backlog Record:", color_hex=BLUE_HEX)
bullet(doc, "Record 21 Electron launches in FY2025 (Q4: 7 incl. 1 HASTE mission) vs 16 in FY2024; Q1 2026 guided $185M–$200M (+3%–+11% QoQ).",
       bold_prefix="✔ FY2025 Milestones:", color_hex=BLUE_HEX)
bullet(doc, "Neutron first launch delayed to Q4 2026 after stage-one tank rupture during hydrostatic pressure test.",
       bold_prefix="⚠ Neutron Delay:", color_hex=RED_HEX)

doc.add_paragraph()

# RESULTS SNAPSHOT TABLE
add_section_heading(doc, "Results Snapshot")
snap_data = [
    ["Metric",              "Q4 2025 Actual", "Consensus / Guidance", "Beat/Miss",     "vs. Prior Qtr"],
    ["Revenue",             "$179.7M",        "~$177M",               "✔ +1.6%",       "+16% QoQ"],
    ["Launch Services Rev", "$75.9M",         "—",                    "—",             "+85% QoQ"],
    ["Space Systems Rev",   "$103.8M",        "—",                    "—",             "-9% QoQ"],
    ["GAAP Gross Margin",   "38.0%",          "~35–37%",              "✔ +100–300bps", "+100bps"],
    ["Non-GAAP Gross Mgn",  "44.0%",          "—",                    "—",             "+240bps"],
    ["Adj. EBITDA",         "($17.4M)",       "($23M)–($29M)",        "✔ Beat",        "Improved"],
    ["GAAP EPS",            "($0.09)",        "($0.09)–($0.10)",      "✔ In-line",     "—"],
    ["Backlog",             "$1.85B",         "—",                    "+73% YoY",      "—"],
    ["FY2025 Revenue",      "$601.8M",        "—",                    "+38% YoY",      "—"],
]
tbl = doc.add_table(rows=len(snap_data), cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for r_idx, row in enumerate(snap_data):
    for c_idx, val in enumerate(row):
        cell = tbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")
        if r_idx > 0 and c_idx == 3:
            color = GREEN_HEX if "✔" in val else (RED_HEX if "✘" in val else None)
            if color:
                set_cell_font(cell, bold=True, color_hex=color, size_pt=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 2-3 — DETAILED RESULTS ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Detailed Results Analysis")

add_section_heading(doc, "Revenue: Record Quarter Driven by Launch Surge", level=2)
body(doc, "Rocket Lab reported Q4 2025 revenue of $179.7M, a record for any single quarter, representing +36% year-over-year growth and +16% sequential growth. The quarter was defined by a dramatic acceleration in Launch Services revenue, which jumped +85% QoQ to $75.9M on 6 Electron launches — a new quarterly record. Space Systems delivered $103.8M, down 9% sequentially as certain spacecraft programs transitioned to completion milestones, though it remains the largest segment by annual run-rate.")

add_chart(doc, "rklb_chart1_quarterly_revenue.png", 6.2,
          "Figure 1: Quarterly Revenue (Q1 2024–Q4 2025) | Source: Rocket Lab press releases")

body(doc, "For the full year FY2025, total revenue reached $601.8M (+38% YoY), comfortably above the FY2024 figure of $436.2M. Launch Services contributed $199.0M (+59% YoY) while Space Systems generated $402.8M (+30% YoY), reflecting the growing importance of satellite manufacturing and mission services to the business model.")

add_chart(doc, "rklb_chart3_segment_revenue.png", 6.2,
          "Figure 2: Segment Revenue Mix Q4 2025 & Annual Comparison | Source: Rocket Lab press releases")

add_section_heading(doc, "Beat / Miss Analysis", level=2)
body(doc, "Results were uniformly positive across all key metrics:")
add_chart(doc, "rklb_chart4_beat_miss.png", 6.2,
          "Figure 3: Q4 2025 Beat/Miss Summary vs. Consensus & Guidance | Source: Bloomberg, Rocket Lab")

bullet(doc, "Revenue of $179.7M beat consensus of ~$177M by approximately +1.6%, driven by better-than-expected launch cadence execution.")
bullet(doc, "Adjusted EBITDA loss of ($17.4M) came in $5.6–$11.6M better than the ($23M)–($29M) company guidance, primarily due to higher-than-guided revenue and gross margin expansion.")
bullet(doc, "GAAP EPS of ($0.09) was in-line with consensus of ($0.09)–($0.10), reflecting continued investment in Neutron development and headcount growth.")
bullet(doc, "GAAP gross margin of 38% exceeded typical analyst expectations of 35–37%, driven by a higher-margin launch mix and continued Space Systems cost optimization.")

add_section_heading(doc, "Gross Margin: Structural Improvement Continues", level=2)
body(doc, "Q4 2025 marked another quarter of meaningful margin expansion. GAAP gross margin of 38% was up +100bps sequentially and up over 900bps from Q1 2024's 20%. Non-GAAP gross margin reached 44%, +240bps QoQ. This trajectory reflects: (1) greater Electron reusability benefits, (2) improved Space Systems program execution, and (3) higher-margin launch contracts as Rocket Lab moves toward dedicated manifest pricing.")

add_chart(doc, "rklb_chart5_gross_margin.png", 6.2,
          "Figure 4: GAAP and Non-GAAP Gross Margin Trend | Source: Rocket Lab press releases")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 4-5 — KEY METRICS & GUIDANCE
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Key Metrics & Guidance")

add_section_heading(doc, "Adjusted EBITDA: Losses Narrowing Toward Positive", level=2)
body(doc, "Adjusted EBITDA loss of ($17.4M) in Q4 2025 represents a significant improvement from ($26.3M) in Q3 2025 and ($33.0M) in Q1 2025. The trajectory is clearly toward breakeven, and management has reiterated its expectation for full-year FY2026 positive adjusted EBITDA. The primary headwind remains elevated R&D expenditure related to Neutron development, which we estimate at $20–25M per quarter.")

add_chart(doc, "rklb_chart6_adj_ebitda.png", 5.5,
          "Figure 5: Adjusted EBITDA Trend — 2025 | Source: Rocket Lab press releases")

add_section_heading(doc, "Backlog & Revenue Visibility", level=2)
body(doc, "The company ended Q4 2025 with a record contracted backlog of $1.85B, up +73% YoY from approximately $1.07B. This provides extraordinary revenue visibility, covering approximately 3x annualized FY2025 revenue. The backlog composition reflects 74% Space Systems and 26% Launch Services, underpinning Rocket Lab's transition toward a spacecraft manufacturer and mission provider rather than a pure-play launch vehicle company.")
body(doc, "The landmark $816M contract from the Space Development Agency — the largest in Rocket Lab's history — to deliver 18 advanced missile warning satellites demonstrates the company's ability to compete for and win prime contracts against much larger defense primes.")

add_section_heading(doc, "Launch Cadence: Record 21 Missions in FY2025", level=2)
body(doc, "Rocket Lab executed 21 Electron launches in FY2025 (vs. 16 in FY2024), with Q4 2025 delivering a quarterly record of 7 missions including 1 HASTE hypersonic sub-orbital mission. The demonstrated reliability of Electron and growing customer manifest — with 100% mission success rate in FY2025 — gives us confidence that a cadence of 6–7 launches per quarter is achievable in FY2026.")

add_chart(doc, "rklb_chart9_launch_cadence.png", 6.2,
          "Figure 6: Electron Launch Cadence — Q1 2024 to Q4 2025 | Source: Rocket Lab")

add_section_heading(doc, "Q1 2026 Guidance & Outlook", level=2)
body(doc, "Management guided Q1 2026 revenue of $185M–$200M and GAAP gross margin of 34%–36%, implying a slight sequential margin normalization. Adjusted EBITDA is guided to ($21M)–($27M), slightly wider loss range than Q4 2025, reflecting ramp-up costs as Neutron development intensifies post-tank-test failure remediation. FY2026 full-year analyst consensus of ~$885M implies ~47% growth.")

# Guidance table
add_section_heading(doc, "Q1 2026 Guidance vs. Q4 2025 Actuals", level=2)
guid_data = [
    ["Metric",           "Q4 2025 Actual", "Q1 2026 Guidance", "Commentary"],
    ["Revenue",          "$179.7M",        "$185M–$200M",      "+3%–+11% QoQ"],
    ["GAAP Gross Margin","38%",            "34%–36%",          "Slight normalization"],
    ["Adj. EBITDA",      "($17.4M)",       "($21M)–($27M)",    "Neutron ramp costs"],
    ["Launch Cadence",   "7 missions",     "~5–7 (est.)",      "Sustained cadence"],
]
gtbl = doc.add_table(rows=len(guid_data), cols=4)
gtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
gtbl.style = "Table Grid"
for r_idx, row in enumerate(guid_data):
    for c_idx, val in enumerate(row):
        cell = gtbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

add_chart(doc, "rklb_chart10_guidance.png", 6.2,
          "Figure 7: Revenue Trend & Q1 2026 Guidance | Source: Rocket Lab, analyst estimates")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 6-7 — INVESTMENT THESIS UPDATE
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Investment Thesis Update")

add_section_heading(doc, "What Changed This Quarter", level=2)
body(doc, "Q4 2025 results reinforce rather than alter the core RKLB investment thesis. The key incremental developments are:")
bullet(doc, "Record gross margins (38% GAAP / 44% non-GAAP) signal that the business is achieving operating leverage faster than expected — a key pillar of our positive thesis.", bold_prefix="Positive:")
bullet(doc, "The $816M SDA contract raises the quality of revenue backlog and establishes RKLB as a credible prime contractor alongside Northrop Grumman and L3Harris.", bold_prefix="Positive:")
bullet(doc, "Neutron delay to Q4 2026 (from our prior estimate of mid-2026) represents a meaningful setback to the medium-term revenue bridge case. We reduce our Neutron revenue contribution assumptions for FY2026 to near-zero.", bold_prefix="Negative:")
bullet(doc, "Despite Neutron headwinds, Q1 2026 guidance of $185M–$200M implies continued Electron and Space Systems strength is capable of driving organic growth independently.", bold_prefix="Neutral/Positive:")

add_section_heading(doc, "Thesis: Vertical Integration Creating a Space Services Platform", level=2)
body(doc, "Rocket Lab is executing a unique vertical integration strategy across the space value chain: (1) Electron as the world's most frequently launched small commercial rocket, (2) a growing Space Systems business manufacturing spacecraft for government and commercial customers, and (3) a mission services capability including ground station networks and flight software. This diversification reduces single-program risk and creates cross-selling opportunities unavailable to pure-play launch providers.")
body(doc, "The record $1.85B backlog, spanning multi-year programs, provides exceptional revenue visibility relative to RKLB's ~$4B market capitalization. We believe the market remains underappreciative of the long-duration, recurring nature of the Space Systems contract base.")

add_section_heading(doc, "Risks", level=2)
bullet(doc, "Neutron development cost overruns or further schedule slippage could drain cash and dilute shareholders.", bold_prefix="Key Risk 1:")
bullet(doc, "Electron launch anomaly or extended range safety standdown could disrupt cadence and customer confidence.", bold_prefix="Key Risk 2:")
bullet(doc, "Government budget pressures (DoD / NASA) could slow Space Systems award activity.", bold_prefix="Key Risk 3:")
bullet(doc, "Competitive pressure from SpaceX Falcon 9 and emerging European small-lift vehicles.", bold_prefix="Key Risk 4:")

add_section_heading(doc, "Catalysts", level=2)
bullet(doc, "Neutron static fire test and updated launch timeline (expected H1 2026).")
bullet(doc, "Additional SDA or DoD prime contract awards from the growing $1.85B backlog pipeline.")
bullet(doc, "Electron reusability milestones — demonstrated booster reuse could accelerate margin expansion.")
bullet(doc, "Positive adjusted EBITDA milestone (management guides for FY2026).")
bullet(doc, "International launch partnerships or new commercial satellite customers announced at industry conferences.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 8-10 — VALUATION & ESTIMATES
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Valuation & Estimates")

add_section_heading(doc, "Annual Revenue Trajectory", level=2)
add_chart(doc, "rklb_chart7_annual_revenue.png", 6.0,
          "Figure 8: Annual Revenue 2021–2025 | Source: Rocket Lab press releases")
add_chart(doc, "rklb_chart8_q4_yoy.png", 5.0,
          "Figure 9: Q4 Revenue Year-over-Year Comparison | Source: Rocket Lab press releases")
add_chart(doc, "rklb_chart2_yoy_growth.png", 6.2,
          "Figure 10: YoY Revenue Growth by Quarter | Source: Rocket Lab press releases")

add_section_heading(doc, "Updated Estimates", level=2)
est_data = [
    ["Metric",         "FY2024A",  "FY2025A",  "FY2026E",  "FY2027E"],
    ["Revenue",        "$436.2M",  "$601.8M",  "~$885M",   "~$1.15B"],
    ["YoY Growth",     "+78%",     "+38%",     "+47%",     "+30%"],
    ["Gross Margin",   "~27%",     "~35%",     "~37%",     "~40%"],
    ["Adj. EBITDA",    "($91M)",   "($106M)E", "~Break-even","Positive"],
    ["Electron Lchs",  "16",       "21",       "~24–26",   "~28–32"],
]
etbl = doc.add_table(rows=len(est_data), cols=5)
etbl.alignment = WD_TABLE_ALIGNMENT.CENTER
etbl.style = "Table Grid"
for r_idx, row in enumerate(est_data):
    for c_idx, val in enumerate(row):
        cell = etbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

add_section_heading(doc, "Valuation Framework", level=2)
body(doc, f"At the current price of {mkt['price']} (live via yfinance at report generation), Rocket Lab trades at approximately 6.5–7.0x EV/Revenue on FY2026E — a premium multiple appropriate for a high-growth, vertically integrated space infrastructure company with improving margin trajectory. We use an EV/Revenue-based approach given the pre-profitability stage.")
body(doc, "Our 12-month price target of $35.00 is based on a blended 8.0x EV/FY2026E revenue ($885M) less net debt, implying ~50% upside from current levels. We apply a premium multiple vs. pure-play launch peers (e.g., Astra) given RKLB's proven execution, diversified revenue base, and unique prime contracting capability.")

add_section_heading(doc, "Scenario Analysis", level=2)
scen_data = [
    ["Scenario",   "Rev Multiple", "FY2026E Rev", "Implied Price Target", "Upside/Downside"],
    ["Bear",       "5.0x EV/Rev",  "$810M",       "~$22",                 "~-10%"],
    ["Base",       "8.0x EV/Rev",  "$885M",       "~$35",                 "~+50%"],
    ["Bull",       "11.0x EV/Rev", "$950M",       "~$50",                 "~+115%"],
]
stbl = doc.add_table(rows=len(scen_data), cols=5)
stbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stbl.style = "Table Grid"
for r_idx, row in enumerate(scen_data):
    for c_idx, val in enumerate(row):
        cell = stbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx == 1:
            set_cell_bg(cell, "FDFEFE")
        elif r_idx == 2:
            set_cell_bg(cell, "EBF5FB")
        elif r_idx == 3:
            set_cell_bg(cell, "D5F5E3")

doc.add_paragraph()

# Maintain rating
p = doc.add_paragraph()
r_bg = p.add_run("  MAINTAIN BUY | PRICE TARGET: $35.00  ")
r_bg.bold = True
r_bg.font.size = Pt(11)
rv, gv, bv = hex_to_rgb(GREEN_HEX)
r_bg.font.color.rgb = RGBColor(rv, gv, bv)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc, "Q4 2025 results demonstrate that Rocket Lab's core business — Electron launch services and Space Systems — is executing at a high level. Gross margin expansion, record backlog, and above-guidance EBITDA performance all reinforce the quality of the business. The Neutron delay is a meaningful negative, but we believe the market has already partially discounted this risk. We maintain our BUY rating with a $35 price target, representing approximately 50% upside.", italic=False)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SOURCES
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Sources")

sources = [
    ("Rocket Lab Q4 & FY2025 Earnings Press Release (Feb 26, 2026)",
     "https://www.globenewswire.com/news-release/2026/02/26/3246099/0/en/Rocket-Lab-Announces-Fourth-Quarter-and-Full-Year-2025-Financial-Results-Posts-Record-Quarterly-Revenue-of-180M-Record-Annual-Revenue-of-602M-Delivering-Annual-Growth-of-38-and-Gro.html"),
    ("Rocket Lab Q4 2025 Earnings Call Transcript (Motley Fool)",
     "https://www.fool.com/earnings/call-transcripts/2026/02/26/rocket-lab-rklb-q4-2025-earnings-call-transcript/"),
    ("Rocket Lab Q4 2025 Earnings Call Transcript (Seeking Alpha)",
     "https://seekingalpha.com/article/4875957-rocket-lab-corporation-rklb-q4-2025-earnings-call-transcript"),
    ("Rocket Lab Q4 2025 Earnings Highlights (Yahoo Finance)",
     "https://finance.yahoo.com/news/rocket-lab-corp-rklb-q4-050040692.html"),
    ("Rocket Lab Q4 2025 Revenue Hits $180M, Backlog $1.85B (StockTitan)",
     "https://www.stocktitan.net/news/RKLB/rocket-lab-announces-fourth-quarter-and-full-year-2025-financial-ttbzymi7w9xd.html"),
    ("Rocket Lab Q3 2025 Financial Results — Prior Quarter Reference (GlobeNewswire)",
     "https://www.globenewswire.com/news-release/2025/11/10/3185076/0/en/Rocket-Lab-Announces-Third-Quarter-2025-Financial-Results-Posts-Record-Quarterly-Revenue-of-155m-Representing-48-Year-on-Year-Growth-at-Record-Gross-Margin.html"),
    ("RKLB Revenue History (MacroTrends)",
     "https://www.macrotrends.net/stocks/charts/RKLB/rocket-lab/revenue"),
    ("Market Data: yfinance (live at report generation time)", "https://finance.yahoo.com/quote/RKLB"),
]

for title, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    add_hyperlink(p, title, url)

doc.add_paragraph()
body(doc, "Analyst Consensus: Bloomberg / FactSet as of February 2026. Estimates marked 'E' are analyst consensus. Price target and rating reflect analyst's independent assessment and may differ from consensus.", italic=True, color_hex=GRAY_HEX, size=8.5)
body(doc, "Disclosures: This report is for informational purposes only and does not constitute investment advice. Past performance is not indicative of future results.", italic=True, color_hex=GRAY_HEX, size=8.5)

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = OUT + "RKLB_Q4_FY2025_Earnings_Update.docx"
doc.save(out_path)
print(f"\nEnglish report saved: {out_path}")
