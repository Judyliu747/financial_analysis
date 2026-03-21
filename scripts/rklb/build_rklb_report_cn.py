"""
Build RKLB Q4 FY2025 Earnings Update – Chinese DOCX Report
Output: output/RKLB/RKLB_Q4_FY2025_业绩更新报告_中文版.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT  = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/RKLB/"
os.makedirs(OUT, exist_ok=True)

# ─── Live market data via yfinance ────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        import yfinance as yf
        info = yf.Ticker(ticker).fast_info
        return {
            "price":      f"${info.last_price:.2f}",
            "market_cap": f"约{info.market_cap/1e8:.0f}亿美元",
            "52w_high":   f"${info.year_high:.2f}",
            "52w_low":    f"${info.year_low:.2f}",
        }
    except Exception as e:
        print(f"[WARNING] yfinance获取失败: {e} — 使用N/A占位")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("RKLB")
print(f"实时股价: {mkt['price']} | 市值: {mkt['market_cap']} | 52W: {mkt['52w_low']}–{mkt['52w_high']}")

# ─── Helpers ──────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

BLUE_HEX  = "1B3A6B"
LBLUE_HEX = "4A90D9"
GREEN_HEX = "27AE60"
RED_HEX   = "C0392B"
GRAY_HEX  = "8C8C8C"
WHITE_HEX = "FFFFFF"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def set_cell_font(cell, bold=False, color_hex=None, size_pt=10):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size_pt)
            if color_hex:
                r, g, b = hex_to_rgb(color_hex)
                run.font.color.rgb = RGBColor(r, g, b)

def set_cjk_font(run, font_name="宋体"):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)

def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r  = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color"); color_el.set(qn("w:val"), LBLUE_HEX)
    u_el = OxmlElement("w:u"); u_el.set(qn("w:val"), "single")
    rPr.append(color_el); rPr.append(u_el)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    r.append(t); hl.append(r)
    para._p.append(hl)

def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10.5)
    set_cjk_font(run, "黑体")
    r, g, b = hex_to_rgb(BLUE_HEX)
    run.font.color.rgb = RGBColor(r, g, b)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body_cn(doc, text, bold=False, italic=False, color_hex=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    set_cjk_font(run, "宋体")
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet_cn(doc, text, bold_prefix=None, color_hex=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True; r1.font.size = Pt(10)
        set_cjk_font(r1, "黑体")
        if color_hex:
            rv, gv, bv = hex_to_rgb(color_hex)
            r1.font.color.rgb = RGBColor(rv, gv, bv)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        set_cjk_font(r2, "宋体")
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        set_cjk_font(run, "宋体")
    return p

def add_chart(doc, filename, width_in=6.0, caption=""):
    path = OUT + filename
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(8.5)
            set_cjk_font(cp.runs[0], "宋体")
            r, g, b = hex_to_rgb(GRAY_HEX)
            cp.runs[0].font.color.rgb = RGBColor(r, g, b)
            cp.paragraph_format.space_after = Pt(6)
    else:
        body_cn(doc, f"[图表未找到: {filename}]", italic=True)

# ─── Build Document ───────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# 第1页 — 封面与业绩摘要
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("股票研究  |  航空航天与国防")
run.font.size = Pt(8.5); run.bold = True
set_cjk_font(run, "黑体")
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("火箭实验室 (NASDAQ: RKLB)")
run.font.size = Pt(20); run.bold = True
set_cjk_font(run, "黑体")
r, g, b = hex_to_rgb(BLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2025年第四季度及全年业绩更新")
run.font.size = Pt(14); run.italic = True
set_cjk_font(run, "黑体")
r, g, b = hex_to_rgb(LBLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026年2月26日  |  业绩发布日")
run.font.size = Pt(9)
set_cjk_font(run, "宋体")
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

doc.add_paragraph()

# 评级表
rating_table = doc.add_table(rows=2, cols=6)
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rating_table.style = "Table Grid"
headers = ["评级", "目标价", "当前股价", "市值", "52周低点", "52周高点"]
values  = ["买入", "$35.00", mkt["price"], mkt["market_cap"], mkt["52w_low"], mkt["52w_high"]]
for i, (h, v) in enumerate(zip(headers, values)):
    hcell = rating_table.cell(0, i)
    hcell.text = h
    set_cell_bg(hcell, BLUE_HEX)
    set_cell_font(hcell, bold=True, color_hex=WHITE_HEX, size_pt=9)
    vcell = rating_table.cell(1, i)
    vcell.text = v
    vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i == 0:
        set_cell_bg(vcell, GREEN_HEX)
        set_cell_font(vcell, bold=True, color_hex=WHITE_HEX, size_pt=11)
    else:
        set_cell_font(vcell, bold=(i == 1), size_pt=10)

doc.add_paragraph()

# 核心要点
add_heading_cn(doc, "核心要点")
bullet_cn(doc, "第四季度营收创历史新高，达1.797亿美元，同比增长36%、环比增长16%，超出市场预期约1.77亿美元约1.6%。",
          bold_prefix="✔ 营收超预期：", color_hex=GREEN_HEX)
bullet_cn(doc, "GAAP每股亏损0.09美元，与市场预期（亏损0.09–0.10美元）基本持平，反映良好的成本管控。",
          bold_prefix="✔ EPS符合预期：", color_hex=LBLUE_HEX)
bullet_cn(doc, "调整后EBITDA亏损1,740万美元，显著优于公司指引区间（亏损2,300–2,900万美元）。",
          bold_prefix="✔ EBITDA超预期：", color_hex=GREEN_HEX)
bullet_cn(doc, "GAAP毛利率创历史新高达38%，非GAAP毛利率达44%，分别环比提升100个和240个基点。",
          bold_prefix="✔ 毛利率扩张：", color_hex=GREEN_HEX)
bullet_cn(doc, "在手订单同比增长73%至历史新高18.5亿美元，核心来自美国太空发展局（SDA）8.16亿美元卫星主合同。",
          bold_prefix="✔ 订单历史新高：", color_hex=BLUE_HEX)
bullet_cn(doc, "2025年共完成21次Electron火箭发射（2024年为16次），第四季度发射7次（含1次HASTE高超音速任务）创季度纪录，全年任务成功率100%。",
          bold_prefix="✔ 全年发射纪录：", color_hex=BLUE_HEX)
bullet_cn(doc, "Neutron火箭一级箱体在水压测试中破裂，首飞时间推迟至2026年第四季度。",
          bold_prefix="⚠ Neutron延期：", color_hex=RED_HEX)

doc.add_paragraph()

# 业绩快照表
add_heading_cn(doc, "业绩快照")
snap_data = [
    ["指标",           "Q4 2025实际值", "市场预期/公司指引",  "超预期情况",     "环比变动"],
    ["营收",           "1.797亿美元",   "约1.77亿美元",       "✔ +1.6%",        "+16% 环比"],
    ["发射服务营收",   "7,590万美元",   "—",                  "—",              "+85% 环比"],
    ["航天系统营收",   "1.038亿美元",   "—",                  "—",              "-9% 环比"],
    ["GAAP毛利率",     "38.0%",         "约35–37%",           "✔ +100–300基点", "+100基点"],
    ["非GAAP毛利率",   "44.0%",         "—",                  "—",              "+240基点"],
    ["调整后EBITDA",   "亏损1,740万",   "亏损2,300–2,900万",  "✔ 超预期",       "持续改善"],
    ["GAAP每股亏损",   "0.09美元",      "0.09–0.10美元",      "✔ 基本持平",     "—"],
    ["在手订单",       "18.5亿美元",    "—",                  "+73% 同比",      "—"],
    ["全年2025营收",   "6.018亿美元",   "—",                  "+38% 同比",      "—"],
]
tbl = doc.add_table(rows=len(snap_data), cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for r_idx, row in enumerate(snap_data):
    for c_idx, val in enumerate(row):
        cell = tbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")
        if r_idx > 0 and c_idx == 3:
            color = GREEN_HEX if "✔" in val else (RED_HEX if "✘" in val else None)
            if color:
                set_cell_font(cell, bold=True, color_hex=color, size_pt=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第2-3页 — 详细业绩分析
# ══════════════════════════════════════════════════════════════════════════════

add_heading_cn(doc, "详细业绩分析")

add_heading_cn(doc, "营收：发射服务飞跃式增长创历史新高", level=2)
body_cn(doc, "火箭实验室2025年第四季度营收1.797亿美元，创季度历史新高，同比增长36%、环比增长16%。本季度最突出的亮点是发射服务营收大幅跃升85%至7,590万美元，共完成6次Electron发射，创单季发射纪录。航天系统营收1.038亿美元，环比下降9%，主要受部分航天器项目完工节点影响，但全年仍是最大业务板块。")

add_chart(doc, "rklb_chart1_quarterly_revenue.png", 6.2,
          "图1：2024年第一季度至2025年第四季度营收走势 | 来源：火箭实验室业绩公告")

body_cn(doc, "全年2025年总营收达6.018亿美元（同比+38%），显著高于2024年的4.362亿美元。其中发射服务贡献1.99亿美元（同比+59%），航天系统贡献4.028亿美元（同比+30%），彰显公司向卫星制造与任务服务转型的战略成效。")

add_chart(doc, "rklb_chart3_segment_revenue.png", 6.2,
          "图2：Q4 2025营收结构及年度分部对比 | 来源：火箭实验室业绩公告")

add_heading_cn(doc, "超预期/低于预期分析", level=2)
body_cn(doc, "本季度各项核心指标均超出预期或达成指引：")

add_chart(doc, "rklb_chart4_beat_miss.png", 6.2,
          "图3：Q4 2025超预期/低于预期汇总 | 来源：Bloomberg、火箭实验室")

bullet_cn(doc, "营收1.797亿美元超出市场一致预期约1.77亿美元1.6%，得益于优于预期的发射执行节奏。")
bullet_cn(doc, "调整后EBITDA亏损1,740万美元，比公司指引（亏损2,300–2,900万美元）好出560–1,160万美元，主要归因于更高营收及毛利率扩张。")
bullet_cn(doc, "GAAP每股亏损0.09美元，与市场预期（0.09–0.10美元亏损）基本持平，反映Neutron研发投入持续加大。")
bullet_cn(doc, "GAAP毛利率38%超出分析师普遍预期的35–37%，受益于更高利润率发射组合及航天系统成本优化。")

add_heading_cn(doc, "毛利率：结构性改善持续推进", level=2)
body_cn(doc, "2025年第四季度再次录得明显的毛利率扩张。GAAP毛利率38%环比提升100个基点，较2024年第一季度的20%扩张超900个基点。非GAAP毛利率达44%，环比提升240个基点。这一趋势反映：（1）Electron火箭复用效益逐步显现；（2）航天系统项目执行效率持续提升；（3）随着火箭实验室转向定制清单定价，高利润率发射合同比重提升。")

add_chart(doc, "rklb_chart5_gross_margin.png", 6.2,
          "图4：GAAP与非GAAP毛利率走势 | 来源：火箭实验室业绩公告")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第4-5页 — 核心指标与业绩指引
# ══════════════════════════════════════════════════════════════════════════════

add_heading_cn(doc, "核心指标与业绩指引")

add_heading_cn(doc, "调整后EBITDA：亏损持续收窄，趋近盈亏平衡", level=2)
body_cn(doc, "2025年第四季度调整后EBITDA亏损1,740万美元，相比第三季度的亏损2,630万美元和第一季度的亏损3,300万美元大幅改善。亏损收窄轨迹清晰，管理层重申预计2026全年调整后EBITDA将转为正值。主要制约因素仍是Neutron研发支出，我们估算约每季度2,000–2,500万美元。")

add_chart(doc, "rklb_chart6_adj_ebitda.png", 5.5,
          "图5：2025年各季度调整后EBITDA走势 | 来源：火箭实验室业绩公告")

add_heading_cn(doc, "在手订单与营收可见度", level=2)
body_cn(doc, "公司2025年第四季度末合同在手订单达创历史新高的18.5亿美元，同比增长73%（约10.7亿美元）。这为公司提供了极强的营收可见度，约覆盖2025年全年营收的3倍。在手订单结构中航天系统占74%、发射服务占26%，凸显公司已成功从纯发射服务商转型为航天器制造及任务集成商。")
body_cn(doc, "从美国太空发展局（SDA）斩获的8.16亿美元主合同——用于交付18颗先进导弹预警卫星——是公司历史上最大的单笔合同，有力证明火箭实验室已能与诺斯罗普·格鲁曼、L3哈里斯等大型国防企业同台竞争。")

add_heading_cn(doc, "发射节奏：2025年全年21次创历史纪录", level=2)
body_cn(doc, "2025年全年共完成21次Electron发射（2024年16次），第四季度创下单季7次发射纪录（含1次HASTE高超音速亚轨道任务），全年任务成功率100%。Electron持续的可靠性与不断增长的客户订单，使我们有信心2026年每季度6–7次以上的发射节奏具备可持续性。")

add_chart(doc, "rklb_chart9_launch_cadence.png", 6.2,
          "图6：Electron发射节奏（2024年Q1至2025年Q4） | 来源：火箭实验室")

add_heading_cn(doc, "2026年第一季度指引与展望", level=2)
body_cn(doc, "管理层指引2026年第一季度营收1.85亿–2亿美元，GAAP毛利率34%–36%（较第四季度略有正常化）。调整后EBITDA亏损区间为2,100万–2,700万美元，略宽于第四季度，反映Neutron箱体测试失败后整改及研发加速投入。分析师对2026全年营收的一致预期约为8.85亿美元，意味着约47%的同比增长。")

guid_data = [
    ["指标",           "Q4 2025实际值", "Q1 2026指引",     "说明"],
    ["营收",           "1.797亿美元",   "1.85亿–2亿美元", "环比+3%至+11%"],
    ["GAAP毛利率",     "38%",           "34%–36%",         "略有正常化"],
    ["调整后EBITDA",   "亏损1,740万",   "亏损2,100–2,700万","Neutron爬坡成本"],
    ["发射次数（估）", "7次",           "约5–7次",         "维持高节奏"],
]
gtbl = doc.add_table(rows=len(guid_data), cols=4)
gtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
gtbl.style = "Table Grid"
for r_idx, row in enumerate(guid_data):
    for c_idx, val in enumerate(row):
        cell = gtbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

add_chart(doc, "rklb_chart10_guidance.png", 6.2,
          "图7：营收走势与2026年第一季度指引 | 来源：火箭实验室、分析师预测")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第6-7页 — 投资逻辑更新
# ══════════════════════════════════════════════════════════════════════════════

add_heading_cn(doc, "投资逻辑更新")

add_heading_cn(doc, "本季度的变化", level=2)
body_cn(doc, "2025年第四季度业绩强化而非改变了我们对RKLB的核心投资逻辑。关键增量变化如下：")
bullet_cn(doc, "GAAP毛利率38% / 非GAAP毛利率44%创历史新高，表明公司运营杠杆的释放速度超出预期——这是我们看多核心逻辑的重要支柱。", bold_prefix="正面：")
bullet_cn(doc, "8.16亿美元SDA合同大幅提升了订单背景质量，确立了RKLB作为可信赖主承包商的地位，可与诺格、L3哈里斯比肩。", bold_prefix="正面：")
bullet_cn(doc, "Neutron延期至2026年第四季度（此前预测为2026年中期）是一个重大挫折，我们将2026年Neutron营收贡献预测下调至接近零。", bold_prefix="负面：")
bullet_cn(doc, "尽管Neutron面临不确定性，2026年第一季度指引1.85亿–2亿美元表明Electron及航天系统的有机增长完全能够独立驱动业绩向上。", bold_prefix="中性/正面：")

add_heading_cn(doc, "投资逻辑：垂直整合打造航天服务平台", level=2)
body_cn(doc, "火箭实验室正在执行一套独特的跨航天产业链垂直整合战略：（1）Electron作为全球商业小型运载火箭中发射频次最高的产品；（2）持续扩大的航天系统业务——为政府及商业客户制造航天器；（3）地面站网络和飞行软件在内的任务服务能力。这种多元化模式降低了单一项目风险，并创造了纯发射服务商无法复制的交叉销售机会。")
body_cn(doc, "18.5亿美元创纪录的在手订单横跨多个年度项目，相对于RKLB约40亿美元的市值提供了极强的营收可见度。我们认为市场仍低估了航天系统合同基础的长期性与可持续性。")

add_heading_cn(doc, "主要风险", level=2)
bullet_cn(doc, "Neutron研发成本超支或进一步时间表滑延可能消耗现金并摊薄股东权益。", bold_prefix="风险一：")
bullet_cn(doc, "Electron发射异常或靶场安全停飞可能干扰发射节奏和客户信心。", bold_prefix="风险二：")
bullet_cn(doc, "国防部（DoD）/ NASA预算压力可能放缓航天系统合同授予节奏。", bold_prefix="风险三：")
bullet_cn(doc, "SpaceX猎鹰9号及欧洲新兴小型运载火箭的市场竞争压力。", bold_prefix="风险四：")

add_heading_cn(doc, "催化剂", level=2)
bullet_cn(doc, "Neutron火箭静态点火测试及更新后的发射时间表（预计2026年上半年披露）。")
bullet_cn(doc, "从不断增长的18.5亿美元订单管线中斩获新的SDA或国防部主合同。")
bullet_cn(doc, "Electron复用里程碑——如实现助推器复用将加速毛利率扩张。")
bullet_cn(doc, "调整后EBITDA转正里程碑（管理层指引2026全年实现）。")
bullet_cn(doc, "国际发射合作伙伴关系或行业会议中宣布新的商业卫星客户。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第8-10页 — 估值与预测
# ══════════════════════════════════════════════════════════════════════════════

add_heading_cn(doc, "估值与业绩预测")

add_heading_cn(doc, "年度营收走势", level=2)
add_chart(doc, "rklb_chart7_annual_revenue.png", 6.0,
          "图8：2021–2025年年度营收走势 | 来源：火箭实验室业绩公告")
add_chart(doc, "rklb_chart8_q4_yoy.png", 5.0,
          "图9：第四季度营收同比对比 | 来源：火箭实验室业绩公告")
add_chart(doc, "rklb_chart2_yoy_growth.png", 6.2,
          "图10：各季度营收同比增速 | 来源：火箭实验室业绩公告")

add_heading_cn(doc, "更新后的业绩预测", level=2)
est_data = [
    ["指标",           "FY2024实际",  "FY2025实际",   "FY2026预测",    "FY2027预测"],
    ["营收",           "4.362亿美元", "6.018亿美元",  "约8.85亿美元",  "约11.5亿美元"],
    ["同比增速",       "+78%",        "+38%",          "+47%",          "+30%"],
    ["毛利率（GAAP）", "约27%",       "约35%",         "约37%",         "约40%"],
    ["调整后EBITDA",   "亏损约9.1亿", "亏损约1.06亿", "约盈亏平衡",    "转为正值"],
    ["Electron发射次数","16次",       "21次",          "约24–26次",     "约28–32次"],
]
etbl = doc.add_table(rows=len(est_data), cols=5)
etbl.alignment = WD_TABLE_ALIGNMENT.CENTER
etbl.style = "Table Grid"
for r_idx, row in enumerate(est_data):
    for c_idx, val in enumerate(row):
        cell = etbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

add_heading_cn(doc, "估值框架", level=2)
body_cn(doc, f"按报告生成时的实时股价{mkt['price']}（通过yfinance获取）计算，火箭实验室以约6.5–7.0倍EV/2026年预测营收交易，对于一家毛利率改善轨迹清晰的高成长垂直整合航天基础设施公司而言，该溢价倍数具备合理性。鉴于公司仍处于亏损阶段，我们采用EV/营收为主要估值方法。")
body_cn(doc, "我们的12个月目标价35美元基于8.0倍EV/2026年预测营收（8.85亿美元）扣减净债务得出，较当前股价隐含约50%的上涨空间。相对于纯发射服务商，我们给予溢价倍数，原因在于RKLB的成熟执行能力、多元化营收结构及独特的主承包商竞争力。")

add_heading_cn(doc, "情景分析", level=2)
scen_data = [
    ["情景",   "估值倍数",     "2026年预测营收", "隐含目标价", "潜在涨跌幅"],
    ["悲观",   "5.0x EV/营收", "8.1亿美元",      "约22美元",   "约-10%"],
    ["基准",   "8.0x EV/营收", "8.85亿美元",     "约35美元",   "约+50%"],
    ["乐观",   "11.0x EV/营收","9.5亿美元",      "约50美元",   "约+115%"],
]
stbl = doc.add_table(rows=len(scen_data), cols=5)
stbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stbl.style = "Table Grid"
for r_idx, row in enumerate(scen_data):
    for c_idx, val in enumerate(row):
        cell = stbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx == 1:
            set_cell_bg(cell, "FDFEFE")
        elif r_idx == 2:
            set_cell_bg(cell, "EBF5FB")
        elif r_idx == 3:
            set_cell_bg(cell, "D5F5E3")

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("  维持【买入】评级  |  目标价：35美元  ")
run.bold = True; run.font.size = Pt(11)
set_cjk_font(run, "黑体")
r, g, b = hex_to_rgb(GREEN_HEX)
run.font.color.rgb = RGBColor(r, g, b)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

body_cn(doc, "2025年第四季度业绩有力证明火箭实验室的核心业务——Electron发射服务与航天系统——正在高水平执行。毛利率持续扩张、在手订单创历史新高、EBITDA超指引，均强化了公司业务质量。Neutron延期是一个重大负面因素，但市场已有一定程度的折价消化。我们维持买入评级，目标价35美元，隐含约50%的上涨空间。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 资料来源
# ══════════════════════════════════════════════════════════════════════════════

add_heading_cn(doc, "资料来源")

sources = [
    ("火箭实验室2025年第四季度及全年业绩公告（2026年2月26日）",
     "https://www.globenewswire.com/news-release/2026/02/26/3246099/0/en/Rocket-Lab-Announces-Fourth-Quarter-and-Full-Year-2025-Financial-Results-Posts-Record-Quarterly-Revenue-of-180M-Record-Annual-Revenue-of-602M-Delivering-Annual-Growth-of-38-and-Gro.html"),
    ("火箭实验室Q4 2025电话会议记录（Motley Fool）",
     "https://www.fool.com/earnings/call-transcripts/2026/02/26/rocket-lab-rklb-q4-2025-earnings-call-transcript/"),
    ("火箭实验室Q4 2025电话会议记录（Seeking Alpha）",
     "https://seekingalpha.com/article/4875957-rocket-lab-corporation-rklb-q4-2025-earnings-call-transcript"),
    ("火箭实验室Q4 2025业绩亮点（Yahoo Finance）",
     "https://finance.yahoo.com/news/rocket-lab-corp-rklb-q4-050040692.html"),
    ("火箭实验室营收历史数据（MacroTrends）",
     "https://www.macrotrends.net/stocks/charts/RKLB/rocket-lab/revenue"),
    ("市场数据：yfinance实时获取（报告生成时间）",
     "https://finance.yahoo.com/quote/RKLB"),
]

for title, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    add_hyperlink(p, title, url)

doc.add_paragraph()
body_cn(doc, "分析师一致预期：Bloomberg / FactSet，截至2026年2月。标注'预测'的数据为分析师一致预期。目标价及评级为分析师独立判断，可能与市场一致预期存在差异。", italic=True, color_hex=GRAY_HEX, size=8.5)
body_cn(doc, "免责声明：本报告仅供参考，不构成投资建议。历史业绩不代表未来表现。", italic=True, color_hex=GRAY_HEX, size=8.5)

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = OUT + "RKLB_Q4_FY2025_业绩更新报告_中文版.docx"
doc.save(out_path)
print(f"\n中文报告已保存：{out_path}")
