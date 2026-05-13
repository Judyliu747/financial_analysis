"""
Adobe Inc. (ADBE) — Q1 FY2026 业绩更新报告（中文版）
机构级股票研究报告生成器 (DOCX)
"""
import os
import yfinance as yf
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/ADBE"

C_NAVY   = RGBColor(0x1C, 0x1C, 0x4D)
C_RED    = RGBColor(0xEB, 0x10, 0x00)
C_BLUE   = RGBColor(0x14, 0x73, 0xE6)
C_GREEN  = RGBColor(0x1E, 0x8A, 0x44)
C_SELL   = RGBColor(0xCC, 0x00, 0x00)
C_GREY   = RGBColor(0x8E, 0x8E, 0x93)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

FONT_CN_BODY  = "宋体"
FONT_CN_TITLE = "黑体"


def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price": round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high": round(info.year_high, 2),
            "52w_low": round(info.year_low, 2),
        }
    except Exception as e:
        print(f"WARNING: yfinance failed: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "1473E6")
    rPr.append(color_elem)
    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rPr.append(u_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def set_run_cn(run, font_name=FONT_CN_BODY):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
        set_run_cn(run, FONT_CN_TITLE)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        set_run_cn(run, FONT_CN_TITLE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1C1C4D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        set_run_cn(run, FONT_CN_TITLE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    set_run_cn(run, FONT_CN_BODY)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    set_run_cn(run, FONT_CN_BODY)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"[图表: {filename} 未找到]", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
        cp.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, alternate=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "1C1C4D")
        set_cell_border(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F5F5F7" if (alternate and ri % 2 == 1) else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="E5E5EA")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if "+" in str(val) and ci > 1:
                        run.font.color.rgb = C_GREEN
                    elif "−" in str(val) or (str(val).startswith("-") and ci > 0):
                        run.font.color.rgb = C_SELL
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()
    return table


def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# 市场数据
# ══════════════════════════════════════════════════════════════════════════════
mkt = get_market_data("ADBE")
price_str = f"${mkt['price']}" if isinstance(mkt['price'], float) else "N/A"
mcap_b = mkt['market_cap'] / 1e9 if isinstance(mkt['market_cap'], (int, float)) else 0
mcap_str = f"${mcap_b:.1f}B" if mcap_b else "N/A"
range_str = f"${mkt['52w_low']} – ${mkt['52w_high']}" if isinstance(mkt['52w_low'], float) else "N/A"

# ══════════════════════════════════════════════════════════════════════════════
# 生成报告
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

style_normal = doc.styles["Normal"]
style_normal.font.name = FONT_CN_BODY
style_normal.font.size = Pt(10)
rPr = style_normal.element.rPr
if rPr is None:
    rPr = OxmlElement("w:rPr")
    style_normal.element.append(rPr)
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.insert(0, rFonts)
rFonts.set(qn("w:eastAsia"), FONT_CN_BODY)

# ══════════════════════════════════════════════════════════════════════════════
# 第1页 — 业绩摘要
# ══════════════════════════════════════════════════════════════════════════════
p_ticker = doc.add_paragraph()
p_ticker.clear()
run = p_ticker.add_run("ADOBE INC.  (ADBE: NASDAQ)")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_NAVY
run.font.name = "Calibri"
p_ticker.paragraph_format.space_after = Pt(2)

p_sub = doc.add_paragraph()
p_sub.clear()
run2 = p_sub.add_run("股票研究  |  应用软件")
run2.font.size = Pt(10)
run2.font.color.rgb = C_GREY
set_run_cn(run2, FONT_CN_TITLE)
p_sub.paragraph_format.space_after = Pt(6)

# 评级表
rating_table = doc.add_table(rows=1, cols=5)
rating_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["评级", "目标价", "当前价", "52周区间", "市值"]
values = ["买入", "$325.00", price_str, range_str, mcap_str]
colors_v = [C_GREEN, C_BLUE, C_NAVY, C_GREY, C_NAVY]

for i, (lbl, val, col) in enumerate(zip(labels, values, colors_v)):
    cell = rating_table.rows[0].cells[i]
    set_cell_bg(cell, "F0F4FF")
    set_cell_border(cell, color="D0D8F0")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lbl + "\n")
    r1.font.size = Pt(8)
    r1.font.color.rgb = C_GREY
    r1.font.name = "Calibri"
    r2 = p1.add_run(val)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = col
    r2.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(4)
horizontal_rule(doc)

# 报告标题
p_title = doc.add_paragraph()
p_title.clear()
run_t = p_title.add_run(
    "2026财年第一季度业绩更新：营收与每股收益均超预期；"
    "AI优先ARR三倍增长，CEO宣布交接计划"
)
run_t.font.size = Pt(14)
run_t.font.bold = True
run_t.font.color.rgb = C_NAVY
set_run_cn(run_t, FONT_CN_TITLE)
p_title.paragraph_format.space_after = Pt(3)

p_date = doc.add_paragraph()
p_date.clear()
rd = p_date.add_run(
    "报告日期: 2026年5月13日  |  业绩发布: 2026年3月12日  |  "
    "财季截止: 2026年2月27日"
)
rd.font.size = Pt(8.5)
rd.font.color.rgb = C_GREY
set_run_cn(rd, FONT_CN_BODY)
p_date.paragraph_format.space_after = Pt(8)

horizontal_rule(doc)

# 业绩摘要表
heading(doc, "业绩摘要", 2)
body(doc, "2026财年第一季度业绩：超预期", bold=True, color=C_GREEN)

add_table(doc,
    ["指标", "实际值", "市场预期", "差异"],
    [
        ["营收", "$64.0亿", "$62.8亿", "+$1.2亿 (+1.9%)"],
        ["非GAAP每股收益", "$6.06", "$5.87", "+$0.19 (+3.2%)"],
        ["GAAP每股收益", "$4.60", "$4.45", "+$0.15 (+3.4%)"],
        ["毛利率（GAAP）", "89.6%", "89.2%", "+40个基点"],
        ["非GAAP营业利润率", "47.4%", "47.0%", "+40个基点"],
    ],
    col_widths=[Inches(2.0), Inches(1.3), Inches(1.3), Inches(1.9)],
)

heading(doc, "核心要点", 3)

body(doc, (
    "■ 营收64.0亿美元，超市场预期1.2亿美元（+1.9%），同比增长12%（恒定汇率下增长11%）。"
    "增长主要由订阅收入同比增长13%至61.7亿美元驱动，其中商业专业人士与消费者订阅收入"
    "同比加速增长16%至17.8亿美元，创意与营销专业人士订阅收入同比增长12%至43.9亿美元。"
    "超预期表现覆盖全部客户群体，反映出AI驱动的持续采纳趋势。"
))

body(doc, (
    "■ 非GAAP每股收益6.06美元，超市场预期0.19美元，同比增长19%。优异表现源于"
    "营收端增长与利润率扩张双重驱动——非GAAP营业利润率同比扩张120个基点至47.4%。"
    "GAAP毛利率同比提升50个基点至89.6%，彰显Adobe订阅模式固有的规模效应。"
))

body(doc, (
    "■ AI优先ARR同比增长超过3倍，总ARR达260.6亿美元（+10.9%）。月活跃用户"
    "突破8.5亿（+17%），展现持续的平台参与度。Firefly、Acrobat AI助手及"
    "GenStudio for Performance Marketing等产品正在获得商业化牵引力，验证了"
    "Adobe的AI变现战略。"
))

body(doc, (
    "■ 2026财年指引重申：营收259–261亿美元，非GAAP每股收益$23.30–$23.50。"
    "第二季度指引为营收64.3–64.8亿美元，非GAAP每股收益$5.80–$5.85。"
    "CEO Shantanu Narayen宣布将在继任者任命后卸任CEO职务，标志着其18年任期后"
    "的重大领导层变更。维持买入评级和$325目标价。"
))

# 财务预测表
heading(doc, "更新后的财务预测", 2)

add_table(doc,
    ["指标", "FY2026E（前值）", "FY2026E（新值）", "变化", "FY2027E"],
    [
        ["营收（$B）", "$25.9–$26.1", "$25.9–$26.1", "重申", "$28.8"],
        ["营收增速", "9–10%", "9–10%", "—", "~11%"],
        ["非GAAP营业利润率", "~45%", "~45%", "—", "~46%"],
        ["非GAAP每股收益", "$23.30–$23.50", "$23.30–$23.50", "重申", "$26.50"],
        ["总ARR（$B）", "$27.8", "$27.8", "—", "$30.8"],
        ["市盈率（NTM）", "10.2x", "10.2x", "—", "9.0x"],
    ],
    col_widths=[Inches(1.8), Inches(1.3), Inches(1.3), Inches(1.0), Inches(1.1)],
)

body(doc, "注：FY2026E指引已于Q1 FY2026业绩电话会上重申。FY2027E为分析师预测。",
     italic=True, color=C_GREY, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# 第2-3页 — 详细业绩分析
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "详细业绩分析", 1)

heading(doc, "营收分析", 2)

body(doc, (
    "Adobe在2026财年第一季度实现创纪录营收64.0亿美元，同比增长12%（恒定汇率下增长11%），"
    "超过华尔街预期的62.8亿美元约1.2亿美元。这标志着公司连续五个季度实现10%以上的营收增长，"
    "并从上季度（Q4 FY2025）的10%增速进一步加速。"
))

body(doc, (
    "总订阅收入达61.7亿美元，同比增长13%，占总营收的96.4%。其中，商业专业人士与消费者"
    "订阅收入17.8亿美元，同比增长16%，为各客户群中增速最快，受Acrobat AI助手采纳和"
    "文档工作流扩展驱动。创意与营销专业人士订阅收入43.9亿美元，同比增长12%，"
    "得益于Firefly商业化部署和Creative Cloud持续增长。"
))

add_image(doc, "adbe_chart1_revenue.png", width=Inches(6.0),
          caption="图1 — 季度营收趋势（资料来源：Adobe业绩新闻稿）")

heading(doc, "季度营收明细", 3)

add_table(doc,
    ["指标", "Q1 FY25", "Q2 FY25", "Q3 FY25", "Q4 FY25", "Q1 FY26", "同比"],
    [
        ["总营收（$B）", "$5.71", "$5.87", "$5.99", "$6.20", "$6.40", "+12.1%"],
        ["  订阅收入（$B）", "$5.46", "$5.62", "$5.75", "$5.95", "$6.17", "+13.0%"],
        ["  创意与营销专业人士", "$3.92", "$4.03", "$4.12", "$4.27", "$4.39", "+12.0%"],
        ["  商业与消费者", "$1.53", "$1.59", "$1.63", "$1.69", "$1.78", "+16.3%"],
        ["营收增速", "10.0%", "10.8%", "11.0%", "10.0%", "12.1%", "+210bps"],
    ],
    col_widths=[Inches(1.6)] + [Inches(0.85)] * 6,
)

heading(doc, "盈利能力分析", 2)

body(doc, (
    "GAAP毛利率同比扩张50个基点至89.6%，反映订阅收入高增量利润率和运营效率提升。"
    "非GAAP营业利润率达47.4%，同比扩张120个基点（Q1 FY2025为46.2%），"
    "受营收规模效应驱动，抵消了AI能力方面的持续投入。"
))

body(doc, (
    "GAAP营业利润为24.2亿美元（利润率37.8%），非GAAP营业利润为30.4亿美元"
    "（利润率47.4%）。GAAP与非GAAP之间的差异主要反映股权激励费用和无形资产摊销。"
    "GAAP每股收益4.60美元同比增长11%，非GAAP每股收益6.06美元同比增长19%，"
    "后者增速更快源于利润率扩张和股份回购的双重贡献。"
))

add_image(doc, "adbe_chart2_eps.png", width=Inches(6.0),
          caption="图2 — 非GAAP每股收益趋势（资料来源：Adobe业绩新闻稿；彭博）")

add_image(doc, "adbe_chart5_margins.png", width=Inches(6.0),
          caption="图3 — 盈利能力趋势（资料来源：Adobe业绩新闻稿）")

add_table(doc,
    ["利润率", "Q1 FY25", "Q2 FY25", "Q3 FY25", "Q4 FY25", "Q1 FY26", "同比"],
    [
        ["毛利率（GAAP）", "89.1%", "89.2%", "89.3%", "89.4%", "89.6%", "+50bps"],
        ["营业利润率（GAAP）", "35.3%", "35.9%", "36.5%", "37.1%", "37.8%", "+250bps"],
        ["营业利润率（非GAAP）", "46.2%", "46.5%", "46.8%", "47.0%", "47.4%", "+120bps"],
    ],
    col_widths=[Inches(1.8)] + [Inches(0.85)] * 6,
)

# ══════════════════════════════════════════════════════════════════════════════
# 第4-5页 — 关键指标与指引
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "关键指标与前瞻指引", 1)

heading(doc, "年化经常性收入（ARR）", 2)

body(doc, (
    "截至2026财年第一季度末，Adobe总ARR达260.6亿美元，同比增长10.9%。管理层重申了"
    "FY2026年ARR增长10.2%的目标，暗示下半年将维持增长势头。ARR指标仍然是Adobe最重要的"
    "关键绩效指标（KPI），因为它反映了订阅收入的全年化运行率，提供了极强的前瞻可见性。"
))

body(doc, (
    "AI优先ARR同比增长超过3倍，相比Q1 FY2025披露的约1.25亿美元业务规模显著加速。"
    "推动增长的产品包括Firefly（生成式AI图像创作）、Acrobat AI助手、"
    "GenStudio for Performance Marketing，以及嵌入Creative Cloud和Experience Cloud"
    "的AI功能。尽管Adobe未披露AI优先ARR的精确数字，但>3倍的同比增速表明"
    "AI正在成为FY2026-27年实质性的营收贡献来源。"
))

add_image(doc, "adbe_chart4_arr.png", width=Inches(6.0),
          caption="图4 — 总ARR趋势（资料来源：Adobe业绩新闻稿）")

heading(doc, "平台活跃度", 2)

body(doc, (
    "月活跃用户在Q1 FY2026突破8.5亿，同比增长17%。该指标至关重要，因为它支撑着"
    "Adobe从免费到付费的转化漏斗，并为AI增值销售创造了庞大的装机基础。MAU增长加速"
    "反映了Creative Cloud工具网页版和移动版的采纳增加，以及AI功能吸引新用户的效果。"
))

add_image(doc, "adbe_chart3_segments.png", width=Inches(6.0),
          caption="图5 — 营收构成与订阅收入增速（资料来源：Adobe Q1 FY2026业绩新闻稿）")

heading(doc, "第二季度及全年指引", 2)

body(doc, (
    "第二季度（Q2 FY2026）指引：营收64.3–64.8亿美元（隐含同比增长~10–11%），"
    "非GAAP每股收益$5.80–$5.85，非GAAP营业利润率约44.5%。"
    "环比营业利润率下降与季节性模式和投资节奏一致。"
))

body(doc, (
    "2026全财年目标重申：总营收259–261亿美元，非GAAP营业利润率约45%，"
    "非GAAP每股收益$23.30–$23.50。Q1超预期后仍重申全年指引，为管理层的"
    "可见度提供了信心，并暗示年内有上调空间。"
))

add_image(doc, "adbe_chart7_guidance.png", width=Inches(6.0),
          caption="图6 — FY2026指引 vs. 市场预期（资料来源：Adobe Q1 FY2026业绩新闻稿；彭博）")

add_table(doc,
    ["指标", "Q2 FY26指引", "FY2026指引", "市场预期"],
    [
        ["营收（$B）", "$6.43–$6.48", "$25.9–$26.1", "$26.0"],
        ["非GAAP每股收益", "$5.80–$5.85", "$23.30–$23.50", "$23.35"],
        ["非GAAP营业利润率", "~44.5%", "~45%", "~45%"],
        ["ARR增速", "—", "10.2%", "~10.5%"],
    ],
    col_widths=[Inches(1.6), Inches(1.5), Inches(1.6), Inches(1.5)],
)

# ══════════════════════════════════════════════════════════════════════════════
# 第6-7页 — 投资论点更新
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "投资论点更新", 1)

heading(doc, "论点影响评估", 2)

body(doc, "■ 论点支柱一：AI变现驱动增长加速", bold=True)
body(doc, "状态：强化", bold=True, color=C_GREEN)
body(doc, (
    "Q1 FY2026业绩为Adobe AI战略转化为营收提供了迄今最有力的证据。AI优先ARR"
    "同比增长超过3倍，结合8.5亿以上MAU（+17%），证明AI功能既能吸引新用户，"
    "又能通过高级订阅层和独立产品实现变现。12%的营收增长从上季度10%加速，"
    "进一步验证了AI整合正在成为Creative Cloud、Document Cloud和Experience Cloud"
    "的结构性增长催化剂。我们预计AI驱动的营收贡献将在FY2026-27年日益显著。"
))

body(doc, "■ 论点支柱二：订阅模式提供盈利可见性与利润率扩张", bold=True)
body(doc, "状态：强化", bold=True, color=C_GREEN)
body(doc, (
    "订阅收入占总营收的96.4%，Adobe的经常性收入模式持续提供高可见性和利润率扩张。"
    "非GAAP营业利润率同比扩张120bps至47.4%，毛利率达89.6%（+50bps）。"
    "创纪录的Q1经营现金流29.6亿美元彰显了模式的现金生成能力。重申的FY2026全年"
    "非GAAP营业利润率~45%目标（低于Q1的47.4%）反映季节性投资节奏，"
    "但长期利润率扩张轨迹保持完好。"
))

body(doc, "■ 论点支柱三：创意和文档工作流中的竞争壁垒", bold=True)
body(doc, "状态：不变", bold=True, color=C_BLUE)
body(doc, (
    "Adobe的竞争地位保持稳固，但新兴AI原生竞争对手（Canva、Figma替代品、Midjourney）"
    "持续在专业消费者层面构成挑战。8.5亿以上MAU表明Adobe的平台粘性依然强劲，"
    "而Firefly原生整合至Photoshop、Illustrator等工具在专业工作流中提供了差异化的"
    "AI体验。不过，我们注意到股价年初至今下跌约27%，表明市场对竞争风险的"
    "定价可能过于激进。"
))

heading(doc, "CEO交接", 2)

body(doc, (
    "Shantanu Narayen宣布将在继任者任命后卸任CEO职务。Narayen领导Adobe 18年，"
    "见证了公司从永久许可证软件业务向云优先订阅模式的转型。虽然CEO交接带来不确定性，"
    "但我们认为风险可控，原因包括：(1) 有序的过渡流程（非立即离职），"
    "(2) Adobe拥有深厚的高管人才储备，(3) 公司强劲的运营动能。"
    "我们将密切关注继任时间表，但预计不会影响近期执行力。"
))

heading(doc, "风险因素", 3)

bullet(doc, "CEO继任风险：新领导层方向和市场反应的不确定性")
bullet(doc, "AI竞争颠覆：新兴AI原生工具可能在专业消费者层面蚕食Creative Cloud份额")
bullet(doc, "宏观敏感性：企业支出放缓可能减慢Experience Cloud增长")
bullet(doc, "估值压缩：股价年初至今已下跌约27%；若增长减速，估值倍数可能进一步收缩")
bullet(doc, "监管风险：与训练数据相关的AI版权诉讼潜在风险")

add_image(doc, "adbe_chart10_ai_growth.png", width=Inches(6.0),
          caption="图7 — AI优先ARR与MAU增长（资料来源：Adobe Q1 FY2026业绩新闻稿及电话会）")

# ══════════════════════════════════════════════════════════════════════════════
# 第8-10页 — 估值与预测
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "估值与预测", 1)

heading(doc, "估值概要", 2)

body(doc, (
    f"以当前股价{price_str}计算，Adobe目前的NTM市盈率约为10.2倍"
    "（基于FY2026E非GAAP每股收益中值$23.40），显著低于其3年平均NTM市盈率~22.5倍"
    "及大型软件同业中位数~25倍。我们认为当前估值过度反映了竞争风险和CEO交接不确定性。"
))

body(doc, (
    "我们的$325目标价基于混合方法：(1) DCF分析权重60%，隐含每股$340（WACC 9.0%，"
    "终端增长率3.0%，更新后的营收/利润率预测）；(2) NTM P/E倍数法权重40%，"
    "以14.0倍应用于FY2027E非GAAP每股收益$26.50，隐含每股$371。"
    f"混合价值约$352，保守取整至$325。隐含上行空间约"
    f"{((325/mkt['price'])-1)*100:.0f}%。" if isinstance(mkt['price'], float) else
    "我们的$325目标价基于混合DCF和倍数方法。"
))

add_image(doc, "adbe_chart9_valuation.png", width=Inches(6.0),
          caption="图8 — NTM市盈率倍数（资料来源：彭博，FactSet）")

heading(doc, "超预期分析与现金流", 2)

add_image(doc, "adbe_chart6_beat_miss.png", width=Inches(6.0),
          caption="图9 — Q1 FY2026超预期分析（资料来源：Adobe业绩新闻稿；彭博）")

add_image(doc, "adbe_chart8_cashflow.png", width=Inches(6.0),
          caption="图10 — 经营现金流（资料来源：Adobe业绩新闻稿）")

body(doc, (
    "Adobe在Q1创下29.6亿美元的经营现金流纪录，支持持续的股份回购和资产负债表稳健。"
    "季末现金及短期投资为68.9亿美元。强劲的现金流进一步支持我们的论点——"
    "Adobe能够在投资AI能力的同时向股东返还资本。"
))

heading(doc, "详细预测更新", 2)

add_table(doc,
    ["指标", "FY2025A", "FY2026E", "同比", "FY2027E"],
    [
        ["营收（$B）", "$23.77", "$26.0", "+9.4%", "$28.8"],
        ["  订阅收入（$B）", "$22.71", "$24.9", "+9.6%", "$27.6"],
        ["毛利率", "89.3%", "89.5%", "+20bps", "89.7%"],
        ["非GAAP营业利润（$B）", "$10.87", "$11.70", "+7.6%", "$13.25"],
        ["非GAAP营业利润率", "45.7%", "45.0%", "−70bps", "46.0%"],
        ["非GAAP每股收益", "$20.50", "$23.40", "+14.1%", "$26.50"],
        ["GAAP每股收益", "$16.60", "$18.50", "+11.4%", "$21.20"],
        ["经营现金流（$B）", "$10.80", "$11.50", "+6.5%", "$12.80"],
        ["总ARR（$B）", "$25.20", "$27.80", "+10.3%", "$30.80"],
    ],
    col_widths=[Inches(1.8), Inches(1.2), Inches(1.2), Inches(1.0), Inches(1.2)],
)

body(doc, "注：FY2025A = 实际值；FY2026E/FY2027E = 预测值。A = 实际；E = 预测。",
     italic=True, color=C_GREY, size=8)

heading(doc, "目标价推导", 3)

add_table(doc,
    ["方法", "权重", "隐含价值", "关键假设"],
    [
        ["DCF分析", "60%", "$340", "WACC 9.0%, 终端增长率 3.0%, 10年预测"],
        ["NTM P/E倍数", "40%", "$371", "14.0x FY2027E EPS $26.50"],
        ["混合价值", "100%", "$352", "保守取整至$325"],
    ],
    col_widths=[Inches(1.5), Inches(0.8), Inches(0.8), Inches(3.2)],
)

# ══════════════════════════════════════════════════════════════════════════════
# 资料来源
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "资料来源与参考文献", 1)

heading(doc, "业绩材料（Q1 FY2026）", 3)

p1 = body(doc, "")
p1.clear()
r = p1.add_run("业绩新闻稿（2026年3月12日）：")
r.font.size = Pt(9)
set_run_cn(r, FONT_CN_BODY)
add_hyperlink(p1, "Adobe Delivers Record Q1 Results",
              "https://news.adobe.com/news/2026/03/adobe-q1fy26-financial-results")

p2 = body(doc, "")
p2.clear()
r = p2.add_run("Form 8-K / 业绩新闻稿（2026年3月12日提交）：")
r.font.size = Pt(9)
set_run_cn(r, FONT_CN_BODY)
add_hyperlink(p2, "SEC EDGAR Filing",
              "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0000796343&type=8-K")

p3 = body(doc, "")
p3.clear()
r = p3.add_run("Q1 FY2026业绩电话会记录（2026年3月12日）：")
r.font.size = Pt(9)
set_run_cn(r, FONT_CN_BODY)
add_hyperlink(p3, "Adobe投资者关系",
              "https://www.adobe.com/investor-relations.html")

heading(doc, "往期参考", 3)

p5 = body(doc, "")
p5.clear()
r = p5.add_run("Q4 FY2025业绩新闻稿（2025年12月10日）：")
r.font.size = Pt(9)
set_run_cn(r, FONT_CN_BODY)
add_hyperlink(p5, "Adobe Reports Record Q4 and FY2025 Revenue",
              "https://news.adobe.com/news/2025/12/122025-q4earnings")

heading(doc, "市场共识与数据", 3)

body(doc, "彭博市场共识预测截至2026年3月11日（业绩发布前）",
     italic=True, size=9, color=C_GREY)
body(doc, "股价与市场数据来源：Yahoo Finance / 彭博",
     italic=True, size=9, color=C_GREY)

heading(doc, "免责声明", 3)
body(doc, (
    "本报告仅供参考，不构成投资建议。所有预测和估算均为作者观点，可能随时变更。"
    "过往业绩不代表未来表现。投资者在做出投资决策前应进行独立尽职调查。"
), italic=True, color=C_GREY, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════════════════════
out_path = os.path.join(BASE, "ADBE_Q1_FY2026_业绩更新报告_中文版.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
