"""
Adobe Inc. (ADBE) — Q1 FY2026 Earnings Update
Institutional equity research report builder (DOCX)
"""
import os
import yfinance as yf
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/ADBE"

C_NAVY   = RGBColor(0x1C, 0x1C, 0x4D)
C_RED    = RGBColor(0xEB, 0x10, 0x00)
C_BLUE   = RGBColor(0x14, 0x73, 0xE6)
C_GREEN  = RGBColor(0x1E, 0x8A, 0x44)
C_SELL   = RGBColor(0xCC, 0x00, 0x00)
C_GREY   = RGBColor(0x8E, 0x8E, 0x93)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price": round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high": round(info.year_high, 2),
            "52w_low": round(info.year_low, 2),
        }
    except Exception as e:
        print(f"WARNING: yfinance failed: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "1473E6")
    rPr.append(color_elem)
    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rPr.append(u_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1C1C4D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"[Chart: {filename} not found]", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
        cp.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, alternate=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "1C1C4D")
        set_cell_border(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F5F5F7" if (alternate and ri % 2 == 1) else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="E5E5EA")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if "+" in str(val) and ci > 1:
                        run.font.color.rgb = C_GREEN
                    elif "−" in str(val) or (str(val).startswith("-") and ci > 0):
                        run.font.color.rgb = C_SELL
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()
    return table


def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# MARKET DATA
# ══════════════════════════════════════════════════════════════════════════════
mkt = get_market_data("ADBE")
price_str = f"${mkt['price']}" if isinstance(mkt['price'], float) else "N/A"
mcap_str = f"${mkt['market_cap'] / 1e9:.1f}B" if isinstance(mkt['market_cap'], (int, float)) else "N/A"
range_str = f"${mkt['52w_low']} – ${mkt['52w_high']}" if isinstance(mkt['52w_low'], float) else "N/A"

# ══════════════════════════════════════════════════════════════════════════════
# BUILD REPORT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — EARNINGS SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
p_ticker = doc.add_paragraph()
p_ticker.clear()
run = p_ticker.add_run("ADOBE INC.  (ADBE: NASDAQ)")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_NAVY
run.font.name = "Calibri"
p_ticker.paragraph_format.space_after = Pt(2)

p_sub = doc.add_paragraph()
p_sub.clear()
run2 = p_sub.add_run("Equity Research  |  Application Software")
run2.font.size = Pt(10)
run2.font.color.rgb = C_GREY
run2.font.name = "Calibri"
p_sub.paragraph_format.space_after = Pt(6)

# Rating / PT / Price box
rating_table = doc.add_table(rows=1, cols=5)
rating_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["Rating", "Price Target", "Current Price", "52-Wk Range", "Market Cap"]
values = ["OUTPERFORM", "$325.00", price_str, range_str, mcap_str]
colors_v = [C_GREEN, C_BLUE, C_NAVY, C_GREY, C_NAVY]

for i, (lbl, val, col) in enumerate(zip(labels, values, colors_v)):
    cell = rating_table.rows[0].cells[i]
    set_cell_bg(cell, "F0F4FF")
    set_cell_border(cell, color="D0D8F0")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lbl + "\n")
    r1.font.size = Pt(8)
    r1.font.color.rgb = C_GREY
    r1.font.name = "Calibri"
    r2 = p1.add_run(val)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = col
    r2.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(4)
horizontal_rule(doc)

# Report title
p_title = doc.add_paragraph()
p_title.clear()
run_t = p_title.add_run(
    "Q1 FY2026 Earnings Update: Record Revenue & EPS Beat; "
    "AI-first ARR Triples as CEO Transition Announced"
)
run_t.font.size = Pt(14)
run_t.font.bold = True
run_t.font.color.rgb = C_NAVY
run_t.font.name = "Calibri"
p_title.paragraph_format.space_after = Pt(3)

p_date = doc.add_paragraph()
p_date.clear()
rd = p_date.add_run(
    "Report Date: May 13, 2026  |  Earnings Released: March 12, 2026  |  "
    "Fiscal Quarter Ended: February 27, 2026"
)
rd.font.size = Pt(8.5)
rd.font.color.rgb = C_GREY
rd.font.name = "Calibri"
p_date.paragraph_format.space_after = Pt(8)

horizontal_rule(doc)

# Earnings Summary Table
heading(doc, "EARNINGS SUMMARY", 2)
body(doc, "Q1 FY2026 RESULTS: BEAT", bold=True, color=C_GREEN)

add_table(doc,
    ["Metric", "Reported", "Consensus", "Variance"],
    [
        ["Revenue", "$6.40B", "$6.28B", "+$120M (+1.9%)"],
        ["Non-GAAP EPS", "$6.06", "$5.87", "+$0.19 (+3.2%)"],
        ["GAAP EPS", "$4.60", "$4.45", "+$0.15 (+3.4%)"],
        ["Gross Margin (GAAP)", "89.6%", "89.2%", "+40bps"],
        ["Non-GAAP Op. Margin", "47.4%", "47.0%", "+40bps"],
    ],
    col_widths=[Inches(2.0), Inches(1.3), Inches(1.3), Inches(1.9)],
)

heading(doc, "KEY TAKEAWAYS", 3)

body(doc, (
    "■ Revenue of $6.40B beat consensus by $120M (+1.9%), representing 12% YoY growth "
    "(11% constant currency). This was driven by strong subscription revenue growth of 13% YoY "
    "to $6.17B, with Business Professionals & Consumers subscription revenue accelerating to "
    "+16% YoY ($1.78B) and Creative & Marketing Professionals growing 12% YoY ($4.39B). "
    "The beat was broad-based across customer groups and reflects continued AI-driven adoption."
))

body(doc, (
    "■ Non-GAAP EPS of $6.06 exceeded consensus of $5.87 by $0.19, up 19% YoY. The "
    "outperformance was fueled by both top-line strength and margin expansion, with non-GAAP "
    "operating margin expanding 120bps YoY to 47.4%. GAAP gross margin also expanded 50bps YoY "
    "to 89.6%, underscoring the inherent scalability of Adobe's subscription model."
))

body(doc, (
    "■ AI-first ARR more than tripled YoY, with total ARR reaching $26.06B (+10.9% YoY). "
    "Monthly active users surpassed 850 million (+17% YoY), demonstrating continued platform "
    "engagement. Products like Firefly, Acrobat AI Assistant, and GenStudio for Performance "
    "Marketing are gaining commercial traction, validating Adobe's AI monetization strategy."
))

body(doc, (
    "■ FY2026 guidance reaffirmed: revenue of $25.9–$26.1B and non-GAAP EPS of "
    "$23.30–$23.50. Q2 FY2026 guided to $6.43–$6.48B in revenue with non-GAAP EPS of "
    "$5.80–$5.85. CEO Shantanu Narayen announced he will transition from the CEO role "
    "after a successor is appointed, marking a significant leadership change after 18 years. "
    "We maintain our OUTPERFORM rating and $325 price target."
))

# Updated Estimates Table
heading(doc, "UPDATED FINANCIAL ESTIMATES", 2)

add_table(doc,
    ["Metric", "FY2026E (Prior)", "FY2026E (New)", "Change", "FY2027E"],
    [
        ["Revenue ($B)", "$25.9 – $26.1", "$25.9 – $26.1", "Reaffirmed", "$28.8"],
        ["Revenue Growth (%)", "9 – 10%", "9 – 10%", "—", "~11%"],
        ["Non-GAAP Op. Margin", "~45%", "~45%", "—", "~46%"],
        ["Non-GAAP EPS ($)", "$23.30 – $23.50", "$23.30 – $23.50", "Reaffirmed", "$26.50"],
        ["Total ARR ($B)", "$27.8", "$27.8", "—", "$30.8"],
        ["P/E (NTM, x)", "10.2x", "10.2x", "—", "9.0x"],
    ],
    col_widths=[Inches(1.8), Inches(1.3), Inches(1.3), Inches(1.0), Inches(1.1)],
)

body(doc, "Note: FY2026E guidance reaffirmed per Q1 FY2026 earnings call. FY2027E are analyst estimates.",
     italic=True, color=C_GREY, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2-3 — DETAILED RESULTS ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "DETAILED RESULTS ANALYSIS", 1)

heading(doc, "Revenue Analysis", 2)

body(doc, (
    "Adobe delivered record Q1 FY2026 revenue of $6.40 billion, up 12% year-over-year "
    "(11% in constant currency), beating the Street consensus of $6.28 billion by approximately "
    "$120 million. This marks five consecutive quarters of 10%+ revenue growth and acceleration "
    "from the 10% growth posted in Q4 FY2025."
))

body(doc, (
    "Total subscription revenue reached $6.17 billion, growing 13% YoY and constituting 96.4% "
    "of total revenue. Within this, Business Professionals & Consumers subscription revenue of "
    "$1.78 billion grew 16% YoY, the fastest growth rate among customer groups, driven by "
    "Acrobat AI Assistant adoption and expanded document workflow capabilities. Creative & Marketing "
    "Professionals subscription revenue of $4.39 billion grew 12% YoY, supported by Firefly's "
    "commercial deployment and continued Creative Cloud strength."
))

add_image(doc, "adbe_chart1_revenue.png", width=Inches(6.0),
          caption="Figure 1 — Quarterly Revenue Progression (Source: Adobe Earnings Releases)")

heading(doc, "Quarterly Revenue Breakdown", 3)

add_table(doc,
    ["Metric", "Q1 FY25", "Q2 FY25", "Q3 FY25", "Q4 FY25", "Q1 FY26", "YoY Chg"],
    [
        ["Total Revenue ($B)", "$5.71", "$5.87", "$5.99", "$6.20", "$6.40", "+12.1%"],
        ["  Subscription ($B)", "$5.46", "$5.62", "$5.75", "$5.95", "$6.17", "+13.0%"],
        ["  C&M Professionals ($B)", "$3.92", "$4.03", "$4.12", "$4.27", "$4.39", "+12.0%"],
        ["  B&P Consumers ($B)", "$1.53", "$1.59", "$1.63", "$1.69", "$1.78", "+16.3%"],
        ["Revenue Growth (%)", "10.0%", "10.8%", "11.0%", "10.0%", "12.1%", "+210bps"],
    ],
    col_widths=[Inches(1.6)] + [Inches(0.85)] * 6,
)

body(doc, "Source: Adobe Earnings Releases, Q1 FY2025 – Q1 FY2026.",
     italic=True, color=C_GREY, size=8)

heading(doc, "Profitability Analysis", 2)

body(doc, (
    "GAAP gross margin expanded 50 basis points YoY to 89.6%, reflecting the high "
    "incremental margins of subscription revenue and operational efficiencies. Non-GAAP "
    "operating margin reached 47.4%, expanding 120bps YoY from 46.2% in Q1 FY2025, driven "
    "by revenue scale benefits offsetting continued investment in AI capabilities."
))

body(doc, (
    "GAAP operating income was $2.42 billion (37.8% margin) and non-GAAP operating income "
    "was $3.04 billion (47.4% margin). The delta between GAAP and non-GAAP primarily reflects "
    "stock-based compensation and amortization of intangible assets. GAAP EPS of $4.60 grew "
    "11% YoY, while non-GAAP EPS of $6.06 grew 19% YoY, with the faster non-GAAP growth "
    "reflecting both margin expansion and share repurchase benefits."
))

add_image(doc, "adbe_chart2_eps.png", width=Inches(6.0),
          caption="Figure 2 — Non-GAAP EPS Progression (Source: Adobe Earnings Releases; Bloomberg)")

add_image(doc, "adbe_chart5_margins.png", width=Inches(6.0),
          caption="Figure 3 — Profitability Trends (Source: Adobe Earnings Releases)")

heading(doc, "Margin Analysis", 3)

add_table(doc,
    ["Margin", "Q1 FY25", "Q2 FY25", "Q3 FY25", "Q4 FY25", "Q1 FY26", "YoY Chg"],
    [
        ["Gross Margin (GAAP)", "89.1%", "89.2%", "89.3%", "89.4%", "89.6%", "+50bps"],
        ["Op. Margin (GAAP)", "35.3%", "35.9%", "36.5%", "37.1%", "37.8%", "+250bps"],
        ["Op. Margin (Non-GAAP)", "46.2%", "46.5%", "46.8%", "47.0%", "47.4%", "+120bps"],
    ],
    col_widths=[Inches(1.8)] + [Inches(0.85)] * 6,
)

body(doc, "Source: Adobe Earnings Releases.", italic=True, color=C_GREY, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 4-5 — KEY METRICS & GUIDANCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "KEY METRICS & GUIDANCE", 1)

heading(doc, "Annualized Recurring Revenue (ARR)", 2)

body(doc, (
    "Total Adobe ARR exiting Q1 FY2026 reached $26.06 billion, growing 10.9% YoY. "
    "Management reaffirmed its FY2026 ARR growth target of 10.2%, implying continued "
    "momentum through the year. The ARR metric remains the single most important KPI "
    "for Adobe, as it captures the full-year run rate of subscription revenue and provides "
    "strong forward visibility."
))

body(doc, (
    "AI-first ARR more than tripled year-over-year, accelerating from the ~$125 million "
    "book of business disclosed in Q1 FY2025. Products driving this include Firefly "
    "(generative AI image creation), Acrobat AI Assistant, GenStudio for Performance Marketing, "
    "and AI-powered features embedded across Creative Cloud and Experience Cloud. While Adobe "
    "does not disclose an exact AI-first ARR figure, the >3x YoY growth rate signals meaningful "
    "commercial traction and positions AI as a material revenue contributor in FY2026–27."
))

add_image(doc, "adbe_chart4_arr.png", width=Inches(6.0),
          caption="Figure 4 — Total ARR Progression (Source: Adobe Earnings Releases)")

heading(doc, "Platform Engagement", 2)

body(doc, (
    "Monthly active users surpassed 850 million in Q1 FY2026, growing 17% YoY. This "
    "engagement metric is critical because it underpins Adobe's freemium-to-paid conversion "
    "funnel and creates a massive installed base for AI upsell opportunities. The acceleration "
    "in MAU growth reflects increased adoption of web-based and mobile versions of Creative "
    "Cloud tools, as well as AI-powered features attracting new users."
))

add_image(doc, "adbe_chart3_segments.png", width=Inches(6.0),
          caption="Figure 5 — Revenue Mix & Subscription Growth by Customer Group (Source: Adobe Q1 FY2026 Earnings Release)")

heading(doc, "Q2 FY2026 & Full-Year Guidance", 2)

body(doc, (
    "For Q2 FY2026, Adobe guided to revenue of $6.43–$6.48 billion (implying ~10–11% "
    "YoY growth), with non-GAAP EPS of $5.80–$5.85 and non-GAAP operating margin of "
    "approximately 44.5%. The sequential decline in operating margin is consistent with "
    "seasonal patterns and investment timing."
))

body(doc, (
    "Full-year FY2026 targets were reaffirmed: total revenue of $25.9–$26.1 billion, "
    "non-GAAP operating margin of approximately 45%, and non-GAAP EPS of $23.30–$23.50. "
    "The reaffirmation following a Q1 beat provides confidence in management's visibility "
    "and suggests potential for upward revisions as the year progresses."
))

add_image(doc, "adbe_chart7_guidance.png", width=Inches(6.0),
          caption="Figure 6 — FY2026 Guidance vs. Street Estimates (Source: Adobe Q1 FY2026 Earnings Release; Bloomberg)")

add_table(doc,
    ["Metric", "Q2 FY26 Guidance", "FY2026 Guidance", "Street Consensus"],
    [
        ["Revenue ($B)", "$6.43 – $6.48", "$25.9 – $26.1", "$26.0"],
        ["Non-GAAP EPS ($)", "$5.80 – $5.85", "$23.30 – $23.50", "$23.35"],
        ["Non-GAAP Op. Margin", "~44.5%", "~45%", "~45%"],
        ["ARR Growth", "—", "10.2%", "~10.5%"],
    ],
    col_widths=[Inches(1.6), Inches(1.5), Inches(1.6), Inches(1.5)],
)

body(doc, "Source: Adobe Q1 FY2026 Earnings Release; Bloomberg consensus as of March 2026.",
     italic=True, color=C_GREY, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 6-7 — UPDATED INVESTMENT THESIS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "UPDATED INVESTMENT THESIS", 1)

heading(doc, "Thesis Impact Assessment", 2)

body(doc, "■ Thesis Pillar 1: AI Monetization Driving Growth Acceleration", bold=True)
body(doc, "Status: STRENGTHENED", bold=True, color=C_GREEN)
body(doc, (
    "Q1 FY2026 results provided the strongest evidence yet that Adobe's AI strategy is "
    "translating to revenue. AI-first ARR more than tripling YoY, combined with 850M+ MAUs "
    "(+17% YoY), demonstrates that AI features are both attracting new users and enabling "
    "monetization through premium tiers and standalone products. The 12% revenue growth "
    "represents an acceleration from Q4 FY2025's 10%, further validating the thesis that "
    "AI integration across Creative Cloud, Document Cloud, and Experience Cloud is becoming "
    "a structural growth catalyst. We expect AI-driven revenue contribution to become "
    "increasingly material in FY2026–27."
))

body(doc, "■ Thesis Pillar 2: Subscription Model Provides Earnings Visibility & Margin Expansion",
     bold=True)
body(doc, "Status: STRENGTHENED", bold=True, color=C_GREEN)
body(doc, (
    "With 96.4% of revenue from subscriptions, Adobe's recurring revenue model continues to "
    "deliver high visibility and margin expansion. Non-GAAP operating margin expanded 120bps "
    "YoY to 47.4%, and gross margin reached 89.6% (+50bps YoY). Record Q1 operating cash flow "
    "of $2.96 billion underscores the model's cash generation power. The reaffirmed FY2026 "
    "non-GAAP operating margin target of ~45% (below Q1's 47.4%) reflects seasonal investment "
    "timing but keeps the long-term margin expansion trajectory intact."
))

body(doc, "■ Thesis Pillar 3: Competitive Moat in Creative and Document Workflows", bold=True)
body(doc, "Status: UNCHANGED", bold=True, color=C_BLUE)
body(doc, (
    "Adobe's competitive position remains robust, though emerging AI-native competitors "
    "(Canva, Figma alternatives, Midjourney) continue to challenge at the prosumer tier. "
    "The 850M+ MAU figure suggests Adobe's platform stickiness remains strong, and the "
    "integration of Firefly natively into Photoshop, Illustrator, and other tools provides "
    "a differentiated AI experience within established professional workflows. However, "
    "we note that the stock's 27% YTD decline suggests the market is pricing in competitive "
    "risk more aggressively than we believe is warranted."
))

heading(doc, "CEO Transition", 2)

body(doc, (
    "Shantanu Narayen announced he will transition from the CEO role after a successor "
    "is appointed. Narayen has led Adobe for 18 years, overseeing the company's transformation "
    "from a perpetual-license software business to a cloud-first subscription powerhouse. "
    "While CEO transitions introduce uncertainty, we view this as manageable given: (1) the "
    "orderly nature of the process (no immediate departure), (2) Adobe's deep bench of senior "
    "executives, and (3) the company's strong operational momentum. We will monitor the "
    "succession timeline closely but do not expect it to impact near-term execution."
))

heading(doc, "Risks", 3)

bullet(doc, "CEO succession risk: uncertainty around new leadership direction and market reception")
bullet(doc, "AI competitive disruption: emerging AI-native tools could erode Creative Cloud share at the prosumer tier")
bullet(doc, "Macro sensitivity: enterprise spending deceleration could slow Experience Cloud growth")
bullet(doc, "Valuation compression: stock has already declined ~27% YTD; further multiple contraction if growth decelerates")
bullet(doc, "Regulatory: potential AI copyright litigation related to training data")

add_image(doc, "adbe_chart10_ai_growth.png", width=Inches(6.0),
          caption="Figure 7 — AI-first ARR & MAU Growth (Source: Adobe Q1 FY2026 Earnings Release & Earnings Call)")

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 8-10 — VALUATION & ESTIMATES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "VALUATION & ESTIMATES", 1)

heading(doc, "Valuation Summary", 2)

body(doc, (
    f"At the current price of {price_str}, Adobe trades at approximately 10.2x NTM P/E "
    "(using FY2026E non-GAAP EPS midpoint of $23.40), which represents a significant "
    "discount to both its 3-year average NTM P/E of ~22.5x and the large-cap software "
    "peer group median of ~25x. We believe the current valuation reflects excessive "
    "discounting of competitive risks and CEO transition uncertainty."
))

body(doc, (
    "Our $325 price target is derived from a blended approach: (1) 60% weight on DCF "
    "analysis implying $340 per share (using a 9.0% WACC, 3.0% terminal growth rate, "
    "and updated revenue/margin estimates), and (2) 40% weight on a 14.0x NTM P/E multiple "
    "applied to our FY2027E non-GAAP EPS of $26.50, implying $371. The blended value of "
    "~$352 is rounded down to $325 to provide a conservative margin of safety. This implies "
    f"~{((325/mkt['price'])-1)*100:.0f}% upside from the current price." if isinstance(mkt['price'], float) else
    "Our $325 price target is derived from a blended DCF and multiple approach."
))

add_image(doc, "adbe_chart9_valuation.png", width=Inches(6.0),
          caption="Figure 8 — NTM P/E Multiple (Source: Bloomberg, FactSet)")

heading(doc, "Beat/Miss Summary & Cash Flow", 2)

add_image(doc, "adbe_chart6_beat_miss.png", width=Inches(6.0),
          caption="Figure 9 — Q1 FY2026 Beat/Miss Analysis (Source: Adobe Earnings Release; Bloomberg)")

add_image(doc, "adbe_chart8_cashflow.png", width=Inches(6.0),
          caption="Figure 10 — Operating Cash Flow (Source: Adobe Earnings Releases)")

body(doc, (
    "Adobe generated record Q1 operating cash flow of $2.96 billion, supporting continued "
    "share repurchases and balance sheet strength. Cash and short-term investments stood at "
    "$6.89 billion at quarter end. The strong cash generation further supports our thesis "
    "that Adobe can simultaneously invest in AI capabilities while returning capital to shareholders."
))

heading(doc, "Updated Detailed Estimates", 2)

add_table(doc,
    ["Metric", "FY2025A", "FY2026E", "YoY Chg", "FY2027E"],
    [
        ["Revenue ($B)", "$23.77", "$26.0", "+9.4%", "$28.8"],
        ["  Subscription ($B)", "$22.71", "$24.9", "+9.6%", "$27.6"],
        ["Gross Margin (%)", "89.3%", "89.5%", "+20bps", "89.7%"],
        ["Non-GAAP Op. Income ($B)", "$10.87", "$11.70", "+7.6%", "$13.25"],
        ["Non-GAAP Op. Margin (%)", "45.7%", "45.0%", "−70bps", "46.0%"],
        ["Non-GAAP EPS ($)", "$20.50", "$23.40", "+14.1%", "$26.50"],
        ["GAAP EPS ($)", "$16.60", "$18.50", "+11.4%", "$21.20"],
        ["Operating Cash Flow ($B)", "$10.80", "$11.50", "+6.5%", "$12.80"],
        ["Total ARR ($B)", "$25.20", "$27.80", "+10.3%", "$30.80"],
    ],
    col_widths=[Inches(1.8), Inches(1.2), Inches(1.2), Inches(1.0), Inches(1.2)],
)

body(doc, "Note: FY2025A = Actual. FY2026E/FY2027E = Estimates. A = Actual; E = Estimate.",
     italic=True, color=C_GREY, size=8)

heading(doc, "Price Target Methodology", 3)

add_table(doc,
    ["Method", "Weight", "Value", "Key Assumptions"],
    [
        ["DCF Analysis", "60%", "$340", "WACC 9.0%, TGR 3.0%, 10-yr forecast"],
        ["NTM P/E Multiple", "40%", "$371", "14.0x FY2027E EPS of $26.50"],
        ["Blended Value", "100%", "$352", "Rounded to $325 (conservative)"],
    ],
    col_widths=[Inches(1.5), Inches(0.8), Inches(0.8), Inches(3.2)],
)

# ══════════════════════════════════════════════════════════════════════════════
# SOURCES SECTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

heading(doc, "SOURCES & REFERENCES", 1)

heading(doc, "Earnings Materials (Q1 FY2026)", 3)

p1 = body(doc, "")
p1.clear()
r = p1.add_run("Earnings Release (March 12, 2026): ")
r.font.size = Pt(9)
r.font.name = "Times New Roman"
add_hyperlink(p1, "Adobe Delivers Record Q1 Results",
              "https://news.adobe.com/news/2026/03/adobe-q1fy26-financial-results")

p2 = body(doc, "")
p2.clear()
r = p2.add_run("Form 8-K / Earnings Press Release (Filed March 12, 2026): ")
r.font.size = Pt(9)
r.font.name = "Times New Roman"
add_hyperlink(p2, "SEC EDGAR Filing",
              "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0000796343&type=8-K")

p3 = body(doc, "")
p3.clear()
r = p3.add_run("Q1 FY2026 Earnings Call Transcript (March 12, 2026): ")
r.font.size = Pt(9)
r.font.name = "Times New Roman"
add_hyperlink(p3, "Adobe Investor Relations",
              "https://www.adobe.com/investor-relations.html")

p4 = body(doc, "")
p4.clear()
r = p4.add_run("Investor Datasheet (Q1 FY2026): ")
r.font.size = Pt(9)
r.font.name = "Times New Roman"
add_hyperlink(p4, "Adobe IR Materials",
              "https://www.adobe.com/investor-relations.html")

heading(doc, "Prior Quarter References", 3)

p5 = body(doc, "")
p5.clear()
r = p5.add_run("Q4 FY2025 Earnings Release (December 10, 2025): ")
r.font.size = Pt(9)
r.font.name = "Times New Roman"
add_hyperlink(p5, "Adobe Reports Record Q4 and FY2025 Revenue",
              "https://news.adobe.com/news/2025/12/122025-q4earnings")

heading(doc, "Consensus & Market Data", 3)

body(doc, "Bloomberg consensus estimates as of March 11, 2026 (pre-earnings)",
     italic=True, size=9, color=C_GREY)
body(doc, "Stock price and market data sourced from Yahoo Finance / Bloomberg",
     italic=True, size=9, color=C_GREY)

heading(doc, "Disclaimer", 3)
body(doc, (
    "This report is for informational purposes only and does not constitute investment advice. "
    "All estimates and projections are the author's own and are subject to change. "
    "Past performance is not indicative of future results. Investors should conduct their own "
    "due diligence before making investment decisions."
), italic=True, color=C_GREY, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = os.path.join(BASE, "ADBE_Q1_FY2026_Earnings_Update.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
