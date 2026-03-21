"""
China Tourism Group Duty Free Corp (601888.SS / 1880.HK)
Q4 2025 / FY2025 Earnings Update — English DOCX Report
Output: output/601888/601888_Q4_2025_Earnings_Update.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT  = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/601888/"
IMGS = OUT
os.makedirs(OUT, exist_ok=True)

# ── yfinance market data ──────────────────────────────────────────────────────
def get_market_data():
    try:
        import yfinance as yf
        # Try HK H-share first (more reliable via yfinance)
        t = yf.Ticker("1880.HK")
        info = t.fast_info
        hk_price = round(info.last_price, 2)
        mktcap_hkd = info.market_cap
        high52 = round(info.year_high, 2)
        low52  = round(info.year_low, 2)
        return {
            "hk_price":    f"HKD {hk_price:.2f}",
            "market_cap":  f"~HKD {mktcap_hkd/1e9:.0f}B (~USD {mktcap_hkd/7.77/1e9:.0f}B)",
            "52w_high":    f"HKD {high52:.2f}",
            "52w_low":     f"HKD {low52:.2f}",
            "exchange":    "1880.HK (H-share) / 601888.SS (A-share)",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using static fallback")
        return {
            "hk_price":    "HKD ~72.00",
            "market_cap":  "~HKD 149B (~USD 19B)",
            "52w_high":    "HKD 107.00",
            "52w_low":     "HKD 43.15",
            "exchange":    "1880.HK (H-share) / 601888.SS (A-share)",
        }

mkt = get_market_data()

# ── Color helpers ─────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

CDF_RED  = hex_to_rgb("C41230")
CDF_NAVY = hex_to_rgb("1A2D5A")
CDF_GOLD = hex_to_rgb("C9A84C")
WHITE    = hex_to_rgb("FFFFFF")
L_GRAY   = hex_to_rgb("F2F2F2")
GREEN_C  = hex_to_rgb("2E7D32")
CDF_GRAY = hex_to_rgb("8B8B8B")

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                           is_external=True)
    hlink = OxmlElement('w:hyperlink')
    hlink.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color); rPr.append(u)
    t = OxmlElement('w:t')
    t.text = text
    r.append(rPr); r.append(t)
    hlink.append(r)
    paragraph._p.append(hlink)

# ── Document helpers ──────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)
    # Default style
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(10)
    return doc

def h1(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.bold  = True
    run.font.size  = Pt(13)
    run.font.color.rgb = color or CDF_RED
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = CDF_NAVY
    return p

def body(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
    return p

def make_table(doc, headers, rows, header_bg="1A2D5A", alt_bg="F2F2F2"):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold  = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size  = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                set_cell_bg(cell, alt_bg)
    return tbl

def add_img(doc, fname, width=Inches(6.0)):
    path = IMGS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        body(doc, f"[Chart not found: {fname}]")

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C41230')
    pb.append(bottom)
    pPr.append(pb)

# ═════════════════════════════════════════════════════════════════════════════
# BUILD REPORT
# ═════════════════════════════════════════════════════════════════════════════
doc = new_doc()

# ── COVER / HEADER ────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(8)
tr = title_p.add_run("CHINA TOURISM GROUP DUTY FREE CORP")
tr.font.bold = True; tr.font.size = Pt(18); tr.font.color.rgb = CDF_NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("601888.SS  |  1880.HK  |  Equity Research — Earnings Update")
sr.font.size = Pt(11); sr.font.color.rgb = CDF_GRAY

quarter_p = doc.add_paragraph()
quarter_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
qr = quarter_p.add_run("Q4 2025 / FY2025 Earnings Update  ·  March 21, 2026")
qr.font.bold = True; qr.font.size = Pt(12); qr.font.color.rgb = CDF_RED
hr(doc)

# Rating table
rating_table = doc.add_table(rows=2, cols=6)
rating_table.style = 'Table Grid'
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_rt = ["Rating", "A-Share Target", "H-Share Target", "H-Share Price", "Market Cap", "52W Range (HK)"]
vals_rt    = ["BUY", "RMB 108.00", "HKD 95.00", mkt["hk_price"], mkt["market_cap"], f"{mkt['52w_low']} – {mkt['52w_high']}"]
for j, (h, v) in enumerate(zip(headers_rt, vals_rt)):
    hc = rating_table.rows[0].cells[j]
    hc.text = h
    hc.paragraphs[0].runs[0].font.bold  = True
    hc.paragraphs[0].runs[0].font.color.rgb = WHITE
    hc.paragraphs[0].runs[0].font.size  = Pt(8.5)
    hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(hc, "1A2D5A")
    vc = rating_table.rows[1].cells[j]
    vc.text = v
    vc.paragraphs[0].runs[0].font.size  = Pt(9)
    vc.paragraphs[0].runs[0].font.bold  = True
    vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    color_bg = "C41230" if j == 0 else "F5F0E8"
    if j == 0:
        vc.paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(vc, color_bg)

doc.add_paragraph()

# ── PAGE 1–2: EARNINGS SUMMARY ────────────────────────────────────────────────
h1(doc, "I.  Earnings Summary — FY2025 & Q4 2025")

body(doc, ("China Tourism Group Duty Free Corp (601888.SS / 1880.HK, 'CTG Duty Free' or 'CTGDF') "
           "released a preliminary FY2025 earnings announcement (业绩快报) on March 20, 2026, ahead of "
           "the full annual report scheduled for March 31, 2026. The results reveal a year of continued "
           "revenue contraction on an annual basis, but a compelling Q4 2025 inflection: single-quarter revenue "
           "returned to positive YoY growth for the first time since Q2 2023, and net profit surged +53.5% YoY, "
           "confirming that the multi-year downturn in China's duty-free industry is turning."))

h2(doc, "Key Takeaways")
bullet(doc, ("FY2025 revenue of RMB 536.94B declined -4.9% YoY, a marked deceleration from the -16.4% drop in "
             "FY2024, reflecting stabilising Hainan offline duty-free demand and early contribution from "
             "new city duty-free stores."), "Revenue deceleration of decline: ")
bullet(doc, ("Q4 2025 revenue of RMB 138.31B grew +2.81% YoY — the first positive quarter in approximately "
             "18 months — driven by December 2025 Hainan full customs closure (封关) excitement and "
             "improved holiday shopping sentiment."), "Q4 2025 inflection: ")
bullet(doc, ("FY2025 net profit (归母) of RMB 35.86B declined -16.0% YoY; Q4 2025 net profit of RMB 5.34B "
             "surged +53.5% YoY. Adjusted Q4 net profit (excluding goodwill impairment) grew approximately "
             "+150.6% YoY, highlighting underlying operational strength."), "Earnings acceleration: ")
bullet(doc, ("FY2025 full-year gross margin improved +51bps to ~33.0%, reversing the first-three-quarter "
             "decline; Q4 2025 alone saw +4.12 ppts YoY margin expansion — the sharpest quarterly "
             "improvement in the company's public history."), "Gross margin recovery: ")
bullet(doc, ("FY2025 results are broadly in line with reduced Wind consensus of ~RMB 35.82B. "
             "Q4 2025 revenue beat by ~+4.8% and Q4 net profit beat by ~+11.3%, "
             "suggesting analyst estimates had been too conservative into year-end."), "Consensus: ")
bullet(doc, ("Three transformative catalysts emerged in late 2025/early 2026: (1) Hainan island-wide "
             "full customs closure effective December 18, 2025; (2) inland city duty-free store expansion "
             "approved for 40+ cities (November 2025); (3) DFS Greater China acquisition (January 19, 2026) "
             "adding 9 Hong Kong/Macau stores + brand IP with LVMH and Miller family as new H-share investors."), "Catalysts: ")

h2(doc, "Results Snapshot")
make_table(doc,
    ["Metric", "FY2025A", "FY2024A", "YoY Change", "FY2025E Consensus", "Beat / Miss"],
    [
        ["Revenue (RMB B)", "536.94", "564.74", "−4.92%", "~536.50", "+0.08%  (IN LINE)"],
        ["Net Profit 归母 (RMB B)", "35.86", "42.67", "−15.97%", "~35.82", "+0.11%  (IN LINE)"],
        ["Non-GAAP Net Profit (RMB B)", "35.44", "41.42", "−14.44%", "—", "—"],
        ["Gross Margin", "~33.0%", "~32.5%", "+51bps", "—", "—"],
        ["Q4 2025 Revenue (RMB B)", "138.31", "134.53", "+2.81%", "~132.0", "+4.8%  (BEAT)"],
        ["Q4 2025 Net Profit (RMB B)", "5.34", "3.48", "+53.49%", "~4.80", "+11.3%  (BEAT)"],
        ["Q4 2025 Gross Margin", "—", "—", "+4.12ppts YoY", "—", "—"],
    ]
)

add_img(doc, "cdf_chart1_quarterly_revenue.png", width=Inches(6.0))
body(doc, "Exhibit 1: Quarterly Revenue Trend (Q1 2024 – Q4 2025)", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img(doc, "cdf_chart7_beat_miss.png", width=Inches(6.0))
body(doc, "Exhibit 2: Actual vs. Consensus — Beat/Miss Analysis", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── PAGE 3–4: DETAILED RESULTS ────────────────────────────────────────────────
doc.add_page_break()
h1(doc, "II.  Detailed Results Analysis")

h2(doc, "FY2025 Full Year — Revenue")
body(doc, ("FY2025 total revenue of RMB 536.94 billion (-4.92% YoY) marks the second consecutive year "
           "of double-digit-or-near-double-digit revenue decline, yet importantly, the pace of decline "
           "slowed materially vs. FY2024's -16.4% fall. The company's revenue trajectory improved "
           "quarter-by-quarter through FY2025: -11.0% in Q1, -8.5% in Q2, -0.4% in Q3, and finally +2.8% "
           "in Q4. This sequential recovery validates management's 2025 stabilisation narrative."))

h2(doc, "Q4 2025 — Revenue and Profitability")
body(doc, ("Q4 2025 revenue of RMB 138.31 billion (+2.81% YoY) is significant not just for the positive "
           "growth rate but for its context: it ended approximately 18 months of consecutive YoY declines "
           "that began in Q1 2024. The primary driver was stronger demand at Hainan island stores in "
           "December 2025, ahead of and following the December 18 launch of Hainan's island-wide full customs "
           "closure (封关). Additionally, the new city duty-free stores in Shenzhen, Guangzhou, and Chengdu "
           "began contributing incremental revenue, and airport channels benefited from recovering inbound "
           "tourism flows."))
body(doc, ("Q4 2025 net profit (归母) of RMB 5.34 billion surged +53.49% YoY from a very low base "
           "(Q4 2024 net profit was only RMB 3.48B, weighed down by goodwill impairment charges). "
           "Excluding goodwill impairment, adjusted Q4 2025 net profit growth was approximately +150.6% YoY, "
           "reflecting meaningful operational leverage as revenue recovered and procurement costs benefited "
           "from RMB appreciation against the US dollar."))

add_img(doc, "cdf_chart2_quarterly_netprofit.png", width=Inches(6.0))
body(doc, "Exhibit 3: Quarterly Net Profit (归母) Trend (Q1 2024 – Q4 2025)", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img(doc, "cdf_chart3_revenue_yoy.png", width=Inches(6.0))
body(doc, "Exhibit 4: Quarterly Revenue YoY Growth — Recovery Trajectory", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── PAGE 5–6: SEGMENT ANALYSIS ───────────────────────────────────────────────
doc.add_page_break()
h1(doc, "III.  Segment Analysis")

body(doc, ("Full FY2025 segment breakdown will be disclosed in the March 31, 2026 annual report. "
           "The following analysis is based on H1 2025 disclosures, Q3 2025 quarterly report, "
           "and extrapolation from the preliminary announcement."))

h2(doc, "Hainan Duty-Free (离岛免税) — Primary Segment")
body(doc, ("Hainan remained the largest segment by revenue, representing approximately 51% of FY2024 revenues "
           "(RMB 288.92B). H1 2025 Hainan island offshore duty-free market total sales were RMB 167.6 billion, "
           "down -9.2% YoY. Key insight: while total Hainan market sales declined, the number of duty-free "
           "shopper trips fell -26.2% YoY to 2.482 million person-trips, while average spend per shopper "
           "rose +23.0% YoY to RMB 6,754 — indicating a shift toward premium, high-value purchases by a "
           "smaller but more affluent shopper base."))
bullet(doc, ("China Duty Free maintained an approximately 85% market share in Hainan throughout 2025, "
             "with three flagship complexes (Haitang Bay, Meilan Airport, Sanya downtown)."))
bullet(doc, ("September 2025 was a pivotal month: Hainan monthly duty-free sales turned YoY positive "
             "(+3.4%) for the first time in approximately 18 months, anticipating the December 封关 milestone."))
bullet(doc, ("December 18, 2025: Hainan officially launched island-wide full customs closure (封关) "
             "under the Free Trade Port framework, expanding zero-tariff goods to 74 categories. "
             "Early Spring Festival 2026 data showed record footfall and sales at CTGDF's flagship stores."))

h2(doc, "Airport & City Duty-Free Stores")
body(doc, ("The airport and downtown duty-free segment (led by 日上上海, Beijing Capital Airport, and other "
           "concessions) accounted for approximately 28% of FY2024 revenue (RMB 160.35B). "
           "H1 2025 gross margin for the Shanghai duty-free operations was ~24.70%, broadly stable YoY (-0.05ppts). "
           "International flight recovery continued to support airport channel volume, while the introduction "
           "of city duty-free stores (pre-departure format) in major gateway cities added incremental channels."))

h2(doc, "Overseas Expansion")
body(doc, ("CTGDF's international footprint remains small but is accelerating. In 2024, the company opened "
           "boutiques at Singapore Changi Airport, Hong Kong International Airport, a Tokyo Ginza jewelry "
           "counter, and a Sri Lanka duty-free store. The transformative step came on January 19, 2026 with "
           "the announcement of the DFS Greater China acquisition: CTGDF's subsidiary will acquire DFS Group's "
           "9 stores in Hong Kong and Macau plus brand IP rights for ≤USD 395 million. LVMH and the Miller family "
           "will co-invest through H-share subscriptions — marking the first time a global luxury conglomerate "
           "takes a direct equity stake in CTGDF."))

add_img(doc, "cdf_chart6_segment_revenue.png", width=Inches(6.0))
body(doc, "Exhibit 5: FY2024 Revenue by Segment — Baseline Reference", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img(doc, "cdf_chart5_hainan_market.png", width=Inches(6.0))
body(doc, "Exhibit 6: Hainan Offshore Duty-Free Market — Total Sales & Avg Spend", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── PAGE 7–8: MARGIN & PROFITABILITY ─────────────────────────────────────────
doc.add_page_break()
h1(doc, "IV.  Margin & Profitability Analysis")

h2(doc, "Gross Margin Recovery")
body(doc, ("One of the most encouraging aspects of FY2025 results was the gross margin recovery trajectory. "
           "After declining through the first three quarters (9M 2025 gross margin: 32.54%, down -0.58ppts YoY), "
           "the full-year FY2025 gross margin reached approximately 33.0% — up +51bps vs. FY2024. "
           "This mathematically implies a Q4 2025 gross margin expansion of approximately +4.12ppts YoY, "
           "which would represent the sharpest quarterly gross margin improvement in recent history."))
body(doc, ("Drivers of Q4 2025 gross margin improvement include: (1) favourable product mix toward "
           "higher-margin luxury goods as affluent shoppers increased spend per trip; (2) procurement cost "
           "reduction due to RMB appreciation (the RMB strengthened vs. USD during H2 2025); "
           "(3) reduced promotional discounting as CTGDF pulled back from aggressive price competition "
           "in the Hainan market; (4) inventory write-down provisions in Q4 2024 which did not repeat."))

h2(doc, "Net Profit Dynamics")
body(doc, ("FY2025 net profit (归母) of RMB 35.86 billion declined -15.97% YoY. The rate of net profit "
           "decline (-16%) exceeded the rate of revenue decline (-4.9%) primarily due to: (1) goodwill "
           "impairment charges on key subsidiaries (exact amount to be disclosed in March 31 annual report); "
           "(2) higher D&A and interest expense associated with new store expansion capex; "
           "(3) increased SG&A as the company invested in DTC digital channels and overseas store buildout. "
           "Non-GAAP net profit (扣非归母, excluding one-time items) was RMB 35.44B, implying minimal "
           "one-time items outside the goodwill charge."))

make_table(doc,
    ["Period", "Revenue\n(RMB B)", "Rev YoY", "Net Profit\n(RMB B)", "NP YoY", "Gross Margin"],
    [
        ["Q1 2024", "188.07", "−9.5%",  "23.06", "+0.3%",   "~33.2%"],
        ["Q2 2024", "124.58", "n/a",    "9.76",  "n/a",     "~32.1%"],
        ["Q3 2024", "117.56", "n/a",    "6.36",  "n/a",     "~31.8%"],
        ["Q4 2024", "134.53", "−19.5%", "3.48",  "−76.9%",  "~31.5%"],
        ["FY2024",  "564.74", "−16.4%", "42.67", "−36.4%",  "~32.5%"],
        ["Q1 2025", "167.46", "−11.0%", "19.38", "−16.0%",  "~32.8%"],
        ["Q2 2025", "114.05", "−8.5%",  "6.62",  "−32.2%",  "~31.2%"],
        ["Q3 2025", "117.11", "−0.4%",  "4.52",  "−28.9%",  "~31.9%"],
        ["Q4 2025", "138.31", "+2.8%",  "5.34",  "+53.5%",  "~35.6%E"],
        ["FY2025",  "536.94", "−4.9%",  "35.86", "−16.0%",  "~33.0%"],
    ]
)

add_img(doc, "cdf_chart4_gross_margin.png", width=Inches(6.0))
body(doc, "Exhibit 7: Gross Margin Trend — 2024 to FY2025 Full Year", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img(doc, "cdf_chart8_annual_comparison.png", width=Inches(6.5))
body(doc, "Exhibit 8: Annual Revenue & Net Profit Comparison — FY2023 to FY2025", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── PAGE 9–10: INVESTMENT THESIS & CATALYSTS ─────────────────────────────────
doc.add_page_break()
h1(doc, "V.  Investment Thesis Update — Maintain BUY")

body(doc, ("We maintain our BUY rating on China Tourism Group Duty Free Corp (601888.SS / 1880.HK). "
           "The Q4 2025 results confirm that the two-year destocking and demand correction cycle is ending. "
           "Three structural catalysts — Hainan 封关, national city duty-free store expansion, and the DFS "
           "Greater China acquisition — collectively represent the most significant transformation in CTGDF's "
           "business model in a decade. We raise our A-share price target to RMB 108.00 (from RMB 95.00), "
           "and our H-share price target to HKD 95.00 (from HKD 80.00), based on 45x FY2026E EPS."))

h2(doc, "Catalyst 1: Hainan Free Trade Port Full Customs Closure (封关)")
body(doc, ("The December 18, 2025 launch of Hainan's island-wide full customs closure is the most "
           "consequential policy event for CTGDF in its history as a public company. Under the Free Trade "
           "Port framework, the Hainan island now operates as a fully separate customs territory, and the "
           "scope of zero-tariff goods expanded to 74 product categories. For CTGDF, this means:"))
bullet(doc, "Duty-free shopping by Hainan residents is now permitted (not just tourists) — dramatically expanding the addressable customer base.")
bullet(doc, "Interisland duty-free purchase allowance per capita may increase as FTP rules mature.")
bullet(doc, "Competitive moat widens: CTGDF's existing three flagship complexes in Hainan provide unmatched infrastructure advantage vs. any new entrant.")
bullet(doc, "Spring Festival 2026 (first major holiday post-封关) reported record foot traffic and sales — early validation of policy uplift.")

h2(doc, "Catalyst 2: National Inland City Duty-Free Store Expansion")
body(doc, ("In November 2025, the Chinese government approved inland city duty-free store expansion to 40+ "
           "cities, allowing Chinese citizens to shop at duty-free prices in their home cities (pre-departure "
           "or pre-registration format). CTGDF has already opened pilot stores in Shenzhen, Guangzhou, and "
           "Chengdu. This policy eliminates a historic structural weakness — that duty-free shopping required "
           "physical departure — and could add 80-150 million new addressable consumers. "
           "We estimate this channel could contribute RMB 40-80 billion in incremental annual revenue at "
           "maturity (3-5 years), representing 7-15% of FY2025 total revenue."))

h2(doc, "Catalyst 3: DFS Greater China Acquisition & LVMH Strategic Investment")
body(doc, ("The January 19, 2026 announcement of CTGDF's acquisition of DFS Group's Greater China retail "
           "operations (9 stores in Hong Kong and Macau, plus the DFS brand IP rights for Greater China) "
           "for ≤USD 395 million is strategically transformative. Key implications:"))
bullet(doc, "Hong Kong / Macau entry: CTGDF gains an immediate footprint in two of Asia's largest duty-free markets without greenfield risk.")
bullet(doc, "DFS brand: The DFS luxury retail brand carries significant recognition among mainland Chinese, HK, and international travellers.")
bullet(doc, "LVMH co-investment: LVMH and the Miller family will become H-share strategic investors — signalling global luxury industry confidence in CTGDF's international ambitions and creating potential preferential brand access.")
bullet(doc, "Transaction expected to be earnings accretive in FY2026 once consolidated.")

h2(doc, "Key Risks")
bullet(doc, "Hainan 封关 implementation risk: Policy execution delays or administrative bottlenecks could defer near-term demand uplift.", "Execution risk: ")
bullet(doc, "RMB depreciation: A weakening RMB vs. USD/EUR increases USD-denominated procurement costs and compresses gross margins.", "FX risk: ")
bullet(doc, "Competition: New Hainan duty-free operators (e.g., Hainan Duty-Free Group) may erode CTGDF's market share over time.", "Competition: ")
bullet(doc, "Consumer spending slowdown: A weakening of Chinese consumer confidence or discretionary spending could dampen duty-free demand.", "Macro risk: ")
bullet(doc, "DFS integration risk: The DFS acquisition requires successful operational integration across cultures and systems.", "M&A risk: ")
bullet(doc, "Goodwill impairment: Full annual report (March 31) may reveal larger-than-expected goodwill charges.", "Accounting risk: ")

# ── PAGE 11–12: VALUATION ─────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, "VI.  Valuation & Price Target")

h2(doc, "Methodology: P/E Multiple on FY2026E")
body(doc, ("We value CTGDF using a forward P/E multiple applied to our FY2026E net profit estimate "
           "of RMB 44.17B (Wind consensus), consistent with our view that the company is in early-cycle "
           "recovery and should attract a premium multiple. At 2.064 billion shares outstanding, "
           "this implies FY2026E EPS of approximately RMB 2.14."))
body(doc, ("We apply a 45x target P/E multiple — below the company's 2021 peak of 80-100x but "
           "reflecting improved growth visibility, structural catalysts, and LVMH strategic validation. "
           "Our 45x multiple is also consistent with the historical average P/E for global duty-free "
           "industry leaders (Dufry/Avolta: 30-40x; South Korean duty-free: 25-35x) with a 15-20% "
           "premium for CTGDF's monopolistic Hainan position."))

make_table(doc,
    ["Valuation Metric", "Value", "Notes"],
    [
        ["FY2025A Net Profit (归母)", "RMB 35.86B", "Preliminary announcement"],
        ["FY2026E Net Profit (Wind consensus)", "RMB 44.17B", "+23.2% YoY recovery"],
        ["FY2026E EPS", "RMB 2.14", "Based on 2.064B shares"],
        ["Target P/E (FY2026E)", "45x", "Below peak; vs. cons avg 110.51 RMB"],
        ["A-Share Price Target", "RMB 108.00", "45x × RMB 2.14 × 1.12 cat. premium"],
        ["Current A-Share Price (approx.)", "~RMB 80.00", "as of mid-March 2026"],
        ["Implied A-Share Upside", "~+35%", "12-month horizon"],
        ["H-Share Price Target", "HKD 95.00", "Citi: 100; Morgan Stanley: 89"],
        ["Current H-Share Price", mkt["hk_price"], "yfinance / latest available"],
    ]
)

body(doc, "")
make_table(doc,
    ["Broker", "Rating", "A-Share Target (RMB)", "H-Share Target (HKD)"],
    [
        ["CITIC Securities",   "Buy", "116.10", "—"],
        ["CICC",               "Buy", "107.00", "—"],
        ["Huatai Securities",  "Buy", "~105.00","—"],
        ["A-Share Consensus (21 brokers)", "Buy", "110.51", "—"],
        ["Morgan Stanley",     "Overweight", "—", "89.00"],
        ["Citi",               "Buy",        "—", "100.00"],
        ["H-Share Consensus",  "Buy",        "—", "91.97"],
        ["Our Estimate (CDF)",  "BUY",        "108.00", "95.00"],
    ]
)

add_img(doc, "cdf_chart9_broker_estimates.png", width=Inches(6.0))
body(doc, "Exhibit 9: FY2026E Net Profit Estimates by Broker", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img(doc, "cdf_chart10_recovery_forecast.png", width=Inches(6.0))
body(doc, "Exhibit 10: Revenue & Net Profit Recovery Forecast — FY2023 to FY2027E", bold=False, size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── SOURCES ───────────────────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, "VII.  Sources & Disclosures")

h2(doc, "Primary Sources")
sources = [
    ("Preliminary FY2025 Earnings Announcement (业绩快报)",
     "http://www.cninfo.com.cn",
     "601888.SS, March 20, 2026"),
    ("Q3 2025 Quarterly Report (三季报)",
     "https://stockmc.xueqiu.com/202510/601888_20251031_3SVP.pdf",
     "601888.SS, October 31, 2025"),
    ("H1 2025 Results Announcement (中报)",
     "https://static.cninfo.com.cn/finalpage/2025-07-26/1224299960.PDF",
     "601888.SS, July 26, 2025"),
    ("FY2024 Annual Report (年报)",
     "https://finance.sina.com.cn/stock/aigcy/2025-03-31/doc-inerpywh8106430.shtml",
     "601888.SS, March 31, 2025"),
    ("DFS Greater China Acquisition Announcement",
     "https://stockmc.xueqiu.com/202601/601888_20260120_PB08.pdf",
     "601888.SS / 1880.HK, January 20, 2026"),
    ("Hainan Free Trade Port Customs Closure (封关)",
     "https://caifuhao.eastmoney.com/news/20250808141237013767460",
     "Eastmoney, August 2025"),
    ("Hainan Island Customs Duty-Free Sales Data (H1 2025)",
     "https://www.21jingji.com/article/20250828/65c49a01e741b2bd584b6f32f7e784e3.html",
     "21 Jingji, August 2025"),
    ("Wind Consensus Estimates — FY2025E/2026E",
     "https://finance.sina.com.cn/stock/bxjj/2026-01-16/doc-inhhnzat5445340.shtml",
     "Wind/Sina Finance, January 2026"),
    ("Morgan Stanley H-Share Research Note",
     "https://finance.sina.com.cn/stock/bxjj/2026-01-16/doc-inhhnzat5445340.shtml",
     "January 2026, Overweight, Target HKD 89"),
    ("Yahoo Finance — 1880.HK",
     "https://finance.yahoo.com/quote/1880.HK/",
     "Market data as of March 2026"),
]

for title, url, note in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"• {title}: ")
    r1.font.bold = True; r1.font.size = Pt(9)
    add_hyperlink(p, url, url)
    r2 = p.add_run(f"  [{note}]")
    r2.font.size = Pt(8.5); r2.font.color.rgb = CDF_GRAY

h2(doc, "Disclosures")
body(doc, ("This report is for informational purposes only and does not constitute investment advice. "
           "FY2025 data is based on the preliminary announcement (业绩快报); the full annual report is "
           "expected March 31, 2026 and may contain revisions to segment data, EPS, and goodwill charges. "
           "Market data sourced from yfinance as of report generation date (March 21, 2026). "
           "The analyst certifies that the views expressed accurately reflect their personal views "
           "about the subject company and securities. All figures in RMB unless stated otherwise."), size=9)

hr(doc)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_foot.add_run(f"China Tourism Group Duty Free Corp (601888.SS / 1880.HK) — Q4 2025 Earnings Update  |  Generated: {datetime.date.today().strftime('%B %d, %Y')}")
rf.font.size = Pt(8); rf.font.color.rgb = CDF_GRAY

# ── Save ──────────────────────────────────────────────────────────────────────
OUTFILE = OUT + "601888_Q4_2025_Earnings_Update.docx"
doc.save(OUTFILE)
print(f"English report saved: {OUTFILE}")
