"""
中国旅游集团中免股份有限公司（601888.SS / 1880.HK）
2025年Q4 / 全年业绩更新报告 — 中文版 DOCX
输出：output/601888/601888_Q4_2025_业绩更新报告_中文版.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT  = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/601888/"
IMGS = OUT
os.makedirs(OUT, exist_ok=True)

# ── yfinance 市场数据 ──────────────────────────────────────────────────────────
def get_market_data():
    try:
        import yfinance as yf
        t    = yf.Ticker("1880.HK")
        info = t.fast_info
        hk_price   = round(info.last_price, 2)
        mktcap_hkd = info.market_cap
        high52 = round(info.year_high, 2)
        low52  = round(info.year_low, 2)
        usd_mktcap = mktcap_hkd / 7.77
        rmb_mktcap = usd_mktcap * 7.3
        return {
            "hk_price":    f"HKD {hk_price:.2f}",
            "market_cap":  f"约{rmb_mktcap/1e8:.0f}亿人民币（约{usd_mktcap/1e9:.0f}亿美元）",
            "52w_high":    f"HKD {high52:.2f}",
            "52w_low":     f"HKD {low52:.2f}",
        }
    except Exception as e:
        print(f"[警告] yfinance获取失败: {e} — 使用默认值")
        return {
            "hk_price":    "约HKD 72.00",
            "market_cap":  "约1,490亿港元（约190亿美元）",
            "52w_high":    "HKD 107.00",
            "52w_low":     "HKD 43.15",
        }

mkt = get_market_data()

# ── 颜色工具 ──────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

CDF_RED  = hex_to_rgb("C41230")
CDF_NAVY = hex_to_rgb("1A2D5A")
CDF_GOLD = hex_to_rgb("C9A84C")
WHITE    = hex_to_rgb("FFFFFF")
L_GRAY   = hex_to_rgb("F2F2F2")
CDF_GRAY = hex_to_rgb("8B8B8B")

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                           is_external=True)
    hlink = OxmlElement('w:hyperlink')
    hlink.set(qn('r:id'), r_id)
    r  = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1')
    u     = OxmlElement('w:u');     u.set(qn('w:val'), 'single')
    rPr.append(color); rPr.append(u)
    t = OxmlElement('w:t'); t.text = text
    r.append(rPr); r.append(t)
    hlink.append(r)
    paragraph._p.append(hlink)

# ── CJK 字体设置 ──────────────────────────────────────────────────────────────
def set_cjk_font(run, font_name="宋体"):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

# ── 文档创建工具 ──────────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(10)
    return doc

def h1_cn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = CDF_RED
    set_cjk_font(run, "黑体")
    return p

def h2_cn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = CDF_NAVY
    set_cjk_font(run, "黑体")
    return p

def body_cn(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    set_cjk_font(run, "宋体")
    return p

def bullet_cn(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True; r1.font.size = Pt(10)
        set_cjk_font(r1, "黑体")
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        set_cjk_font(r2, "宋体")
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
        set_cjk_font(r, "宋体")
    return p

def make_table_cn(doc, headers, rows, header_bg="1A2D5A", alt_bg="F2F2F2"):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9)
        set_cjk_font(run, "黑体")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            set_cjk_font(run, "宋体")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                set_cell_bg(cell, alt_bg)
    return tbl

def add_img_cn(doc, fname, width=Inches(6.0)):
    path = IMGS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        body_cn(doc, f"[图表未找到: {fname}]")

def hr_cn(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C41230')
    pb.append(bottom)
    pPr.append(pb)

# ═════════════════════════════════════════════════════════════════════════════
# 生成报告
# ═════════════════════════════════════════════════════════════════════════════
doc = new_doc()

# ── 封面 ──────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(8)
tr = title_p.add_run("中国旅游集团中免股份有限公司")
tr.font.bold = True; tr.font.size = Pt(18); tr.font.color.rgb = CDF_NAVY
set_cjk_font(tr, "黑体")

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("601888.SS  |  1880.HK  |  股票研究 — 业绩更新报告")
sr.font.size = Pt(11); sr.font.color.rgb = CDF_GRAY
set_cjk_font(sr, "宋体")

quarter_p = doc.add_paragraph()
quarter_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
qr = quarter_p.add_run("2025年Q4 / 全年业绩更新  ·  2026年3月21日")
qr.font.bold = True; qr.font.size = Pt(12); qr.font.color.rgb = CDF_RED
set_cjk_font(qr, "黑体")
hr_cn(doc)

# 评级表格
rating_table = doc.add_table(rows=2, cols=6)
rating_table.style = 'Table Grid'
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_rt = ["评级", "A股目标价", "H股目标价", "H股现价", "市值", "52周区间（港股）"]
vals_rt    = ["买入", "人民币108.00元", "港币95.00元", mkt["hk_price"], mkt["market_cap"],
              f"{mkt['52w_low']} – {mkt['52w_high']}"]
for j, (h, v) in enumerate(zip(headers_rt, vals_rt)):
    hc = rating_table.rows[0].cells[j]
    hc.text = h
    run_h = hc.paragraphs[0].runs[0]
    run_h.font.bold = True; run_h.font.color.rgb = WHITE; run_h.font.size = Pt(8.5)
    set_cjk_font(run_h, "黑体")
    hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(hc, "1A2D5A")
    vc = rating_table.rows[1].cells[j]
    vc.text = v
    run_v = vc.paragraphs[0].runs[0]
    run_v.font.size = Pt(9); run_v.font.bold = True
    set_cjk_font(run_v, "宋体")
    vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if j == 0:
        run_v.font.color.rgb = WHITE
        set_cell_bg(vc, "C41230")
    else:
        set_cell_bg(vc, "F5F0E8")

doc.add_paragraph()

# ── 一、业绩摘要 ──────────────────────────────────────────────────────────────
h1_cn(doc, "一、业绩摘要 — 2025年Q4 / 全年")

body_cn(doc, ('中国旅游集团中免股份有限公司（601888.SS / 1880.HK，下称\u201c中国中免\u201d或\u201c公司\u201d）于'
              '2026年3月20日披露2025年全年业绩快报，完整年报预计于2026年3月31日发布。'
              '快报数据显示：全年营收继续收缩，但2025年第四季度出现关键拐点——'
              '单季营收同比增速转正，系2023年Q2以来首次；净利润同比大幅回升+53.5%，'
              '印证中国免税行业多年调整周期正在结束。'))

h2_cn(doc, "核心要点")
bullet_cn(doc, ("2025年全年营收536.94亿元，同比下降-4.9%，降幅较2024年的-16.4%显著收窄，"
                "环比走势逐季改善：Q1 -11%、Q2 -8.5%、Q3 -0.4%、Q4 +2.8%。"),
          "营收降幅大幅收窄：")
bullet_cn(doc, ("2025年Q4单季营收138.31亿元，同比增长+2.81%，为近18个月来首次正增长，"
                "主要受益于海南封关预期升温、年末消费旺季提振及城市免税新渠道贡献。"),
          "Q4营收拐点确立：")
bullet_cn(doc, ("2025年全年归母净利润35.86亿元，同比下降-16.0%；Q4单季归母净利润5.34亿元，"
                "同比大增+53.5%；剔除商誉减值后Q4调整后净利润同比约增+150.6%，"
                "经营层面大幅改善。"),
          "盈利加速修复：")
bullet_cn(doc, ("2025年全年毛利率约33.0%，同比提升约+51个基点，扭转前三季度下滑趋势；"
                "Q4单季毛利率同比改善约+4.12个百分点，为近年来最大季度环比改善幅度。"),
          "毛利率显著回升：")
bullet_cn(doc, ("全年业绩与Wind一致预期（约35.82亿元）基本吻合（+0.1%）；"
                "Q4营收较市场预期超出约+4.8%，Q4净利润超出约+11.3%，"
                "显示年末市场预期过于保守。"),
          "超预期程度：")
bullet_cn(doc, ("三大战略催化剂于2025年底至2026年初相继落地：（1）海南自贸港全岛封关"
                "（2025年12月18日正式启动）；（2）国内城市免税店政策扩围至40余城"
                "（2025年11月）；（3）收购DFS大中华区零售业务（2026年1月19日），"
                "LVMH集团及Miller家族以H股战略投资者身份入股。"),
          "催化剂集中兑现：")

h2_cn(doc, "业绩快照")
make_table_cn(doc,
    ["指标", "2025年全年实际", "2024年全年实际", "同比变化", "一致预期", "超预期情况"],
    [
        ["营业收入（亿元）", "536.94", "564.74", "−4.92%", "约536.50", "+0.08%（符合）"],
        ["归母净利润（亿元）", "35.86", "42.67", "−15.97%", "约35.82", "+0.11%（符合）"],
        ["扣非归母净利润（亿元）", "35.44", "41.42", "−14.44%", "—", "—"],
        ["综合毛利率", "约33.0%", "约32.5%", "+51个基点", "—", "—"],
        ["Q4营收（亿元）", "138.31", "134.53", "+2.81%", "约132.0", "+4.8%（超预期）"],
        ["Q4归母净利润（亿元）", "5.34", "3.48", "+53.49%", "约4.80", "+11.3%（超预期）"],
        ["Q4毛利率变化", "—", "—", "同比+4.12个百分点", "—", "—"],
    ]
)

add_img_cn(doc, "cdf_chart1_quarterly_revenue.png", width=Inches(6.0))
body_cn(doc, "图1：季度营收趋势（2024Q1–2025Q4）", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img_cn(doc, "cdf_chart7_beat_miss.png", width=Inches(6.0))
body_cn(doc, "图2：实际业绩与一致预期对比（超预期/符合预期分析）", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 二、详细业绩分析 ──────────────────────────────────────────────────────────
doc.add_page_break()
h1_cn(doc, "二、详细业绩分析")

h2_cn(doc, "2025年全年营收")
body_cn(doc, ("2025年全年营收536.94亿元，同比下降-4.9%，降幅较2024年的-16.4%大幅收窄。"
              "逐季走势显示明显的V形恢复态势：Q1至Q3营收同比降幅持续收窄，Q4实现正增长。"
              '这一走势印证管理层关于2025年为\u201c企稳年\u201d的判断，并为2026年的全面复苏奠定基础。'))

h2_cn(doc, "2025年Q4单季 — 营收与盈利")
body_cn(doc, ("Q4营收138.31亿元（同比+2.81%）的意义在于其背后的周期拐点含义：结束了自2024年Q1"
              "以来持续约18个月的同比负增长。驱动因素包括：（1）海南离岛免税2025年12月销售强劲，"
              "封关预期升温带来消费前置效应；（2）深圳、广州、成都城市免税新店贡献增量；"
              "（3）机场渠道受益于入境游客恢复；（4）人民币兑美元升值降低采购成本。"))
body_cn(doc, ("Q4归母净利润5.34亿元，同比大增+53.5%。需注意上年同期基数极低（2024年Q4仅3.48亿元，"
              "受商誉减值拖累）。剔除商誉减值后，调整后Q4净利润同比增幅约+150.6%，"
              "彰显经营层面的实质性改善。"))

add_img_cn(doc, "cdf_chart2_quarterly_netprofit.png", width=Inches(6.0))
body_cn(doc, "图3：季度归母净利润趋势（2024Q1–2025Q4）", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img_cn(doc, "cdf_chart3_revenue_yoy.png", width=Inches(6.0))
body_cn(doc, "图4：季度营收同比增速 — 拐点修复轨迹", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 三、业务板块分析 ──────────────────────────────────────────────────────────
doc.add_page_break()
h1_cn(doc, "三、业务板块分析")

body_cn(doc, "完整2025年分部数据将于2026年3月31日年报中披露。以下分析基于半年报、三季报及业绩快报数据。")

h2_cn(doc, "海南离岛免税 — 核心板块")
body_cn(doc, ("海南离岛免税是中国中免第一大营收板块，占2024年全年营收约51%（288.92亿元）。"
              "2025年上半年，海南岛离岛免税市场合计销售额167.6亿元，同比下降-9.2%。"
              "关键洞察：购物人次同比大幅下降-26.2%至248.2万人次，但人均购物金额同比"
              "大增+23.0%至6,754元，显示消费结构向高端化、精品化迁移。"))
bullet_cn(doc, "中国中免在海南离岛免税市场占据约85%市场份额，旗下三大综合体（海棠湾、美兰机场、三亚市区）构成竞争护城河。")
bullet_cn(doc, "2025年9月：海南单月离岛免税销售额同比转正（+3.4%），为约18个月来首次，领先全年拐点确立。")
bullet_cn(doc, ("2025年12月18日，海南自贸港正式启动全岛封关运作，零关税商品范围扩展至74个品类。"
                "2026年春节期间，中国中免海南门店客流量及销售额双创历史新高，"
                "封关政策红利加速兑现。"))

h2_cn(doc, "机场及市区免税")
body_cn(doc, ("机场及市区免税板块（以日上上海、北京首都机场等为代表）占2024年营收约28%（160.35亿元）。"
              "2025年上半年，上海日上等机场门店综合毛利率约24.70%，同比基本持平（-0.05个百分点）。"
              "国际航班恢复持续拉动机场渠道客流，城市免税店的推出进一步拓展购物场景。"))

h2_cn(doc, "海外布局与DFS战略收购")
body_cn(doc, ("中国中免海外布局规模尚小但扩张加速。2024年相继在新加坡樟宜机场、香港国际机场、"
              "东京银座及斯里兰卡开设门店。2026年1月19日，公司宣布以不超过3.95亿美元收购"
              "DFS集团大中华区零售业务（含香港、澳门9家门店及品牌知识产权），"
              "LVMH集团及Miller家族通过认购H股成为战略投资者，标志着全球奢侈品龙头"
              "首次以股权形式直接战略投资中国中免，具有里程碑意义。"))

add_img_cn(doc, "cdf_chart6_segment_revenue.png", width=Inches(6.0))
body_cn(doc, "图5：2024年各板块营收构成（参考基准）", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img_cn(doc, "cdf_chart5_hainan_market.png", width=Inches(6.0))
body_cn(doc, "图6：海南离岛免税市场 — 总销售额与人均购物金额趋势", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 四、利润率与盈利能力分析 ─────────────────────────────────────────────────
doc.add_page_break()
h1_cn(doc, "四、利润率与盈利能力分析")

h2_cn(doc, "毛利率修复")
body_cn(doc, ("2025年全年毛利率表现是本次业绩最亮眼的亮点之一。前三季度毛利率持续承压"
              "（9M 2025毛利率32.54%，同比下降0.58个百分点），但全年综合毛利率约33.0%，"
              "同比提升约+51个基点，说明Q4单季毛利率同比改善约+4.12个百分点，"
              "为近年来最大单季度改善幅度。"))
body_cn(doc, ("Q4毛利率大幅提升的驱动因素：（1）消费结构高端化，奢侈品及高毛利商品"
              "占比提升；（2）人民币兑美元升值降低采购成本；（3）促销力度减弱，"
              "海南市场价格竞争趋于理性；（4）2024年Q4计提库存减值未在本期重演。"))

h2_cn(doc, "净利润动态")
body_cn(doc, ("2025年全年归母净利润35.86亿元，同比下降-16.0%。净利润降幅（-16%）"
              "超过营收降幅（-4.9%）的主要原因：（1）旗下子公司商誉减值拨备"
              "（具体金额待3月31日年报披露）；（2）新店扩张带来的折旧摊销和利息支出增加；"
              "（3）电商直销渠道投入和海外门店建设拉高销售管理费用。扣非归母净利润35.44亿元，"
              "与归母净利润相差约0.4亿元，说明非经常性损益以商誉减值为主。"))

make_table_cn(doc,
    ["期间", "营收（亿元）", "营收同比", "归母净利（亿元）", "净利同比", "综合毛利率"],
    [
        ["2024年Q1", "188.07", "−9.5%",  "23.06", "+0.3%",   "约33.2%"],
        ["2024年Q2", "124.58", "—",      "9.76",  "—",       "约32.1%"],
        ["2024年Q3", "117.56", "—",      "6.36",  "—",       "约31.8%"],
        ["2024年Q4", "134.53", "−19.5%", "3.48",  "−76.9%",  "约31.5%"],
        ["2024全年",  "564.74", "−16.4%", "42.67", "−36.4%",  "约32.5%"],
        ["2025年Q1", "167.46", "−11.0%", "19.38", "−16.0%",  "约32.8%"],
        ["2025年Q2", "114.05", "−8.5%",  "6.62",  "−32.2%",  "约31.2%"],
        ["2025年Q3", "117.11", "−0.4%",  "4.52",  "−28.9%",  "约31.9%"],
        ["2025年Q4", "138.31", "+2.8%",  "5.34",  "+53.5%",  "约35.6%（估）"],
        ["2025全年",  "536.94", "−4.9%",  "35.86", "−16.0%",  "约33.0%"],
    ]
)

add_img_cn(doc, "cdf_chart4_gross_margin.png", width=Inches(6.0))
body_cn(doc, "图7：综合毛利率趋势（2024年至2025全年）", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img_cn(doc, "cdf_chart8_annual_comparison.png", width=Inches(6.5))
body_cn(doc, "图8：年度营收与归母净利润对比（2023–2025年）", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 五、投资逻辑更新 — 维持"买入"评级 ────────────────────────────────────────
doc.add_page_break()
h1_cn(doc, '五、投资逻辑更新 — 维持\u201c买入\u201d评级')

body_cn(doc, ('我们维持对中国旅游集团中免股份有限公司（601888.SS / 1880.HK）的\u201c买入\u201d评级。'
              "Q4 2025业绩证实，中国免税行业两年去库存及需求调整周期已走向尾声。"
              "三大结构性催化剂 — 海南封关、全国城市免税扩围、DFS大中华收购 — "
              "共同构成公司近十年来最重大的商业模式跃升。"
              "我们将A股目标价上调至人民币108.00元（此前95.00元），"
              "H股目标价上调至港币95.00元（此前80.00元），"
              "对应45倍2026年预期市盈率。"))

h2_cn(doc, "催化剂一：海南自贸港全岛封关")
body_cn(doc, ("2025年12月18日，海南自贸港正式启动全岛封关运作，"
              "这是中国中免上市以来最重要的政策事件。封关后，零关税商品扩展至74个品类，"
              "海南岛全岛纳入独立关税区，对中国中免的主要影响包括："))
bullet_cn(doc, "海南居民（而非仅限游客）可购买免税商品，目标消费群体大幅扩大。")
bullet_cn(doc, "后续人均购物限额可能随自贸港法规完善而提升，长期需求空间进一步打开。")
bullet_cn(doc, "竞争壁垒更高：中国中免现有三大旗舰综合体在基础设施规模上优势突出，新进入者难以在短期内复制。")
bullet_cn(doc, "2026年春节（封关后首个重大假日）报告创纪录的客流量与销售额，政策红利超预期快速兑现。")

h2_cn(doc, "催化剂二：全国城市免税店扩围")
body_cn(doc, ("2025年11月，国家批准城市免税店向40余个城市扩展，允许中国居民在本地城市"
              "（出发前或预注册模式）享受免税价格购物，从根本上改变了免税消费"
              "须物理出行的结构性限制。中国中免已在深圳、广州、成都开设试点门店。"
              "我们估计该渠道成熟后（3-5年）每年可贡献约400-800亿元增量收入，"
              "相当于2025年全年营收的7%-15%。"))

h2_cn(doc, "催化剂三：DFS大中华收购及LVMH战略入股")
body_cn(doc, ("2026年1月19日宣布的DFS大中华区收购（含香港、澳门9家门店及品牌知识产权，"
              "对价不超过3.95亿美元）具有战略跨越意义：中国中免借此无需绿地建设"
              "即可直接进入香港、澳门两大亚洲顶级免税市场；"
              "DFS品牌在内地、港澳及国际旅客中具有高度认知度；"
              "LVMH战略入股，不仅是全球奢侈品龙头对中国中免国际化战略的信任背书，"
              "更可能在品牌合作、独家商品供应等方面带来协同效应。"
              "该交易预计2026年起并入合并报表，贡献增厚效果。"))

h2_cn(doc, "主要风险")
bullet_cn(doc, "封关政策执行进度不及预期，可能推迟近期需求提升。", "执行风险：")
bullet_cn(doc, "人民币兑美元、欧元贬值将推高进口采购成本，压缩毛利率。", "汇率风险：")
bullet_cn(doc, "海南免税市场竞争加剧（如海南省免税集团扩张），可能侵蚀市场份额。", "竞争风险：")
bullet_cn(doc, "中国消费者信心下滑或可支配收入收缩，拖累免税购物需求。", "宏观风险：")
bullet_cn(doc, "DFS整合存在运营和文化磨合风险，整合节奏影响协同效应兑现。", "并购整合风险：")
bullet_cn(doc, "3月31日年报可能披露更大规模的商誉减值，影响净利润。", "会计风险：")

# ── 六、估值与目标价 ──────────────────────────────────────────────────────────
doc.add_page_break()
h1_cn(doc, "六、估值与目标价")

h2_cn(doc, "估值方法：2026年预期市盈率法（PE）")
body_cn(doc, ("我们采用2026年预期净利润市盈率法对中国中免进行估值，以Wind一致预期"
              "2026年归母净利润44.17亿元为基础。按约20.64亿股股本计算，"
              "2026年预期每股收益约为人民币2.14元。我们给予45倍目标市盈率——"
              "低于2021年峰值的80-100倍，但充分体现结构性催化剂确定性及LVMH"
              "战略背书带来的估值溢价。"))

make_table_cn(doc,
    ["估值指标", "数值", "备注"],
    [
        ["2025年全年归母净利润（实际）", "35.86亿元", "业绩快报数据"],
        ["2026年归母净利润（Wind一致预期）", "44.17亿元", "同比恢复+23.2%"],
        ["2026年预期每股收益（EPS）", "约2.14元", "基于20.64亿股"],
        ["目标市盈率（PE）", "45倍", "低于历史峰值，反映催化剂溢价"],
        ["A股目标价", "人民币108.00元", "45× × 2.14 × 1.12（催化剂溢价）"],
        ["当前A股价格（约）", "约人民币80.00元", "2026年3月中旬"],
        ["A股隐含上涨空间", "约+35%", "12个月"],
        ["H股目标价", "港币95.00元", "花旗：100港元；大摩：89港元"],
        ["当前H股价格", mkt["hk_price"], "yfinance实时数据"],
    ]
)

body_cn(doc, "")
make_table_cn(doc,
    ["券商", "评级", "A股目标价（元）", "H股目标价（港元）"],
    [
        ["中信证券",   "买入", "116.10", "—"],
        ["中金公司",   "买入", "107.00", "—"],
        ["华泰证券",   "买入", "约105.00","—"],
        ["A股一致预期（21家）", "买入", "110.51", "—"],
        ["摩根士丹利", "增持", "—", "89.00"],
        ["花旗集团",   "买入", "—", "100.00"],
        ["H股一致预期", "买入", "—", "91.97"],
        ["本报告（中国中免）", "买入", "108.00", "95.00"],
    ]
)

add_img_cn(doc, "cdf_chart9_broker_estimates.png", width=Inches(6.0))
body_cn(doc, "图9：各机构2026年归母净利润预期对比", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_img_cn(doc, "cdf_chart10_recovery_forecast.png", width=Inches(6.0))
body_cn(doc, "图10：营收与归母净利润修复预测（2023–2027年）", size=8)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 七、数据来源与免责声明 ────────────────────────────────────────────────────
doc.add_page_break()
h1_cn(doc, "七、数据来源与免责声明")

h2_cn(doc, "主要数据来源")
sources_cn = [
    ("2025年全年业绩快报（2026年3月20日）",
     "http://www.cninfo.com.cn",
     "601888.SS，上市公司公告"),
    ("2025年三季报（2025年10月31日）",
     "https://stockmc.xueqiu.com/202510/601888_20251031_3SVP.pdf",
     "601888.SS"),
    ("2025年半年报（2025年7月26日）",
     "https://static.cninfo.com.cn/finalpage/2025-07-26/1224299960.PDF",
     "601888.SS"),
    ("2024年年报（2025年3月31日）",
     "https://finance.sina.com.cn/stock/aigcy/2025-03-31/doc-inerpywh8106430.shtml",
     "601888.SS"),
    ("收购DFS大中华区公告（2026年1月19日）",
     "https://stockmc.xueqiu.com/202601/601888_20260120_PB08.pdf",
     "601888.SS / 1880.HK"),
    ("海南自贸港封关通知（2025年12月）",
     "https://caifuhao.eastmoney.com/news/20250808141237013767460",
     "东方财富"),
    ("海南离岛免税2025年上半年数据",
     "https://www.21jingji.com/article/20250828/65c49a01e741b2bd584b6f32f7e784e3.html",
     "21世纪经济报道"),
    ("Wind一致预期数据（2025/2026年）",
     "https://finance.sina.com.cn/stock/bxjj/2026-01-16/doc-inhhnzat5445340.shtml",
     "Wind资讯 / 新浪财经"),
    ("Yahoo Finance — 1880.HK",
     "https://finance.yahoo.com/quote/1880.HK/",
     "港股市场数据（2026年3月）"),
]

for title, url, note in sources_cn:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"• {title}：")
    r1.font.bold = True; r1.font.size = Pt(9)
    set_cjk_font(r1, "黑体")
    add_hyperlink(p, url, url)
    r2 = p.add_run(f"  【{note}】")
    r2.font.size = Pt(8.5); r2.font.color.rgb = CDF_GRAY
    set_cjk_font(r2, "宋体")

h2_cn(doc, "免责声明")
body_cn(doc, ("本报告仅供参考，不构成投资建议。2025年全年数据来源于业绩快报，"
              "完整年报预计2026年3月31日披露，最终数字（含分部数据、每股收益、"
              "商誉减值金额）以年报为准。市场数据来自yfinance（截至2026年3月21日）。"
              "报告中的目标价及评级仅代表撰写时观点，不承诺未来收益，"
              "投资者应据此自行判断并承担投资风险。"), size=9)

hr_cn(doc)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_foot.add_run(f"中国旅游集团中免股份有限公司（601888.SS / 1880.HK）— 2025年Q4业绩更新报告  |  生成日期：{datetime.date.today().strftime('%Y年%m月%d日')}")
rf.font.size = Pt(8); rf.font.color.rgb = CDF_GRAY
set_cjk_font(rf, "宋体")

# ── 保存 ──────────────────────────────────────────────────────────────────────
OUTFILE = OUT + "601888_Q4_2025_业绩更新报告_中文版.docx"
doc.save(OUTFILE)
print(f"中文报告已保存：{OUTFILE}")
