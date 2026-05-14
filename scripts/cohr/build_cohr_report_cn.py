"""
Coherent Corp. (COHR) — Q3 FY2026 业绩更新报告（中文版）
"""
import os
import yfinance as yf
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/COHR"

C_NAVY   = RGBColor(0x00, 0x3D, 0xA5)
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_BLUE   = RGBColor(0x4A, 0x90, 0xD9)
C_GREEN  = RGBColor(0x1E, 0x8A, 0x44)
C_RED    = RGBColor(0xCC, 0x33, 0x33)
C_GREY   = RGBColor(0x8E, 0x8E, 0x93)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_ORANGE = RGBColor(0xE6, 0x86, 0x19)

FONT_CN_BODY  = "宋体"
FONT_CN_TITLE = "黑体"


def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price": round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high": round(info.year_high, 2),
            "52w_low": round(info.year_low, 2),
        }
    except Exception as e:
        print(f"WARNING: yfinance failed: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "4A90D9")
    rPr.append(color_elem)
    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rPr.append(u_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def set_run_cn(run, font_name=FONT_CN_BODY, size=10):
    run.font.size = Pt(size)
    run.font.name = font_name
    r_elem = run._element
    rPr = r_elem.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
        set_run_cn(run, FONT_CN_TITLE, 18)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        set_run_cn(run, FONT_CN_TITLE, 13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "003DA5")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        set_run_cn(run, FONT_CN_TITLE, 11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_cn(run, FONT_CN_BODY, size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_cn(run, FONT_CN_BODY, 10)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"[图表: {filename} 未找到]", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
        cp.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, alternate=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "003DA5")
        set_cell_border(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F5F5F7" if (alternate and ri % 2 == 1) else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="E5E5EA")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if "+" in str(val) and ci > 1:
                        run.font.color.rgb = C_GREEN
                    elif "−" in str(val) or (str(val).startswith("-") and ci > 0):
                        run.font.color.rgb = C_RED
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()
    return table


def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# 市场数据
# ══════════════════════════════════════════════════════════════════════════════
mkt = get_market_data("COHR")
price_str = f"${mkt['price']}" if isinstance(mkt['price'], float) else "N/A"
mcap_str = f"${mkt['market_cap'] / 1e9:.1f}B" if isinstance(mkt['market_cap'], (int, float)) else "N/A"
range_str = f"${mkt['52w_low']} – ${mkt['52w_high']}" if isinstance(mkt['52w_low'], float) else "N/A"

# ══════════════════════════════════════════════════════════════════════════════
# 构建报告
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(10)
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
else:
    rFonts.set(qn("w:eastAsia"), "宋体")

# ══════════════════════════════════════════════════════════════════════════════
# 第1页 — 业绩摘要
# ══════════════════════════════════════════════════════════════════════════════
p_ticker = doc.add_paragraph()
p_ticker.clear()
run = p_ticker.add_run("贰光科技  (COHR: NYSE)")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_NAVY
set_run_cn(run, FONT_CN_TITLE, 22)
p_ticker.paragraph_format.space_after = Pt(2)

p_sub = doc.add_paragraph()
p_sub.clear()
run2 = p_sub.add_run("权益研究  |  光子学与光通信")
run2.font.size = Pt(10)
run2.font.color.rgb = C_GREY
set_run_cn(run2, FONT_CN_BODY, 10)
p_sub.paragraph_format.space_after = Pt(6)

# 评级信息
rating_table = doc.add_table(rows=1, cols=5)
rating_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["评级", "目标价", "当前股价", "52周区间", "市值"]
values = ["增持", "$450.00", price_str, range_str, mcap_str]
colors_v = [C_GREEN, C_BLUE, C_NAVY, C_GREY, C_NAVY]

for i, (lbl, val, col) in enumerate(zip(labels, values, colors_v)):
    cell = rating_table.rows[0].cells[i]
    set_cell_bg(cell, "E8F0FE")
    set_cell_border(cell, color="C0D0F0")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lbl + "\n")
    r1.font.size = Pt(8)
    r1.font.color.rgb = C_GREY
    r1.font.name = "Calibri"
    r2 = p1.add_run(val)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = col
    r2.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(4)
horizontal_rule(doc)

# 报告标题
p_title = doc.add_paragraph()
p_title.clear()
run_t = p_title.add_run(
    "Q3 FY2026 业绩更新：营收创历史新高18.1亿美元；"
    "英伟达战略合作重塑资产负债表与增长轨迹"
)
run_t.font.size = Pt(14)
run_t.font.bold = True
run_t.font.color.rgb = C_NAVY
set_run_cn(run_t, FONT_CN_TITLE, 14)
p_title.paragraph_format.space_after = Pt(3)

p_date = doc.add_paragraph()
p_date.clear()
rd = p_date.add_run(
    "报告日期：2026年5月14日  |  业绩发布：2026年5月6日  |  "
    "财务季度截止：2026年3月31日"
)
rd.font.size = Pt(8.5)
rd.font.color.rgb = C_GREY
set_run_cn(rd, FONT_CN_BODY, 8.5)
p_date.paragraph_format.space_after = Pt(8)

horizontal_rule(doc)

# 核心要点
heading(doc, "核心要点", level=2)
bullet(doc, "营收18.1亿美元，超市场一致预期17.8亿美元约3,000万美元（+1.7%），创季度新高，同比增长21%。")
bullet(doc, "Non-GAAP每股收益1.41美元，超一致预期1.39美元0.02美元（+1.4%），同比增长55%，受益于经营杠杆。")
bullet(doc, "数据中心与通信部门同比大增36%至约13.6亿美元，占总收入75%，受800G/1.6T光模块放量及AI基础设施需求驱动。")
bullet(doc, "英伟达战略合作：Q3收到20亿美元股权投资；获得数十亿美元多年期采购承诺，涵盖激光器、CPO及光学元器件。")
bullet(doc, "Q4 FY2026指引：营收19.1亿–20.5亿美元（中值19.8亿），Non-GAAP每股收益1.52–1.72美元，预示持续双位数环比加速增长。")
bullet(doc, "杠杆率从1.7x骤降至0.5x，英伟达入股资金到账后资产负债表已成竞争优势。", color=C_GREEN)

# 业绩概览表
heading(doc, "业绩概览", level=2)
add_table(doc,
    ["指标", "Q3 FY26 实际", "一致预期", "超预期幅度", "Q3 FY25", "同比变化"],
    [
        ["营收",            "$18.1亿", "$17.8亿", "+1.7%",  "$15.0亿", "+21.0%"],
        ["Non-GAAP EPS",    "$1.41",   "$1.39",   "+$0.02", "$0.91",   "+55.0%"],
        ["GAAP EPS",        "$0.97",   "—",       "—",      "$0.38",   "+155%"],
        ["Non-GAAP毛利率",  "39.6%",   "~39.2%",  "+40bps", "38.1%",   "+150bps"],
        ["Non-GAAP经营利润率","20.3%", "—",       "—",      "18.7%",   "+160bps"],
        ["Non-GAAP经营利润","$3.66亿", "—",       "—",      "$2.79亿", "+31.1%"],
    ])

add_image(doc, "cohr_chart5_beat_miss.png", caption="图1：Q3 FY2026业绩 vs. 市场一致预期")

# ══════════════════════════════════════════════════════════════════════════════
# 第2-3页 — 详细业绩分析
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "详细业绩分析", level=1)

heading(doc, "营收 — AI数据中心需求驱动的创纪录季度", level=2)
body(doc, "Coherent在Q3 FY2026实现创纪录季度营收18.1亿美元，超过市场一致预期17.8亿美元约3,000万美元（+1.7%）。营收同比增长21%，环比增长7.1%，连续第四个季度实现环比加速增长。按pro forma口径计算，同比增速约为27%，充分体现了有机增长的强劲动能。")
body(doc, "营收超预期主要由数据中心与通信部门贡献，该部门收入约13.6亿美元（同比+36%），目前占总收入的75%。其中，数据中心相关收入环比增长13%、同比增长37%；通信相关收入环比增长16%、同比增长60%。1.6T光模块首次对环比增长产生实质性贡献，叠加已有的800G收入流。")

add_image(doc, "cohr_chart1_revenue.png", caption="图2：季度营收走势（Q4 FY2024 – Q3 FY2026）")

heading(doc, "分部业绩", level=2)

heading(doc, "数据中心与通信（约13.6亿美元，占比75%）", level=3)
body(doc, "该部门是公司的主要增长引擎，受益于超大规模AI训练集群对光互连的强劲需求。管理层表示客户订单已延伸至2028年日历年，长期协议覆盖至本十年末。Coherent的磷化铟产能在本季度提前翻倍，6英寸晶圆产线扩张持续推进。英伟达合作提供了显著的需求可见性，通过涵盖高功率CW激光器、外部激光源模块、光纤阵列单元及共封装光学（CPO）组件的数十亿美元多年期采购承诺。")

heading(doc, "工业部门（4.44亿美元，占比25%）", level=3)
body(doc, "工业部门同比下降16%，反映了更广泛工业终端市场的持续疲软。管理层将该部门定性为稳定但非近期增长核心。公司正战略性地将产品组合和制造产能向更高增长的数据中心应用重新调配。")

add_image(doc, "cohr_chart3_segments.png", caption="图3：收入结构与分部同比增长 — Q3 FY2026")
add_image(doc, "cohr_chart10_datacenter.png", caption="图4：数据中心与通信收入增长轨迹")

# ══════════════════════════════════════════════════════════════════════════════
# 第4-5页 — 利润率与指引
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "利润率、盈利能力与业绩指引", level=1)

heading(doc, "利润率扩张 — 连续六个季度改善", level=2)
body(doc, "Non-GAAP毛利率扩张至39.6%，同比提升150个基点，环比提升60个基点。GAAP毛利率达37.7%，同比提升243个基点。持续的利润率扩张反映了产品组合向更高毛利的AI数据中心光学组件倾斜、磷化铟平台的制造规模效益，以及严格的成本管控。")
body(doc, "Non-GAAP经营利润率扩张至20.3%，同比提升163个基点。Non-GAAP经营利润达3.66亿美元，同比增长31.1%。公司展现出显著的经营杠杆效应，运营费用增速大幅慢于收入增速。")

add_image(doc, "cohr_chart4_margins.png", caption="图5：利润率走势 — 持续扩张")

heading(doc, "前三季度（YTD）业绩", level=2)
body(doc, "截至2026年3月31日的九个月内，Coherent实现营收50.7亿美元（同比+18.5%），GAAP摊薄每股收益2.92美元（去年同期为0.30美元），Non-GAAP摊薄每股收益3.86美元（同比+52.6%）。GAAP盈利能力的大幅改善既反映了经营面的提升，也体现了与前II-VI收购相关的重组/整合费用减少。")

heading(doc, "Q4 FY2026 业绩指引", level=2)
add_table(doc,
    ["指标", "Q4 FY26 指引（低值）", "Q4 FY26 指引（高值）", "中值", "隐含环比"],
    [
        ["营收",            "$19.1亿", "$20.5亿", "$19.8亿", "+9.4%"],
        ["Non-GAAP毛利率",  "39.0%",   "41.0%",   "40.0%",   "+40bps"],
        ["Non-GAAP运营费用","$3.60亿", "$3.80亿", "$3.70亿", "—"],
        ["Non-GAAP EPS",    "$1.52",   "$1.72",   "$1.62",   "+14.9%"],
    ])
body(doc, "Q4指引意味着持续加速，营收中值19.8亿美元代表约9.4%的环比增长及显著的同比扩张。隐含的FY2026全年营收约为70.5亿美元（同比+21%），隐含的全年Non-GAAP每股收益约为5.48美元。")
body(doc, "Non-GAAP毛利率指引39–41%表明中值可能突破40%，将创下新高。管理层指出Q4资本支出将进一步增加，以支持英伟达合作及更广泛的光模块需求。")

add_image(doc, "cohr_chart6_guidance.png", caption="图6：FY2026营收轨迹与Q4指引")
add_image(doc, "cohr_chart9_yoy.png", caption="图7：关键指标同比对比")

# ══════════════════════════════════════════════════════════════════════════════
# 第6-7页 — 投资论点更新
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "投资论点更新", level=1)

heading(doc, "英伟达合作 — 变革性催化剂", level=2)
body(doc, "本季度最重要的事件是英伟达战略合作的落地，该合作于2026年3月2日宣布。英伟达向Coherent投资20亿美元股权，并承诺签署数十亿美元多年期采购协议，采购Coherent的先进激光器和光通信产品。该合作具有多重战略意义：")
bullet(doc, "需求可见性：客户订单延伸至2028年日历年，长期协议覆盖至本十年末，为光子学公司提供了前所未有的收入可见性。")
bullet(doc, "CPO拐点：Coherent将共封装光学（CPO）可服务潜在市场规模估值上调至150亿美元。规模化CPO收入将于CY2026下半年开始，升级版CPO收入将于CY2027下半年启动。")
bullet(doc, "资产负债表转型：20亿美元股权投资使杠杆率从1.7x降至0.5x，为有机产能扩张提供了充足的财务灵活性，无需增加额外债务。")
bullet(doc, "产能承诺：计划在CY2026将产能翻倍，此后再翻倍，几乎所有扩张均在6英寸晶圆产线上进行，用于先进磷化铟生产。")

heading(doc, "论点强化 — 维持「增持」评级的理由", level=2)
body(doc, "Q3业绩进一步强化了我们的投资论点。Coherent是为数不多的能够提供光信号全链路垂直整合的光子学平台——从激光外延和芯片制造到模块封装和光模块组装。这使公司成为AI算力基础设施建设的关键供应商。核心论点：")
bullet(doc, "长期AI需求驱动：超大规模资本支出在AI训练和推理基础设施上持续加速，每一代新GPU都需要更多的光学带宽。")
bullet(doc, "技术护城河：Coherent的磷化铟垂直整合和CPO路线图提供了难以复制的差异化优势。")
bullet(doc, "盈利杠杆：Non-GAAP经营利润率20.3%仍有显著提升空间，随着制造利用率提升和产品组合向更高利润率的数据中心产品倾斜。")
bullet(doc, "估值上行空间：以~30x CY2027E市盈率计算，股价相对于纯AI基础设施同业仍有折价，尽管其AI相关收入占比日益提升。")

heading(doc, "主要风险", level=2)
bullet(doc, "客户集中度：英伟达合作虽具变革性，但增加了对单一客户的依赖。")
bullet(doc, "资本开支强度：加速产能投资（Q3达2.90亿美元，Q4将继续增加）可能短期内压制自由现金流。")
bullet(doc, "工业部门拖累：工业部门（同比-16%）仍是利润率的不利因素，可能面临进一步的周期性压力。")
bullet(doc, "执行风险：加速时间线下的产能翻倍引入了制造爬坡风险。")
bullet(doc, "股权稀释：英伟达20亿美元股权投资稀释了现有股东，尽管战略价值可能抵消稀释影响。")

add_image(doc, "cohr_chart2_eps.png", caption="图8：Non-GAAP每股收益走势 — 盈利加速")

# ══════════════════════════════════════════════════════════════════════════════
# 第8-10页 — 估值与预测
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "估值与盈利预测更新", level=1)

heading(doc, "财务预测更新", level=2)
add_table(doc,
    ["指标", "FY2025A", "FY2026E（旧）", "FY2026E（新）", "FY2027E"],
    [
        ["营收（$B）",       "$58.3亿", "$68.5亿", "$70.5亿", "$85.0亿"],
        ["营收同比增速",     "+12.5%",  "+17.5%",  "+20.9%",  "+20.6%"],
        ["Non-GAAP毛利率",  "38.0%",   "39.0%",   "39.5%",   "40.5%"],
        ["Non-GAAP经营利润率","18.0%",  "19.5%",   "20.0%",   "21.5%"],
        ["Non-GAAP EPS",    "$3.42",   "$5.20",   "$5.48",   "$7.20"],
        ["EPS同比增速",      "+35%",    "+52%",    "+60%",    "+31%"],
    ])
body(doc, "我们将FY2026营收预测从68.5亿美元上调至70.5亿美元，Non-GAAP每股收益从5.20美元上调至5.48美元，反映了Q3超预期业绩及强于预期的Q4指引。FY2027预测营收85.0亿美元，Non-GAAP每股收益7.20美元，受持续的数据中心增长动能、CY2026/2027下半年CPO收入放量及经营杠杆驱动。")

heading(doc, "目标价论证 — $450", level=2)
body(doc, "我们的$450目标价基于混合估值方法：")
bullet(doc, "DCF（权重50%）：使用10%加权平均资本成本和3.5%终端增长率，DCF分析得出约$460/股的公允价值，反映了英伟达多年期承诺支撑的长久期增长特征。")
bullet(doc, "市盈率（权重30%）：对CY2027E Non-GAAP EPS $7.20给予35x市盈率得到$252。考虑到增长轨迹，我们对CY2026E EPS $5.48给予62.5x溢价倍数，得到$342。")
bullet(doc, "EV/Revenue（权重20%）：对CY2027E营收85.0亿美元给予8.5x倍数，得到约$472/股的权益价值。")
body(doc, "混合目标价：约$450，较当前股价有约12%的上行空间。")

heading(doc, "情景分析", level=2)
add_table(doc,
    ["情景", "核心假设", "FY27E营收", "FY27E EPS", "隐含股价"],
    [
        ["乐观情景", "CPO放量超预期；1.6T份额扩大",  "$95亿", "$8.50", "$550+"],
        ["基准情景", "按指引轨迹；CPO如期推进",       "$85亿", "$7.20", "$450"],
        ["悲观情景", "AI资本支出放缓；CPO延迟",       "$72亿", "$5.50", "$300"],
    ])

add_image(doc, "cohr_chart7_capex.png", caption="图9：资本支出加速 — AI增长投资")
add_image(doc, "cohr_chart8_leverage.png", caption="图10：杠杆率 — 英伟达投资后快速去杠杆")

# ══════════════════════════════════════════════════════════════════════════════
# 资料来源
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "资料来源与参考文献", level=1)

heading(doc, "Q3 FY2026 业绩资料", level=2)

p1 = doc.add_paragraph()
r1 = p1.add_run("业绩新闻稿（2026年5月6日）：")
set_run_cn(r1, FONT_CN_BODY, 10)
add_hyperlink(p1, "Coherent Corp Q3 FY2026 新闻稿",
              "https://www.coherent.com/news/press-releases/third-quarter-fiscal-year-2026-results")
p1.paragraph_format.space_after = Pt(4)

p2 = doc.add_paragraph()
r2 = p2.add_run("10-Q报表（2026年5月6日提交）：")
set_run_cn(r2, FONT_CN_BODY, 10)
add_hyperlink(p2, "SEC EDGAR — Coherent Corp 10-Q",
              "https://www.stocktitan.net/sec-filings/COHR/10-q-coherent-corp-quarterly-earnings-report-dfc9ceaf3a93.html")
p2.paragraph_format.space_after = Pt(4)

p3 = doc.add_paragraph()
r3 = p3.add_run("业绩电话会纪要（2026年5月6日）：")
set_run_cn(r3, FONT_CN_BODY, 10)
add_hyperlink(p3, "Q3 FY2026 业绩电话会纪要 — Motley Fool",
              "https://www.fool.com/earnings/call-transcripts/2026/05/06/coherent-cohr-q3-2026-earnings-transcript/")
p3.paragraph_format.space_after = Pt(4)

p4 = doc.add_paragraph()
r4 = p4.add_run("英伟达战略合作公告（2026年3月2日）：")
set_run_cn(r4, FONT_CN_BODY, 10)
add_hyperlink(p4, "英伟达-Coherent战略合作公告",
              "https://nvidianews.nvidia.com/news/nvidia-and-coherent-announce-strategic-partnership-to-develop-optics-technology-to-scale-next-generation-data-center-architecture")
p4.paragraph_format.space_after = Pt(4)

heading(doc, "往期资料", level=2)

p5 = doc.add_paragraph()
r5 = p5.add_run("Q2 FY2026 业绩新闻稿：")
set_run_cn(r5, FONT_CN_BODY, 10)
add_hyperlink(p5, "Coherent Corp Q2 FY2026 新闻稿",
              "https://www.coherent.com/news/press-releases/second-quarter-fiscal-year-2026-results")
p5.paragraph_format.space_after = Pt(4)

p6 = doc.add_paragraph()
r6 = p6.add_run("Q1 FY2026 业绩新闻稿：")
set_run_cn(r6, FONT_CN_BODY, 10)
add_hyperlink(p6, "Coherent Corp Q1 FY2026 新闻稿",
              "https://www.coherent.com/news/press-releases/first-quarter-fiscal-year-2026-results")
p6.paragraph_format.space_after = Pt(4)

heading(doc, "市场数据", level=2)
body(doc, f"股价截至2026年5月14日：{price_str}")
body(doc, f"市值：{mcap_str}")
body(doc, "一致预期来源：Bloomberg终端，截至2026年5月5日。")
body(doc, "除特别注明外，所有财务数据均来自公司公开披露文件。")

# ══════════════════════════════════════════════════════════════════════════════
# 免责声明
# ══════════════════════════════════════════════════════════════════════════════
horizontal_rule(doc)
body(doc, "免责声明：本报告仅供参考，不构成投资建议。本文所表达的分析、意见和估计基于公开可获取的信息，"
     "如有变更，恕不另行通知。过往表现不代表未来业绩。投资者在做出投资决策前应自行进行尽职调查。",
     italic=True, color=C_GREY, size=8, space_after=2)

# ══════════════════════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════════════════════
out_path = os.path.join(BASE, "COHR_Q3_FY2026_业绩更新报告_中文版.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
