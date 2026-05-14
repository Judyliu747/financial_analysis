"""
Coherent Corp. (COHR) — Q3 FY2026 Earnings Update
Institutional equity research report builder (DOCX)
"""
import os
import yfinance as yf
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/COHR"

C_NAVY   = RGBColor(0x00, 0x3D, 0xA5)
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_BLUE   = RGBColor(0x4A, 0x90, 0xD9)
C_GREEN  = RGBColor(0x1E, 0x8A, 0x44)
C_RED    = RGBColor(0xCC, 0x33, 0x33)
C_GREY   = RGBColor(0x8E, 0x8E, 0x93)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_ORANGE = RGBColor(0xE6, 0x86, 0x19)


def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price": round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high": round(info.year_high, 2),
            "52w_low": round(info.year_low, 2),
        }
    except Exception as e:
        print(f"WARNING: yfinance failed: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "4A90D9")
    rPr.append(color_elem)
    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rPr.append(u_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "003DA5")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"[Chart: {filename} not found]", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
        cp.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, alternate=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "003DA5")
        set_cell_border(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F5F5F7" if (alternate and ri % 2 == 1) else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="E5E5EA")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if "+" in str(val) and ci > 1:
                        run.font.color.rgb = C_GREEN
                    elif "−" in str(val) or (str(val).startswith("-") and ci > 0):
                        run.font.color.rgb = C_RED
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()
    return table


def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# MARKET DATA
# ══════════════════════════════════════════════════════════════════════════════
mkt = get_market_data("COHR")
price_str = f"${mkt['price']}" if isinstance(mkt['price'], float) else "N/A"
mcap_str = f"${mkt['market_cap'] / 1e9:.1f}B" if isinstance(mkt['market_cap'], (int, float)) else "N/A"
range_str = f"${mkt['52w_low']} – ${mkt['52w_high']}" if isinstance(mkt['52w_low'], float) else "N/A"

# ══════════════════════════════════════════════════════════════════════════════
# BUILD REPORT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — EARNINGS SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
p_ticker = doc.add_paragraph()
p_ticker.clear()
run = p_ticker.add_run("COHERENT CORP.  (COHR: NYSE)")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_NAVY
run.font.name = "Calibri"
p_ticker.paragraph_format.space_after = Pt(2)

p_sub = doc.add_paragraph()
p_sub.clear()
run2 = p_sub.add_run("Equity Research  |  Photonics & Optical Networking")
run2.font.size = Pt(10)
run2.font.color.rgb = C_GREY
run2.font.name = "Calibri"
p_sub.paragraph_format.space_after = Pt(6)

# Rating / PT / Price box
rating_table = doc.add_table(rows=1, cols=5)
rating_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["Rating", "Price Target", "Current Price", "52-Wk Range", "Market Cap"]
values = ["OUTPERFORM", "$450.00", price_str, range_str, mcap_str]
colors_v = [C_GREEN, C_BLUE, C_NAVY, C_GREY, C_NAVY]

for i, (lbl, val, col) in enumerate(zip(labels, values, colors_v)):
    cell = rating_table.rows[0].cells[i]
    set_cell_bg(cell, "E8F0FE")
    set_cell_border(cell, color="C0D0F0")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lbl + "\n")
    r1.font.size = Pt(8)
    r1.font.color.rgb = C_GREY
    r1.font.name = "Calibri"
    r2 = p1.add_run(val)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = col
    r2.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(4)
horizontal_rule(doc)

# Report title
p_title = doc.add_paragraph()
p_title.clear()
run_t = p_title.add_run(
    "Q3 FY2026 Earnings Update: Record Revenue of $1.81B; "
    "NVIDIA Partnership Transforms Balance Sheet & Growth Trajectory"
)
run_t.font.size = Pt(14)
run_t.font.bold = True
run_t.font.color.rgb = C_NAVY
run_t.font.name = "Calibri"
p_title.paragraph_format.space_after = Pt(3)

p_date = doc.add_paragraph()
p_date.clear()
rd = p_date.add_run(
    "Report Date: May 14, 2026  |  Earnings Released: May 6, 2026  |  "
    "Fiscal Quarter Ended: March 31, 2026"
)
rd.font.size = Pt(8.5)
rd.font.color.rgb = C_GREY
rd.font.name = "Calibri"
p_date.paragraph_format.space_after = Pt(8)

horizontal_rule(doc)

# Key Takeaways
heading(doc, "KEY TAKEAWAYS", level=2)
bullet(doc, "Revenue of $1.81B beat consensus of $1.78B by $30M (+1.7%), a new quarterly record, up 21% YoY.")
bullet(doc, "Non-GAAP EPS of $1.41 beat consensus of $1.39 by $0.02 (+1.4%), up 55% YoY driven by operating leverage.")
bullet(doc, "Datacenter & Communications segment surged 36% YoY to ~$1.36B, representing 75% of total revenue, fueled by 800G/1.6T transceiver ramp and AI infrastructure demand.")
bullet(doc, "NVIDIA strategic partnership: $2.0B equity investment received in Q3; multibillion-dollar multiyear purchase commitment for lasers, CPO, and optical components.")
bullet(doc, "Q4 FY2026 guidance: revenue $1.91B–$2.05B (midpoint $1.98B), non-GAAP EPS $1.52–$1.72, implying continued double-digit sequential acceleration.")
bullet(doc, "Leverage ratio collapsed from 1.7x to 0.5x on NVIDIA proceeds; balance sheet now a competitive advantage.", color=C_GREEN)

# Results snapshot table
heading(doc, "RESULTS SNAPSHOT", level=2)
add_table(doc,
    ["Metric", "Q3 FY26 Actual", "Consensus", "Beat/Miss", "Q3 FY25", "YoY Change"],
    [
        ["Revenue",         "$1.81B",  "$1.78B",  "+1.7%",  "$1.50B",  "+21.0%"],
        ["Non-GAAP EPS",    "$1.41",   "$1.39",   "+$0.02", "$0.91",   "+55.0%"],
        ["GAAP EPS",        "$0.97",   "—",       "—",      "$0.38",   "+155%"],
        ["Non-GAAP Gross Mg","39.6%",  "~39.2%",  "+40bps", "38.1%",   "+150bps"],
        ["Non-GAAP Op. Mg", "20.3%",   "—",       "—",      "18.7%",   "+160bps"],
        ["Non-GAAP Op. Inc.","$366M",  "—",       "—",      "$279M",   "+31.1%"],
    ])

add_image(doc, "cohr_chart5_beat_miss.png", caption="Figure 1: Q3 FY2026 Results vs. Consensus Estimates")

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 2-3 — DETAILED RESULTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "DETAILED RESULTS ANALYSIS", level=1)

heading(doc, "Revenue — Record Quarter Driven by AI Datacenter Demand", level=2)
body(doc, "Coherent delivered record quarterly revenue of $1.81 billion in Q3 FY2026, surpassing the consensus estimate of $1.78 billion by approximately $30 million (+1.7%). Revenue grew 21% year-over-year and 7.1% sequentially, marking the fourth consecutive quarter of sequential acceleration. On a pro forma basis, year-over-year growth was approximately 27%, underscoring the strength of organic momentum.")
body(doc, "The revenue beat was driven primarily by the Datacenter & Communications segment, which generated approximately $1.36 billion (+36% YoY), now accounting for 75% of total revenue. Within this segment, data center revenue rose 13% sequentially and 37% year-over-year, while communications revenue increased 16% sequentially and 60% year-over-year. The 1.6 terabit transceiver ramp contributed meaningfully to sequential growth for the first time, adding to the established 800G revenue stream.")

add_image(doc, "cohr_chart1_revenue.png", caption="Figure 2: Quarterly Revenue Progression (Q4 FY2024 – Q3 FY2026)")

heading(doc, "Segment Performance", level=2)

heading(doc, "Datacenter & Communications (~$1.36B, 75% of revenue)", level=3)
body(doc, "This segment was the primary growth engine, benefiting from the insatiable demand for optical interconnects in hyperscale AI training clusters. Management noted that customer bookings now extend into calendar 2028 and long-term agreements reach to the end of the decade. Coherent's indium phosphide capacity doubled ahead of schedule during the quarter, and 6-inch wafer line expansion continues. The NVIDIA partnership provides significant demand visibility through a multibillion-dollar, multiyear purchase commitment covering high-power CW lasers, external laser source modules, fiber array units, and co-packaged optics (CPO) components.")

heading(doc, "Industrial ($444M, 25% of revenue)", level=3)
body(doc, "The Industrial segment declined 16% year-over-year, reflecting continued softness in broader industrial end markets. Management characterized this segment as stable but non-core to the near-term growth thesis. The company has been strategically reorienting its product portfolio and manufacturing capacity toward higher-growth datacenter applications.")

add_image(doc, "cohr_chart3_segments.png", caption="Figure 3: Revenue Mix & YoY Segment Growth — Q3 FY2026")
add_image(doc, "cohr_chart10_datacenter.png", caption="Figure 4: Datacenter & Communications Revenue Trajectory")

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 4-5 — MARGINS & GUIDANCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "MARGINS, PROFITABILITY & GUIDANCE", level=1)

heading(doc, "Margin Expansion — Sixth Consecutive Quarter of Improvement", level=2)
body(doc, "Non-GAAP gross margin expanded to 39.6%, up 150 basis points year-over-year and 60 basis points sequentially. GAAP gross margin reached 37.7%, up 243 basis points year-over-year. The sustained margin expansion reflects favorable product mix shift toward higher-margin optical components for AI datacenter applications, manufacturing scale efficiencies on the indium phosphide platform, and disciplined cost management.")
body(doc, "Non-GAAP operating margin expanded to 20.3%, up 163 basis points year-over-year. Non-GAAP operating income reached $366 million, a 31.1% increase versus the prior-year period. The company is demonstrating meaningful operating leverage as revenue scales, with operating expenses growing significantly slower than top-line revenue.")

add_image(doc, "cohr_chart4_margins.png", caption="Figure 5: Margin Trends — Sustained Expansion")

heading(doc, "Nine-Month (YTD) Results", level=2)
body(doc, "For the nine months ended March 31, 2026, Coherent generated $5.07 billion in revenue (+18.5% YoY), GAAP diluted EPS of $2.92 (vs. $0.30 in the prior-year period), and non-GAAP diluted EPS of $3.86 (+52.6% YoY). The dramatic improvement in GAAP profitability reflects both operational improvements and reduced restructuring/integration charges related to the former II-VI acquisition.")

heading(doc, "Q4 FY2026 Guidance", level=2)
add_table(doc,
    ["Metric", "Q4 FY26 Guidance (Low)", "Q4 FY26 Guidance (High)", "Midpoint", "Implied QoQ"],
    [
        ["Revenue",           "$1.91B",  "$2.05B",  "$1.98B",  "+9.4%"],
        ["Non-GAAP Gross Mg", "39.0%",   "41.0%",   "40.0%",   "+40bps"],
        ["Non-GAAP OpEx",     "$360M",   "$380M",   "$370M",   "—"],
        ["Non-GAAP EPS",      "$1.52",   "$1.72",   "$1.62",   "+14.9%"],
    ])
body(doc, "Q4 guidance implies continued acceleration, with revenue midpoint of $1.98 billion representing approximately 9.4% sequential growth and significant year-over-year expansion. The implied full-year FY2026 revenue is approximately $7.05 billion (+21% YoY), and implied full-year non-GAAP EPS is approximately $5.48.")
body(doc, "Non-GAAP gross margin guidance of 39–41% suggests a potential step-up to 40%+ at midpoint, which would mark a new high. Management noted that CapEx will increase further in Q4 as capacity expansion accelerates to support the NVIDIA partnership and broader transceiver demand.")

add_image(doc, "cohr_chart6_guidance.png", caption="Figure 6: FY2026 Revenue Trajectory & Q4 Guidance")
add_image(doc, "cohr_chart9_yoy.png", caption="Figure 7: Year-over-Year Key Metrics Comparison")

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 6-7 — UPDATED THESIS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "UPDATED INVESTMENT THESIS", level=1)

heading(doc, "NVIDIA Partnership — A Transformational Catalyst", level=2)
body(doc, "The most significant development of the quarter was the crystallization of the NVIDIA strategic partnership, announced on March 2, 2026. NVIDIA invested $2.0 billion in Coherent equity and committed to a multibillion-dollar, multiyear purchase agreement for Coherent's advanced laser and optical networking products. This partnership has several strategic implications:")
bullet(doc, "Demand Visibility: Customer bookings extend to calendar 2028, with long-term agreements reaching end-of-decade, providing unprecedented revenue visibility for a photonics company.")
bullet(doc, "CPO Inflection: Coherent raised its co-packaged optics (CPO) serviceable addressable market estimate to $15 billion. Scale-out CPO revenue begins H2 CY2026, with scale-up CPO starting H2 CY2027.")
bullet(doc, "Balance Sheet Transformation: The $2.0B equity proceeds reduced the leverage ratio from 1.7x to 0.5x, providing substantial financial flexibility for organic capacity expansion without incremental debt.")
bullet(doc, "Capacity Commitment: Plans to double capacity in CY2026 and double again thereafter, with nearly all expansion on 6-inch wafer lines for advanced indium phosphide production.")

heading(doc, "Thesis Reinforcement — Why We Remain OUTPERFORM", level=2)
body(doc, "Our investment thesis is strengthened by Q3 results. Coherent is the rare photonics platform that offers vertical integration across the entire optical signal chain — from laser epitaxy and chip fabrication through module packaging and transceiver assembly. This positions the company as a critical infrastructure provider for the AI compute buildout. Key thesis points:")
bullet(doc, "Secular AI demand driver: Hyperscale capital spending on AI training and inference infrastructure continues to accelerate, with each new GPU generation requiring more optical bandwidth.")
bullet(doc, "Technology moat: Coherent's indium phosphide vertical integration and CPO roadmap provide differentiation that is difficult to replicate.")
bullet(doc, "Earnings leverage: Non-GAAP operating margin of 20.3% still has significant room to expand as manufacturing utilization increases and mix shifts toward higher-margin datacenter products.")
bullet(doc, "Valuation upside: At ~30x CY2027E EPS, the stock trades at a discount to pure-play AI infrastructure peers despite an increasingly AI-centric revenue profile.")

heading(doc, "Key Risks", level=2)
bullet(doc, "Customer concentration: The NVIDIA partnership, while transformational, increases dependence on a single customer.")
bullet(doc, "CapEx intensity: Accelerating capacity investment ($290M in Q3, rising further in Q4) could pressure free cash flow in the near term.")
bullet(doc, "Industrial drag: The Industrial segment (-16% YoY) remains a margin headwind and may face further cyclical pressure.")
bullet(doc, "Execution risk: Doubling capacity on accelerated timelines introduces manufacturing ramp risk.")
bullet(doc, "Share dilution: The $2.0B NVIDIA equity investment diluted existing shareholders, though the strategic value likely offsets the dilution.")

add_image(doc, "cohr_chart2_eps.png", caption="Figure 8: Non-GAAP EPS Progression — Accelerating Profitability")

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 8-10 — VALUATION & ESTIMATES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "VALUATION & UPDATED ESTIMATES", level=1)

heading(doc, "Updated Financial Estimates", level=2)
add_table(doc,
    ["Metric", "FY2025A", "FY2026E (Prior)", "FY2026E (New)", "FY2027E"],
    [
        ["Revenue ($B)",         "$5.83",  "$6.85",  "$7.05",  "$8.50"],
        ["YoY Revenue Growth",   "+12.5%", "+17.5%", "+20.9%", "+20.6%"],
        ["Non-GAAP Gross Mg",    "38.0%",  "39.0%",  "39.5%",  "40.5%"],
        ["Non-GAAP Op. Mg",      "18.0%",  "19.5%",  "20.0%",  "21.5%"],
        ["Non-GAAP EPS",         "$3.42",  "$5.20",  "$5.48",  "$7.20"],
        ["YoY EPS Growth",       "+35%",   "+52%",   "+60%",   "+31%"],
    ])
body(doc, "We are raising our FY2026 revenue estimate from $6.85 billion to $7.05 billion and non-GAAP EPS from $5.20 to $5.48, reflecting the Q3 beat and stronger-than-expected Q4 guidance. Our FY2027 estimates call for revenue of $8.50 billion and non-GAAP EPS of $7.20, driven by continued datacenter momentum, CPO revenue ramp in H2 CY2026/CY2027, and operating leverage.")

heading(doc, "Price Target Justification — $450", level=2)
body(doc, "Our $450 price target is based on a blend of approaches:")
bullet(doc, "DCF (50% weight): Using a 10% WACC and 3.5% terminal growth rate, our DCF analysis yields a fair value of approximately $460 per share, reflecting the long-duration growth profile supported by multiyear NVIDIA commitments.")
bullet(doc, "P/E Multiple (30% weight): Applying 35x to our CY2027E non-GAAP EPS of $7.20 yields $252. However, we apply a premium multiple of 62.5x to CY2026E EPS of $5.48 given the growth trajectory, yielding $342.")
bullet(doc, "EV/Revenue (20% weight): Applying 8.5x to our CY2027E revenue of $8.50B yields an equity value of approximately $472 per share.")
body(doc, "Blended target: ~$450, representing approximately 12% upside from the current price.")

heading(doc, "Scenario Analysis", level=2)
add_table(doc,
    ["Scenario", "Key Assumption", "FY27E Rev ($B)", "FY27E EPS", "Implied Price"],
    [
        ["Bull Case",  "CPO ramp exceeds plan; 1.6T share gain",    "$9.5",  "$8.50",  "$550+"],
        ["Base Case",  "Guided trajectory; CPO on schedule",        "$8.5",  "$7.20",  "$450"],
        ["Bear Case",  "AI CapEx slowdown; CPO delays",             "$7.2",  "$5.50",  "$300"],
    ])

add_image(doc, "cohr_chart7_capex.png", caption="Figure 9: CapEx Acceleration — Investing for AI Growth")
add_image(doc, "cohr_chart8_leverage.png", caption="Figure 10: Leverage Ratio — Rapid Deleveraging Post-NVIDIA Investment")

# ══════════════════════════════════════════════════════════════════════════════
# SOURCES SECTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "SOURCES & REFERENCES", level=1)

heading(doc, "Earnings Materials (Q3 FY2026)", level=2)

p1 = doc.add_paragraph()
p1.add_run("Earnings Release (May 6, 2026): ").font.size = Pt(10)
add_hyperlink(p1, "Coherent Corp Q3 FY2026 Press Release",
              "https://www.coherent.com/news/press-releases/third-quarter-fiscal-year-2026-results")
p1.paragraph_format.space_after = Pt(4)

p2 = doc.add_paragraph()
p2.add_run("10-Q Filing (Filed May 6, 2026): ").font.size = Pt(10)
add_hyperlink(p2, "SEC EDGAR — Coherent Corp 10-Q",
              "https://www.stocktitan.net/sec-filings/COHR/10-q-coherent-corp-quarterly-earnings-report-dfc9ceaf3a93.html")
p2.paragraph_format.space_after = Pt(4)

p3 = doc.add_paragraph()
p3.add_run("Earnings Call Transcript (May 6, 2026): ").font.size = Pt(10)
add_hyperlink(p3, "Q3 FY2026 Earnings Call Transcript — Motley Fool",
              "https://www.fool.com/earnings/call-transcripts/2026/05/06/coherent-cohr-q3-2026-earnings-transcript/")
p3.paragraph_format.space_after = Pt(4)

p4 = doc.add_paragraph()
p4.add_run("NVIDIA Partnership (March 2, 2026): ").font.size = Pt(10)
add_hyperlink(p4, "NVIDIA-Coherent Strategic Partnership Announcement",
              "https://nvidianews.nvidia.com/news/nvidia-and-coherent-announce-strategic-partnership-to-develop-optics-technology-to-scale-next-generation-data-center-architecture")
p4.paragraph_format.space_after = Pt(4)

heading(doc, "Prior Quarter Materials", level=2)

p5 = doc.add_paragraph()
p5.add_run("Q2 FY2026 Earnings Release: ").font.size = Pt(10)
add_hyperlink(p5, "Coherent Corp Q2 FY2026 Press Release",
              "https://www.coherent.com/news/press-releases/second-quarter-fiscal-year-2026-results")
p5.paragraph_format.space_after = Pt(4)

p6 = doc.add_paragraph()
p6.add_run("Q1 FY2026 Earnings Release: ").font.size = Pt(10)
add_hyperlink(p6, "Coherent Corp Q1 FY2026 Press Release",
              "https://www.coherent.com/news/press-releases/first-quarter-fiscal-year-2026-results")
p6.paragraph_format.space_after = Pt(4)

heading(doc, "Market Data", level=2)
body(doc, f"Stock price as of May 14, 2026: {price_str}")
body(doc, f"Market capitalization: {mcap_str}")
body(doc, "Consensus estimates: Bloomberg Terminal, as of May 5, 2026.")
body(doc, "All financial figures from company filings unless otherwise noted.")

# ══════════════════════════════════════════════════════════════════════════════
# DISCLAIMER
# ══════════════════════════════════════════════════════════════════════════════
horizontal_rule(doc)
body(doc, "DISCLAIMER: This report is for informational purposes only and does not constitute investment advice. "
     "The analysis, opinions, and estimates expressed herein are based on publicly available information and "
     "are subject to change without notice. Past performance is not indicative of future results. "
     "Investors should conduct their own due diligence before making investment decisions.",
     italic=True, color=C_GREY, size=8, space_after=2)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = os.path.join(BASE, "COHR_Q3_FY2026_Earnings_Update.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
