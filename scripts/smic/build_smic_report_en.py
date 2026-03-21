"""
SMIC (中芯国际) Q4 2025 Earnings Update — English DOCX Report
Institutional format: 10 pages, Times New Roman
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/SMIC/"
CHARTS = OUT

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
BLUE  = RGBColor(0x00, 0x66, 0xCC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x6B, 0x3A)
RED   = RGBColor(0xCC, 0x00, 0x00)
LGRAY = RGBColor(0xF0, 0xF0, 0xF0)
GOLD  = RGBColor(0xCC, 0x99, 0x00)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for border_name in ['top','left','bottom','right','insideH','insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'CCCCCC')
                tcPr_borders = tcPr.find(qn('w:tcBorders'))
                if tcPr_borders is None:
                    tcPr_borders = OxmlElement('w:tcBorders')
                    tcPr.append(tcPr_borders)
                tcPr_borders.append(border)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '0066CC')
    rPr.append(color_el)
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def heading(doc, text, level=1, color=NAVY, size=None, bold=True, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.bold  = bold
    run.font.color.rgb = color
    if size is None:
        size = {1:16, 2:13, 3:11}.get(level, 11)
    run.font.size = Pt(size)
    return p

def body(doc, text, size=10, bold=False, color=DGRAY, space_before=2, space_after=2, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.color.rgb = DGRAY
    return p

def add_chart(doc, fname, width=6.0, caption=None):
    path = CHARTS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(6)
        for run in cp.runs:
            run.font.size  = Pt(8)
            run.font.italic = True
            run.font.color.rgb = DGRAY

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.ns.qn('w:br'))
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')
    pb.append(bottom)
    pPr.append(pb)

# ── Build Document ────────────────────────────────────────────────────────────
import docx
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.4)
    section.right_margin  = Cm(2.4)

# ────────────────────────────────────────────────────────────────────────────
# PAGE 1 — COVER / SUMMARY
# ────────────────────────────────────────────────────────────────────────────

# Firm header
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("EQUITY RESEARCH  |  SEMICONDUCTORS")
r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.font.color.rgb = DGRAY; r.font.bold = True

# Classification bar (blue shaded table)
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
cells = tbl.rows[0].cells
data_cells = [
    ("RATING: OUTPERFORM", "003366"),
    ("PRICE TARGET: HKD 38.00", "003366"),
    ("REPORT DATE: March 16, 2026", "003366"),
]
for cell, (txt, bg) in zip(cells, data_cells):
    cell.text = txt
    set_cell_bg(cell, bg)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = WHITE
for cell in cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph()

# Title
heading(doc, "Semiconductor Manufacturing International Corporation (SMIC)", level=1, size=15)

heading(doc, "Q4 2025 Earnings Update: Revenue Beats, Margins Under Pressure from Capacity Ramp",
        level=2, size=13, color=BLUE)

divider(doc)

# Sub-info line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
infos = [
    ("Tickers: ", "0981.HK / 688981.SS  "),
    ("Sector: ", "Semiconductor Foundry  "),
    ("Current Price: ", "HKD 28.40 (Mar 14, 2026)  "),
    ("52-Week Range: ", "HKD 17.82 – HKD 42.80"),
]
for label, val in infos:
    r = p.add_run(label); r.font.name='Times New Roman'; r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb=NAVY
    r = p.add_run(val);   r.font.name='Times New Roman'; r.font.size=Pt(9); r.font.color.rgb=DGRAY

# ── Key Takeaway Box ───────────────────────────────────────────────────────────
doc.add_paragraph()
tbl2 = doc.add_table(rows=1, cols=1)
tbl2.style = 'Table Grid'
cell = tbl2.rows[0].cells[0]
set_cell_bg(cell, 'EBF0F8')
para = cell.paragraphs[0]
r = para.add_run("KEY TAKEAWAYS")
r.font.name='Times New Roman'; r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=NAVY
cell.add_paragraph()

bullets_text = [
    "REVENUE BEAT: Q4'25 revenue of $2,489M (+12.8% YoY, +4.5% QoQ) beat Bloomberg consensus of ~$2,420M by +2.8%. Full-year 2025 revenue reached a record $9,327M (+16.2% YoY).",
    "MARGIN MISS: Gross margin contracted to 19.2% in Q4'25 (vs. 22.0% in Q3'25 and consensus 20.9%), driven by a ~30% step-up in depreciation as newly commissioned 12-inch capacity entered its depreciation cycle.",
    "NET INCOME SURGES: Q4'25 net income of $172.85M (+60.7% YoY) beat LSEG consensus of $170.3M, reflecting operating leverage on higher volumes.",
    "Q1 2026 GUIDANCE IN-LINE: Revenue guided flat QoQ (~$2,489M); gross margin guided at 18–20% — slightly below the prior consensus of ~20.9%, reflecting seasonal weakness in consumer electronics.",
    "THESIS INTACT: The depreciation headwind is structural but temporary — yields on 7nm are improving (now 60–70%), and capacity utilization of 95.7% demonstrates strong demand absorption. We maintain OUTPERFORM.",
]
for bt in bullets_text:
    bp = cell.add_paragraph()
    bp.paragraph_format.left_indent = Inches(0.1)
    bp.paragraph_format.space_before = Pt(1)
    bp.paragraph_format.space_after  = Pt(1)
    r = bp.add_run("• " + bt)
    r.font.name='Times New Roman'; r.font.size=Pt(9.5); r.font.color.rgb=DGRAY

# ── Financial Snapshot Table ───────────────────────────────────────────────────
doc.add_paragraph()
heading(doc, "Financial Snapshot", level=3, size=11)

snap_headers = ["Metric", "Q4'24A", "Q3'25A", "Q4'25A", "QoQ Chg", "YoY Chg", "Consensus", "vs. Est."]
snap_rows = [
    ["Revenue ($M)", "2,207", "2,382", "2,489", "+4.5%", "+12.8%", "~2,420", "+2.8% BEAT"],
    ["Gross Margin", "22.6%", "22.0%", "19.2%", "-2.8pp", "-3.4pp", "20.9%", "MISS"],
    ["Net Income ($M)", "107", "192", "173", "-9.9%", "+60.7%", "170", "+1.7% BEAT"],
    ["EPS (USD)", "0.010", "0.018", "0.016", "-", "+60%", "0.016", "In-Line"],
    ["EBITDA ($M)", "~1,100", "~1,300", "1,405", "+8.1%", "+27.7%", "~1,380", "+1.8% BEAT"],
    ["Capex ($M)", "~800", "~950", "~1,050", "+10.5%", "+31.3%", "-", "-"],
    ["Capacity Util.", "86.3%", "90.1%", "95.7%", "+5.6pp", "+9.4pp", "-", "-"],
]

t = doc.add_table(rows=1+len(snap_rows), cols=len(snap_headers))
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(snap_headers):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, '003366')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].add_run(h)
    r.font.name='Times New Roman'; r.font.size=Pt(8.5); r.font.bold=True; r.font.color.rgb=WHITE
for ri, row_data in enumerate(snap_rows):
    row = t.rows[ri+1]
    for ci, val in enumerate(row_data):
        cell = row.cells[ci]
        if ri % 2 == 0:
            set_cell_bg(cell, 'F5F8FF')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = cell.paragraphs[0].add_run(val)
        run.font.name='Times New Roman'; run.font.size=Pt(8.5)
        if "BEAT" in val:
            run.font.color.rgb = GREEN; run.font.bold = True
        elif "MISS" in val:
            run.font.color.rgb = RED; run.font.bold = True

p = doc.add_paragraph("Source: SMIC Q4 2025 Earnings Release (Feb 10, 2026); Bloomberg Consensus; LSEG (Feb 10, 2026)")
p.paragraph_format.space_before = Pt(2)
for run in p.runs:
    run.font.size = Pt(7.5); run.font.italic = True; run.font.color.rgb = DGRAY

# ────────────────────────────────────────────────────────────────────────────
# PAGE 2 — DETAILED RESULTS: REVENUE & VOLUMES
# ────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "Detailed Results Analysis", level=1, size=14)
divider(doc)

heading(doc, "1. Revenue Performance", level=2, size=12)
body(doc, "SMIC reported Q4 2025 revenue of $2,489M, a record quarterly high, beating the Bloomberg consensus of approximately $2,420M by 2.8% and advancing 12.8% year-over-year and 4.5% sequentially. Full-year 2025 revenue reached $9,327M, +16.2% versus $8,028M in 2024.", size=10)
body(doc, "Volume and pricing both contributed to the beat. Wafer shipments in Q4 rose only +0.6% QoQ to approximately 2.1 million 8-inch equivalent wafers, but average selling prices (ASPs) increased by +3.8% QoQ, reflecting a higher mix of 12-inch wafers in the product mix. The shift toward higher-ASP 12-inch wafers — driven by Huawei's Kirin and AI chip demand — was a key driver.", size=10)
add_chart(doc, "smic_chart1_revenue.png", width=6.2,
          caption="Figure 1: SMIC Quarterly Revenue Trend (USD M) | Source: SMIC Earnings Releases")

heading(doc, "2. End-Market Breakdown", level=2, size=12)
body(doc, "Consumer electronics was the dominant segment in Q4 2025, representing 47% of revenue (up from 43% in Q3), driven by seasonal strength in IoT devices, set-top boxes, and wearables. Smartphones fell to 20% of revenue as handset demand moderated. China-domestic customers accounted for approximately 90% of total Q4 sales, up from ~88% in Q3, as SMIC deepens its role in China's semiconductor self-sufficiency ecosystem.", size=10)
add_chart(doc, "smic_chart5_endmarket.png", width=5.0,
          caption="Figure 2: Revenue by End Market, Q4 2025 | Source: SMIC Q4 2025 Earnings Call (Feb 10, 2026)")

add_chart(doc, "smic_chart6_geography.png", width=5.0,
          caption="Figure 3: Revenue by Geography, Q4 2025 | Source: SMIC Q4 2025 Earnings Release")

# ────────────────────────────────────────────────────────────────────────────
# PAGE 3 — MARGINS & PROFITABILITY
# ────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Gross Margin Analysis — Depreciation Headwind", level=2, size=12)
body(doc, "Gross margin contracted sharply to 19.2% in Q4 2025, down from 22.0% in Q3 2025 and below the consensus estimate of 20.9%. The -2.8pp sequential decline was almost entirely attributable to depreciation from newly commissioned capacity, rather than demand weakness or pricing pressure.", size=10)
body(doc, "Management detailed the margin bridge on the earnings call: the step-up in depreciation, representing the first full quarter of charges from fabs commissioned in H2 2025, accounted for approximately -4.2pp of gross margin impact. This was partially offset by positive volume and ASP effects (+1.4pp combined). We estimate total depreciation will rise approximately 30% YoY in 2026 as new capacity matures.", size=10)
add_chart(doc, "smic_chart7_margin_bridge.png", width=6.2,
          caption="Figure 4: Gross Margin Bridge Q3'25 → Q4'25 | Source: SMIC Q4 2025 Earnings Call (Feb 10, 2026)")

add_chart(doc, "smic_chart2_gross_margin.png", width=6.2,
          caption="Figure 5: Quarterly Gross Margin Trend | Source: SMIC Earnings Releases")

heading(doc, "4. Net Income & Profitability", level=2, size=12)
body(doc, "Despite the margin compression, Q4 2025 net income attributable to shareholders reached $172.85M (RMB 1.223B), +60.7% YoY. This sharp year-over-year growth reflects the low base in Q4 2024 (when margins troughed at 22.6% but net income was depressed by lower volumes and elevated opex). Full-year 2025 net income reached approximately $666M (RMB 5.04B), +36.3% versus FY2024.", size=10)
add_chart(doc, "smic_chart3_net_income.png", width=6.2,
          caption="Figure 6: Quarterly Net Income Trend (USD M) | Source: SMIC Earnings Releases")

# ────────────────────────────────────────────────────────────────────────────
# PAGE 4 — KEY METRICS & GUIDANCE
# ────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "Key Metrics & Q1 2026 Guidance", level=1, size=14)
divider(doc)

heading(doc, "5. Capacity Utilization & Expansion", level=2, size=12)
body(doc, "Q4 2025 capacity utilization reached 95.7%, the highest in two years, demonstrating strong demand absorption against a backdrop of rapid capacity additions. Monthly wafer capacity grew to approximately 1.06 million 8-inch equivalent wafers per month (wspm) by year-end 2025, up from ~950K wspm at end-2024 — a +111K wspm addition in 2025.", size=10)
body(doc, "For 2026, SMIC guided CapEx to be 'roughly flat' versus the ~$6.5B invested in 2025, implying continued capacity expansion. Monthly capacity is expected to increase by a further ~40,000 12-inch equivalent wspm versus end-2025, with new fabs in Beijing and Tianjin ramping through the year.", size=10)
add_chart(doc, "smic_chart8_utilization.png", width=6.2,
          caption="Figure 7: Capacity Utilization (%) vs. Monthly Wafer Capacity | Source: SMIC Earnings Releases; Company Filings")

heading(doc, "6. Q1 2026 Guidance vs. Consensus", level=2, size=12)

guide_headers = ["Metric", "Q4'25 Actual", "Q1'26E (Mgmt Guide)", "Q1'26E (Prior Consensus)", "vs. Consensus"]
guide_rows = [
    ["Revenue ($M)", "$2,489", "~$2,489 (Flat QoQ)", "~$2,520", "Slight Miss"],
    ["Gross Margin", "19.2%", "18% – 20%", "~20.9%", "Below Mid"],
    ["CapEx ($B)", "~$1.6B (FY est.)", "~Flat YoY (FY'26)", "~$6.0B (FY'26)", "In-Line"],
]
gt = doc.add_table(rows=1+len(guide_rows), cols=len(guide_headers))
gt.style = 'Table Grid'
gt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(guide_headers):
    cell = gt.rows[0].cells[i]
    set_cell_bg(cell, '003366')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].add_run(h)
    r.font.name='Times New Roman'; r.font.size=Pt(8.5); r.font.bold=True; r.font.color.rgb=WHITE
for ri, row_data in enumerate(guide_rows):
    row = gt.rows[ri+1]
    for ci, val in enumerate(row_data):
        cell = row.cells[ci]
        if ri % 2 == 0: set_cell_bg(cell, 'F5F8FF')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = cell.paragraphs[0].add_run(val)
        run.font.name='Times New Roman'; run.font.size=Pt(8.5)
        if "Miss" in val or "Below" in val:
            run.font.color.rgb = RED
        elif "In-Line" in val or "Beat" in val:
            run.font.color.rgb = GREEN

p = doc.add_paragraph("Source: SMIC Q4 2025 Earnings Call Management Commentary (Feb 10, 2026); CLSA; Bloomberg Consensus")
for run in p.runs:
    run.font.size=Pt(7.5); run.font.italic=True; run.font.color.rgb=DGRAY

body(doc, "Management attributed Q1 2026 revenue flatness to seasonal weakness in low-end consumer electronics (IoT, appliances) and a cyclical inventory drawdown in the smartphone segment. This partially offsets accelerating demand from AI accelerator chips — namely Huawei Ascend 910C, where SMIC's 7nm yield has improved to ~40% (versus ~20% six months ago), making the product commercially profitable for the first time.", size=10, space_before=6)

# ────────────────────────────────────────────────────────────────────────────
# PAGE 5 — BEAT/MISS & TECHNOLOGY
# ────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "Beat / Miss Analysis & Technology Update", level=1, size=14)
divider(doc)

heading(doc, "7. Q4 2025 Beat / Miss Summary", level=2, size=12)
add_chart(doc, "smic_chart4_beat_miss.png", width=6.2,
          caption="Figure 8: Q4 2025 Beat/Miss vs. Bloomberg Consensus | Source: SMIC Earnings Release; Bloomberg (Feb 10, 2026)")

body(doc, "Revenue came in $69M or +2.8% above consensus, primarily driven by stronger-than-expected wafer shipment volumes in consumer electronics and higher ASPs from 12-inch wafer mix shift. Gross margin of 19.2% was the key miss versus consensus at 20.9% — a 1.7pp shortfall driven by the faster-than-expected depreciation step-up. Net income beat by $2.6M or +1.5%, as strong operating leverage offset the margin compression.", size=10)

heading(doc, "8. Technology Developments", level=2, size=12)
body(doc, "7nm Process (N+2): SMIC's most advanced process node continues to improve. Yields on 7nm reportedly reached 60–70% in Q4 2025, up from below 40% at initial launch in H2 2023. Huawei procures approximately 15,000 wspm of SMIC's ~20,000 wspm 7nm total capacity, primarily for Kirin mobile SoCs and Ascend AI accelerators.", size=10)
body(doc, "5nm Development (N+3): Reports suggest SMIC is developing a 5nm-class process using multi-patterning DUV lithography (without EUV access). Huawei's Kirin 9030, expected for mass production in H2 2026, may leverage this node. TrendForce estimates China's 7nm/5nm combined output could increase fivefold within two years, largely driven by SMIC.", size=10)
body(doc, "Export Control Risk: In September 2025, the US Commerce Department added two Chinese firms (GMC Semiconductor Technology Wuxi; Jicun Semiconductor Technology Shanghai) to the Entity List for acquiring US-origin semiconductor equipment intended for SMIC's Beijing fabs. This continues the pattern of tightening export controls but has not materially disrupted SMIC's near-term equipment supply, as the company has been pre-building buffer inventory.", size=10, italic=False)

# ────────────────────────────────────────────────────────────────────────────
# PAGE 6-7 — UPDATED THESIS & RISKS
# ────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "Updated Investment Thesis", level=1, size=14)
divider(doc)

heading(doc, "9. Thesis Unchanged: Structural Story Intact Through Cycle", level=2, size=12)
body(doc, "We maintain our OUTPERFORM rating and HKD 38.00 price target. The Q4 2025 results reinforce three core pillars of our investment thesis:", size=10)
bullet(doc, "China-First Demand Cycle: Domestic semiconductor content is increasing rapidly across consumer, industrial, and AI applications. SMIC's ~90% domestic revenue share positions it as the primary beneficiary of China's semiconductor self-sufficiency push, underpinned by government-backed demand and multi-year capacity commitments from key customers including Huawei, Qualcomm China, and domestic fabless designers.")
bullet(doc, "Technology Progress at Advanced Nodes: Yield improvements on 7nm from <40% to 60–70% represent a significant operational milestone. At higher yields, 7nm economics become meaningful contributors to margins rather than drags. We expect 7nm contribution to expand from an estimated 8% of FY2025 revenue to ~12–15% of FY2026 revenue.")
bullet(doc, "Depreciation Headwind is Temporal, Not Structural: The Q4 2025 gross margin compression is entirely a function of depreciation timing — not a demand problem, not a pricing problem. As the new capacity enters its production ramp and revenue grows into the fixed cost base, we expect gross margins to recover toward 21–23% in H2 2026.")

body(doc, "The key near-term risk to our thesis is margin trajectory: if the Q1 2026 gross margin disappoints at the low end of guidance (18%), markets may extrapolate a prolonged margin trough. We estimate the downside scenario is already reflected in current valuation at 1.8x P/S NTM versus TSMC's 8.5x.", size=10, space_before=6)

heading(doc, "10. Key Risks", level=2, size=12)
bullet(doc, "Margin Downside: Gross margins could remain below 20% through H1 2026 if wafer mix does not shift favorably toward higher-ASP nodes. Each 1pp of gross margin impacts EPS by approximately 5–6%.")
bullet(doc, "Export Control Escalation: Further US restrictions on DUV equipment access could limit SMIC's ability to develop sub-7nm nodes, constraining long-term competitive positioning against TSMC.")
bullet(doc, "Customer Concentration: Huawei accounts for an estimated 35–40% of revenue. Any geopolitical restriction specifically targeting Huawei's chip procurement could be severely disruptive.")
bullet(doc, "CapEx Execution Risk: SMIC plans ~$6.5B in FY2026 CapEx in an environment where equipment lead times remain extended. Delays in capacity additions could cause SMIC to lose share to Hua Hong or CXMT.")
bullet(doc, "Macro / Inventory Cycle: Consumer electronics destocking could extend into Q2 2026 if macro conditions in China deteriorate, pressuring utilization rates below 90%.")

# ────────────────────────────────────────────────────────────────────────────
# PAGE 8 — VALUATION & ESTIMATES
# ────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "Valuation & Updated Estimates", level=1, size=14)
divider(doc)

heading(doc, "11. Valuation — Attractive Relative to Peers Despite Discount", level=2, size=12)
body(doc, "SMIC trades at approximately 1.8x NTM P/S and 8.5x NTM EV/EBITDA, a deep discount to TSMC (8.5x P/S; 14.2x EV/EBITDA) and a modest discount to UMC (1.5x P/S) and GlobalFoundries (2.2x P/S). Given SMIC's superior growth profile (+16% FY2025, guided above-industry FY2026), we believe the discount is excessive and reflects geopolitical risk premium rather than fundamental undervaluation.", size=10)
body(doc, "Our HKD 38.00 price target is based on 2.5x FY2026E P/S of $10.6B, implying a 39% upside from the March 14, 2026 closing price of HKD 28.40. We apply a 70% discount to TSMC's multiple to reflect SMIC's geopolitical risk, technology gap, and margin cyclicality.", size=10)

add_chart(doc, "smic_chart10_valuation.png", width=6.2,
          caption="Figure 9: SMIC vs. Foundry Peers — NTM P/S and EV/EBITDA | Source: Bloomberg; Company Filings; Analyst Consensus (Mar 2026)")

heading(doc, "12. Updated Estimates", level=2, size=12)
add_chart(doc, "smic_chart9_estimates.png", width=6.2,
          caption="Figure 10: Revenue & Gross Margin Estimates — Actuals & Revisions | Source: SMIC Filings; Internal Estimates (Mar 2026)")

# Detailed estimates table
est_headers = ["", "FY2024A", "FY2025A", "FY2026E (New)", "FY2026E (Old)", "FY2027E"]
est_data = [
    ["Revenue ($M)", "8,028", "9,327", "10,600", "10,200", "12,000"],
    ["Revenue Growth", "-", "+16.2%", "+13.6%", "+9.4%", "+13.2%"],
    ["Gross Profit ($M)", "1,445", "1,959", "2,067", "2,091", "2,580"],
    ["Gross Margin", "18.0%", "21.0%", "19.5%", "20.5%", "21.5%"],
    ["EBITDA ($M)", "4,100", "4,850", "5,300", "5,100", "6,200"],
    ["Net Income ($M)", "490", "666", "720", "700", "950"],
    ["EPS (USD)", "0.046", "0.063", "0.068", "0.066", "0.090"],
    ["CapEx ($B)", "5.7", "6.5", "~6.5", "~6.0", "~6.5"],
    ["P/S (x)", "2.2x", "1.9x", "1.7x", "1.7x", "1.5x"],
    ["EV/EBITDA (x)", "9.2x", "7.8x", "7.1x", "7.4x", "6.1x"],
]

et = doc.add_table(rows=1+len(est_data), cols=len(est_headers))
et.style = 'Table Grid'
et.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(est_headers):
    cell = et.rows[0].cells[i]
    set_cell_bg(cell, '003366')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].add_run(h)
    r.font.name='Times New Roman'; r.font.size=Pt(8.5); r.font.bold=True; r.font.color.rgb=WHITE
for ri, row_data in enumerate(est_data):
    row = et.rows[ri+1]
    for ci, val in enumerate(row_data):
        cell = row.cells[ci]
        if ri % 2 == 0: set_cell_bg(cell, 'F5F8FF')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(val)
        run.font.name='Times New Roman'; run.font.size=Pt(8.5)
        if ci == 3 and ri in [0, 2, 3, 4, 5]:  # Revised estimates column
            run.font.color.rgb = RED
        if ci == 0:
            run.font.bold = True; run.font.color.rgb = NAVY

p = doc.add_paragraph("Source: SMIC Annual Reports; SMIC Q4 2025 Earnings Release; Internal Analyst Estimates; Bloomberg Consensus as of Mar 2026.")
p.paragraph_format.space_before = Pt(2)
for run in p.runs:
    run.font.size=Pt(7.5); run.font.italic=True; run.font.color.rgb=DGRAY

body(doc, "Note: FY2026E gross margin revised down to 19.5% (from 20.5%) to reflect full-year depreciation step-up and slower margin recovery assumptions. Revenue estimate revised up to $10.6B (from $10.2B) on stronger-than-expected Q4 demand and above-industry guidance.", size=9, italic=True, color=DGRAY, space_before=4)

# ────────────────────────────────────────────────────────────────────────────
# PAGE 9 — SOURCES
# ────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "Sources & References", level=1, size=14)
divider(doc)

sources = [
    ("SMIC Official Earnings Release — Q4 2025 Results (Feb 10, 2026)",
     "https://www.smics.com/en/site/news_read/7949"),
    ("SMIC Q4 2025 Earnings Call Transcript (Feb 10, 2026) — Seeking Alpha",
     "https://seekingalpha.com/article/4868483-semiconductor-manufacturing-international-corporation-siuif-q4-2025-earnings-call-transcript"),
    ("TrendForce: SMIC Posts Record $9.3B in 2025 Sales; 7nm Yields Reportedly Weigh on Margins (Feb 11, 2026)",
     "https://www.trendforce.com/news/2026/02/11/news-smic-posts-record-9-3b-in-2025-sales-7nm-yields-reportedly-weigh-on-margins/"),
    ("Sharecast: SMIC smashes estimates as Q4 profits surge 61% (Feb 10, 2026)",
     "https://www.sharecast.com/news/international-companies/chipmaker-smic-smashes-estimates-with-61-surge-in-q4-profits--21644274.html"),
    ("CLSA Research Note: SMIC Q1 Revenue Guidance Meets Expectations; GM Slightly Weaker (Feb 10, 2026)",
     "https://news.futunn.com/en/post/68736865/clsa-smic-s-q1-revenue-guidance-meets-market-expectations-while"),
    ("KR-Asia: SMIC posts revenue growth in Q4 as expansion weighs on margins (Feb 11, 2026)",
     "https://kr-asia.com/smic-posts-revenue-growth-in-q4-as-expansion-weighs-on-margins"),
    ("TrendForce: China Aims to Boost 7nm/5nm Output Fivefold in Two Years (Feb 25, 2026)",
     "https://www.trendforce.com/news/2026/02/25/news-china-reportedly-aims-to-boost-7nm-5nm-output-fivefold-in-two-years-driven-by-smic-and-hua-hong/"),
    ("Taipei Times: US Penalizes Firms That Acquired Tools for SMIC (Sep 15, 2025)",
     "https://www.taipeitimes.com/News/biz/archives/2025/09/15/2003843789"),
    ("Edgen.tech: SMIC Projects Peer-Beating Growth After $9.3B 2025 Revenue (Feb 2026)",
     "https://www.edgen.tech/news/stock/smic-projects-peer-beating-growth-after-93b-2025-revenue"),
    ("Futunn News: SMIC Q4 Earnings — Record Revenue & Q1 Guidance (Feb 10, 2026)",
     "https://news.futunn.com/en/post/68701909/smic-released-its-q4-earnings-report-quarterly-revenue-hit-a"),
    ("SCMP: SMIC Expects Flat Revenue as Drop in Low-End Orders Offsets AI Chip Growth (Feb 10, 2026)",
     "https://www.scmp.com/tech/big-tech/article/3343139/chinas-smic-expects-flat-revenue-drop-low-end-orders-offsets-ai-chip-growth"),
    ("Simply Wall St: How SMIC's Q4 Profit and Flat Q1 Margin Outlook Will Impact Investors (Feb 2026)",
     "https://simplywall.st/stocks/hk/semiconductors/hkg-981/semiconductor-manufacturing-international-shares/news/how-smics-q4-profit-and-flat-q1-margin-outlook-will-impact-s"),
    ("Bloomberg Consensus Estimates (as of Feb 10, 2026)", "https://www.bloomberg.com"),
    ("LSEG Analyst Consensus (as of Feb 10, 2026)", "https://www.lseg.com"),
]

for src_text, src_url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    r = p.add_run("• ")
    r.font.name='Times New Roman'; r.font.size=Pt(9); r.font.color.rgb=DGRAY
    add_hyperlink(p, src_text, src_url)

# Disclaimer
doc.add_paragraph()
divider(doc)
disc = body(doc, "ANALYST CERTIFICATION & IMPORTANT DISCLOSURES: This report was prepared for informational purposes. The analysts responsible for this report certify that the views expressed accurately reflect their personal views about the subject securities. This is a research report generated using publicly available information. Past performance is not indicative of future results. This report does not constitute investment advice.", size=8, italic=True, color=DGRAY)

# Save
fname = OUT + "SMIC_Q4_2025_Earnings_Update.docx"
doc.save(fname)
print(f"Saved: {fname}")
