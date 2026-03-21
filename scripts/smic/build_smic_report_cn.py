"""
SMIC (中芯国际) Q4 2025 业绩更新报告 — 中文版 DOCX
机构格式：宋体正文，黑体标题
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/SMIC/"
CHARTS = OUT

# ── 颜色 ──────────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
BLUE  = RGBColor(0x00, 0x66, 0xCC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x6B, 0x3A)
RED   = RGBColor(0xCC, 0x00, 0x00)

# ── 工具函数 ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '0066CC')
    rPr.append(color_el)
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def heading_cn(doc, text, level=1, color=NAVY, size=None, space_before=6, space_after=4):
    """黑体标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.bold = True
    run.font.color.rgb = color
    if size is None:
        size = {1:16, 2:13, 3:11}.get(level, 11)
    run.font.size = Pt(size)
    # 设置中文字体
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), '黑体')
    rPr.insert(0, rFonts)
    return p

def body_cn(doc, text, size=10.5, bold=False, color=DGRAY, space_before=3, space_after=3, italic=False):
    """宋体正文"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.name   = '宋体'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.insert(0, rFonts)
    return p

def bullet_cn(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run("• " + text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.font.color.rgb = DGRAY
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.insert(0, rFonts)
    return p

def add_chart(doc, fname, width=6.0, caption=None):
    path = CHARTS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(6)
        for run in cp.runs:
            run.font.size   = Pt(8)
            run.font.italic = True
            run.font.color.rgb = DGRAY

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')
    pb.append(bottom)
    pPr.append(pb)

def make_table(doc, headers, rows, hdr_bg='003366', alt_bg='F5F8FF', font_size=9):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, hdr_bg)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(h)
        r.font.name='黑体'; r.font.size=Pt(font_size); r.font.bold=True; r.font.color.rgb=WHITE
        rPr = r._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '黑体')
        rPr.insert(0, rFonts)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 0: set_cell_bg(cell, alt_bg)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name='宋体'; run.font.size=Pt(font_size)
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rPr.insert(0, rFonts)
            if "超预期" in str(val) or "优于" in str(val):
                run.font.color.rgb = GREEN; run.font.bold = True
            elif "低于" in str(val) or "不及" in str(val):
                run.font.color.rgb = RED; run.font.bold = True
            elif ci == 0:
                run.font.bold = True; run.font.color.rgb = NAVY
    return t

def source_note(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.size=Pt(7.5); run.font.italic=True; run.font.color.rgb=DGRAY
        run.font.name='宋体'

# ── 正式构建文档 ──────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.4)
    section.right_margin  = Cm(2.4)

# ════════════════════════════════════════════════════════════════════
# 第一页：封面与摘要
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run("股票研究  |  半导体行业")
r.font.name='黑体'; r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb=DGRAY

# 评级栏
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
cells = tbl.rows[0].cells
data_cells = [
    ("评级：跑赢大市", "003366"),
    ("目标价：港币 38.00 元", "003366"),
    ("报告日期：2026年3月16日", "003366"),
]
for cell, (txt, bg) in zip(cells, data_cells):
    cell.text = txt
    set_cell_bg(cell, bg)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name='黑体'; run.font.size=Pt(9); run.font.bold=True; run.font.color.rgb=WHITE

doc.add_paragraph()

# 主标题
heading_cn(doc, "中芯国际集成电路制造有限公司（SMIC）", level=1, size=15)
heading_cn(doc, "2025年第四季度业绩更新：营收超预期，折旧压力致毛利率承压",
           level=2, size=13, color=BLUE)
divider(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
infos = [
    ("股票代码：", "0981.HK / 688981.SS  "),
    ("行业：", "半导体晶圆代工  "),
    ("当前股价：", "港币 28.40 元（2026年3月14日）  "),
    ("52周区间：", "港币 17.82 – 42.80 元"),
]
for label, val in infos:
    r = p.add_run(label); r.font.name='黑体'; r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb=NAVY
    r = p.add_run(val);   r.font.name='宋体'; r.font.size=Pt(9); r.font.color.rgb=DGRAY

# 核心结论框
doc.add_paragraph()
tbl2 = doc.add_table(rows=1, cols=1)
tbl2.style = 'Table Grid'
cell = tbl2.rows[0].cells[0]
set_cell_bg(cell, 'EBF0F8')
para = cell.paragraphs[0]
r = para.add_run("核心结论")
r.font.name='黑体'; r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=NAVY
rPr = r._r.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:eastAsia'), '黑体')
rPr.insert(0, rFonts)

cell.add_paragraph()
bullets_cn = [
    "营收超预期：2025Q4营收24.89亿美元（同比+12.8%，环比+4.5%），超彭博市场一致预期约24.20亿美元的2.8%。全年2025年营收创历史新高，达93.27亿美元（同比+16.2%）。",
    "毛利率低于预期：Q4毛利率降至19.2%（较Q3的22.0%下降2.8个百分点），低于市场一致预期20.9%。主因为新投产12英寸产能进入折旧周期，折旧费用同比预计增约30%，属周期性而非结构性问题。",
    "净利润大幅增长：Q4归母净利润1.7285亿美元（+60.7% YoY），小幅超过LSEG一致预期1.703亿美元，体现了较强的经营杠杆效应。",
    "2026Q1指引符合预期：营收指引环比持平（约24.89亿美元）；毛利率指引18%–20%，略低于市场此前预期的约20.9%，主要反映消费电子旺季退去后的季节性走弱。",
    "投资逻辑不变：折旧压力为阶段性，7nm良率已提升至60-70%，Q4产能利用率达95.7%，需求旺盛。维持【跑赢大市】评级。",
]
for bt in bullets_cn:
    bp = cell.add_paragraph()
    bp.paragraph_format.left_indent = Inches(0.1)
    bp.paragraph_format.space_before = Pt(1)
    bp.paragraph_format.space_after  = Pt(1)
    r = bp.add_run("• " + bt)
    r.font.name='宋体'; r.font.size=Pt(9.5); r.font.color.rgb=DGRAY
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)

# 财务数据摘要表
doc.add_paragraph()
heading_cn(doc, "财务数据摘要", level=3, size=11)

snap_headers = ["指标", "2024Q4实", "2025Q3实", "2025Q4实", "环比变动", "同比变动", "市场预期", "与预期对比"]
snap_rows = [
    ["营收（百万美元）", "2,207", "2,382", "2,489", "+4.5%", "+12.8%", "~2,420", "超预期+2.8%"],
    ["毛利率", "22.6%", "22.0%", "19.2%", "-2.8pp", "-3.4pp", "20.9%", "低于预期"],
    ["归母净利润（百万美元）", "107", "192", "173", "-9.9%", "+60.7%", "170", "超预期+1.7%"],
    ["每股盈利（美元）", "0.010", "0.018", "0.016", "-", "+60%", "0.016", "符合预期"],
    ["EBITDA（百万美元）", "~1,100", "~1,300", "1,405", "+8.1%", "+27.7%", "~1,380", "超预期+1.8%"],
    ["产能利用率", "86.3%", "90.1%", "95.7%", "+5.6pp", "+9.4pp", "-", "-"],
]
make_table(doc, snap_headers, snap_rows, font_size=8.5)
source_note(doc, "资料来源：中芯国际2025Q4业绩公告（2026年2月10日）；彭博一致预期；LSEG（2026年2月10日）")

# ════════════════════════════════════════════════════════════════════
# 第二页：营收与下游市场分析
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading_cn(doc, "详细业绩分析", level=1, size=14)
divider(doc)

heading_cn(doc, "一、营收表现", level=2, size=12)
body_cn(doc, "中芯国际2025Q4营收24.89亿美元，创单季历史新高，超彭博一致预期24.20亿美元约2.8%，同比增长12.8%，环比增长4.5%。全年2025年营收93.27亿美元，较2024年80.28亿美元增长16.2%。")
body_cn(doc, "出货量与售价双双贡献超预期表现。Q4晶圆出货量（按8英寸当量折算）环比仅增长约0.6%，但平均销售价格（ASP）环比提升3.8%，反映12英寸晶圆在产品结构中占比提升。受益于华为麒麟芯片及AI加速器需求拉动，12英寸高ASP产品占比持续扩大，是本季营收超预期的主要驱动力。")
add_chart(doc, "smic_chart1_revenue.png", width=6.2,
          caption="图1：中芯国际季度营收趋势（百万美元）| 资料来源：中芯国际各季度业绩公告")

heading_cn(doc, "二、下游应用市场拆分", level=2, size=12)
body_cn(doc, "2025Q4消费电子仍是最大下游市场，占营收47%（Q3为43%），受益于IoT设备、机顶盒及可穿戴产品季节性旺季备货。智能手机收入占比降至20%，反映手机终端需求趋于平稳。境内客户营收占比约90%，较Q3的约88%进一步提升，体现中芯国际在国内半导体自主可控产业链中的核心地位持续强化。")
add_chart(doc, "smic_chart5_endmarket.png", width=5.0,
          caption="图2：2025Q4下游应用市场收入占比 | 资料来源：中芯国际2025Q4业绩说明会（2026年2月10日）")
add_chart(doc, "smic_chart6_geography.png", width=5.0,
          caption="图3：2025Q4收入地区分布 | 资料来源：中芯国际2025Q4业绩公告")

# ════════════════════════════════════════════════════════════════════
# 第三页：毛利率与盈利分析
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading_cn(doc, "三、毛利率分析——折旧压力剖析", level=2, size=12)
body_cn(doc, "2025Q4毛利率降至19.2%，较Q3的22.0%下滑2.8个百分点，低于市场一致预期20.9%。管理层在业绩说明会上明确表示：此次毛利率下滑几乎完全由新增产能进入折旧周期所致，并非需求疲软或定价压力。")
body_cn(doc, "据管理层在说明会上的分析：H2 2025年新投产的晶圆厂于Q4进入首个完整折旧季度，合计对毛利率造成约-4.2个百分点的拖累。出货量增加和ASP提升合计贡献正向约+1.4个百分点的对冲效应。预计2026年全年折旧费用同比增加约30%，折旧压力将贯穿全年。")
add_chart(doc, "smic_chart7_margin_bridge.png", width=6.2,
          caption="图4：毛利率拆解桥接图 Q3'25 → Q4'25 | 资料来源：中芯国际2025Q4业绩说明会（2026年2月10日）")
add_chart(doc, "smic_chart2_gross_margin.png", width=6.2,
          caption="图5：季度毛利率走势 | 资料来源：中芯国际历季业绩公告")

heading_cn(doc, "四、净利润与盈利能力", level=2, size=12)
body_cn(doc, "尽管毛利率承压，2025Q4归母净利润达1.7285亿美元（约12.23亿人民币），同比+60.7%，略超LSEG一致预期1.703亿美元。净利润大幅增长主要源于2024Q4低基数效应（彼时净利润仅约1.07亿美元）以及更高出货量带来的营运杠杆效应。全年2025年归母净利润约6.66亿美元（约50.41亿人民币），同比+36.3%。")
add_chart(doc, "smic_chart3_net_income.png", width=6.2,
          caption="图6：季度归母净利润走势（百万美元）| 资料来源：中芯国际历季业绩公告")

# ════════════════════════════════════════════════════════════════════
# 第四页：关键指标与业绩指引
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading_cn(doc, "关键指标与2026年一季度业绩指引", level=1, size=14)
divider(doc)

heading_cn(doc, "五、产能利用率与扩产进展", level=2, size=12)
body_cn(doc, "2025Q4产能利用率达95.7%，为近两年最高水平，充分验证了旺盛的市场需求可以有效吸收公司持续投入的新增产能。2025年末月产能（按8英寸当量）约达106万片/月，较2024年末约95万片/月净增约11.1万片/月。")
body_cn(doc, "就2026年产能规划，管理层指引全年资本支出【基本持平】2025年约65亿美元，预计2026年末月产能较2025年末再增约4万片12英寸当量/月，北京和天津新厂将于全年陆续爬产。")
add_chart(doc, "smic_chart8_utilization.png", width=6.2,
          caption="图7：产能利用率（%）与月产能走势 | 资料来源：中芯国际历季业绩公告及公司文件")

heading_cn(doc, "六、2026年一季度指引与市场预期对比", level=2, size=12)
guide_headers = ["指标", "2025Q4实际", "2026Q1指引（管理层）", "2026Q1预期（市场）", "与市场预期对比"]
guide_rows = [
    ["营收（百万美元）", "2,489", "~2,489（环比持平）", "~2,520", "略低于预期"],
    ["毛利率", "19.2%", "18% – 20%", "~20.9%", "低于区间中值"],
    ["资本支出（十亿美元）", "~1.6（季度估）", "全年~持平2025", "全年~60亿", "符合预期"],
]
make_table(doc, guide_headers, guide_rows, font_size=8.5)
source_note(doc, "资料来源：中芯国际2025Q4业绩说明会（2026年2月10日）；里昂证券；彭博一致预期")

body_cn(doc, "管理层将Q1 2026营收环比持平归因于消费电子低端产品（IoT、家电）季节性去库存，以及手机板块短期库存调整。上述负面因素与AI加速器芯片需求加速增长形成对冲——中芯国际为华为昇腾910C（7nm制程）代工的良率已从约半年前的20%提升至当前约40%，产品已实现商业化盈利。")

# ════════════════════════════════════════════════════════════════════
# 第五页：与预期对比及技术进展
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading_cn(doc, "与预期对比分析及技术进展", level=1, size=14)
divider(doc)

heading_cn(doc, "七、2025Q4超预期/不及预期汇总", level=2, size=12)
add_chart(doc, "smic_chart4_beat_miss.png", width=6.2,
          caption="图8：2025Q4实际业绩与彭博一致预期对比 | 资料来源：中芯国际业绩公告；彭博（2026年2月10日）")
body_cn(doc, "营收超出一致预期6900万美元（+2.8%），主要由消费电子出货量好于预期以及12英寸晶圆占比提升推动。毛利率19.2%为本季主要不及预期项，低于一致预期20.9%约1.7个百分点，主因折旧加速度超出市场预判。归母净利润超出LSEG预期约260万美元（+1.5%），得益于出货量增长带来的强劲经营杠杆。")

heading_cn(doc, "八、技术节点进展", level=2, size=12)
body_cn(doc, "7纳米制程（N+2）：中芯国际最先进制程节点持续改善，2025Q4良率据报道已达60–70%，远高于2023年下半年量产初期的40%以下。华为约采购中芯国际7nm约2万片/月产能中的1.5万片/月，主要用于麒麟系列手机SoC及昇腾AI加速器。")
body_cn(doc, "5纳米制程开发（N+3）：有报道指出中芯国际正在开发采用多重曝光DUV光刻（无需EUV）的5nm级制程，华为麒麟9030 SoC（预计2026年下半年量产）或将采用该节点。集邦咨询（TrendForce）预计，中国7nm/5nm合计产能在未来两年有望增长五倍，主要依托中芯国际。")
body_cn(doc, "出口管制风险：2025年9月，美国商务部将两家中国企业（无锡GMC半导体科技、上海积村半导体科技）列入实体清单，原因是其为中芯国际北京晶圆厂采购美国管制设备。此类持续收紧的出口管制措施对中芯国际近期设备供应影响有限（公司已构建缓冲库存），但长期制约其发展先进节点的能力。")

# ════════════════════════════════════════════════════════════════════
# 第六页：投资逻辑与风险
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading_cn(doc, "更新投资逻辑", level=1, size=14)
divider(doc)

heading_cn(doc, "九、投资逻辑不变：结构性成长逻辑穿越周期", level=2, size=12)
body_cn(doc, "维持【跑赢大市】评级，目标价港币38.00元。2025Q4业绩进一步印证我们投资逻辑的三大核心支柱：")
bullet_cn(doc, "国产替代主线受益者：中国消费、工业及AI应用领域半导体本土化需求持续高增长，中芯国际约90%境内营收占比使其成为自主可控产业链的核心受益者，背后有政府政策支撑和华为、高通（中国）等大客户的多年产能承诺保障。")
bullet_cn(doc, "先进制程良率持续改善：7nm良率从40%以下提升至60–70%，是重要的经营里程碑。随着良率提高，7nm制程将从前期的拖累因素转变为盈利贡献项，预计7nm营收占比将从2025全年约8%提升至2026年约12–15%。")
bullet_cn(doc, "折旧压力为阶段性而非结构性：Q4毛利率下滑完全源于折旧时序问题，需求面和定价面均无问题。随着新产能进入量产爬坡期、营收规模持续扩大，预计毛利率将于2026年下半年回升至21–23%区间。")
body_cn(doc, "近期核心风险在于毛利率走势：若2026Q1毛利率落至指引区间下限（18%），市场可能担忧毛利率低迷延续。但我们认为当前1.8x NTM P/S的估值（远低于台积电的8.5x）已充分反映地缘政治风险溢价，下行空间有限。")

heading_cn(doc, "十、主要风险因素", level=2, size=12)
bullet_cn(doc, "毛利率超预期下行：若晶圆产品结构未能按预期向高ASP制程迁移，2026上半年毛利率或持续低于20%。毛利率每下降1个百分点对应EPS影响约5–6%。")
bullet_cn(doc, "出口管制升级：若美国进一步限制DUV设备供应，将制约中芯国际发展7nm以下制程的能力，削弱其与台积电的长期竞争力。")
bullet_cn(doc, "大客户集中度风险：华为预估占营收35–40%。如出台专门针对华为芯片采购的地缘政治限制，将对中芯国际造成重大冲击。")
bullet_cn(doc, "资本支出执行风险：在设备交期仍然偏长的背景下，中芯国际计划2026全年投入约65亿美元资本支出，扩产延误可能导致份额流失至华虹半导体或长鑫存储。")
bullet_cn(doc, "宏观与库存周期：若中国宏观经济走弱，消费电子去库存周期或延伸至2026年二季度，拖累产能利用率跌破90%。")

# ════════════════════════════════════════════════════════════════════
# 第七页：估值与盈利预测
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading_cn(doc, "估值与盈利预测更新", level=1, size=14)
divider(doc)

heading_cn(doc, "十一、估值——相对同业具吸引力", level=2, size=12)
body_cn(doc, "中芯国际当前交易于约1.8倍NTM市销率（P/S）和8.5倍NTM EV/EBITDA，相对台积电（P/S 8.5倍；EV/EBITDA 14.2倍）存在显著折让，亦低于联华电子（P/S 1.5倍）和格芯（P/S 2.2倍）的估值中枢。鉴于中芯国际的成长性显著优于同业（2025年+16%，2026年指引高于行业平均），我们认为当前折让程度过大，主要反映地缘政治风险溢价而非基本面折让。")
body_cn(doc, "目标价港币38.00元基于2026年预期营收106亿美元的2.5倍P/S，相对2026年3月14日收盘价港币28.40元约有39%上行空间。我们在台积电估值基础上给予70%折让，以反映中芯国际的地缘政治风险、技术代差和毛利率周期性。")
add_chart(doc, "smic_chart10_valuation.png", width=6.2,
          caption="图9：中芯国际与晶圆代工同业估值对比（NTM P/S与EV/EBITDA）| 资料来源：彭博；公司文件；分析师一致预期（2026年3月）")

heading_cn(doc, "十二、盈利预测更新", level=2, size=12)
add_chart(doc, "smic_chart9_estimates.png", width=6.2,
          caption="图10：营收与毛利率预测——实际值与修订后预测 | 资料来源：中芯国际文件；内部预测（2026年3月）")

est_headers_cn = ["指标", "2024年实际", "2025年实际", "2026年预测（新）", "2026年预测（旧）", "2027年预测"]
est_data_cn = [
    ["营收（百万美元）", "8,028", "9,327", "10,600", "10,200", "12,000"],
    ["营收增速", "-", "+16.2%", "+13.6%", "+9.4%", "+13.2%"],
    ["毛利润（百万美元）", "1,445", "1,959", "2,067", "2,091", "2,580"],
    ["毛利率", "18.0%", "21.0%", "19.5%", "20.5%", "21.5%"],
    ["EBITDA（百万美元）", "4,100", "4,850", "5,300", "5,100", "6,200"],
    ["归母净利润（百万美元）", "490", "666", "720", "700", "950"],
    ["每股盈利（美元）", "0.046", "0.063", "0.068", "0.066", "0.090"],
    ["资本支出（十亿美元）", "5.7", "6.5", "~6.5", "~6.0", "~6.5"],
    ["市销率（倍）", "2.2x", "1.9x", "1.7x", "1.7x", "1.5x"],
    ["EV/EBITDA（倍）", "9.2x", "7.8x", "7.1x", "7.4x", "6.1x"],
]
make_table(doc, est_headers_cn, est_data_cn, font_size=8.5)
source_note(doc, "资料来源：中芯国际年报；2025Q4业绩公告；内部分析师预测；彭博一致预期（截至2026年3月）")
body_cn(doc, "注：2026年预测毛利率由20.5%下调至19.5%，以体现全年折旧增长及毛利率回升节奏慢于前期假设；营收预测由102亿美元上调至106亿美元，体现Q4需求强于预期及管理层高于行业增速的指引。", size=9, italic=True, color=DGRAY)

# ════════════════════════════════════════════════════════════════════
# 第八页：资料来源
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading_cn(doc, "资料来源与参考文献", level=1, size=14)
divider(doc)

sources_cn = [
    ("中芯国际2025年第四季度及全年业绩公告（2026年2月10日）",
     "https://www.smics.com/en/site/news_read/7949"),
    ("中芯国际2025Q4业绩说明会电话录音文字记录（2026年2月10日）— Seeking Alpha",
     "https://seekingalpha.com/article/4868483-semiconductor-manufacturing-international-corporation-siuif-q4-2025-earnings-call-transcript"),
    ("集邦咨询（TrendForce）：中芯国际2025年销售额创纪录达93亿美元，7nm良率拖累利润率（2026年2月11日）",
     "https://www.trendforce.com/news/2026/02/11/news-smic-posts-record-9-3b-in-2025-sales-7nm-yields-reportedly-weigh-on-margins/"),
    ("Sharecast：中芯国际Q4利润激增61%，超越市场预期（2026年2月10日）",
     "https://www.sharecast.com/news/international-companies/chipmaker-smic-smashes-estimates-with-61-surge-in-q4-profits--21644274.html"),
    ("里昂证券研究点评：中芯国际Q1营收指引符合预期，毛利率指引略弱（2026年2月10日）",
     "https://news.futunn.com/en/post/68736865/clsa-smic-s-q1-revenue-guidance-meets-market-expectations-while"),
    ("KR-Asia：中芯国际Q4营收增长，产能扩张拖累利润率（2026年2月11日）",
     "https://kr-asia.com/smic-posts-revenue-growth-in-q4-as-expansion-weighs-on-margins"),
    ("集邦咨询（TrendForce）：中国7nm/5nm产能规模两年内增五倍（2026年2月25日）",
     "https://www.trendforce.com/news/2026/02/25/news-china-reportedly-aims-to-boost-7nm-5nm-output-fivefold-in-two-years-driven-by-smic-and-hua-hong/"),
    ("台北时报：美国制裁为中芯国际采购设备的相关企业（2025年9月15日）",
     "https://www.taipeitimes.com/News/biz/archives/2025/09/15/2003843789"),
    ("Edgen.tech：中芯国际2025年营收93亿美元后预计高于同业增速（2026年2月）",
     "https://www.edgen.tech/news/stock/smic-projects-peer-beating-growth-after-93b-2025-revenue"),
    ("富途牛牛：中芯国际2025Q4业绩发布——季度营收创历史新高（2026年2月10日）",
     "https://news.futunn.com/en/post/68701909/smic-released-its-q4-earnings-report-quarterly-revenue-hit-a"),
    ("南华早报：中芯国际预计营收持平，低端订单下滑被AI芯片需求对冲（2026年2月10日）",
     "https://www.scmp.com/tech/big-tech/article/3343139/chinas-smic-expects-flat-revenue-drop-low-end-orders-offsets-ai-chip-growth"),
    ("彭博分析师一致预期（截至2026年2月10日）", "https://www.bloomberg.com"),
    ("LSEG（路孚特）分析师一致预期（截至2026年2月10日）", "https://www.lseg.com"),
]

for src_text, src_url in sources_cn:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    r = p.add_run("• ")
    r.font.name='宋体'; r.font.size=Pt(9); r.font.color.rgb=DGRAY
    add_hyperlink(p, src_text, src_url)

doc.add_paragraph()
divider(doc)
disc_p = doc.add_paragraph()
disc_p.paragraph_format.space_before = Pt(4)
r = disc_p.add_run("免责声明：本报告基于公开信息，仅供参考，不构成投资建议。所有预测及估值均为分析师独立判断，不代表任何机构的官方立场。投资涉及风险，过往业绩不代表未来表现。")
r.font.name='宋体'; r.font.size=Pt(8); r.font.italic=True; r.font.color.rgb=DGRAY

# 保存
fname = OUT + "SMIC_Q4_2025_业绩更新报告_中文版.docx"
doc.save(fname)
print(f"已保存：{fname}")
