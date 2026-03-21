"""
Alibaba Group (BABA) — Q3 FY2026 (Calendar Q4 2025) Earnings Update
English DOCX report builder
Output: output/BABA/BABA_Q3_FY2026_Earnings_Update.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

OUT  = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/BABA/"
os.makedirs(OUT, exist_ok=True)

# ── yfinance market data ───────────────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        import yfinance as yf
        info = yf.Ticker(ticker).fast_info
        return {
            "price":      f"${info.last_price:.2f}",
            "market_cap": f"~${info.market_cap/1e9:.0f}B",
            "52w_high":   f"${info.year_high:.2f}",
            "52w_low":    f"${info.year_low:.2f}",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using N/A")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("BABA")
print(f"Market data: {mkt}")

# ── Color helpers ──────────────────────────────────────────────────────────────
def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

ALI_ORANGE  = hex_to_rgb("FF6A00")
ALI_NAVY    = hex_to_rgb("1F2D5C")
CLOUD_BLUE  = hex_to_rgb("0070C0")
GREEN_ACC   = hex_to_rgb("00A651")
RED_ACC     = hex_to_rgb("E8192C")
GRAY_LIGHT  = hex_to_rgb("D9D9D9")
WHITE       = hex_to_rgb("FFFFFF")
BLACK       = hex_to_rgb("000000")

# ── Document helpers ───────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

def set_font(run, name="Times New Roman", size=10, bold=False,
             color=None, italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=ALI_NAVY, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p

def body(doc, text, size=9.5, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic)
    return p

def bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def set_cell_bg(cell, color_tuple):
    r, g, b = color_tuple
    hex_color = f"{r:02X}{g:02X}{b:02X}"
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_run.append(new_t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_img(doc, fname, width=6.0):
    path = OUT + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        body(doc, f"[Chart not found: {fname}]", italic=True)

# ── TABLE HELPER ──────────────────────────────────────────────────────────────
def make_table(doc, headers, rows,
               hdr_bg=ALI_NAVY, hdr_fg=WHITE, alt_bg=None,
               col_widths=None, hdr_size=8.5, row_size=8.5):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, hdr_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=hdr_size, bold=True, color=hdr_fg)

    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        bg = alt_bg if (alt_bg and i % 2 == 1) else WHITE
        for j, cell_val in enumerate(row_data):
            cell = row.cells[j]
            if bg != WHITE:
                set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_val))
            set_font(run, size=row_size)

    if col_widths:
        for j, w in enumerate(col_widths):
            set_col_width(table, j, w)

    return table

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER & EARNINGS SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

# Report title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ALIBABA GROUP HOLDING LIMITED (NYSE: BABA | HKEX: 9988)")
set_font(run, size=15, bold=True, color=ALI_NAVY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Q3 FY2026 (Calendar Q4 2025) Earnings Update")
set_font(run2, size=13, bold=True, color=ALI_ORANGE)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Quarter Ended December 31, 2025 | Reported March 18, 2026")
set_font(run3, size=9.5, italic=True, color=tuple(hex_to_rgb("666666")))

doc.add_paragraph()

# Rating / snapshot table
make_table(doc,
    headers=["Rating", "Price Target", "Current Price", "Market Cap", "52-Week Range"],
    rows=[["BUY", "$120.00", mkt["price"], mkt["market_cap"],
           f"{mkt['52w_low']} – {mkt['52w_high']}"]],
    hdr_bg=ALI_NAVY, hdr_fg=WHITE,
    col_widths=[2.5, 2.5, 2.5, 2.5, 3.5],
    hdr_size=9, row_size=9,
)
doc.add_paragraph()

add_heading(doc, "KEY TAKEAWAYS", level=1, color=ALI_NAVY, size=11)
bullets_key = [
    "Revenue MISS: Q3 FY2026 total revenue of RMB 284.8B ($40.7B) came in ~2% below consensus of ~RMB 290.7B, driven by slower international commerce and ongoing divestiture headwinds.",
    "Cloud BEAT: Cloud Intelligence Group revenue surged +36% YoY to RMB 43.3B ($6.2B), beating estimates by ~8%; AI product revenue maintained triple-digit growth for 10 consecutive quarters.",
    "Profitability MISS: Non-GAAP EPS of RMB 7.09/ADS ($1.01) missed consensus ~RMB 9.00; Adjusted EBITA fell -57% YoY to RMB 23.4B — reflecting deliberate heavy investment in AI infrastructure and quick commerce.",
    "Quick Commerce Surge: Taobao Instant Commerce (quick commerce) revenue +56% YoY — Alibaba's fastest-growing domestic commerce segment, though margin-dilutive.",
    "User Momentum Strong: 88VIP members surpassed 59M (double-digit YoY growth); Qwen AI consumer app exceeded 300M monthly active users — validating AI consumer strategy.",
    "Deliberate Investment Cycle: Management reaffirmed aggressive reinvestment in AI infrastructure, Qwen model, and quick commerce logistics. Near-term margin pressure is intentional and expected to persist into Q4 FY2026.",
    "No Formal Guidance Issued: Alibaba declined to provide quantitative Q4 FY2026 revenue or earnings guidance; Cloud momentum and AI demand remain strong qualitative signals.",
    "Stock Reaction: Shares fell ~5% in premarket on March 19 — market punished profitability miss; longer-term AI/cloud upside story intact.",
]
for b in bullets_key:
    bullet(doc, b, size=9.5)

doc.add_paragraph()

# ── Results Snapshot Table ─────────────────────────────────────────────────────
add_heading(doc, "RESULTS SNAPSHOT — Q3 FY2026 vs. Consensus Estimates", level=1, color=ALI_NAVY, size=11)
snap_hdrs  = ["Metric", "Actual", "Consensus Est.", "Beat / Miss", "YoY Change"]
snap_rows  = [
    ["Total Revenue (RMB B)", "¥284.8B", "~¥290.7B", "MISS  –2.0%", "+2% reported\n(+9% LL)"],
    ["Total Revenue (USD)", "$40.7B", "~$41.5B", "MISS  –2.0%", "+2% YoY"],
    ["Adj. EBITA (RMB B)", "¥23.4B", "~¥30.0B", "MISS  –22%", "–57% YoY"],
    ["Non-GAAP EPS/ADS (RMB)", "¥7.09", "~¥9.00", "MISS  –21%", "Significant decline"],
    ["Non-GAAP EPS/ADS (USD)", "$1.01", "~$1.29", "MISS  –22%", "Significant decline"],
    ["Cloud Revenue (RMB B)", "¥43.3B", "~¥40.0B", "BEAT  +8.3%", "+36% YoY"],
    ["Free Cash Flow (RMB B)", "¥11.3B", "~¥20.0B", "MISS  –44%", "–71% YoY"],
    ["Gross Margin", "~39.7%", "~40.5%", "MISS  –80bps", "~–70bps YoY"],
]
make_table(doc, snap_hdrs, snap_rows,
           hdr_bg=ALI_NAVY, hdr_fg=WHITE,
           alt_bg=hex_to_rgb("F5F7FA"),
           col_widths=[4.0, 2.5, 2.5, 2.5, 3.0],
           hdr_size=8.5, row_size=8.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2-3 — DETAILED RESULTS ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "DETAILED RESULTS ANALYSIS", level=1, color=ALI_NAVY, size=13)

add_heading(doc, "Revenue — Topline Softer on Divested Businesses", level=2, color=ALI_ORANGE, size=11)
body(doc, "Alibaba reported Q3 FY2026 (October–December 2025) total revenue of RMB 284.8B ($40.7B), a +2% YoY increase on a reported basis. Excluding the impact of divested businesses — including the Sun Art hypermarket chain and Intime department store — organic (like-for-like) revenue growth was approximately +9% YoY, meaningfully above the headline figure. The ~2% reported revenue miss versus consensus (~RMB 290.7B) was principally attributable to slower growth in the International Digital Commerce segment and ongoing structural drag from the divested retail assets.")

body(doc, "Revenue grew across all retained segments on an organic basis, with Cloud Intelligence being the clear standout. Quick commerce (Taobao Instant Commerce) posted the strongest domestic growth at +56% YoY, though scale remains limited (~¥20.8B) relative to core commerce. China Wholesale (1688.com) grew +5% YoY to ¥6.9B, driven by small and mid-sized enterprise demand.")

add_img(doc, "baba_chart1_quarterly_revenue.png", width=5.8)
body(doc, "Chart 1: Alibaba Group — Quarterly Revenue (RMB Billions), Q4 FY24–Q3 FY26", size=8, italic=True)
doc.add_paragraph()

# ── Segment Detail ─────────────────────────────────────────────────────────────
add_heading(doc, "Segment Revenue Breakdown — Q3 FY2026", level=2, color=ALI_ORANGE, size=11)
seg_hdrs = ["Business Segment", "Q3 FY26 Revenue", "YoY Growth", "Key Commentary"]
seg_rows = [
    ["Cloud Intelligence Group", "¥43.3B ($6.2B)", "+36% YoY", "AI product revenue: triple-digit growth for 10 consec. quarters; Cloud market share 36% in China"],
    ["Intl. Digital Commerce (AIDC)", "¥32.4B ($4.6B)", "+3% YoY", "Deceleration from +32% year-ago; Lazada drag partially offset by AliExpress; loss narrowing"],
    ["Taobao Instant Commerce", "¥20.8B ($3.0B)", "+56% YoY", "Fastest-growing domestic segment; heavy logistics investment weigh on margins"],
    ["China Wholesale (1688.com)", "¥6.9B ($1.0B)", "+5% YoY", "SME-driven B2B commerce; steady growth"],
    ["Cainiao, Media & Others", "¥67.3B ($9.6B)", "–25% YoY", "Reported decline due to Sun Art/Intime divestitures; Cainiao growth offset"],
    ["Total", "¥284.8B ($40.7B)", "+2% (rpt'd)\n+9% LL", "Organic growth healthy; reported dragged by divestitures"],
]
make_table(doc, seg_hdrs, seg_rows, hdr_bg=ALI_NAVY, hdr_fg=WHITE,
           alt_bg=hex_to_rgb("F5F7FA"),
           col_widths=[3.8, 2.8, 2.5, 5.4],
           hdr_size=8.5, row_size=8.5)
doc.add_paragraph()

add_img(doc, "baba_chart3_segment_revenue.png", width=5.8)
body(doc, "Chart 2: Q3 FY2026 Revenue by Business Segment (RMB Billions)", size=8, italic=True)

doc.add_page_break()

# ── Cloud Deep-Dive ────────────────────────────────────────────────────────────
add_heading(doc, "Cloud Intelligence Group — The Crown Jewel (+36%)", level=2, color=CLOUD_BLUE, size=11)
body(doc, "Cloud Intelligence Group delivered another quarter of accelerating growth, with revenue of RMB 43.3B ($6.2B), up +36% YoY — materially ahead of consensus estimates of ~¥40.0B (~+8% beat). This marks a remarkable re-acceleration from the near-stagnation of FY2024 (when cloud grew only +3–7%), and reflects the powerful tailwind from AI-driven infrastructure demand across Chinese enterprises.")

body(doc, "AI product revenue sustained triple-digit YoY growth for the 10th consecutive quarter. Alibaba's Qwen model family has emerged as a leading frontier model in China, with the Qwen consumer app surpassing 300M monthly active users. Cloud market share in China climbed to ~36%, up for three consecutive quarters. External revenue across FY2026 (through February 2026) surpassed CNY 100 billion — a landmark milestone that validates Alibaba Cloud's enterprise traction.")

body(doc, "The cloud growth narrative is now firmly intact: Alibaba's AI strategy is differentiated, its model capabilities are competitive globally, and enterprise adoption is accelerating. Cloud is transforming from a cost center into Alibaba's most valuable growth engine. We see cloud growing 30–40% through FY2027 if AI adoption continues at current pace.")

add_img(doc, "baba_chart2_cloud_revenue.png", width=5.8)
body(doc, "Chart 3: Cloud Intelligence Group — Quarterly Revenue & YoY Growth", size=8, italic=True)

doc.add_paragraph()
add_heading(doc, "International Digital Commerce (AIDC) — Structural Reset", level=2, color=ALI_ORANGE, size=11)
body(doc, "AIDC revenue of RMB 32.4B ($4.6B) grew only +3% YoY, a sharp deceleration from the +32% pace recorded in the year-ago Q3 FY2025 quarter. The slowdown reflects two factors: (1) Lazada (Southeast Asia) experienced continued operational restructuring and a reduction in loss-making GMV, partially offsetting gains from AliExpress; and (2) AliExpress is intentionally shifting from a volume-driven to a margin-driven model.")

body(doc, "The positive read-through is AIDC's significant profitability improvement — losses narrowed substantially YoY through logistics optimization and a more disciplined approach to international expansion. Management appears to be prioritizing quality of international revenue over quantity, consistent with the overall portfolio rationalization approach (Intime, Sun Art divestitures).")

add_img(doc, "baba_chart8_aidc_revenue.png", width=5.8)
body(doc, "Chart 4: International Digital Commerce (AIDC) — Revenue & YoY Growth Trend", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4-5 — KEY METRICS & PROFITABILITY ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "KEY METRICS & PROFITABILITY ANALYSIS", level=1, color=ALI_NAVY, size=13)

add_heading(doc, "The Investment Cycle — Deliberate Profitability Sacrifice", level=2, color=ALI_ORANGE, size=11)
body(doc, "The defining feature of Alibaba's Q3 FY2026 results is not what it missed, but why. All profitability metrics fell sharply — not due to structural erosion, but intentional reinvestment:")
bullets_invest = [
    "Adjusted EBITA: RMB 23.4B (–57% YoY) — compressed by AI infrastructure capex, quick commerce logistics buildout, and enhanced consumer experience spending.",
    "Non-GAAP EPS/ADS: RMB 7.09 ($1.01) vs consensus ~¥9.00 — missed by ~21%; below the prior quarter's ¥4.36 though that was the cycle trough.",
    "Free Cash Flow: RMB 11.3B (–71% YoY) — reflecting capex intensity; management has not disclosed specific capex figures.",
    "GAAP Net Income: RMB 15.6B (–66% YoY) — includes non-cash charges and investment mark-to-market.",
    "Gross Margin: ~39.7% — modestly lower YoY, driven by quick commerce logistics costs and cloud infrastructure buildout.",
]
for b in bullets_invest:
    bullet(doc, b, size=9.5)

add_img(doc, "baba_chart4_adj_ebita.png", width=5.8)
body(doc, "Chart 5: Adjusted EBITA vs Revenue YoY Growth — Investment Cycle Analysis", size=8, italic=True)
doc.add_paragraph()

add_img(doc, "baba_chart5_eps.png", width=5.8)
body(doc, "Chart 6: Non-GAAP EPS per ADS (RMB) — Q4 FY25 to Q3 FY26", size=8, italic=True)
doc.add_paragraph()

# ── Profitability metrics table ────────────────────────────────────────────────
add_heading(doc, "Profitability Metrics — Last 4 Quarters", level=2, color=ALI_ORANGE, size=11)
prof_hdrs = ["Metric", "Q4 FY25", "Q1 FY26", "Q2 FY26", "Q3 FY26", "Q3 FY26 YoY"]
prof_rows = [
    ["Total Revenue (¥B)", "236.5", "247.7", "247.8", "284.8", "+2% (rpt'd)"],
    ["Adj. EBITA (¥B)", "46.4", "38.8", "9.1", "23.4", "–57%"],
    ["Adj. EBITA Margin", "~19.6%", "~15.7%", "~3.7%", "~8.2%", "–7pp YoY"],
    ["Non-GAAP EPS/ADS (¥)", "12.52", "14.75", "4.36", "7.09", "~–55%"],
    ["Free Cash Flow (¥B)", "49.0", "41.6", "22.5", "11.3", "–71%"],
    ["Gross Margin (%)", "~42.5%", "~41.9%", "~40.2%", "~39.7%", "–80bps"],
]
make_table(doc, prof_hdrs, prof_rows, hdr_bg=ALI_NAVY, hdr_fg=WHITE,
           alt_bg=hex_to_rgb("F5F7FA"),
           col_widths=[4.0, 2.0, 2.0, 2.0, 2.0, 2.5],
           hdr_size=8, row_size=8)
doc.add_paragraph()

add_img(doc, "baba_chart10_fcf_ebita.png", width=5.8)
body(doc, "Chart 7: Adj. EBITA vs. Free Cash Flow — Investment Cycle Deep-Dive", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 5-6 — USER METRICS & CONSUMER STRATEGY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "USER METRICS & CONSUMER STRATEGY", level=1, color=ALI_NAVY, size=13)

add_heading(doc, "88VIP — High-Value Consumer Engine at 59M Members", level=2, color=ALI_ORANGE, size=11)
body(doc, "Alibaba's premium 88VIP membership program surpassed 59 million members in Q3 FY2026, recording double-digit YoY growth for the quarter. 88VIP members represent the highest-spending cohort within Taobao and Tmall — their ARPU is significantly above average, and they tend to anchor cross-category purchasing behavior across Alibaba's ecosystem.")

body(doc, "The continued expansion of 88VIP validates Alibaba's premiumization strategy in China. Despite a challenging consumer environment in China, the company has successfully grown its high-quality user base, suggesting that its enhanced product selection, AI-driven personalization, and loyalty benefits are resonating with the upper-middle consumer tier.")

add_img(doc, "baba_chart9_88vip_members.png", width=5.8)
body(doc, "Chart 8: 88VIP Premium Members — Quarterly Growth (Millions)", size=8, italic=True)
doc.add_paragraph()

add_heading(doc, "Qwen AI App — 300M Monthly Active Users", level=2, color=ALI_ORANGE, size=11)
body(doc, "Alibaba's Qwen AI consumer app surpassed 300 million monthly active users as of Q3 FY2026 — a landmark milestone that positions it among the top-tier global AI consumer applications alongside OpenAI's ChatGPT and Google's Gemini. The Qwen model family is also gaining significant enterprise traction via Alibaba Cloud's Model-as-a-Service offerings.")

body(doc, "Qwen's 300M MAU milestone reinforces the thesis that Alibaba is building a durable AI moat. The flywheel between consumer AI usage data, model training improvements, and enterprise Cloud demand is increasingly visible. T-Head AI chips have now shipped cumulatively over 470,000 units through February 2026, reducing Alibaba's dependence on Nvidia GPUs and improving cloud AI margin economics over time.")

add_heading(doc, "Beat/Miss Visual Summary", level=2, color=ALI_ORANGE, size=11)
add_img(doc, "baba_chart6_beat_miss.png", width=5.8)
body(doc, "Chart 9: Q3 FY2026 Beat/Miss vs. Consensus Estimates — Key Metrics", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 7 — INVESTMENT THESIS UPDATE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "INVESTMENT THESIS UPDATE", level=1, color=ALI_NAVY, size=13)

add_heading(doc, "Thesis Intact — Investment Cycle Creates Compelling Entry Point", level=2, color=ALI_ORANGE, size=11)
body(doc, "We maintain our BUY rating on Alibaba (BABA) with a price target of $120.00 per ADS. The Q3 FY2026 results were mixed on the surface — meaningful profitability misses — but the underlying strategic picture remains compelling. Alibaba is executing a deliberate, management-sanctioned reinvestment cycle, prioritizing long-term competitive positioning in AI and quick commerce over short-term EPS delivery. This creates temporary but significant near-term multiple compression that, in our view, is an opportunity for long-duration investors.")

add_heading(doc, "What's Changed", level=2, color=ALI_ORANGE, size=11)
body(doc, "Positive Developments:")
bullets_pos = [
    "Cloud re-acceleration to +36% YoY is faster than our prior 30% forecast; AI product revenue sustaining triple-digit growth for 10 quarters is structural, not cyclical.",
    "88VIP at 59M members demonstrates Alibaba's high-value consumer relationships are deepening, not eroding — a key bull thesis datapoint.",
    "Qwen at 300M MAU creates a consumer AI moat that could monetize directly (subscriptions) and indirectly (data flywheel for Cloud model training).",
    "International commerce (AIDC) losses narrowing materially — Alibaba is shifting from loss-funding international expansion to a disciplined, quality-focused approach.",
    "T-Head AI chip shipments (470K cumulative) reduce GPU procurement risk and should improve cloud AI economics over time.",
]
for b in bullets_pos:
    bullet(doc, b, size=9.5)

doc.add_paragraph()
body(doc, "Areas of Caution:")
bullets_neg = [
    "Profitability compression (Adj. EBITA –57%, FCF –71%) is deeper and more protracted than initially expected; timeline for recovery into Q4 FY2026 remains unclear.",
    "No formal guidance provided — management's deliberate opacity on investment levels creates uncertainty for investors modeling near-term earnings.",
    "AIDC growth deceleration (from +32% to +3%) reflects execution challenges in Southeast Asia via Lazada; this remains a key watch item.",
    "China macro remains soft — domestic consumer sentiment has not fully recovered, and the quick commerce buildout is a bet on sustained engagement.",
    "Stock –5% post-earnings reinforces that the market is penalizing profitability misses regardless of strategic rationale.",
]
for b in bullets_neg:
    bullet(doc, b, size=9.5)

doc.add_paragraph()
add_heading(doc, "Key Catalysts to Watch", level=2, color=ALI_ORANGE, size=11)
cat_hdrs = ["Catalyst", "Timeline", "Bull Case", "Bear Case"]
cat_rows = [
    ["Cloud growth rate trajectory", "Q4 FY26 (May 2026)", "Maintains 35%+ — AI enterprise adoption accelerates", "+36% decelerates — AI spending cyclical pause"],
    ["Profitability recovery", "Q4 FY26–Q1 FY27", "Adj. EBITA rebounds to ¥35B+ as investments mature", "Investment cycle extends — EPS recovery delayed to FY2028"],
    ["Qwen monetization", "FY2027", "Cloud premium tier / Qwen subscriptions drive new revenue", "Qwen remains free; monetization path unclear"],
    ["AIDC turnaround (Lazada)", "H1 FY2027", "Lazada restructuring complete; AIDC growth re-accelerates to 15%+", "Lazada continues to drag; AIDC growth stays single-digit"],
    ["China macro recovery", "CY2026", "Consumer spending recovery lifts GMV and merchant advertising", "Prolonged China slowdown; quick commerce fails to scale profitably"],
    ["Regulatory / geopolitical", "Ongoing", "No incremental regulatory headwinds", "US-China tech restrictions impact cloud enterprise deals"],
]
make_table(doc, cat_hdrs, cat_rows, hdr_bg=ALI_NAVY, hdr_fg=WHITE,
           alt_bg=hex_to_rgb("F5F7FA"),
           col_widths=[3.5, 2.5, 4.0, 4.0],
           hdr_size=8.5, row_size=8.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 8-10 — VALUATION & ESTIMATES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "VALUATION & UPDATED ESTIMATES", level=1, color=ALI_NAVY, size=13)

add_heading(doc, "Valuation Framework — Sum-of-the-Parts", level=2, color=ALI_ORANGE, size=11)
body(doc, "Alibaba's multi-segment structure warrants a sum-of-the-parts (SOTP) valuation approach. We apply differentiated multiples to each segment reflecting their growth profiles and comparable peer valuations. Cloud Intelligence is valued on a revenue multiple consistent with high-growth cloud peers; China Commerce on earnings/EBITA; International Commerce at a discount given profitability drag.")

sotp_hdrs = ["Segment", "Metric", "FY27E Value", "Multiple", "Segment Value"]
sotp_rows = [
    ["Cloud Intelligence Group", "FY27E Rev: ~¥210B", "¥210B", "6.0x Rev", "~¥1,260B ($180B)"],
    ["China Core Commerce (T&T)", "FY27E Adj. EBITA: ~¥180B", "¥180B", "10x EBITA", "~¥1,800B ($257B)"],
    ["International Commerce (AIDC)", "FY27E Rev: ~¥155B", "¥155B", "2.0x Rev", "~¥310B ($44B)"],
    ["Quick Commerce (Taobao Instant)", "Growth option", "–", "Option value", "~¥70B ($10B)"],
    ["Cainiao / Media / Other", "FY27E Rev: ~¥280B", "¥280B", "0.5x Rev", "~¥140B ($20B)"],
    ["Investments / Net Cash", "Net cash + investments", "~¥500B", "1.0x", "~¥500B ($71B)"],
    ["Total Enterprise Value", "", "", "", "~¥4,080B ($583B)"],
    ["Implied Price per ADS", "", "", "", "~$115–$125 per ADS"],
]
make_table(doc, sotp_hdrs, sotp_rows, hdr_bg=ALI_NAVY, hdr_fg=WHITE,
           alt_bg=hex_to_rgb("F5F7FA"),
           col_widths=[3.8, 2.8, 2.8, 2.5, 3.0],
           hdr_size=8.5, row_size=8.5)
doc.add_paragraph()

body(doc, "Our SOTP analysis yields an implied ADS value of $115–$125. We set our 12-month price target at $120, representing ~40–50% upside from current levels (approximately $80 per ADS). The discount reflects: (1) continued near-term EPS compression risk; (2) China geopolitical uncertainty; and (3) execution risk on the investment cycle timeline.")

add_heading(doc, "Updated Financial Estimates", level=2, color=ALI_ORANGE, size=11)
est_hdrs = ["Metric", "FY2025A", "FY2026E", "FY2027E", "FY28E"]
est_rows = [
    ["Total Revenue (¥B)", "¥941.2B", "¥1,025B", "¥1,140B", "¥1,265B"],
    ["Revenue Growth YoY", "+7%", "+9%", "+11%", "+11%"],
    ["Cloud Revenue (¥B)", "¥117.3B", "¥159.1B", "¥211B", "¥265B"],
    ["Cloud YoY Growth", "+20%", "+36%", "+33%", "+26%"],
    ["Adj. EBITA (¥B)", "~¥175B", "~¥90B", "~¥155B", "~¥200B"],
    ["Adj. EBITA Margin", "~18.6%", "~8.8%", "~13.6%", "~15.8%"],
    ["Non-GAAP EPS/ADS (¥)", "~¥50.0", "~¥28.0", "~¥52.0", "~¥67.0"],
    ["Non-GAAP P/E (at $80/ADS)", "~11x", "~20x", "~11x", "~8.5x"],
]
make_table(doc, est_hdrs, est_rows, hdr_bg=ALI_NAVY, hdr_fg=WHITE,
           alt_bg=hex_to_rgb("F5F7FA"),
           col_widths=[4.0, 2.5, 2.5, 2.5, 2.5],
           hdr_size=8.5, row_size=8.5)
doc.add_paragraph()

add_heading(doc, "Gross Margin & Cloud Growth Trends", level=2, color=ALI_ORANGE, size=11)
add_img(doc, "baba_chart7_gross_margin.png", width=5.8)
body(doc, "Chart 10: Gross Margin Trend (%) — Q4 FY24 to Q3 FY26", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 10 — RISKS & SOURCES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "KEY RISKS", level=1, color=ALI_NAVY, size=13)
risk_hdrs = ["Risk Category", "Description", "Severity"]
risk_rows = [
    ["Investment Cycle Duration", "Management may continue heavy spending longer than market expects, delaying EPS recovery beyond FY2027.", "HIGH"],
    ["China Macro / Consumer", "Soft Chinese consumer sentiment and real estate weakness could limit GMV recovery; quick commerce economics unproven at scale.", "MEDIUM-HIGH"],
    ["Geopolitical / ADR Risk", "US-listed ADR remains subject to PCAOB oversight risk; HKEX listing partially mitigates but US delistin remains a tail risk.", "MEDIUM"],
    ["Cloud Competition", "Huawei, Tencent, Baidu, and ByteDance all aggressively competing in China cloud; pricing pressure could slow cloud margin recovery.", "MEDIUM"],
    ["International Commerce", "Lazada SE Asia restructuring remains incomplete; execution of AIDC quality-over-quantity pivot could depress reported growth for multiple quarters.", "MEDIUM"],
    ["Regulatory", "China tech regulation remains active; platform economy rules could impose new compliance costs or restrict certain business practices.", "LOW-MEDIUM"],
    ["Qwen Monetization", "Qwen at 300M MAU is impressive but not yet meaningfully monetized; model capability races could erode Qwen's competitive differentiation.", "LOW-MEDIUM"],
]
make_table(doc, risk_hdrs, risk_rows, hdr_bg=RED_ACC, hdr_fg=WHITE,
           alt_bg=hex_to_rgb("FFF5F5"),
           col_widths=[3.5, 8.5, 2.5],
           hdr_size=8.5, row_size=8.5)

doc.add_paragraph()

# ── Sources ────────────────────────────────────────────────────────────────────
add_heading(doc, "SOURCES & DISCLOSURES", level=1, color=ALI_NAVY, size=11)
body(doc, "All financial data sourced from Alibaba Group official filings and investor relations materials:", size=9)

sources = [
    ("Alibaba Q3 FY2026 Earnings Press Release (March 18, 2026)",
     "https://www.businesswire.com/news/home/20260318501558/en/Alibaba-Group-Announces-December-Quarter-2025-Results"),
    ("CNBC: Alibaba revenue misses estimates in December quarter as net income drops 66%",
     "https://www.cnbc.com/2026/03/19/alibaba-december-quarter-earnings-ai-investment.html"),
    ("Alibaba Q2 FY2026 Earnings Press Release (November 24, 2025)",
     "https://www.businesswire.com/news/home/20251124757764/en/Alibaba-Group-Announces-September-Quarter-2025-Results-and-Interim-Results-for-the-Six-Months-Ended-September-30-2025"),
    ("Alibaba Q4 FY2025 & Full Year FY2025 Results (May 14, 2025)",
     "https://www.businesswire.com/news/home/20250514856295/en/Alibaba-Group-Announces-March-Quarter-2025-and-Fiscal-Year-2025-Results"),
    ("Alibaba Investor Relations — BABA SEC Filings (EDGAR)",
     "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=BABA&type=20-F&dateb=&owner=include&count=10"),
    ("Alibaba AliViews: Eddie Wu on Q3 FY2026 Earnings",
     "https://www.alizila.com/aliviews-eddie-wu-on-alibabas-q3-earnings/"),
    ("Alpha Spread: Alibaba Misses Revenue Estimates as Net Income Plunges 66%",
     "https://www.alphaspread.com/market-news/earnings/alibaba-misses-revenue-estimates-as-net-income-plunges-66-ai-and-cloud-see-strong-growth"),
    ("AlphaStreet: Key metrics from Alibaba Group BABA Q3 2025 earnings",
     "https://news.alphastreet.com/key-metrics-from-alibaba-groups-baba-q3-2025-earnings-results/"),
]

for title, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run("• ")
    set_font(run, size=9)
    add_hyperlink(p, url, title)

doc.add_paragraph()
disc_text = ("ANALYST DISCLOSURES: This report is prepared for informational purposes only. "
             "The rating (BUY) and price target ($120.00) represent the analyst's current view based "
             "on publicly available information as of March 21, 2026. This is not investment advice. "
             "Market data sourced via yfinance. All financial figures in Chinese Renminbi (RMB/CNY) "
             "unless otherwise stated; USD conversion at approximately ¥7.0/USD.")
body(doc, disc_text, size=8, italic=True)

# ── Save ────────────────────────────────────────────────────────────────────────
outfile = OUT + "BABA_Q3_FY2026_Earnings_Update.docx"
doc.save(outfile)
print(f"\nEnglish report saved: {outfile}")
