"""
阿里巴巴集团 (BABA) — FY2026 Q3（2025自然年Q4）业绩更新
中文版 DOCX 报告生成
输出: output/BABA/BABA_Q3_FY2026_业绩更新报告_中文版.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/BABA/"
os.makedirs(OUT, exist_ok=True)

# ── yfinance 市场数据 ──────────────────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        import yfinance as yf
        info = yf.Ticker(ticker).fast_info
        return {
            "price":      f"${info.last_price:.2f}",
            "market_cap": f'约{info.market_cap/1e8:.0f}亿美元',
            "52w_high":   f"${info.year_high:.2f}",
            "52w_low":    f"${info.year_low:.2f}",
        }
    except Exception as e:
        print(f"[警告] yfinance 获取失败: {e} — 使用 N/A")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("BABA")
print(f"市场数据: {mkt}")

# ── 颜色工具 ──────────────────────────────────────────────────────────────────
def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

ALI_ORANGE = hex_to_rgb("FF6A00")
ALI_NAVY   = hex_to_rgb("1F2D5C")
CLOUD_BLUE = hex_to_rgb("0070C0")
GREEN_ACC  = hex_to_rgb("00A651")
RED_ACC    = hex_to_rgb("E8192C")
GRAY_LIGHT = hex_to_rgb("D9D9D9")
WHITE      = hex_to_rgb("FFFFFF")
GRAY_TEXT  = hex_to_rgb("666666")

# ── 文档工具 ──────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

def set_cjk_font(run, cjk_name="宋体"):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cjk_name)

def set_font_cn(run, size=10, bold=False, color=None, italic=False,
                latin="Times New Roman", cjk="宋体"):
    run.font.name   = latin
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_cjk_font(run, cjk)

def add_heading_cn(doc, text, level=1, color=ALI_NAVY, size=14, cjk="黑体"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font_cn(run, size=size, bold=True, color=color, cjk=cjk)
    return p

def body_cn(doc, text, size=9.5, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font_cn(run, size=size, italic=italic)
    return p

def bullet_cn(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font_cn(run, size=size)
    return p

def set_cell_bg(cell, color_tuple):
    r, g, b = color_tuple
    hex_color = f"{r:02X}{g:02X}{b:02X}"
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_hyperlink_cn(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_run.append(new_t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_img(doc, fname, width=6.0):
    path = OUT + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        body_cn(doc, f"[图表未找到: {fname}]", italic=True)

def make_table_cn(doc, headers, rows,
                  hdr_bg=ALI_NAVY, hdr_fg=WHITE, alt_bg=None,
                  col_widths=None, hdr_size=8.5, row_size=8.5,
                  cjk="宋体"):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, hdr_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font_cn(run, size=hdr_size, bold=True, color=hdr_fg, cjk=cjk)

    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        bg  = alt_bg if (alt_bg and i % 2 == 1) else WHITE
        for j, cell_val in enumerate(row_data):
            cell = row.cells[j]
            if bg != WHITE:
                set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_val))
            set_font_cn(run, size=row_size, cjk=cjk)

    if col_widths:
        for j, w in enumerate(col_widths):
            set_col_width(table, j, w)

    return table

# ═══════════════════════════════════════════════════════════════════════════════
# 第一页 — 封面与业绩摘要
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("阿里巴巴集团控股有限公司（纽交所：BABA | 港交所：9988）")
set_font_cn(run, size=15, bold=True, color=ALI_NAVY, cjk="黑体")

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("FY2026 Q3（2025自然年Q4）业绩更新报告")
set_font_cn(run2, size=13, bold=True, color=ALI_ORANGE, cjk="黑体")

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("报告期间：2025年10月–12月 | 发布日期：2026年3月18日 | 报告日期：2026年3月21日")
set_font_cn(run3, size=9.5, italic=True, color=GRAY_TEXT)

doc.add_paragraph()

# 评级表
make_table_cn(doc,
    headers=["评级", "目标价", "当前股价", "市值", "52周区间"],
    rows=[["买入（BUY）", "120.00美元", mkt["price"], mkt["market_cap"],
           f'{mkt["52w_low"]} – {mkt["52w_high"]}']],
    hdr_bg=ALI_NAVY, hdr_fg=WHITE,
    col_widths=[2.5, 2.5, 2.5, 3.0, 3.5],
    hdr_size=9, row_size=9,
)
doc.add_paragraph()

add_heading_cn(doc, "核心要点", size=11)
bullets = [
    "营收略低于预期：FY2026 Q3总营收人民币2848亿元（约401亿美元），同比+2%（剔除资产处置后有机增速+9%），低于市场预期约2%（市场预期约2907亿元）。",
    "云业务强劲超预期：云智能集团营收同比大增+36%至人民币433亿元，超出市场预期约8%；AI产品营收已连续10个季度实现三位数增长。",
    "利润承压（主动投资导致）：经调整EBITA同比下降57%至人民币234亿元；Non-GAAP每ADS摊薄EPS为人民币7.09元（约1.01美元），显著低于市场预期约人民币9.00元。",
    "即时零售高速增长：淘宝闪购（即时零售）营收同比+56%，为国内增速最快业务板块，但持续加大物流投入拖累利润。",
    "用户生态健康：88VIP超级会员突破5900万，同比双位数增长；通义千问（Qwen）AI应用月活用户超3亿，验证AI消费端战略。",
    "主动投资周期：管理层重申将持续加大AI基础设施、大模型及即时零售物流的投入，利润短期承压可能延续至FY2026 Q4。",
    "未提供量化业绩指引：公司未披露下季度营收或利润数字指引，但定性表态云业务增长势头持续。",
    "股价反应：结果发布后股价盘前下跌约5%，反映市场对利润超预期偏差的负面情绪；中长期AI/云成长逻辑不变。",
]
for b in bullets:
    bullet_cn(doc, b)

doc.add_paragraph()

add_heading_cn(doc, "业绩快照 — FY2026 Q3 vs 市场预期", size=11)
snap_hdrs = ["指标", "实际值", "市场预期", "超/低预期", "同比变化"]
snap_rows = [
    ["总营收（亿元）", "2848亿", "约2907亿", "低预期 –2.0%", "+2%（有机+9%）"],
    ["总营收（亿美元）", "约407亿", "约415亿", "低预期 –2.0%", "+2%"],
    ["经调整EBITA（亿元）", "234亿", "约300亿", "低预期 –22%", "–57%"],
    ["Non-GAAP每ADS EPS（人民币）", "7.09元", "约9.00元", "低预期 –21%", "大幅下降"],
    ["Non-GAAP每ADS EPS（美元）", "$1.01", "约$1.29", "低预期 –22%", "大幅下降"],
    ["云智能集团营收（亿元）", "433亿", "约400亿", "超预期 +8.3%", "+36%"],
    ["自由现金流（亿元）", "113亿", "约200亿", "低预期 –44%", "–71%"],
    ["毛利率", "约39.7%", "约40.5%", "低预期 –80bps", "约–70bps"],
]
make_table_cn(doc, snap_hdrs, snap_rows,
              hdr_bg=ALI_NAVY, hdr_fg=WHITE,
              alt_bg=hex_to_rgb("F5F7FA"),
              col_widths=[3.8, 2.5, 2.5, 2.8, 2.9],
              hdr_size=8.5, row_size=8.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第二至三页 — 详细业绩分析
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_cn(doc, "详细业绩分析", size=13)

add_heading_cn(doc, "营收——资产处置拖累报告增速，有机增长健康", size=11, color=ALI_ORANGE)
body_cn(doc, "阿里巴巴FY2026 Q3（2025年10月至12月）总营收为人民币2848亿元（约407亿美元），报告口径同比增长+2%。剔除已处置资产影响（包括大润发超市及银泰百货），有机/同口径营收增速约+9%，远高于报告增速。相较于市场预期约2907亿元，实际营收低约2%，主要原因为国际数字商业板块（AIDC）增速放缓，以及已处置零售资产的结构性拖累。")

body_cn(doc, "在保留业务口径下，各板块均实现有机增长，其中云智能集团表现最为亮眼。淘宝闪购（即时零售）同比+56%，为国内最快增长板块。中国批发（1688.com）同比+5%至人民币69亿元，受中小企业需求驱动。")

add_img(doc, "baba_chart1_quarterly_revenue.png", width=5.8)
body_cn(doc, "图1：阿里巴巴集团季度营收趋势（人民币十亿元）", size=8, italic=True)
doc.add_paragraph()

add_heading_cn(doc, "分板块营收——FY2026 Q3", size=11, color=ALI_ORANGE)
seg_hdrs = ["业务板块", "FY2026 Q3营收", "同比增速", "核心亮点"]
seg_rows = [
    ["云智能集团", "433亿元（约62亿美元）", "+36%", "AI产品连续10季三位数增长；云市场份额达36%"],
    ["国际数字商业（AIDC）", "324亿元（约46亿美元）", "+3%", "由+32%大幅放缓；来赞达拖累；亏损明显收窄"],
    ["淘宝闪购（即时零售）", "208亿元（约30亿美元）", "+56%", "增速最快；物流重投资拖累利润"],
    ["中国批发（1688.com）", "69亿元（约10亿美元）", "+5%", "中小企业B2B平台；稳健增长"],
    ["菜鸟、娱乐及其他", "673亿元（约96亿美元）", "–25%", "大润发/银泰处置致报告下降；剔除后持平"],
    ["合计", "2848亿元（约407亿美元）", "+2%（有机+9%）", "有机增长健康；处置资产拖累报告增速"],
]
make_table_cn(doc, seg_hdrs, seg_rows,
              hdr_bg=ALI_NAVY, hdr_fg=WHITE,
              alt_bg=hex_to_rgb("F5F7FA"),
              col_widths=[3.5, 3.2, 2.5, 5.3],
              hdr_size=8.5, row_size=8.5)
doc.add_paragraph()

add_img(doc, "baba_chart3_segment_revenue.png", width=5.8)
body_cn(doc, "图2：FY2026 Q3分业务板块营收构成（人民币十亿元）", size=8, italic=True)

doc.add_page_break()

add_heading_cn(doc, "云智能集团——核心增长引擎，同比+36%", size=11, color=CLOUD_BLUE)
body_cn(doc, "云智能集团FY2026 Q3实现营收人民币433亿元（约62亿美元），同比增速+36%，显著超出市场预期约400亿元（超预期约+8%）。自FY2024年仅+3%的增速，到当前+36%的加速，背后是AI驱动的企业级云需求的强劲崛起。AI产品营收连续10个季度实现三位数同比增长，阿里云在中国的市场份额上升至约36%，连续三个季度保持份额提升。")

body_cn(doc, "通义千问（Qwen）大模型家族已成为中国最具竞争力的前沿模型之一，其消费端应用月活用户突破3亿，平头哥AI芯片累计出货量截至2026年2月超过47万颗，有助于降低对英伟达GPU的依赖并改善云AI毛利经济模型。")

add_img(doc, "baba_chart2_cloud_revenue.png", width=5.8)
body_cn(doc, "图3：云智能集团季度营收及同比增速", size=8, italic=True)
doc.add_paragraph()

add_heading_cn(doc, "国际数字商业（AIDC）——结构性重塑中", size=11, color=ALI_ORANGE)
body_cn(doc, "AIDC营收人民币324亿元（约46亿美元），同比仅+3%，较去年同期+32%大幅放缓。放缓原因：（1）来赞达（Lazada，东南亚）持续进行业务重组，主动收缩亏损性GMV；（2）速卖通（AliExpress）正从规模导向转向利润导向模式。积极方面是AIDC亏损大幅收窄，体现管理层国际业务精细化运营的成效。")

add_img(doc, "baba_chart8_aidc_revenue.png", width=5.8)
body_cn(doc, "图4：国际数字商业（AIDC）季度营收及同比增速", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第四至五页 — 利润与关键指标分析
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_cn(doc, "利润与关键指标分析", size=13)

add_heading_cn(doc, "主动投资周期——以短期利润换长期竞争力", size=11, color=ALI_ORANGE)
body_cn(doc, 'FY2026 Q3业绩的核心特征，并非在于"miss什么"，而在于"为何而miss"。所有利润指标均大幅下滑，但这是管理层主动战略选择，而非经营结构恶化：')
bullets_inv = [
    "经调整EBITA：人民币234亿元（同比–57%），受AI基础设施资本开支、即时零售物流建设及用户体验优化投入压缩。",
    "Non-GAAP每ADS摊薄EPS：人民币7.09元（约1.01美元），低于市场预期约9.00元，偏差约–21%；较Q2 FY26低点（4.36元）有所回升。",
    "自由现金流：人民币113亿元（同比–71%），反映密集的资本开支节奏；公司尚未披露具体年度资本支出数字。",
    "GAAP净利润：人民币156亿元（同比–66%），含非现金项目及投资公允价值变动影响。",
    "毛利率：约39.7%，同比小幅收窄，主要受即时零售物流成本及云基础设施建设期摊销影响。",
]
for b in bullets_inv:
    bullet_cn(doc, b)

add_img(doc, "baba_chart4_adj_ebita.png", width=5.8)
body_cn(doc, "图5：经调整EBITA与营收同比增速——投资周期分析", size=8, italic=True)
doc.add_paragraph()

add_img(doc, "baba_chart5_eps.png", width=5.8)
body_cn(doc, "图6：Non-GAAP每ADS EPS（人民币）——FY2025 Q4至FY2026 Q3", size=8, italic=True)
doc.add_paragraph()

add_heading_cn(doc, "利润指标——近四季度", size=11, color=ALI_ORANGE)
prof_hdrs = ["指标", "FY25 Q4", "FY26 Q1", "FY26 Q2", "FY26 Q3", "FY26 Q3同比"]
prof_rows = [
    ["总营收（亿元）", "2365亿", "2477亿", "2478亿", "2848亿", "+2%（有机+9%）"],
    ["经调整EBITA（亿元）", "464亿", "388亿", "91亿", "234亿", "–57%"],
    ["经调整EBITA利润率", "约19.6%", "约15.7%", "约3.7%", "约8.2%", "约–7pp"],
    ["Non-GAAP每ADS EPS（人民币）", "12.52元", "14.75元", "4.36元", "7.09元", "约–55%"],
    ["自由现金流（亿元）", "490亿", "416亿", "225亿", "113亿", "–71%"],
    ["毛利率（%）", "约42.5%", "约41.9%", "约40.2%", "约39.7%", "约–80bps"],
]
make_table_cn(doc, prof_hdrs, prof_rows,
              hdr_bg=ALI_NAVY, hdr_fg=WHITE,
              alt_bg=hex_to_rgb("F5F7FA"),
              col_widths=[4.0, 2.0, 2.0, 2.0, 2.0, 2.5],
              hdr_size=8, row_size=8)
doc.add_paragraph()

add_img(doc, "baba_chart10_fcf_ebita.png", width=5.8)
body_cn(doc, "图7：经调整EBITA与自由现金流对比——投资周期深度剖析", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第五至六页 — 用户指标与超/低预期可视化
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_cn(doc, "用户指标与消费端战略", size=13)

add_heading_cn(doc, "88VIP超级会员突破5900万", size=11, color=ALI_ORANGE)
body_cn(doc, "88VIP超级会员在FY2026 Q3突破5900万，同比实现双位数增长。88VIP会员是淘宝天猫消费者中ARPU（客单价）最高的群体，其跨品类购买行为强劲，是驱动GMV和品牌广告收入的核心用户群体。")

body_cn(doc, "在中国消费市场相对低迷的背景下，88VIP高质量用户群持续扩张，表明阿里巴巴通过优质选品、AI个性化推荐及会员权益体系成功留住高价值消费者。这是我们投资逻辑的核心验证点之一。")

add_img(doc, "baba_chart9_88vip_members.png", width=5.8)
body_cn(doc, "图8：88VIP超级会员季度增长（百万人）", size=8, italic=True)
doc.add_paragraph()

add_heading_cn(doc, "通义千问（Qwen）——月活用户突破3亿", size=11, color=ALI_ORANGE)
body_cn(doc, 'FY2026 Q3，阿里巴巴旗下通义千问AI应用月活用户突破3亿，位居全球头部AI消费应用行列。Qwen大模型家族也通过阿里云"模型即服务（MaaS）"在企业端获得快速渗透，平头哥AI芯片累计出货量截至2026年2月已超47万颗。')

body_cn(doc, "Qwen 3亿MAU里程碑验证了阿里巴巴在消费端AI的竞争壁垒。消费数据飞轮、模型训练优化与云端企业需求之间的正向循环日益清晰。我们认为Qwen可通过订阅制（B端+C端）实现未来货币化。")

add_heading_cn(doc, "超/低预期可视化", size=11, color=ALI_ORANGE)
add_img(doc, "baba_chart6_beat_miss.png", width=5.8)
body_cn(doc, "图9：FY2026 Q3各核心指标相对市场预期超/低情况", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第七页 — 投资逻辑更新
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_cn(doc, "投资逻辑更新", size=13)

add_heading_cn(doc, "维持买入——投资周期创造优质买点", size=11, color=ALI_ORANGE)
body_cn(doc, '我们维持对阿里巴巴（BABA）的"买入"评级，目标价120美元/ADS。FY2026 Q3业绩表面看利润指标偏差较大，但深层战略逻辑依然坚实。公司正在执行一轮经管理层充分认可的主动再投资周期——以牺牲短期EPS为代价，换取AI和即时零售赛道的长期竞争优势。这在短期内造成明显的估值压缩，但在我们看来，恰恰是长线投资者的介入机会。')

add_heading_cn(doc, "新变化梳理", size=11, color=ALI_ORANGE)
body_cn(doc, "积极进展：")
bullets_pos = [
    "云智能集团加速至+36%，快于我们此前30%的预测；AI产品连续10季三位数增长具有结构性，非周期性。",
    "88VIP突破5900万会员，高价值用户关系持续深化，是我们核心看多论据之一。",
    "Qwen月活3亿，消费端AI护城河显现，未来可通过订阅直接货币化，也可通过云端间接受益（数据飞轮）。",
    "AIDC亏损大幅收窄，国际业务由亏损扩张转向精细化运营，质量改善明显。",
    "平头哥AI芯片累计出货47万颗，逐步降低对英伟达GPU的依赖，有望改善未来云AI毛利经济性。",
]
for b in bullets_pos:
    bullet_cn(doc, b)

doc.add_paragraph()
body_cn(doc, "需关注的风险点：")
bullets_neg = [
    "利润压缩（EBITA –57%，FCF –71%）深度和持续时间超出预期；FY2026 Q4利润回升时间表尚不明确。",
    "未提供量化指引，管理层对投资力度的主动淡化处理增加了短期盈利建模难度。",
    "AIDC增速由+32%降至+3%，来赞达东南亚运营重组仍在进行，是重要跟踪项。",
    "中国宏观消费端尚未完全复苏，即时零售大规模盈利能力待验证。",
    "股价业绩后下跌约5%，市场对利润miss惩罚明显，与战略逻辑形成短期背离。",
]
for b in bullets_neg:
    bullet_cn(doc, b)

doc.add_paragraph()
add_heading_cn(doc, "重要催化剂跟踪", size=11, color=ALI_ORANGE)
cat_hdrs = ["催化剂", "时间窗口", "乐观情景", "悲观情景"]
cat_rows = [
    ["云业务增速持续性", "FY2026 Q4（2026年5月）", "维持35%+，AI企业化渗透加速", "+36%开始回落，AI支出进入周期性停顿"],
    ["利润率回升路径", "FY2026 Q4–FY2027 Q1", "EBITA回升至350亿+，投资成效显现", "投资周期延长，EPS复苏推迟至FY2028"],
    ["Qwen货币化进展", "FY2027", "云高端订阅+Qwen会员制，形成新营收来源", "Qwen保持免费，货币化路径不清晰"],
    ["AIDC/来赞达转型", "FY2027 H1", "来赞达重组完成，AIDC增速回升至15%+", "来赞达持续拖累，AIDC个位数增长延续"],
    ["中国消费复苏", "2026年全年", "消费复苏带动GMV提升，广告回暖", "中国经济持续低迷，即时零售盈利难"],
    ["监管/地缘政治", "持续关注", "无新增监管压力", "中美科技限制影响云端企业客户获取"],
]
make_table_cn(doc, cat_hdrs, cat_rows,
              hdr_bg=ALI_NAVY, hdr_fg=WHITE,
              alt_bg=hex_to_rgb("F5F7FA"),
              col_widths=[3.5, 2.8, 3.8, 4.4],
              hdr_size=8.5, row_size=8.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第八至十页 — 估值与财务预测
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_cn(doc, "估值与财务预测更新", size=13)

add_heading_cn(doc, "估值框架——分部加总法（SOTP）", size=11, color=ALI_ORANGE)
body_cn(doc, "鉴于阿里巴巴多板块业务结构，我们采用分部加总法（SOTP）进行估值，对各板块分别赋予与其增长特征及可比公司匹配的估值倍数。云智能集团按高增长云业务营收倍数估值；中国核心商业按EBITA倍数估值；国际商业因利润尚未转正而折价处理。")

sotp_hdrs = ["业务板块", "估值基础", "FY27E基础值", "估值倍数", "对应价值"]
sotp_rows = [
    ["云智能集团", "FY27E营收：约2100亿", "2100亿元", "6.0x营收", "约1.26万亿元（约1800亿美元）"],
    ["中国核心商业（淘宝天猫）", "FY27E调整EBITA：约1800亿", "1800亿元", "10x EBITA", "约1.80万亿元（约2570亿美元）"],
    ["国际商业（AIDC）", "FY27E营收：约1550亿", "1550亿元", "2.0x营收", "约3100亿元（约440亿美元）"],
    ["即时零售（淘宝闪购）", "成长期权", "——", "期权价值", "约700亿元（约100亿美元）"],
    ["菜鸟、娱乐及其他", "FY27E营收：约2800亿", "2800亿元", "0.5x营收", "约1400亿元（约200亿美元）"],
    ["净现金及投资", "账面净现金+长期股权投资", "约5000亿元", "1.0x", "约5000亿元（约710亿美元）"],
    ["企业价值合计", "", "", "", "约4.08万亿元（约5830亿美元）"],
    ["对应每ADS价值", "", "", "", "约115–125美元"],
]
make_table_cn(doc, sotp_hdrs, sotp_rows,
              hdr_bg=ALI_NAVY, hdr_fg=WHITE,
              alt_bg=hex_to_rgb("F5F7FA"),
              col_widths=[3.5, 3.0, 2.8, 2.5, 3.7],
              hdr_size=8.5, row_size=8.5)
doc.add_paragraph()

body_cn(doc, "SOTP分析显示每ADS内在价值约115–125美元，我们设定12个月目标价为120美元，较当前价格（约80美元）隐含约40–50%上涨空间。折扣因素包括：（1）近期EPS压缩风险持续；（2）中美地缘政治不确定性；（3）投资周期时间线的执行风险。")

add_heading_cn(doc, "财务预测更新", size=11, color=ALI_ORANGE)
est_hdrs = ["指标", "FY2025A", "FY2026E", "FY2027E", "FY2028E"]
est_rows = [
    ["总营收（亿元）", "9412亿", "约1.025万亿", "约1.14万亿", "约1.265万亿"],
    ["营收增速", "+7%", "+9%", "+11%", "+11%"],
    ["云智能营收（亿元）", "1173亿", "约1591亿", "约2110亿", "约2650亿"],
    ["云业务增速", "+20%", "+36%", "+33%", "+26%"],
    ["经调整EBITA（亿元）", "约1750亿", "约900亿", "约1550亿", "约2000亿"],
    ["经调整EBITA利润率", "约18.6%", "约8.8%", "约13.6%", "约15.8%"],
    ["Non-GAAP每ADS EPS（人民币）", "约50元", "约28元", "约52元", "约67元"],
    ["Non-GAAP市盈率（按$80/ADS）", "约11x", "约20x", "约11x", "约8.5x"],
]
make_table_cn(doc, est_hdrs, est_rows,
              hdr_bg=ALI_NAVY, hdr_fg=WHITE,
              alt_bg=hex_to_rgb("F5F7FA"),
              col_widths=[4.0, 2.5, 2.5, 2.5, 2.5],
              hdr_size=8.5, row_size=8.5)
doc.add_paragraph()

add_heading_cn(doc, "毛利率与云业务增速趋势", size=11, color=ALI_ORANGE)
add_img(doc, "baba_chart7_gross_margin.png", width=5.8)
body_cn(doc, "图10：毛利率趋势（%）——FY2024 Q4至FY2026 Q3", size=8, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第十页 — 风险提示与参考资料
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_cn(doc, "主要风险提示", size=13)
risk_hdrs = ["风险类别", "描述", "风险等级"]
risk_rows = [
    ["投资周期持续性", "管理层可能持续重投资超出市场预期，EPS恢复时间线延迟至FY2028以后。", "高"],
    ["中国宏观/消费端", "中国消费情绪疲软，即时零售大规模盈利能力尚待验证。", "中高"],
    ["ADR监管/退市风险", "美股ADR面临PCAOB审查风险；港股上市部分对冲，但尾部风险仍存。", "中"],
    ["云业务竞争", "华为、腾讯、百度、字节跳动均在中国云市场激烈竞争，价格战可能拖缓云利润回升。", "中"],
    ["国际商业（来赞达）", "来赞达东南亚重组尚未完成，AIDC增速可能持续低迷多个季度。", "中"],
    ["监管合规", "中国科技监管持续，平台经济规则演变可能带来新合规成本或业务限制。", "中低"],
    ["Qwen货币化", "Qwen 3亿MAU令人印象深刻，但商业化路径尚不清晰，竞争对手快速追赶可能侵蚀优势。", "中低"],
]
make_table_cn(doc, risk_hdrs, risk_rows,
              hdr_bg=RED_ACC, hdr_fg=WHITE,
              alt_bg=hex_to_rgb("FFF5F5"),
              col_widths=[3.5, 9.0, 2.0],
              hdr_size=8.5, row_size=8.5)
doc.add_paragraph()

# ── 参考资料 ───────────────────────────────────────────────────────────────────
add_heading_cn(doc, "参考资料与信息披露", size=11)
body_cn(doc, "所有财务数据来源于阿里巴巴集团官方公告及投资者关系材料：", size=9)

sources_cn = [
    ("阿里巴巴集团FY2026 Q3业绩新闻稿（2026年3月18日）",
     "https://www.businesswire.com/news/home/20260318501558/en/Alibaba-Group-Announces-December-Quarter-2025-Results"),
    ("CNBC：阿里巴巴十二月季度营收未达预期，净利润同比下降66%",
     "https://www.cnbc.com/2026/03/19/alibaba-december-quarter-earnings-ai-investment.html"),
    ("阿里巴巴集团FY2026 Q2业绩新闻稿（2025年11月24日）",
     "https://www.businesswire.com/news/home/20251124757764/en/Alibaba-Group-Announces-September-Quarter-2025-Results-and-Interim-Results-for-the-Six-Months-Ended-September-30-2025"),
    ("阿里巴巴集团FY2025 Q4及全年业绩新闻稿（2025年5月14日）",
     "https://www.businesswire.com/news/home/20250514856295/en/Alibaba-Group-Announces-March-Quarter-2025-and-Fiscal-Year-2025-Results"),
    ("阿里巴巴EDGAR SEC文件（20-F年报）",
     "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=BABA&type=20-F&dateb=&owner=include&count=10"),
    ("AliViews：吴泳铭就FY2026 Q3业绩发表评述",
     "https://www.alizila.com/aliviews-eddie-wu-on-alibabas-q3-earnings/"),
    ("Alpha Spread：阿里巴巴营收未达预期，净利润骤降66%",
     "https://www.alphaspread.com/market-news/earnings/alibaba-misses-revenue-estimates-as-net-income-plunges-66-ai-and-cloud-see-strong-growth"),
    ("AlphaStreet：阿里巴巴BABA Q3 2025业绩关键指标解读",
     "https://news.alphastreet.com/key-metrics-from-alibaba-groups-baba-q3-2025-earnings-results/"),
]

for title, url in sources_cn:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run("• ")
    set_font_cn(run, size=9)
    add_hyperlink_cn(p, url, title)

doc.add_paragraph()
disc = ("【分析师声明】本报告仅供参考，不构成投资建议。评级（买入）与目标价（120美元）均基于截至2026年3月21日"
        "的公开信息。市场数据通过yfinance动态获取。所有财务数据以人民币（RMB/CNY）为单位，"
        "除特别说明外；美元换算汇率约为7.0元/美元。")
body_cn(doc, disc, size=8, italic=True)

# ── 保存 ───────────────────────────────────────────────────────────────────────
outfile = OUT + "BABA_Q3_FY2026_业绩更新报告_中文版.docx"
doc.save(outfile)
print(f"\n中文报告已保存：{outfile}")
