"""
Energy Fuels Inc. (UUUU) — Q4 & FY2025 业绩更新报告（中文版）
机构股票研究报告生成器（DOCX）
运行：python build_uuuu_report_cn.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/UUUU"

C_NAVY   = RGBColor(0x1C, 0x1C, 0x4D)
C_BLUE   = RGBColor(0x00, 0x71, 0xE3)
C_GREEN  = RGBColor(0x1E, 0x8A, 0x44)
C_RED    = RGBColor(0xCC, 0x00, 0x00)
C_GREY   = RGBColor(0x8E, 0x8E, 0x93)
C_LTBLUE = RGBColor(0xE8, 0xF2, 0xFF)
C_AMBER  = RGBColor(0xE6, 0x7E, 0x22)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

BODY_FONT    = "宋体"
HEADING_FONT = "黑体"


# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)

def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.name = HEADING_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1C1C4D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text, bold=False, italic=False, color=None, size=10.5, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"【图表：{filename} 未找到 — 请先运行 build_uuuu_charts.py】", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
            run.font.name = BODY_FONT
        cp.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows, col_widths=None, alternate=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "1C1C4D")
        set_cell_border(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = C_WHITE
            run.font.size = Pt(9)
            run.font.name = HEADING_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F0F4FF" if (alternate and ri % 2 == 0) else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = BODY_FONT
                run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def page_break(doc):
    doc.add_page_break()

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
def build_report_cn():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── 封面 / 标题 ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("股票研究  |  铀矿及关键矿产")
    run.font.size = Pt(9)
    run.font.color.rgb = C_GREY
    run.font.bold = True
    run.font.name = HEADING_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
    p.paragraph_format.space_after = Pt(2)

    heading(doc, "Energy Fuels Inc.（股票代码：UUUU / EFR.TSX）", level=1)

    p = doc.add_paragraph()
    run = p.add_run("2025年第四季度及全年业绩更新  |  评级：买入（市场一致）  |  目标价：$27.25（H.C. Wainwright）")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = C_BLUE
    run.font.name = HEADING_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run("报告日期：2026年3月17日  |  业绩发布日：2026年2月26–27日  |  当前股价：约$18.67  |  市值：约$45.1亿")
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_GREY
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    p.paragraph_format.space_after = Pt(8)

    divider(doc)

    # ── 投资摘要 ────────────────────────────────────────────────────────────────
    heading(doc, "投资摘要", level=2)
    body(doc,
         "Energy Fuels（UUUU）于2026年2月26–27日发布了2025年第四季度及全年业绩。"
         "公司在运营层面实现突破——铀矿开采量达172万磅U3O8（超出指引上限约11%），"
         "精制铀产量达101.5万磅，并完成7亿美元可转换票据发行，"
         "将营运资金推升至约9.274亿美元。"
         "第四季度营收为2,710万美元，全年营收为6,590万美元"
         "（2024年为7,810万美元），下降主要因肯尼亚Kwale重矿砂（HMS）业务停产。"
         "全年净亏损扩大至8,610万美元（每股亏损0.38美元），"
         "反映公司在铀矿、稀土元素（REE）、重矿砂及医疗同位素多条业务线上的激进扩张投入。"
    )
    body(doc,
         "正面亮点：铀矿单位成本从$53/磅降至$43/磅；签署六份长期供应合同（履约至2032年）；"
         "镝氧化物（Dy2O3）中试生产通过韩国汽车制造商认证；"
         "二期REE可行性研究显示净现值（NPV）为19亿美元（总投资4.1亿美元）；"
         "管理层交接安排落实（Ross Bhappu于2026年4月15日接任CEO）。"
         "负面方面：第四季度每股亏损（$0.08）较市场一致预期（$0.07）差$0.01。",
         space_after=8
    )

    # 关键指标概览
    heading(doc, "关键指标一览", level=3)
    add_table(doc,
        headers=["指标", "2025年Q4", "2024年Q4", "2025年全年", "2024年全年"],
        rows=[
            ["总营收（百万美元）", "$27.1", "$39.9", "$65.9", "$78.1"],
            ["铀矿营收（百万美元）", "约$27.1", "约$12.6", "$48.2", "$37.9"],
            ["重矿砂营收（百万美元）", "—", "约$27.3", "$15.8", "$39.9"],
            ["净利润/（亏损）（百万美元）", "（约$20.9）", "（$28.8）", "（$86.1）", "（$47.8）"],
            ["每股收益（基本）", "（$0.08）", "（$0.19）", "（$0.38）", "（$0.28）"],
            ["每股收益 vs 一致预期", "差$0.01（未达预期）", "差$0.12（未达预期）", "—", "差$0.18（未达预期）"],
            ["营收 vs 一致预期", "超出$0.1M（超预期）", "—", "—", "—"],
            ["铀矿销售量（千磅）", "360", "约150", "650", "450"],
            ["加权平均实现价格（美元/磅）", "$74.93", "约$84.0", "$74.21", "$87.2"],
            ["营运资金（百万美元）", "—", "—", "$927.4", "$170.9"],
        ],
        col_widths=[2.3, 1.1, 1.1, 1.1, 1.1]
    )
    body(doc, "资料来源：Energy Fuels新闻稿（2026年2月26日；2025年12月29日）；Zacks；H.C. Wainwright研究报告。", italic=True, color=C_GREY, size=8.5)
    body(doc, "", space_after=10)

    add_image(doc, "uuuu_chart1_revenue.png", caption="图1：UUUU各业务板块季度营收（2024年Q1–2025年Q4）")

    # ── 第一节：2025年第四季度详细业绩 ─────────────────────────────────────────
    page_break(doc)
    heading(doc, "一、2025年第四季度详细业绩", level=2)

    heading(doc, "营收与盈利", level=3)
    body(doc,
         "2025年Q4总营收为2,710万美元，超出Zacks一致预期（2,700万美元）约0.38%（+10万美元）。"
         "相较于2024年Q4的3,990万美元，下降主要由于Kwale重矿砂矿山（2024年底完成回填封矿）贡献为零。"
         "铀浓缩物收入占本季度几乎全部营收。"
    )
    bullet(doc, "Q4 2025铀矿营收：约2,700万美元（36万磅 × $74.93/磅加权平均实现价格）")
    bullet(doc, "重矿砂营收：极少/零（Kwale HMS已停产）")
    bullet(doc, "五氧化二钒营收：零（持有90.5万磅V2O5库存，待市场好转后出售）")
    bullet(doc, "REE/NdPr氧化物营收：中试阶段，Q4 2025无商业化REE销售收入入账")
    body(doc, "", space_after=4)

    heading(doc, "每股收益与净利润", level=3)
    body(doc,
         "Q4 2025净亏损约2,090万美元，即每股亏损（$0.08），"
         "略差于市场一致预期（$0.07）——每股差$0.01。"
         "相较于2024年Q4的每股亏损（$0.19），盈利质量显著改善，"
         "反映铀矿单位经济效益的提升，尽管企业层面的可转换票据发行成本有所上升。"
    )
    bullet(doc, "Q4 2025 EPS：（$0.08）实际 vs（$0.07）预期  →  未达预期$0.01")
    bullet(doc, "2024年Q4 EPS：（$0.19）  |  2025年Q3 EPS：（$0.07）  |  2025年Q2 EPS：（$0.10）")
    body(doc, "", space_after=4)

    heading(doc, "铀矿业务——2025年第四季度", level=3)
    body(doc,
         "铀矿板块是Q4 2025业绩的核心驱动力。White Mesa磨矿厂于Q4 2025启动传统矿石处理流程"
         "（预计持续至2026年Q2），工厂展示了月均约25万磅的产能，"
         "其中2025年12月单月精制U3O8产量即达35万磅。"
    )
    add_table(doc,
        headers=["指标", "Q4 2025", "Q3 2025", "Q2 2025", "Q1 2025"],
        rows=[
            ["铀矿销售量（千磅）", "360", "240", "约150", "约50"],
            ["加权平均实现价格（美元/磅）", "$74.93", "约$69.6（估）", "约$79.0（估）", "约$80.0（估）"],
            ["铀矿营收（百万美元）", "约$27.0", "约$10.4", "约$4.2", "约$7.0"],
            ["销售成本/磅（近似值）", "约$43", "约$43", "约$47", "约$50"],
        ],
        col_widths=[2.5, 1.2, 1.2, 1.2, 1.2]
    )
    body(doc, "", space_after=6)

    add_image(doc, "uuuu_chart2_uranium_ops.png", caption="图2：铀矿销售量及每磅实现价格（2024年Q1–2025年Q4）")

    heading(doc, "五氧化二钒（V2O5）", level=3)
    body(doc,
         "截至2025年Q3，公司持有90.5万磅精制五氧化二钒库存。"
         "公司选择在市场价格达到管理层目标价之前不出售，"
         "2025年全年及2024年全年钒矿销售收入均为零。"
         "钒库存代表着对未来钒价上涨的实物期权，当前钒价约为$6–7/磅。"
    )

    heading(doc, "稀土元素（REE）——2025年第四季度里程碑", level=3)
    body(doc,
         "2025年Q4及全年均未确认商业化REE销售收入（分离稀土氧化物仍处中试阶段）。"
         "但公司取得以下重要里程碑："
    )
    bullet(doc, "截至2025年Q3，White Mesa磨矿厂共生产29千克纯度99.9%的镝（Dy）氧化物")
    bullet(doc, "Dy氧化物通过韩国主要汽车制造商认证（用于电动汽车驱动电机永磁体）")
    bullet(doc, "1.2吨NdPr氧化物成功制造约3.0吨稀土永磁体，驱动约1,500辆电动汽车")
    bullet(doc, "铽（Tb）氧化物首批量产目标为2026年Q1初")
    bullet(doc, "截至2025年Q3库存：37,000千克精制NdPr氧化物；9,000千克Sm+重稀土碳酸盐")
    body(doc, "", space_after=4)

    add_image(doc, "uuuu_chart3_net_income.png", caption="图3：季度净利润/（亏损）及每股收益（2024年Q1–2025年Q4）")
    add_image(doc, "uuuu_chart4_beat_miss.png", caption="图4：2025年各季度实际业绩 vs 一致预期")

    # ── 第二节：2025年全年业绩 ──────────────────────────────────────────────────
    page_break(doc)
    heading(doc, "二、2025年全年业绩", level=2)

    heading(doc, "利润表摘要", level=3)
    add_table(doc,
        headers=["项目", "2025年全年", "2024年全年", "变动幅度", "备注"],
        rows=[
            ["总营收（百万美元）", "$65.9", "$78.1", "–15.6%", "Kwale HMS停产影响"],
            ["  铀浓缩物", "$48.2", "$37.9", "+27.2%", "65万磅@$74.21/磅"],
            ["  重矿砂", "$15.8", "$39.9", "–60.4%", "Kwale末年/Donald待产"],
            ["  五氧化二钒", "$0", "$0", "—", "90.5万磅存库未售"],
            ["  稀土/其他", "约$1.9", "约$0.3", "—", "中试规模/极少"],
            ["毛利润（百万美元）", "约$20.3", "约$28.2", "–28.0%", "铀矿毛利率约42%"],
            ["经营亏损（百万美元）", "（约$58）", "（约$19）", "—", "含管理费用+勘探+折旧"],
            ["净亏损（百万美元）", "（$86.1）", "（$47.8）", "–80.1%", "含可转换票据发行费用"],
            ["每股净亏损", "（$0.38）", "（$0.28）", "–35.7%", "摊薄股数约2.26亿股"],
        ],
        col_widths=[2.3, 1.0, 1.0, 1.0, 2.2]
    )
    body(doc, "资料来源：Energy Fuels新闻稿，2026年2月26日（美通社）。", italic=True, color=C_GREY, size=8.5)
    body(doc, "", space_after=6)

    heading(doc, "铀矿板块——2025年全年", level=3)
    add_table(doc,
        headers=["指标", "2025年全年", "2024年全年"],
        rows=[
            ["铀矿销售量（千磅）", "650", "450"],
            ["加权平均实现价格（美元/磅）", "$74.21", "$87.2"],
            ["  现货销售", "35万磅@$76.90/磅", "25万磅@$91.51/磅"],
            ["  长期合同销售", "30万磅@$71.06/磅", "20万磅@$75.13/磅"],
            ["铀矿营收（百万美元）", "$48.2", "$37.9"],
            ["销售成本/磅（加权平均）", "约$43", "约$53"],
            ["铀矿毛利率", "约42%", "约56%"],
            ["开采矿量（百万磅U3O8）", "1.72", "1.10"],
            ["精制/成品产量（百万磅）", "1.015", "0.55"],
            ["长期供应合同数量", "6份（履约至2032年）", "4份"],
        ],
        col_widths=[3.0, 2.0, 2.0]
    )
    body(doc, "", space_after=4)

    heading(doc, "铀矿库存（截至2025年12月31日）", level=3)
    add_table(doc,
        headers=["类别", "磅数（U3O8）", "备注"],
        rows=[
            ["开采矿石/矿化材料", "1,240,000磅", "地下及地表矿堆"],
            ["在制品", "130,000磅", "位于White Mesa磨矿厂"],
            ["成品U3O8", "810,000磅", "可交付/出售"],
            ["合计", "约2,180,000磅", "约为2024年底约98万磅的2.2倍"],
        ],
        col_widths=[2.8, 1.8, 2.9]
    )
    body(doc, "", space_after=8)

    add_image(doc, "uuuu_chart5_annual_revenue.png", caption="图5：各业务板块年度营收（2023–2025年）")
    add_image(doc, "uuuu_chart6_uranium_production.png", caption="图6：铀矿开采、精制及库存（磅数U3O8）")

    heading(doc, "资产负债表（2025年12月31日）", level=3)
    add_table(doc,
        headers=["项目", "2025年12月31日", "2024年12月31日"],
        rows=[
            ["营运资金", "$927.4M", "$170.9M"],
            ["现金及现金等价物", "$64.7M", "$38.6M"],
            ["有价证券", "$797.1M", "$80.9M"],
            ["现金+有价证券合计", "约$861.8M", "约$119.5M"],
            ["总资产", "$14.1亿", "约$4.7亿"],
            ["总负债", "$7.293亿", "约$0（无债务）"],
            ["可转换优先票据", "$7亿（票息0.75%）", "—"],
            ["净股权", "约$6.8亿", "约$4.7亿"],
        ],
        col_widths=[2.8, 1.8, 1.8]
    )
    body(doc,
         "7亿美元可转换票据（2025年10月，票息0.75%，超额认购7倍以上，2031年到期）"
         "彻底改变了公司的资产负债结构。初始转换价格：$20.34/股；"
         "附加上限期权（capped call）后，有效稀释触发价升至$30.70/股（相当于发行价的100%溢价）。",
         space_after=6
    )

    add_image(doc, "uuuu_chart7_cost_structure.png", caption="图7：铀矿成本结构 vs 实现价格")
    add_image(doc, "uuuu_chart8_balance_sheet.png", caption="图8：资产负债表关键指标")

    # ── 第三节：铀矿及稀土市场背景 ────────────────────────────────────────────
    page_break(doc)
    heading(doc, "三、铀矿及稀土市场背景", level=2)

    heading(doc, "2025年铀矿现货价格走势", level=3)
    body(doc,
         "2025年铀矿现货市场整体偏弱、区间震荡，全年交易价格在$63/磅（4月低点）至$82–83/磅"
         "（9月高点）之间，相较2024年1月约$106/磅的峰值大幅回调。"
         "2025年价格下行的主要驱动因素包括：（1）2023–2024年累积多头头寸平仓；"
         "（2）铀浓缩低料位（underfeeding）和公用事业库存去化带来的二级供应；"
         "（3）特朗普政府政策不确定性影响市场情绪。"
    )
    body(doc,
         "与此形成对比的是，长期合同价格从约$80/磅稳步升至约$86/磅（2025年全年），"
         "反映公用事业对多年期燃料供应安全的强劲需求。"
         "这一「现货弱、长期强」的分化格局实际上有利于Energy Fuels——"
         "公司通过长期合同锁定销售（2025年全年30万磅@$71.06/磅），同时积累库存以备未来交付。",
         space_after=6
    )
    add_table(doc,
        headers=["时期", "U3O8现货价（美元/磅）", "长期合同价（美元/磅）", "市场评述"],
        rows=[
            ["2024年1月", "约$106", "约$68", "后Cameco/Kazatom供应忧虑推至峰值"],
            ["2024年Q2", "约$85–90", "约$73", "逐步回调"],
            ["2024年Q4", "约$78–82", "约$78", "现货价与长期价趋于收敛"],
            ["2025年Q1（4月低点）", "约$63–65", "约$80", "18个月低位；长期价>现货价"],
            ["2025年Q3（9月高点）", "约$82–83", "约$84", "部分反弹"],
            ["2025年Q4", "约$70–76", "约$85", "年末走软"],
            ["2026年1月", "约$98–100", "约$87", "快速反弹；接近近两年新高"],
        ],
        col_widths=[1.8, 1.7, 1.7, 2.3]
    )
    body(doc, "资料来源：UxC、TradeTech、Sprott铀矿2026年展望、ANS核能新闻。", italic=True, color=C_GREY, size=8.5)
    body(doc, "", space_after=6)

    add_image(doc, "uuuu_chart11_uranium_market.png", caption="图11：铀矿现货价格 vs 长期合同价格走势（2024–2026年1月）")

    heading(doc, "铀矿市场结构性展望", level=3)
    bullet(doc, "全球铀矿供应缺口约1.3亿磅/年；新反应堆建设及数据中心电力需求推动需求持续增长")
    bullet(doc, "2024年全球新增6台核反应堆并网（主要来自中国）；全球在建反应堆逾60台")
    bullet(doc, "美国《2024年禁止俄罗斯铀进口法案》在结构上有利于美国本土铀矿生产商")
    bullet(doc, "2026年1月铀期货升至约$100/磅——近两年新高")
    bullet(doc, "Energy Fuels是美国仅有的3家活跃铀矿生产商之一（另外两家为Cameco-Cigar和enCore Energy）")
    body(doc, "", space_after=6)

    heading(doc, "稀土（REE）市场背景", level=3)
    body(doc,
         "西方稀土供应链仍高度依赖中国——中国控制全球约85%的稀土加工产能。"
         "美国及盟友对本土稀土分离的需求因国防/电动汽车永磁体需求而加速增长。"
         "Energy Fuels的White Mesa磨矿厂是中国以外整个西半球唯一能够商业化生产"
         "分离重稀土氧化物（镝、铽）的设施。"
    )
    bullet(doc, "NdPr氧化物价格：约$50–60/千克（2025年）；Dy氧化物：约$300–360/千克；Tb氧化物：约$1,400–1,600/千克")
    bullet(doc, "美国能源部（DOE）和国防部（DOD）积极资助本土稀土供应链建设")
    bullet(doc, "Energy Fuels NdPr氧化物通过三星SDI类型永磁体制造商（韩国）资质认证")
    body(doc, "", space_after=8)

    # ── 第四节：2026年指引 ──────────────────────────────────────────────────────
    page_break(doc)
    heading(doc, "四、2026年业绩指引", level=2)

    heading(doc, "铀矿产量与销售计划", level=3)
    add_table(doc,
        headers=["指标", "2026年指引", "2025年实际", "变动"],
        rows=[
            ["开采矿量（百万磅U3O8）", "2.0–2.5", "1.72", "+16%至+45%"],
            ["精制/成品产量（百万磅）", "1.5–2.5", "1.015", "+48%至+146%"],
            ["铀矿销售量（百万磅）", "1.5–2.0", "0.65", "+131%至+208%"],
            ["成品库存加权平均成本（美元/磅）", "目标：30美元初段", "约$43", "降幅>25%"],
            ["Pinyon Plain矿山处理成本（美元/磅）", "$23–30", "约$43（混合）", "大幅下降"],
            ["  其中：开采/运输", "$10–14/磅", "—", "—"],
            ["  其中：磨矿", "$13–16/磅", "—", "—"],
            ["长期供应合同", "6份（履约至2032年）", "6份", "稳定"],
        ],
        col_widths=[2.7, 1.8, 1.5, 1.5]
    )
    body(doc,
         "管理层指引2026年铀矿销售量接近翻倍，支撑因素包括：（1）Pinyon Plain高品位矿石继续在White Mesa处理；"
         "（2）可能重启La Sal和Pandora矿山；（3）进入2026年时持有218万磅精制+在制铀矿库存。"
         "预期单位成本降至「30美元初段」，意味着在当前铀价（$80–100/磅）下毛利率可达55%以上。",
         space_after=6
    )

    heading(doc, "稀土元素（REE）2026年目标", level=3)
    bullet(doc, "一期（NdPr产自独居石）：一期B/C阶段全产能下，NdPr年产量约850–1,000吨（目标2027年达产）")
    bullet(doc, "重稀土商业化分离：2026年Q4在White Mesa完成Dy、Tb（及可能Sm）商业规模分离产能建设")
    bullet(doc, "铽氧化物首批千克：目标2026年Q1完成")
    bullet(doc, "二期可行性研究（2026年1月发布）：资本开支4.1亿美元；年产5,513吨NdPr+165吨Dy+48吨Tb；NPV 19亿美元；IRR 33%")
    bullet(doc, "二期许可申请：2026/2027年开始；投产目标2028年底至2029年初")
    body(doc, "", space_after=4)

    heading(doc, "资本支出与战略布局", level=3)
    bullet(doc, "2026年资本支出：未明确指引；重点为White Mesa稀土回路升级及矿山开发")
    bullet(doc, "Donald项目（澳大利亚，49% JV合资/Astron）：最终投资决定（FID）预计2026年Q1；总资本开支约5.2亿澳元；获Export Finance Australia 8,000万澳元债务融资意向函")
    bullet(doc, "Donald项目生产时间：2027年H2；为White Mesa提供重稀土原料")
    bullet(doc, "Vara Mada项目（马达加斯加）：税后NPV 18亿美元（10%折现率），矿山寿命约38年；与政府就投资协议谈判中")
    bullet(doc, "医疗同位素：镭-226（Ra-226）研发中试设施预计2026年投用；商业化生产目标2028年")
    bullet(doc, "管理层交接：Ross Bhappu（总裁→CEO）自2026年4月15日起生效；Mark Chalmers退休并保留2年顾问身份")
    body(doc, "", space_after=6)

    add_image(doc, "uuuu_chart9_ree_roadmap.png", caption="图9：稀土扩张各阶段NPV与EBITDA路线图")

    # ── 第五节：业绩发布电话会议要点 ────────────────────────────────────────────
    page_break(doc)
    heading(doc, "五、业绩发布电话会议要点（2026年2月27日）", level=2)

    heading(doc, "White Mesa磨矿厂——运营情况", level=3)
    body(doc,
         "White Mesa磨矿厂（位于犹他州Blanding）是Energy Fuels运营的核心，"
         "也是美国目前唯一在产的传统铀矿磨矿厂，"
         "同时是西半球唯一能商业化生产分离重稀土氧化物的设施（中国除外）。"
    )
    bullet(doc, "磨矿厂年产许可证：800万磅U3O8")
    bullet(doc, "传统矿石处理流程：2025年Q4启动，计划运行至2026年Q2")
    bullet(doc, "峰值产能：2025年12月单月精制U3O8产量达35万磅")
    bullet(doc, "月均产能：约25万磅")
    bullet(doc, "稀土中试回路：镝氧化物纯度99.9%；铽氧化物即将生产")
    bullet(doc, "成本改善：单位COGS从$53/磅（2024年）→$43/磅（2025年）→目标30美元初段（2026年）")
    body(doc, "", space_after=4)

    heading(doc, "Pinyon Plain矿山——亚利桑那州", level=3)
    body(doc,
         "Pinyon Plain是铀矿产量增长的主要驱动力，据管理层表述，其矿石品位居美国铀矿历史最高之列。"
    )
    bullet(doc, "2025年全年开采矿量：约153万磅U3O8（平均品位约1.62% eU3O8）")
    bullet(doc, "矿山寿命及资源量依然较大；持续进行开发钻探")
    bullet(doc, "预可研报告（PFS）更新中——可能改善项目经济指标并扩大资源量估算")
    bullet(doc, "2026年预期处理成本：$23–30/磅总成本（开采/运输$10–14/磅 + 磨矿$13–16/磅）")
    body(doc, "", space_after=4)

    heading(doc, "La Sal矿山群及其他矿山——犹他州/科罗拉多州", level=3)
    bullet(doc, "La Sal矿山群（犹他州）：2025年全年贡献约19万磅含U3O8开采量")
    bullet(doc, "Pandora矿山（科罗拉多州）：在产但产能较低")
    bullet(doc, "Whirlwind矿山（科罗拉多州）：备用/保养维护状态")
    bullet(doc, "管理层指引：2026年将根据200–250万磅目标的需要，视情况重启其他矿山")
    body(doc, "", space_after=4)

    heading(doc, "稀土分离设施——战略愿景", level=3)
    body(doc, "管理层在此次电话会议上就REE扩张路线图给出了迄今最为详尽的规划：")
    bullet(doc, "一期（现状）：从独居石中分离NdPr；一期B/C全规模下约850–1,000吨NdPr/年")
    bullet(doc, "一期商业规模Dy/Tb/Sm分离：目标2026年Q4在White Mesa完成投产")
    bullet(doc, "二期可行性研究（2026年1月发布）：资本开支4.1亿美元，40年项目周期；年产5,513吨NdPr+165吨Dy+48吨Tb")
    bullet(doc, "二期NPV：单独19亿美元；结合Vara Mada（马达加斯加）原料后37亿美元")
    bullet(doc, "二期+Vara Mada年度EBITDA：约7.65亿美元（前15年）")
    bullet(doc, "二期税后IRR：33%；每股NPV：$7.96（单独二期）至$15.26（二期+Vara Mada）")
    bullet(doc, "Donald项目（澳大利亚）：提供重稀土原料——FID计划2026年Q1，2027年H2投产")
    body(doc, "", space_after=4)

    heading(doc, "重大战略公告与并购", level=3)
    bullet(doc, "Australian Strategic Materials（ASM，ASX上市）收购：Energy Fuels同意收购ASM，新增下游稀土金属合金生产能力，从而构建西方世界唯一的完整关键矿产供应链")
    bullet(doc, "2025年Q4签署两份新增长期铀供应合同（现共持有6份，履约至2032年）")
    bullet(doc, "Export Finance Australia为Donald项目出具非约束性8,000万澳元债务融资意向函")
    bullet(doc, "与Vulcan Elements成立合资公司，共同推进美国稀土永磁体供应安全")
    bullet(doc, "管理层交接：Mark Chalmers退休，Ross Bhappu于2026年4月15日接任CEO")
    body(doc, "", space_after=8)

    # ── 第六节：分析师反应 ──────────────────────────────────────────────────────
    page_break(doc)
    heading(doc, "六、分析师反应与目标价（业绩发布后）", level=2)

    add_table(doc,
        headers=["分析机构", "评级", "目标价", "变动", "核心观点"],
        rows=[
            ["H.C. Wainwright", "买入", "$27.25", "上调（原$26.75）", "铀矿产量加速增长；Pinyon Plain预可研催化剂；REE中试里程碑降低一期风险"],
            ["市场一致（6–8家机构）", "强力买入", "均值$23.08（区间$15.50–$34.00）", "—", "7家买入 / 1家持有 / 0家卖出（截至2026年3月9日）"],
            ["隐含上涨空间（均值目标价）", "—", "约+24%（vs当前价）", "—", "基于$23.08均值目标价vs当前股价$18.67"],
            ["隐含上涨空间（HC W.目标价）", "—", "约+46%（vs当前价）", "—", "基于$27.25 vs $18.67当前股价"],
        ],
        col_widths=[1.8, 0.8, 1.7, 1.2, 2.8]
    )
    body(doc, "", space_after=4)

    body(doc,
         "尽管Q4 2025每股收益较预期差$0.01，但营收小幅超出预期，"
         "加之2026年指引强劲（铀矿销售量1.5–200万磅，相较2025年的65万磅增加逾倍），"
         "总体市场反应较为积极。二期REE可行性研究净现值19亿–37亿美元，"
         "被分析师视为当前市值（约45亿美元）下尚未得到充分定价的核心资产。"
         "ASM收购增添下游合金产能，被认为在战略上完善了西方关键矿产供应链，"
         "并可能吸引政府（DOE/DOD）签订采购协议及提供融资支持。",
         space_after=6
    )

    add_image(doc, "uuuu_chart10_price_targets.png", caption="图10：分析师目标价 vs 当前股价（2026年3月）")

    # ── 第七节：股价与市值 ───────────────────────────────────────────────────────
    heading(doc, "七、股价与市场数据", level=2)

    add_table(doc,
        headers=["指标", "数值"],
        rows=[
            ["当前股价（2026年3月16日）", "约$18.67"],
            ["市值", "约$45.1亿"],
            ["52周区间", "$3.20 – $27.90"],
            ["交易所", "美国证交所（UUUU）/ 多伦多证交所（EFR）"],
            ["流通股数（约）", "约2.42亿股"],
            ["企业价值（约）", "约$43.5亿（净现金约1.62亿美元，扣除可转换票据后）"],
            ["EV / 2025年营收", "约66倍（成长期溢价）"],
            ["EV / 2026年预期营收（市场一致）", "约17–20倍（按175万磅@$80/磅=约1.4亿美元营收估算）"],
            ["下次业绩发布日期", "约2026年5月7日（2026年Q1）"],
        ],
        col_widths=[3.0, 4.5]
    )
    body(doc, "", space_after=4)
    body(doc,
         "UUUU股价从2025年4月低点（约$3.20）大幅上涨近400%，至52周高位$27.90，"
         "主要驱动因素为7亿美元可转换票据发行、稀土里程碑突破及铀矿产量加速。"
         "自2026年1月高点以来，股价回调至约$18.67，"
         "主要因铀矿现货价格走软及大盘波动加大。"
         "当前约45亿美元市值远超当期盈利水平，"
         "但内嵌了REE二期+Vara Mada项目37亿美元NPV带来的巨大期权价值。",
         space_after=8
    )

    # ── 第八节：历史季度数据 ────────────────────────────────────────────────────
    page_break(doc)
    heading(doc, "八、历史季度数据——趋势分析", level=2)

    add_table(doc,
        headers=["季度", "营收（百万美元）", "净利润/（亏损）（百万美元）", "每股收益（美元）", "铀矿销量（千磅）", "铀价（美元/磅）", "铀矿营收（百万美元）"],
        rows=[
            ["2024年Q1", "$25.3", "+$3.6", "+$0.02", "300", "$84.4", "$25.3"],
            ["2024年Q2", "约$16.2", "约（$9.8）", "（$0.06）", "约0", "—", "约$0"],
            ["2024年Q3", "约$23.7", "约（$12.8）", "（$0.07）", "约0", "—", "约$0"],
            ["2024年Q4", "约$39.9", "（$28.8）", "（$0.19）", "约150", "约$84", "约$12.6"],
            ["2024年全年", "$78.1", "（$47.8）", "（$0.28）", "450", "$87.2", "$37.9"],
            ["2025年Q1", "$16.9", "（$26.3）", "（$0.13）", "约50", "约$80", "约$7.0"],
            ["2025年Q2", "约$12.7", "（$22.2）", "（$0.10）", "约150", "约$79", "约$4.2"],
            ["2025年Q3", "$17.7", "（$16.7）", "（$0.07）", "240", "约$70", "$10.4"],
            ["2025年Q4", "$27.1", "约（$20.9）", "（$0.08）", "360", "$74.93", "约$27.0"],
            ["2025年全年", "$65.9", "（$86.1）", "（$0.38）", "650", "$74.21", "$48.2"],
        ],
        col_widths=[1.0, 1.2, 1.5, 0.9, 1.0, 1.0, 1.0]
    )
    body(doc, "资料来源：Energy Fuels季度业绩新闻稿（2024年Q1–Q4，2025年Q1–Q4）；Zacks；StockTitan SEC文件。", italic=True, color=C_GREY, size=8.5)
    body(doc, "", space_after=4)

    body(doc,
         "季度数据揭示的主要趋势：（1）2024年Q2和Q3铀矿销售量为零，公司在此期间积累库存；"
         "（2）铀矿销售量在2025年下半年大幅加速，2025年Q4成为公司历史上单季铀矿营收最高的季度；"
         "（3）2025年全年单位经济效益持续改善（每磅销售成本逐季下降）；"
         "（4）重矿砂营收随Kwale停产趋向归零；"
         "（5）2025年净亏损轨迹反映的是高强度投资/产能爬坡成本，而非结构性经营亏损。",
         space_after=8
    )

    # ── 第九节：主要风险 ────────────────────────────────────────────────────────
    heading(doc, "九、主要风险因素", level=2)

    heading(doc, "下行风险", level=3, color=C_RED)
    bullet(doc, "铀矿价格风险：现货价格持续波动；若跌破$60/磅将显著压缩毛利率")
    bullet(doc, "执行风险：二期REE资本开支（4.1亿美元）需要融资——依赖资本市场、政府支持或采购协议")
    bullet(doc, "股权稀释风险：7亿美元可转换票据在$20.34/股转换价格下可产生约3,440万股潜在稀释")
    bullet(doc, "监管/许可风险：Vara Mada（马达加斯加）需与政府谈判；Donald项目仍需环境审批")
    bullet(doc, "稀土价格风险：NdPr、Dy、Tb价格受中国供应政策及全球电动汽车需求波动影响")
    bullet(doc, "管理层交接风险：CEO换届（Bhappu接任Chalmers）带来执行连续性不确定性")
    body(doc, "", space_after=4)

    heading(doc, "上行催化剂", level=3, color=C_GREEN)
    bullet(doc, "铀矿价格反弹至$100+/磅——在2026年计划销售量150–200万磅的基础上，具有显著的业绩弹性")
    bullet(doc, "DOE/DOD签订REE采购合同——政府采购将降低二期建设融资风险")
    bullet(doc, "稀土二期最终投资决定（FID）及融资安排落地公告")
    bullet(doc, "Donald项目（澳大利亚）确认FID并正式启动开采")
    bullet(doc, "钒库存变现——90.5万磅V2O5按当前$6–7/磅可产生约500–700万美元收入")
    bullet(doc, "Pinyon Plain预可研报告更新，改善资源量估算或矿山寿命延长")
    bullet(doc, "医疗同位素（Ra-226）商业协议签署——为关键矿产战略提供进一步多元化")
    body(doc, "", space_after=8)

    # ── 免责声明与资料来源 ─────────────────────────────────────────────────────
    divider(doc)
    heading(doc, "资料来源与免责声明", level=3, color=C_GREY)
    sources_cn = [
        "Energy Fuels公司新闻稿：《2025年业绩及2026年指引》（2026年2月26日）——investors.energyfuels.com",
        "Energy Fuels公司新闻稿：《年终铀矿产销量超出指引》（2025年12月29日）",
        "Energy Fuels公司新闻稿：《2025年Q3业绩》（2025年11月3日）",
        "Energy Fuels公司新闻稿：《2025年Q2业绩》（2025年8月6日）",
        "Energy Fuels公司新闻稿：《2025年Q1业绩》（2025年5月7日）",
        "美通社：《Energy Fuels发布2025年业绩及2026年指引》（2026年2月26日）",
        "H.C. Wainwright研究报告：目标价上调至$27.25（2026年2月27日）",
        "Yahoo Finance / Zacks：Q4 2025业绩摘要及一致预期数据",
        "GuruFocus / Investing.com：Q4 2025业绩电话会议要点",
        "Sprott：《2026年铀矿展望》；《铀矿的双重市场》",
        "Nasdaq.com：《2026年铀矿价格预测：主要趋势》",
        "TipRanks：《Energy Fuels业绩电话会议要点：大胆的增长转型》",
        "StockTitan：UUUU SEC文件（10-Q）——2025年Q3营收1,770万美元",
        "MarketBeat / WallStreetZen / StockAnalysis：分析师一致预期数据",
        "ANS核能新闻：2025年铀矿现货价格数据",
    ]
    for s in sources_cn:
        bullet(doc, s, color=C_GREY)
    body(doc, "", space_after=4)
    body(doc,
         "免责声明：本报告仅供参考，不构成投资建议或任何买卖证券的邀约。"
         "所有财务数据均来源于公开文件及新闻稿。"
         "标注「估」的数据为基于部分信息推算所得，可能与官方最终公布数据存在差异。"
         "过往业绩不代表未来表现。",
         italic=True, color=C_GREY, size=8.5
    )

    # ── 保存 ───────────────────────────────────────────────────────────────────
    out_path = os.path.join(BASE, "UUUU_Q4_FY2025_业绩更新报告_中文版.docx")
    doc.save(out_path)
    print(f"已保存：{out_path}")


if __name__ == "__main__":
    build_report_cn()
