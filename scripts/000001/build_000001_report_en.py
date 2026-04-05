"""
build_000001_report_en.py
Generates English DOCX earnings update report for
平安银行 (Ping An Bank, 000001.SZ)
Q4 2025 / FY2025 Earnings Update
Output: output/000001/000001_Q4_2025_Earnings_Update.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis"
OUT  = os.path.join(BASE, "output", "000001")
os.makedirs(OUT, exist_ok=True)

CHART_PATHS = {i: os.path.join(OUT, f"pab_chart{i}_{n}.png") for i, n in {
    1: "quarterly_revenue",
    2: "quarterly_netprofit",
    3: "nim_trend",
    4: "asset_quality",
    5: "capital_ratios",
    6: "retail_wealth",
    7: "beat_miss",
    8: "annual_comparison",
    9: "broker_targets",
    10: "income_structure",
}.items()}


# ── Market Data ────────────────────────────────────────────────────────────────
def get_market_data():
    try:
        import yfinance as yf
        t    = yf.Ticker("000001.SZ")
        info = t.fast_info
        price  = round(info.last_price, 2)
        mktcap = info.market_cap
        high52 = round(info.year_high, 2)
        low52  = round(info.year_low,  2)
        return {
            "price":      f"RMB {price:.2f}",
            "market_cap": f"~RMB {mktcap/1e9:.0f}B (~USD {mktcap/7.25/1e9:.0f}B)",
            "52w_high":   f"RMB {high52:.2f}",
            "52w_low":    f"RMB {low52:.2f}",
            "exchange":   "000001.SZ (A-share, Shenzhen)",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using static fallback")
        return {
            "price":      "RMB ~10.87",
            "market_cap": "~RMB 211B (~USD 29B)",
            "52w_high":   "RMB ~13.80",
            "52w_low":    "RMB ~9.20",
            "exchange":   "000001.SZ (A-share, Shenzhen)",
        }


mkt   = get_market_data()
TODAY = datetime.date.today().strftime("%B %d, %Y")


# ── Colors ─────────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

PAB_RED   = hex_to_rgb("D0222A")
PAB_ORANGE= hex_to_rgb("F07018")
PAB_NAVY  = hex_to_rgb("003366")
PAB_GOLD  = hex_to_rgb("C8960C")
WHITE_C   = hex_to_rgb("FFFFFF")
L_GRAY    = hex_to_rgb("F0F4F8")
M_GRAY    = hex_to_rgb("8B8B8B")
GREEN_C   = hex_to_rgb("2E7D32")
RED_C     = hex_to_rgb("C41230")


# ── Document Helpers ───────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def add_hyperlink(paragraph, url, text, color=None):
    part      = paragraph.part
    r_id      = part.relate_to(url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run   = OxmlElement("w:r")
    rPr       = OxmlElement("w:rPr")
    rStyle    = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if color:
        clr = OxmlElement("w:color")
        clr.set(qn("w:val"), color.lstrip("#"))
        rPr.append(clr)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before       = Pt(14)
    p.paragraph_format.space_after        = Pt(4)
    p.paragraph_format.keep_with_next     = True
    run = p.add_run(text.upper())
    run.bold = True; run.font.size = Pt(12); run.font.color.rgb = PAB_NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "003366")
    pBdr.append(bot); pPr.append(pBdr)
    return p


def h2(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = PAB_RED
    return p


def body(doc, text, bold=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Times New Roman"; run.bold = bold
    return p


def bullet(doc, text, bold=False, color=None):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10); run.bold = bold
    if color: run.font.color.rgb = color
    return p


def hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot); pPr.append(pBdr)


def add_img(doc, path, caption, width=Inches(6.0)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)
    else:
        body(doc, f"[Chart not found: {path}]")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(8.5); run.font.color.rgb = M_GRAY; run.italic = True
    cap.paragraph_format.space_after = Pt(8)


def make_table(doc, headers, rows, hdr_bg="003366", alt_bg="F0F4F8"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hr_row.cells[i]
        set_cell_bg(cell, hdr_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE_C
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return tbl


# ════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

# ── Cover ──────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("Ping An Bank (000001.SZ)")
tr.bold = True; tr.font.size = Pt(16); tr.font.color.rgb = PAB_NAVY
tr.font.name = "Times New Roman"

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("Q4 2025 / FY2025 Earnings Update — NIM Stabilization Signals Recovery; Maintain BUY")
sr.bold = True; sr.font.size = Pt(11); sr.font.color.rgb = PAB_RED
sr.font.name = "Times New Roman"

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dp.add_run(f"Equity Research  |  April 5, 2026  |  {mkt['exchange']}")
dr.font.size = Pt(9); dr.font.color.rgb = M_GRAY; dr.font.name = "Times New Roman"

hr(doc)

# ── Rating Table ───────────────────────────────────────────────────────────────
rt = doc.add_table(rows=1, cols=6)
rt.style = "Table Grid"; rt.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs_r = ["Rating", "Price Target\n(A-share)", "Current Price",
          "Market Cap", "52W High/Low", "Shares Out."]
vals_r  = ["BUY ↑", "RMB 14.00",
           mkt["price"], mkt["market_cap"],
           f"{mkt['52w_high']} / {mkt['52w_low']}", "~19.4B shares"]
row_r   = rt.rows[0]
for i, (hdr, val) in enumerate(zip(hdrs_r, vals_r)):
    cell = row_r.cells[i]
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    h_run = p.add_run(hdr + "\n")
    h_run.bold = True; h_run.font.size = Pt(8); h_run.font.color.rgb = WHITE_C
    v_run = p.add_run(val)
    v_run.font.size = Pt(9.5)
    v_run.font.color.rgb = PAB_GOLD if i < 2 else WHITE_C
    v_run.bold = (i < 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph()

# ── I. Executive Summary ───────────────────────────────────────────────────────
h1(doc, "I. Executive Summary")
body(doc,
     "Ping An Bank (000001.SZ) reported FY2025 full-year results on March 20, 2026, delivering "
     "net profit attributable to shareholders of RMB 42.63B (-4.2% YoY), in line with consensus "
     "expectations of ~RMB 42.2B. Operating revenue of RMB 131.44B declined -10.4% YoY, consistent "
     "with the bank's ongoing structural revenue pressure. The headline story, however, is inflection: "
     "Q4 2025 net interest income grew +2.76% YoY — the first positive YoY NII growth in multiple "
     "quarters — and the full-year NIM declined only 9bps to 1.78% vs. FY2024's -51bps compression. "
     "Asset quality improved, with the NPL formation rate declining to 1.63% (-17bps YoY). "
     "Management declared the 'most difficult period has passed' at the March 23 analyst conference.")

h2(doc, "Key Takeaways")
bullet(doc, "Revenue in-line; net profit slight beat: FY2025 NP of RMB 42.63B vs. "
            "consensus ~RMB 42.2B (+1.0% beat). Revenue of RMB 131.44B "
            "was slightly below consensus ~RMB 132.0B (-0.4%).", bold=True, color=GREEN_C)
bullet(doc, "NIM stabilization — the critical inflection: Full-year NIM of 1.78% "
            "(-9bps YoY vs. FY2024's -51bps) reflects sharp deceleration in spread compression. "
            "Q4 NII turned positive YoY (+2.76%) for the first time in several quarters, "
            "driven by a 47bps decline in average funding costs to 1.67%.")
bullet(doc, "Asset quality improving: NPL ratio stable at 1.05% (-1bp YoY); NPL formation rate "
            "fell 17bps to 1.63%; personal loan NPL improved -16bps to 1.23%. "
            "Provision coverage declined 30pp to 220.88%, as elevated provisions normalize.")
bullet(doc, "Capital strengthening: CET1 rose 24bps to 9.36% YoY; Tier 1 up 80bps "
            "to 11.49% — driven by Q4 retained earnings and reduced risk-weighted asset growth.")
bullet(doc, "Retail transformation on track: Retail banking net profit surged +828% YoY "
            "to RMB 2.68B (FY2024: RMB 289M). Private banking clients grew +9.1% to "
            "105,600; wealth management fee income +15.8% to RMB 5.06B.")
bullet(doc, "Dividend: Total FY2025 dividend of RMB 0.596/share (含税), "
            "payout ratio 28.83%; dividend yield ~5.5% at current price. "
            "Management targets 20–30% payout ratio, with aspiration to increase in 2026.")
bullet(doc, "Maintain BUY with PT RMB 14.00, implying 0.56x FY2026E BPS of ~RMB 24.9 "
            "and ~29% upside. Management's '2026 return to growth' target is credible given "
            "NIM stabilization and declining credit costs.")

doc.add_paragraph()

# Results Snapshot Table
h2(doc, "Results Snapshot — FY2025 vs. Consensus")
make_table(doc,
    headers=["Metric", "FY2025 Actual", "Consensus (Pre-Results)", "Beat/Miss", "YoY Change"],
    rows=[
        ["Operating Revenue (RMB B)",    "131.44", "~132.0",  "-0.4% miss ↓",  "-10.4%"],
        ["Net Interest Income (RMB B)",  "88.02",  "~88.5",   "-0.5% miss ↓",  "-5.8%"],
        ["Net Profit Attr. (RMB B)",     "42.63",  "~42.2",   "+1.0% beat ✓",  "-4.2%"],
        ["EPS (RMB)",                    "2.07",   "~2.04",   "+1.5% beat ✓",  "-4.2%"],
        ["NIM (%)",                      "1.78%",  "~1.76%",  "+2bps beat ✓",  "-9bps"],
        ["NPL Ratio (%)",                "1.05%",  "~1.07%",  "-2bps beat ✓",  "-1bp"],
        ["Provision Coverage (%)",       "220.88%","~225%",   "-4pp miss ↓",   "-29.8pp"],
        ["CET1 Capital Ratio (%)",       "9.36%",  "~9.20%",  "+16bps beat ✓", "+24bps"],
        ["Cost-to-Income (%)",           "29.06%", "~29.0%",  "In Line",        "+1.4pp"],
        ["ROE (%)",                      "9.15%",  "~9.2%",   "In Line",        "-0.93pp"],
    ],
    hdr_bg="003366", alt_bg="F0F4F8"
)
add_img(doc, CHART_PATHS[7], "Exhibit 1: FY2025 Actual vs. Pre-Results Consensus (Key Metrics)", width=Inches(5.5))
hr(doc)

# ── II. Detailed Results ───────────────────────────────────────────────────────
h1(doc, "II. Detailed Quarterly & Full-Year Results")

h2(doc, "Revenue — Cyclical Low; Q4 Shows Sequential Stabilization")
body(doc,
     "FY2025 operating revenue of RMB 131.44B declined -10.4% YoY (FY2024: RMB 146.70B), "
     "the third consecutive year of double-digit revenue contraction driven primarily by NIM "
     "compression and declining non-interest income. On a quarterly basis, Q4 2025 revenue of "
     "~RMB 30.77B declined -12.4% YoY, but sequential performance was more benign: Q1-Q4 2025 "
     "revenue was RMB 33.71B / 35.68B / 31.28B / 30.77B, showing relative stability in H2 "
     "after Q2's temporary non-interest income boost.")
body(doc,
     "Net interest income of RMB 88.02B (-5.8% YoY) was the primary revenue drag, though the "
     "pace of decline sharply decelerated vs. FY2024's implied ~-10% NII decline. Critically, "
     "Q4 NII grew +2.76% YoY — a key inflection point attributable to: (1) liability cost "
     "improvement (average funding cost fell 47bps to 1.67%); (2) slower asset repricing "
     "pressure as LPR cuts moderated; and (3) deposit mix improvement from high-cost time "
     "deposits toward lower-cost current/wealth product structures.")

add_img(doc, CHART_PATHS[1], "Exhibit 2: Quarterly Operating Revenue (Q1 2024–Q4 2025, RMB Billion)", width=Inches(6.0))
add_img(doc, CHART_PATHS[10], "Exhibit 3: Revenue Mix — Net Interest vs. Non-Interest Income (FY2022–FY2025)", width=Inches(5.5))

h2(doc, "Net Profit — Stable Decline; Q4 Heavy Provisioning as Expected")
body(doc,
     "FY2025 net profit attributable to shareholders of RMB 42.63B declined -4.2% YoY "
     "(FY2024: RMB 44.51B), a slightly narrowing pace vs. FY2024's -4.2%. The Q4 standalone "
     "quarter implied net profit of ~RMB 4.29B was exceptionally weak (~-10% YoY vs Q4 2024's "
     "RMB 4.78B), reflecting the bank's customary year-end practice of front-loading provisioning "
     "charges into Q4. Provision charges (信用减值损失) declined -17.9% YoY to RMB 40.57B for FY2025, "
     "reflecting improving underlying credit quality and reduced need for large reserve builds.")
body(doc,
     "Full-year credit cost fell to 1.38% (-18bps YoY), providing a meaningful tailwind to "
     "bottom-line profitability in FY2026E as this normalization continues. The bank's pre-provision "
     "operating profit (PPOP) declined ~-11.9% YoY, indicating that revenue contraction remains the "
     "primary challenge — but as NIM stabilizes and credit costs normalize, we expect net profit to "
     "return to positive growth in FY2026E.")

add_img(doc, CHART_PATHS[2], "Exhibit 4: Quarterly Net Profit with YoY Growth Labels (Q1 2024–Q4 2025)", width=Inches(6.0))
hr(doc)

# ── III. Key Metrics ───────────────────────────────────────────────────────────
h1(doc, "III. Key Banking Metrics")

h2(doc, "Net Interest Margin (NIM) — Critical Inflection Confirmed")
body(doc,
     "The FY2025 NIM of 1.78% represents only -9bps of decline vs. FY2024's 1.87%, "
     "a dramatic improvement over FY2024's -51bps compression (FY2023: 2.38%). "
     "This deceleration reflects: (1) average deposit cost declined 42bps to 1.65%, "
     "benefiting from regulatory deposit rate cuts and mix shift; (2) slower LPR-linked "
     "asset repricing (only one 10bps LPR cut in H2 2025 vs. multiple cuts in 2024); "
     "and (3) incremental benefit from corporate loan growth (+3.5% YoY) vs. "
     "declining personal loans (-2.3% YoY, primarily mortgage). "
     "Management declared NIM 'essentially stabilized' at the March 2026 investor conference, "
     "with FY2026E NIM consensus of ~1.77% — implying near-flat from FY2025.")

add_img(doc, CHART_PATHS[3], "Exhibit 5: NIM Trend (FY2022–FY2025 Actual; FY2026E Consensus)", width=Inches(5.5))

h2(doc, "Asset Quality — Continuing Improvement; Real Estate Remains Watchpoint")
body(doc,
     "Overall asset quality continued to improve in FY2025. The NPL ratio edged down 1bp to "
     "1.05%, with personal loan NPL declining 16bps to 1.23% — signaling that retail credit "
     "normalization (post-COVID consumer stress) is progressing. The NPL formation rate declined "
     "17bps to 1.63% and the overdue 90d+ deviation ratio of 0.56x (below 1.0x is favorable) "
     "indicates strong forward coverage of potential new NPLs.")
body(doc,
     "Corporate NPL rose 17bps to 0.87%, driven by real estate exposure where the NPL ratio "
     "deteriorated 43bps to 2.22% (FY2024: 1.79%). Real estate remains the key credit watchpoint, "
     "though the bank has been actively de-risking: real estate loans as a percentage of total "
     "corporate loans have declined significantly over 2023–2025. Provision coverage of 220.88% "
     "(-29.8pp YoY) is above the regulatory minimum and reflects deliberate normalization as "
     "underlying credit quality improves and provisioning needs decline.")

add_img(doc, CHART_PATHS[4], "Exhibit 6: NPL Ratio & Provision Coverage Ratio (H1 2024 – FY2025)", width=Inches(5.5))

h2(doc, "Capital — Well-Capitalized; Organic Growth Sufficient")
body(doc,
     "CET1 capital ratio improved 24bps to 9.36% at FY2025 year-end (FY2024: 9.12%), "
     "and the total capital ratio rose 66bps to 13.77%. The improvement was driven by "
     "retained earnings and controlled risk-weighted asset growth (+0.5% loan growth). "
     "Management confirmed no rights offering is planned; organic capital generation is "
     "sufficient to support moderate 2026 business expansion. A higher capital base supports "
     "the bank's ability to maintain the 20–30% dividend payout target.")

add_img(doc, CHART_PATHS[5], "Exhibit 7: Capital Adequacy Ratios (FY2024 vs. FY2025)", width=Inches(5.0))

make_table(doc,
    headers=["Metric", "FY2023A", "FY2024A", "H1 2025A", "FY2025A", "FY2026E"],
    rows=[
        ["Revenue (RMB B)",             "164.70", "146.70", "69.39",  "131.44", "~135.2"],
        ["Net Profit (RMB B)",          "46.46",  "44.51",  "24.87",  "42.63",  "~43.36"],
        ["NIM (%)",                     "2.38",   "1.87",   "1.80",   "1.78",   "~1.77"],
        ["NPL Ratio (%)",               "1.06",   "1.06",   "1.10",   "1.05",   "~1.03"],
        ["Provision Coverage (%)",      "277.6",  "250.71", "246.6",  "220.88", "~215"],
        ["CET1 (%)",                    "8.84",   "9.12",   "9.20",   "9.36",   "~9.5"],
        ["ROE (%)",                     "11.4",   "10.08",  "9.52",   "9.15",   "~8.9"],
        ["Cost-to-Income (%)",          "26.3",   "~27.7",  "28.3",   "29.06",  "~27.7"],
        ["EPS (RMB)",                   "2.25",   "2.16",   "1.21",   "2.07",   "~2.19"],
        ["BVPS (RMB)",                  "21.1",   "~22.4",  "~23.1",  "~24.0",  "~24.9"],
    ],
    hdr_bg="003366", alt_bg="F0F4F8"
)
hr(doc)

# ── IV. Retail Transformation ──────────────────────────────────────────────────
h1(doc, "IV. Retail Banking Transformation Progress")
body(doc,
     "Ping An Bank's multi-year retail transformation remains the central long-term investment "
     "thesis. FY2025 marked a breakthrough in retail profitability: retail banking net profit "
     "surged +828% YoY to RMB 2.68B (FY2024: RMB 289M), increasing its share of total bank "
     "net profit to 6.3% from 0.6% — reflecting the cyclical trough passage in retail credit costs. "
     "Management estimates the retail transformation is '~70% complete, ~30% remaining.'")

h2(doc, "Wealth Management — A Bright Spot")
bullet(doc, "Total retail AUM: RMB 4.24 trillion (+1.1% YoY vs. RMB 4.19 trillion in FY2024) — "
            "modest growth constrained by equity market volatility but demonstrates resilient client relationships.")
bullet(doc, "Private banking clients: +9.1% to 105,600; private banking AUM: RMB 1.991 trillion (+0.8% YoY). "
            "Private banking is the highest-margin retail segment and the key long-term growth driver.")
bullet(doc, "Wealth management fee income: +15.8% YoY to RMB 5.06B — driven by insurance +53.3%, "
            "mutual funds +8.9%, and wealth products +8.8%. Fee income is growing even as total "
            "non-interest income declines, indicating structural quality improvement.")
bullet(doc, "Pocket Bank (口袋银行) MAU ~40M — a key digital distribution advantage for fee income.")
bullet(doc, "Strategic integration with Ping An Insurance: Cross-selling bancassurance products "
            "drove insurance commission income +53.3% YoY; the insurance-bank synergy model "
            "is differentiating vs. non-conglomerate peers.")

add_img(doc, CHART_PATHS[6], "Exhibit 8: Retail AUM, Wealth Clients & Private Banking Clients (FY2023–FY2025)", width=Inches(5.5))
hr(doc)

# ── V. Investment Thesis ───────────────────────────────────────────────────────
h1(doc, "V. Investment Thesis Update — Maintain BUY")
body(doc,
     "We maintain our BUY rating with a 12-month price target of RMB 14.00 (0.56x FY2026E "
     "BPS ~RMB 24.9), representing ~29% upside from current levels. The investment thesis "
     "is predicated on: (1) NIM stabilization as the most critical near-term catalyst; "
     "(2) declining credit costs as retail NPL formation peaks; (3) retail transformation "
     "driving sustained fee income growth; and (4) a ~5.5% dividend yield providing "
     "downside protection at current valuations (~0.44x P/B).")

h2(doc, "Bull Case Catalysts")
bullet(doc, "NIM inflection: Q4 NII turned positive YoY for the first time in multiple quarters. "
            "If Q1 2026 NII continues to grow YoY, revenue growth will return in 2026 — "
            "the single most important catalyst for re-rating.")
bullet(doc, "Credit cost normalization: NPL formation rate of 1.63% is declining. As retail "
            "credit stress (personal loans, credit cards, mortgages) continues to normalize, "
            "provisions can decline further from RMB 40.57B in FY2025 to ~RMB 36–38B in FY2026E, "
            "directly boosting bottom-line growth even if PPOP is flat.")
bullet(doc, "Retail profit contribution expansion: Retail banking at 6.3% of total net profit "
            "in FY2025 vs. 0.6% in FY2024 demonstrates the transformation is delivering. "
            "Targeting 25–30% retail contribution by FY2027E as consumer credit quality normalizes.")
bullet(doc, "Valuation floor at ~0.44x P/B: Historically low vs. Ping An Bank's 5-year average "
            "of ~0.85x P/B. A re-rating to 0.6x P/B alone (still below historical average) "
            "implies RMB 15/share upside from current levels.")
bullet(doc, "Ping An Group parent support: Strong backing from China's largest insurance group "
            "provides funding, insurance cross-sell synergies, and systemic importance protection.")

h2(doc, "Key Risks")
bullet(doc, "Revenue pressure continuation: If LPR is cut further in 2026 or deposit repricing "
            "is slower than expected, NIM could compress again, delaying revenue recovery.")
bullet(doc, "Real estate NPL deterioration: Corporate NPL rose (mainly real estate at 2.22%). "
            "A broader real estate credit event could force fresh provision builds, "
            "potentially reversing the credit cost improvement trend.")
bullet(doc, "Macro slowdown: Consumption softness and SME stress could re-accelerate retail NPL formation, "
            "particularly in personal consumption loans and credit card receivables.")
bullet(doc, "Capital adequacy pressure: CET1 at 9.36% is above minimum but leaves limited buffer "
            "for rapid loan growth or unexpected losses. Dividend aspirations could be constrained.")
bullet(doc, "Competition from state-owned banks: Large state banks (ICBC, CCB, ABC) are aggressively "
            "competing in wealth management and corporate lending — Ping An Bank's competitive moat "
            "is narrowing in these segments.")
hr(doc)

# ── VI. Valuation ──────────────────────────────────────────────────────────────
h1(doc, "VI. Valuation & Updated Estimates")

h2(doc, "Price Target: RMB 14.00 (0.56x FY2026E BPS)")
body(doc,
     "Our 12-month price target of RMB 14.00 is based on 0.56x FY2026E book value per share "
     "of ~RMB 24.9 (consensus). This P/B multiple is justified by our FY2026E ROE estimate "
     "of ~8.9%, with a required return of ~10% (COE), implying a justified P/B of ~0.9x via "
     "Gordon Growth Model — we apply a 40% discount to intrinsic value to reflect execution "
     "risk on the retail transformation and macro uncertainty, yielding 0.56x P/B. "
     "The consensus target price of RMB 14.25 (25 institutions) is consistent with our target.")
body(doc,
     "Bull scenario (0.70x FY2026E BPS): RMB 17.4/share — requires full NIM stabilization "
     "and accelerating retail profitability. Bear scenario (0.40x FY2026E BPS): "
     "RMB 10.0/share — implies renewed credit deterioration and further NIM compression.")

make_table(doc,
    headers=["Metric", "FY2023A", "FY2024A", "FY2025A", "FY2026E", "FY2027E"],
    rows=[
        ["Revenue (RMB B)",           "164.70", "146.70", "131.44", "~135.2", "~139.0"],
        ["Revenue Growth (%)",        "-8.5",   "-10.9",  "-10.4",  "~+2.9",  "~+2.8"],
        ["Net Profit (RMB B)",        "46.46",  "44.51",  "42.63",  "~43.36", "~44.40"],
        ["NP Growth (%)",             "+2.1",   "-4.2",   "-4.2",   "~+1.7",  "~+2.4"],
        ["EPS (RMB)",                 "2.25",   "2.16",   "2.07",   "~2.19",  "~2.26"],
        ["BVPS (RMB)",               "~21.1",  "~22.4",  "~24.0",  "~24.9",  "~26.0"],
        ["P/B at RMB 10.87 (x)",     "0.52x",  "0.49x",  "0.45x",  "0.44x",  "0.42x"],
        ["P/B at PT RMB 14.00 (x)",  "0.66x",  "0.63x",  "0.58x",  "0.56x",  "0.54x"],
        ["P/E at current price (x)", "4.83x",  "5.03x",  "5.25x",  "4.96x",  "4.81x"],
        ["Dividend Yield (%) at current", "~5.2%", "~5.1%", "~5.5%", "~5.5%", "~5.7%"],
    ],
    hdr_bg="003366", alt_bg="F0F4F8"
)
body(doc, "Note: FY2026E/2027E are Wind consensus estimates. BVPS estimates based on projected retained earnings.")
add_img(doc, CHART_PATHS[8], "Exhibit 9: Annual Revenue & Net Profit (FY2022–FY2027E)", width=Inches(6.0))
add_img(doc, CHART_PATHS[9], "Exhibit 10: Broker FY2026E EPS Estimates & Price Targets", width=Inches(6.0))
hr(doc)

# ── VII. Sources ───────────────────────────────────────────────────────────────
h1(doc, "VII. Sources & References")
body(doc, "All data as of April 5, 2026. FY2025 Annual Report officially released March 20, 2026.")

src_data = [
    ("Ping An Bank FY2025 Annual Report Released (Sina Finance, March 20, 2026)",
     "https://finance.sina.com.cn/jryx/bank/2026-03-20/doc-inhrrzut6659056.shtml"),
    ("Ping An Bank FY2025 Net Profit RMB 42.63B; Dividend 5.96元/10 shares (Sina Finance)",
     "https://finance.sina.com.cn/roll/2026-03-20/doc-inhrshar6590029.shtml"),
    ("Ping An Bank FY2025 Annual Report Analyst Review — Galaxy Securities (StockStar, March 24)",
     "https://stock.stockstar.com/JC2026032400000016.shtml"),
    ("Ping An Bank FY2025 Annual Report Deep Dive (Sina Finance, March 31, 2026)",
     "https://finance.sina.com.cn/wm/2026-03-31/doc-inhsxkrq2161128.shtml"),
    ("Ping An Bank 2026 Analyst Conference: 'Full Force Return to Growth' (Sina Finance, March 23)",
     "https://finance.sina.com.cn/roll/2026-03-23/doc-inhryxfp3966939.shtml"),
    ("Ping An Bank: Consensus PT RMB 14.25; Valuation at Historical Low (Sina Finance, March 22)",
     "https://finance.sina.com.cn/roll/2026-03-22/doc-inhrvxax5011786.shtml"),
    ("Reading Ping An Bank FY2025 Annual Report: The Most Difficult Period Has Passed (Sina Finance)",
     "https://finance.sina.com.cn/stock/bxjj/2026-03-25/doc-inhsecsi9934000.shtml"),
    ("Ping An Bank 2025: First Step Out of the Deep Water Zone (Sina Finance, March 25, 2026)",
     "https://finance.sina.com.cn/wm/2026-03-25/doc-inhsfmms3417265.shtml"),
    ("21st Century Business Herald: Annual Report — Dividend RMB 10B, NIM -9bps (March 20, 2026)",
     "https://www.21jingji.com/article/20260320/herald/5ee8bf7a9198ff98e638898d5820cc17.html"),
    ("Ping An Bank H1 2025 Net Profit -3.9%; NIM 1.80% (Wall Street CN, August 2025)",
     "https://wallstreetcn.com/articles/3753917"),
    ("Wind / Tonghuashun F10: FY2026E Consensus Estimates (000001.SZ)",
     "https://basic.10jqka.com.cn/000001/worth.html"),
    ("Retail Transformation 30% Remaining: Ping An Bank 2026 Recovery Methodology Forming (STCN)",
     "https://www.stcn.com/article/detail/3693988.html"),
]

for title, url in src_data:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    add_hyperlink(p, url, title, color="#003366")

doc.add_paragraph()
disc = doc.add_paragraph()
disc_run = disc.add_run(
    "DISCLAIMER: This report is for informational purposes only and does not constitute investment "
    "advice. All financial data sourced from public company filings and consensus estimates. "
    "Market data via yfinance as of report date. FY2026E/2027E figures are analyst consensus estimates."
)
disc_run.font.size = Pt(8); disc_run.font.color.rgb = M_GRAY; disc_run.italic = True
disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT, "000001_Q4_2025_Earnings_Update.docx")
doc.save(out_path)
print(f"✅ English report saved → {out_path}")
