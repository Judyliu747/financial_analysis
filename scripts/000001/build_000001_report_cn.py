"""
build_000001_report_cn.py
生成平安银行（000001.SZ）Q4 2025 / FY2025 业绩更新报告（中文版）
输出: output/000001/000001_Q4_2025_业绩更新报告_中文版.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── 路径 ─────────────────────────────────────────────────────────────────────
BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis"
OUT  = os.path.join(BASE, "output", "000001")
os.makedirs(OUT, exist_ok=True)

CHART_PATHS = {i: os.path.join(OUT, f"pab_chart{i}_{n}.png") for i, n in {
    1: "quarterly_revenue",
    2: "quarterly_netprofit",
    3: "nim_trend",
    4: "asset_quality",
    5: "capital_ratios",
    6: "retail_wealth",
    7: "beat_miss",
    8: "annual_comparison",
    9: "broker_targets",
    10: "income_structure",
}.items()}


# ── 市场数据 ─────────────────────────────────────────────────────────────────
def get_market_data():
    try:
        import yfinance as yf
        t    = yf.Ticker("000001.SZ")
        info = t.fast_info
        price  = round(info.last_price, 2)
        mktcap = info.market_cap
        high52 = round(info.year_high, 2)
        low52  = round(info.year_low,  2)
        return {
            "price":      f"{price:.2f}元",
            "market_cap": f"约{mktcap/1e8:.0f}亿元（约{mktcap/7.25/1e9:.0f}亿美元）",
            "52w_high":   f"{high52:.2f}元",
            "52w_low":    f"{low52:.2f}元",
            "exchange":   "000001.SZ（深交所A股）",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using static fallback")
        return {
            "price":      "约10.87元",
            "market_cap": "约2,109亿元（约291亿美元）",
            "52w_high":   "约13.80元",
            "52w_low":    "约9.20元",
            "exchange":   "000001.SZ（深交所A股）",
        }


mkt = get_market_data()
TODAY = datetime.date.today().strftime("%Y年%m月%d日")


# ── 颜色 ─────────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

PAB_RED   = hex_to_rgb("D0222A")
PAB_ORANGE= hex_to_rgb("F07018")
PAB_NAVY  = hex_to_rgb("003366")
PAB_GOLD  = hex_to_rgb("C8960C")
WHITE_C   = hex_to_rgb("FFFFFF")
L_GRAY    = hex_to_rgb("F0F4F8")
M_GRAY    = hex_to_rgb("8B8B8B")
GREEN_C   = hex_to_rgb("2E7D32")
RED_C     = hex_to_rgb("C41230")


# ── 文档辅助 ──────────────────────────────────────────────────────────────────
def set_cjk_font(run, font_name="宋体"):
    rPr    = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def add_hyperlink(paragraph, url, text, color=None):
    part  = paragraph.part
    r_id  = part.relate_to(url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run   = OxmlElement("w:r")
    rPr       = OxmlElement("w:rPr")
    rStyle    = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if color:
        clr = OxmlElement("w:color")
        clr.set(qn("w:val"), color.lstrip("#"))
        rPr.append(clr)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def h1_cn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before    = Pt(14)
    p.paragraph_format.space_after     = Pt(4)
    p.paragraph_format.keep_with_next  = True
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(12.5); run.font.color.rgb = PAB_NAVY
    run.font.name = "Times New Roman"
    set_cjk_font(run, "黑体")
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "003366")
    pBdr.append(bot); pPr.append(pBdr)
    return p


def h2_cn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = PAB_RED
    run.font.name = "Times New Roman"
    set_cjk_font(run, "黑体")
    return p


def body_cn(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.first_line_indent = Pt(20)
    run = p.add_run(text)
    run.font.size = Pt(10.5); run.font.name = "Times New Roman"; run.bold = bold
    set_cjk_font(run, "宋体")
    return p


def bullet_cn(doc, text, bold=False, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5); run.font.name = "Times New Roman"; run.bold = bold
    set_cjk_font(run, "宋体")
    if color: run.font.color.rgb = color
    return p


def hr_cn(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot); pPr.append(pBdr)


def add_img_cn(doc, path, caption, width=Inches(6.0)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)
    else:
        body_cn(doc, f"【图表未找到：{path}】")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(8.5); run.font.color.rgb = M_GRAY; run.italic = True
    run.font.name = "Times New Roman"
    set_cjk_font(run, "宋体")
    cap.paragraph_format.space_after = Pt(8)


def make_table_cn(doc, headers, rows, hdr_bg="003366", alt_bg="F0F4F8"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hr_row.cells[i]
        set_cell_bg(cell, hdr_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE_C
        run.font.name = "Times New Roman"
        set_cjk_font(run, "黑体")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9); run.font.name = "Times New Roman"
            set_cjk_font(run, "宋体")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return tbl


# ════════════════════════════════════════════════════════════════════════════
# 构建文档
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

# ── 封面 ──────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("平安银行股份有限公司（000001.SZ）")
tr.bold = True; tr.font.size = Pt(16); tr.font.color.rgb = PAB_NAVY
tr.font.name = "Times New Roman"; set_cjk_font(tr, "黑体")

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("2025年四季度 / FY2025 业绩更新 — 净息差企稳信号确立，重回增长在望，维持买入")
sr.bold = True; sr.font.size = Pt(11); sr.font.color.rgb = PAB_RED
sr.font.name = "Times New Roman"; set_cjk_font(sr, "黑体")

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dp.add_run(f"股票研究  |  2026年4月5日  |  {mkt['exchange']}")
dr.font.size = Pt(9); dr.font.color.rgb = M_GRAY
dr.font.name = "Times New Roman"; set_cjk_font(dr, "宋体")

hr_cn(doc)

# ── 评级信息表 ────────────────────────────────────────────────────────────────
rt = doc.add_table(rows=1, cols=6)
rt.style = "Table Grid"; rt.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ["评级", "目标价\n（A股）", "当前股价", "总市值", "52周高/低", "总股本"]
vals  = ["买入 ↑", "人民币 14.00元", mkt["price"], mkt["market_cap"],
         f"{mkt['52w_high']} / {mkt['52w_low']}", "约194亿股"]
row_r = rt.rows[0]
for i, (hdr, val) in enumerate(zip(hdrs, vals)):
    cell = row_r.cells[i]
    set_cell_bg(cell, "003366")
    p    = cell.paragraphs[0]
    h_r  = p.add_run(hdr + "\n")
    h_r.bold = True; h_r.font.size = Pt(8); h_r.font.color.rgb = WHITE_C
    h_r.font.name = "Times New Roman"; set_cjk_font(h_r, "黑体")
    v_r  = p.add_run(val)
    v_r.font.size = Pt(9); v_r.font.color.rgb = PAB_GOLD if i < 2 else WHITE_C
    v_r.bold = (i < 2); v_r.font.name = "Times New Roman"; set_cjk_font(v_r, "宋体")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph()

# ── 一、业绩摘要 ──────────────────────────────────────────────────────────────
h1_cn(doc, "一、业绩摘要")
body_cn(doc,
    "平安银行股份有限公司（下称\u201c平安银行\u201d或\u201c公司\u201d）于2026年3月20日正式发布"
    "2025年年度报告。FY2025全年实现归母净利润426.33亿元，同比下降4.2%，与Wind一致预期约422亿元"
    "基本符合（超预期约+1.0%）；营业收入1,314.42亿元，同比下降10.4%，略低于一致预期约1,320亿元。"
    "核心亮点在于净息差（NIM）企稳拐点明确：2025年四季度净利息收入同比增长+2.76%，为连续多季度"
    "下滑后首次转正；全年NIM仅下行9个基点至1.78%，较2024年的-51个基点大幅收窄。"
    "管理层在3月23日业绩发布会上宣称\u201c最难的时候已经过去\u201d，并明确2026年目标为\u201c全力以赴重回增长\u201d。"
)

h2_cn(doc, "核心要点")
bullet_cn(doc,
    "盈利基本符合预期：FY2025归母净利润426.33亿元（一致预期约422亿元，超预期+1.0%）；"
    "EPS 2.07元。营业收入1,314.42亿元，略低于预期（逊预期-0.4%）。",
    bold=True, color=GREEN_C)
bullet_cn(doc,
    "净息差拐点确立 — 最关键变量：全年NIM 1.78%（-9bps，FY2024为-51bps），降幅大幅收窄；"
    "Q4净利息收入同比+2.76%，为多个季度以来首次转正；负债端成本改善（平均资金成本降47bps至1.67%）是主因。")
bullet_cn(doc,
    "资产质量持续改善：不良率1.05%（同比-1bp）；不良生成率降至1.63%（-17bps）；"
    "个人贷款不良率1.23%（-16bps），零售信用风险正在消退。")
bullet_cn(doc,
    "资本充足率提升：核心一级资本充足率9.36%（+24bps），一级资本充足率11.49%（+80bps）。")
bullet_cn(doc,
    "零售转型成效显现：零售条线净利润大幅增长+828%至26.83亿元（2024年仅2.89亿元），"
    "占总利润比重从0.6%提升至6.3%；财富管理手续费收入+15.8%至50.61亿元。")
bullet_cn(doc,
    "分红：FY2025合计每10股派息5.96元（含税），派息率28.83%；"
    "按当前股价约10.87元计算，股息率约5.5%，提供较强安全边际。")
bullet_cn(doc,
    "维持\u201c买入\u201d评级，目标价14.00元（0.56倍FY2026E每股净资产约24.9元），"
    "潜在上涨空间约29%；Wind 25家机构一致预期目标价14.25元。",
    bold=True)

doc.add_paragraph()

h2_cn(doc, "业绩快照 — FY2025实际 vs 市场预期")
make_table_cn(doc,
    headers=["指标", "FY2025实际", "市场一致预期（发布前）", "超/逊预期", "同比变化"],
    rows=[
        ["营业收入（亿元）",       "1,314.42", "约1,320",  "-0.4% 略逊 ↓", "-10.4%"],
        ["净利息收入（亿元）",     "880.21",   "约885",    "-0.5% 略逊 ↓", "-5.8%"],
        ["归母净利润（亿元）",     "426.33",   "约422",    "+1.0% 超预期 ✓","-4.2%"],
        ["EPS（元）",             "2.07",     "约2.04",   "+1.5% 超预期 ✓","-4.2%"],
        ["净息差NIM（%）",        "1.78%",    "约1.76%",  "+2bps 超预期 ✓","-9bps"],
        ["不良贷款率（%）",       "1.05%",    "约1.07%",  "-2bps 超预期 ✓","-1bp"],
        ["拨备覆盖率（%）",       "220.88%",  "约225%",   "-4pp 略逊 ↓",   "-29.8pp"],
        ["核心一级资本充足率（%）","9.36%",    "约9.20%",  "+16bps 超预期 ✓","+24bps"],
        ["成本收入比（%）",       "29.06%",   "约29.0%",  "基本符合",       "+1.4pp"],
        ["ROE（%）",             "9.15%",    "约9.2%",   "基本符合",       "-0.93pp"],
    ],
    hdr_bg="003366", alt_bg="F0F4F8"
)
add_img_cn(doc, CHART_PATHS[7], "图1：FY2025实际业绩 vs 预公布前市场一致预期", width=Inches(5.5))

hr_cn(doc)

# ── 二、季度与全年业绩详细分析 ────────────────────────────────────────────────
h1_cn(doc, "二、季度与全年业绩详细分析")

h2_cn(doc, "营业收入 — 承压但出现边际改善信号")
body_cn(doc,
    "FY2025营业收入1,314.42亿元，同比下降10.4%（FY2024：1,466.95亿元），为连续第三年两位数降幅，"
    "主要受净息差持续收窄及非息收入下滑拖累。季度维度看，Q4 2025营收约307.74亿元，同比-12.4%，"
    "但净利息收入Q4同比+2.76%为核心亮点。四季度营收分别为337.09 / 356.76 / 312.83 / 307.74亿元，"
    "下半年总体趋于稳定。"
)
body_cn(doc,
    "净利息收入880.21亿元（-5.8% YoY）是最大拖累项，但降幅较FY2024明显收窄。"
    "非利息净收入434.21亿元（-18.5% YoY）主要受债券投资收益下滑影响。"
    "其中，财富管理手续费收入50.61亿元（+15.8%）逆势增长，体现业务结构持续优化。"
)
add_img_cn(doc, CHART_PATHS[1], "图2：季度营业收入走势（2024Q1–2025Q4，亿元）", width=Inches(6.0))
add_img_cn(doc, CHART_PATHS[10], "图3：营业收入结构 — 利息收入 vs 非利息收入（FY2022–FY2025）", width=Inches(5.5))

h2_cn(doc, "归母净利润 — 四季度集中计提拨备，全年基本符合预期")
body_cn(doc,
    "FY2025归母净利润426.33亿元，同比下降4.2%（FY2024：445.08亿元）。Q4隐含单季净利润约42.94亿元，"
    "大幅低于Q3的134.69亿元，反映银行年末集中计提拨备的惯常操作，符合历史规律。"
    "全年信用减值损失405.67亿元（-17.9% YoY），拨备计提压力持续下降，"
    "未来随不良生成率趋势性下行，拨备节约对利润的贡献有望持续释放。"
)
body_cn(doc,
    "全年信贷成本降至1.38%（-18bps YoY），是底线利润的重要支撑。"
    "拨备前利润（PPOP）同比下降约11.9%，表明营收仍是核心挑战，但净息差企稳后"
    "PPOP压力将趋于减轻，FY2026E净利润有望实现约+1.7%的温和正增长。"
)
add_img_cn(doc, CHART_PATHS[2], "图4：季度归母净利润走势及同比增速（2024Q1–2025Q4）", width=Inches(6.0))

hr_cn(doc)

# ── 三、核心经营指标分析 ──────────────────────────────────────────────────────
h1_cn(doc, "三、核心银行经营指标分析")

h2_cn(doc, "净息差（NIM）— 关键拐点已经到来")
body_cn(doc,
    "FY2025全年净息差1.78%，较FY2024的1.87%下行仅9个基点，较FY2024年的-51bps大幅收窄，"
    "亦显著优于FY2023年-48bps的压缩幅度。Q4净利息收入同比+2.76%为多季度内首次正增长，"
    "标志着息差最困难阶段已经过去。驱动因素：（一）负债成本大幅改善，全年平均资金成本降47bps至1.67%，"
    "存款成本降42bps至1.65%；（二）2025年下半年贷款重定价压力趋缓（LPR调降次数减少）；"
    "（三）存款期限结构改善，定期存款占比下降。"
)
body_cn(doc,
    "管理层明确表示净息差\u201c基本趋于稳定\u201d，2026年目标净息差约1.77%（Wind一致预期）。"
    "负债端成本改善仍有一定空间，有助于对冲资产端利率下行压力。"
)
add_img_cn(doc, CHART_PATHS[3], "图5：净息差（NIM）走势（FY2022–FY2025实际；FY2026E一致预期）", width=Inches(5.5))

h2_cn(doc, "资产质量 — 零售信用风险持续出清，对公房地产为主要关注点")
body_cn(doc,
    "FY2025整体资产质量持续改善。不良贷款率1.05%（-1bp），个人贷款不良率显著改善至1.23%（-16bps），"
    "印证零售信用风险正在消退。不良生成率降至1.63%（-17bps）；逾期90天以上贷款偏离度仅0.56倍，"
    "说明认定标准严格，未来新增不良压力可控。"
)
body_cn(doc,
    "主要关注点：对公不良率上升17bps至0.87%，其中房地产贷款不良率升至2.22%（+43bps）。"
    "但公司已持续压缩房地产贷款敞口，2025年房地产贷款占对公贷款比重已明显下降。"
    "拨备覆盖率220.88%（-29.8pp），仍远超监管最低要求，覆盖能力充足。"
)
add_img_cn(doc, CHART_PATHS[4], "图6：不良贷款率 & 拨备覆盖率（H1 2024–FY2025）", width=Inches(5.5))

h2_cn(doc, "资本充足率 — 有机积累，无需外部融资")
body_cn(doc,
    "FY2025末核心一级资本充足率9.36%（+24bps），一级资本充足率11.49%（+80bps），"
    "资本充足率13.77%（+66bps），资本实力稳步提升。管理层明确表示不计划进行配股或外部融资，"
    "依靠内生资本积累即可支持2026年业务适度扩张及分红安排。"
)
add_img_cn(doc, CHART_PATHS[5], "图7：资本充足率（FY2024 vs FY2025）", width=Inches(5.0))

make_table_cn(doc,
    headers=["指标", "FY2023A", "FY2024A", "H1 2025A", "FY2025A", "FY2026E"],
    rows=[
        ["营业收入（亿元）",        "1,647.0", "1,467.0", "693.9",  "1,314.4", "约1,352"],
        ["归母净利润（亿元）",      "464.6",   "445.1",   "248.7",  "426.3",   "约433.6"],
        ["净息差NIM（%）",         "2.38",    "1.87",    "1.80",   "1.78",    "约1.77"],
        ["不良贷款率（%）",        "1.06",    "1.06",    "1.10",   "1.05",    "约1.03"],
        ["拨备覆盖率（%）",        "277.6",   "250.71",  "246.6",  "220.88",  "约215"],
        ["核心一级资本充足率（%）", "8.84",    "9.12",    "9.20",   "9.36",    "约9.5"],
        ["ROE（%）",              "11.4",    "10.08",   "9.52",   "9.15",    "约8.9"],
        ["成本收入比（%）",        "26.3",    "约27.7",  "28.3",   "29.06",   "约27.7"],
        ["EPS（元）",             "2.25",    "2.16",    "1.21",   "2.07",    "约2.19"],
        ["每股净资产BVPS（元）",   "约21.1",  "约22.4",  "约23.1", "约24.0",  "约24.9"],
    ],
    hdr_bg="003366", alt_bg="F0F4F8"
)

hr_cn(doc)

# ── 四、零售转型进展 ──────────────────────────────────────────────────────────
h1_cn(doc, "四、零售银行转型进展")
body_cn(doc,
    "零售业务转型是平安银行核心长期投资逻辑。FY2025年，零售条线净利润大幅增长+828%至26.83亿元"
    "（FY2024：2.89亿元），占全行净利润比重从0.6%提升至6.3%，标志着零售信贷风险成本周期性高点"
    "已经过去。管理层表示零售转型\u201c完成约70%，剩余约30%\u201d。"
)

h2_cn(doc, "财富管理 — 亮点业务")
bullet_cn(doc,
    "零售AUM（管理资产总额）：4.24万亿元（+1.1% YoY），"
    "受权益市场波动影响增速温和，但客户关系稳固。")
bullet_cn(doc,
    "私行客户：10.56万人（+9.1% YoY）；私行AUM：1.9913万亿元（+0.8% YoY）。"
    "私行是零售最高盈利段，是中长期核心增长驱动。")
bullet_cn(doc,
    "财富管理手续费收入：50.61亿元（+15.8% YoY），其中保险佣金+53.3%，"
    "公募基金+8.9%，理财产品+8.8%，收入质量持续提升。",
    bold=True)
bullet_cn(doc,
    "口袋银行APP月活用户约4,000万，是手续费收入增长的核心数字分发渠道。")
bullet_cn(doc,
    "保险+银行协同效应：与平安集团保险板块深度联动，银保佣金收入+53.3%，"
    "集团综合金融生态形成显著差异化竞争优势。")

add_img_cn(doc, CHART_PATHS[6], "图8：零售AUM、财富客户及私行客户增长趋势（FY2023–FY2025）", width=Inches(5.5))

hr_cn(doc)

# ── 五、投资逻辑更新 ──────────────────────────────────────────────────────────
h1_cn(doc, '五、投资逻辑更新 — 维持\u201c买入\u201d评级')
body_cn(doc,
    '维持\u201c买入\u201d评级，目标价14.00元（0.56倍FY2026E每股净资产约24.9元，'
    "Wind一致预期），较当前股价约10.87元具备约29%上涨空间。"
    "核心逻辑：（一）净息差企稳为最重要近期催化剂；（二）信贷成本下降释放利润增长空间；"
    "（三）零售转型带动手续费收入持续增长；（四）约5.5%股息率提供下行保护，估值处历史低位（约0.44倍市净率）。"
)

h2_cn(doc, "多头逻辑（牛市情形）")
bullet_cn(doc,
    "息差拐点：Q4净利息收入同比+2.76%，若Q1 2026延续正增长，营收将于2026年转正，"
    "为股价重估最重要的单一催化剂。")
bullet_cn(doc,
    "信贷成本正常化：不良生成率降至1.63%且趋势下行，预计FY2026E拨备节约释放约25亿元以上利润，"
    "即使PPOP基本持平，归母净利润仍可实现正增长。",
    bold=True)
bullet_cn(doc,
    "零售转型利润贡献扩大：FY2025零售利润占比从0.6%跃升至6.3%，"
    "目标FY2027E达25%至30%，为长期盈利中枢上移的核心驱动力。")
bullet_cn(doc,
    "历史低位估值：当前市净率约0.44倍，较历史5年均值约0.85倍折价约48%。"
    "若估值修复至0.60倍，对应股价约14.9元，较当前上涨约37%。")
bullet_cn(doc,
    "平安集团背书：中国最大保险集团强大股东背景提供资本支持、"
    "保险渠道协同及系统性重要性保护，是差异化护城河。")

h2_cn(doc, "主要风险")
bullet_cn(doc,
    "营收压力持续：若2026年LPR再降息或存款重定价慢于预期，NIM可能再度承压，拖延营收复苏进程。")
bullet_cn(doc,
    "对公房地产信用风险：房地产不良率已升至2.22%，若行业信用事件扩大，或需新增计提拨备，"
    "对冲信贷成本改善红利。")
bullet_cn(doc,
    "宏观经济下行：消费低迷与中小企业经营压力可能重新加快零售不良生成，"
    "尤其是消费贷和信用卡应收账款。")
bullet_cn(doc,
    "资本充足率缓冲有限：核心一级资本充足率9.36%虽满足监管要求，但在贷款快速扩张或"
    "意外损失情形下，提升分红比例的空间将受到制约。")
bullet_cn(doc,
    "国有大行竞争加剧：工农中建交等国有大行在财富管理和对公贷款市场主动进攻，"
    "平安银行在这些领域的竞争护城河面临收窄压力。")

hr_cn(doc)

# ── 六、估值分析 ──────────────────────────────────────────────────────────────
h1_cn(doc, "六、估值分析与盈利预测更新")

h2_cn(doc, "目标价：14.00元（0.56倍FY2026E市净率）")
body_cn(doc,
    "12个月目标价14.00元，基于0.56倍FY2026E每股净资产约24.9元（Wind一致预期）。"
    "估值依据：参考DDM/GGM模型，以ROE约8.9%除以权益成本约10.0%，"
    "内在合理P/B约0.9倍，考虑零售转型执行风险及宏观不确定性，给予40%折扣，"
    "得出目标P/B约0.56倍，与Wind 25家机构平均目标价14.25元高度一致。"
)
body_cn(doc,
    "牛市情形（0.70倍P/B）：目标价约17.4元，对应业绩全面复苏及净息差明显回升。"
    "熊市情形（0.40倍P/B）：目标价约10.0元，对应营收持续萎缩及新一轮资产质量恶化。"
)

make_table_cn(doc,
    headers=["指标", "FY2023A", "FY2024A", "FY2025A", "FY2026E", "FY2027E"],
    rows=[
        ["营业收入（亿元）",          "1,647.0", "1,467.0", "1,314.4", "约1,352", "约1,390"],
        ["营收增速（%）",             "-8.5",    "-10.9",   "-10.4",   "约+2.9",  "约+2.8"],
        ["归母净利润（亿元）",        "464.6",   "445.1",   "426.3",   "约433.6", "约444.0"],
        ["净利润增速（%）",           "+2.1",    "-4.2",    "-4.2",    "约+1.7",  "约+2.4"],
        ["EPS（元）",                "2.25",    "2.16",    "2.07",    "约2.19",  "约2.26"],
        ["每股净资产BVPS（元）",     "约21.1",  "约22.4",  "约24.0",  "约24.9",  "约26.0"],
        ["当前股价对应P/B（倍）",    "0.52×",   "0.49×",   "0.45×",   "0.44×",   "0.42×"],
        ["目标价14元对应P/B（倍）",  "0.66×",   "0.63×",   "0.58×",   "0.56×",   "0.54×"],
        ["当前股价对应P/E（倍）",    "4.83×",   "5.03×",   "5.25×",   "4.96×",   "4.81×"],
        ["股息率（按当前股价）",     "约5.2%",  "约5.1%",  "约5.5%",  "约5.5%",  "约5.7%"],
    ],
    hdr_bg="003366", alt_bg="F0F4F8"
)
body_cn(doc, "注：FY2026E/2027E为Wind一致预期；BVPS基于留存收益预测；股息率基于约10.87元当前价格。")
add_img_cn(doc, CHART_PATHS[8], "图9：年度营收与归母净利润走势（FY2022–FY2027E）", width=Inches(6.0))
add_img_cn(doc, CHART_PATHS[9], "图10：主要券商FY2026E目标价 & EPS预测汇总", width=Inches(6.0))

hr_cn(doc)

# ── 七、数据来源 ──────────────────────────────────────────────────────────────
h1_cn(doc, "七、数据来源与参考资料")
body_cn(doc, "所有数据截至2026年4月5日；FY2025年报于2026年3月20日正式发布。")

src_data = [
    ("平安银行2025年年报发布：全年营收1314.42亿元（新浪财经，2026-03-20）",
     "https://finance.sina.com.cn/jryx/bank/2026-03-20/doc-inhrrzut6659056.shtml"),
    ("平安银行归母净利润426.33亿元，拟每10股派息5.96元（新浪财经，2026-03-20）",
     "https://finance.sina.com.cn/roll/2026-03-20/doc-inhrshar6590029.shtml"),
    ("平安银行2025年年报业绩点评：单季利息净收入正增（银河证券，2026-03-24）",
     "https://stock.stockstar.com/JC2026032400000016.shtml"),
    ("平安银行2025年年报解读（新浪财经，2026-03-31）",
     "https://finance.sina.com.cn/wm/2026-03-31/doc-inhsxkrq2161128.shtml"),
    ("平安银行业绩会：2026年将全力以赴重回增长（新浪财经，2026-03-23）",
     "https://finance.sina.com.cn/roll/2026-03-23/doc-inhryxfp3966939.shtml"),
    ("读懂平安银行2025年报：最难的时候过去了（新浪财经，2026-03-25）",
     "https://finance.sina.com.cn/stock/bxjj/2026-03-25/doc-inhsecsi9934000.shtml"),
    ("平安银行2025：走出深水区的第一步（新浪财经，2026-03-25）",
     "https://finance.sina.com.cn/wm/2026-03-25/doc-inhsfmms3417265.shtml"),
    ("21世纪经济报道：年报出炉，分红百亿，净息差降9bps（2026-03-20）",
     "https://www.21jingji.com/article/20260320/herald/5ee8bf7a9198ff98e638898d5820cc17.html"),
    ("Wind / 同花顺F10：FY2026E一致预期（000001.SZ）",
     "https://basic.10jqka.com.cn/000001/worth.html"),
    ("战略转型进度条剩余30%！平安银行2026重回增长方法论（证券时报）",
     "https://www.stcn.com/article/detail/3693988.html"),
    ("平安银行接受境内外投资者调研（新浪财经，2026-03-25）",
     "https://finance.sina.com.cn/stock/aigc/jgdy/2026-03-25/doc-inhseyvw3553593.shtml"),
    ("平安银行H1 2025净利微降3.9%，净息差收窄至1.80%（华尔街见闻，2025年8月）",
     "https://wallstreetcn.com/articles/3753917"),
]

for title, url in src_data:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    add_hyperlink(p, url, title, color="#003366")

doc.add_paragraph()
disc = doc.add_paragraph()
disc_run = disc.add_run(
    "免责声明：本报告仅供参考，不构成任何投资建议。所有财务数据来源于公司公告及公开信息，"
    "市场数据通过yfinance实时获取。FY2026E/2027E为分析师一致预期，不代表实际结果。"
)
disc_run.font.size = Pt(8.5); disc_run.font.color.rgb = M_GRAY; disc_run.italic = True
disc_run.font.name = "Times New Roman"; set_cjk_font(disc_run, "宋体")
disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── 保存 ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT, "000001_Q4_2025_业绩更新报告_中文版.docx")
doc.save(out_path)
print(f"\u2705 中文报告已生成 → {out_path}")
