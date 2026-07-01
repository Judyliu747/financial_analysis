#!/usr/bin/env python3
"""Generate SpaceX (SPCX) Morning Note — 2026年6月18日 (Chinese / 中文版)."""

import os
import yfinance as yf
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "output", "SPCX")
os.makedirs(OUT, exist_ok=True)

SONG = "宋体"
HEI  = "黑体"

# ── Market Data (dynamic via yfinance) ───────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        info = yf.Ticker(ticker).fast_info
        return {
            "price": round(info.last_price, 2), "market_cap": info.market_cap,
            "52w_high": round(info.year_high, 2), "52w_low": round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

UNIVERSE = ["SPCX", "RKLB", "NBIS", "NVDA"]
mkt = {tk: get_market_data(tk) for tk in UNIVERSE}

def px(tk):  return f"${mkt[tk]['price']}" if mkt[tk]['price'] != 'N/A' else 'N/A'
def cap(tk):
    m = mkt[tk]['market_cap']
    if m == 'N/A': return 'N/A'
    return f"{m/1e12:.2f}万亿美元" if m >= 1e12 else f"{m/1e9:.0f}亿美元"
def rng(tk):
    d = mkt[tk]
    if d['52w_low'] == 'N/A': return 'N/A'
    return f"${d['52w_low']}–${d['52w_high']}"

SPCX_INIT = 160.95
SPCX_IPO  = 135.00
spcx_now  = mkt["SPCX"]["price"]
PT        = 150.0
if spcx_now != 'N/A':
    chg_init = (spcx_now / SPCX_INIT - 1) * 100
    chg_ipo  = (spcx_now / SPCX_IPO - 1) * 100
    dn_pt    = (PT / spcx_now - 1) * 100
    ps_now   = 94 * (mkt["SPCX"]["market_cap"] / 2.11e12)
else:
    chg_init = chg_ipo = dn_pt = ps_now = float('nan')

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def set_cn(run, cn=SONG, en="Times New Roman", size=11):
    run.font.name = en; run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = parse_xml(f'<w:rFonts {nsdecls("w")}/>'); rPr.insert(0, rF)
    rF.set(qn("w:eastAsia"), cn)

def add_p(doc, text, bold=False, font_size=11, color=None, alignment=None,
          space_after=6, space_before=0, heading=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn(run, cn=HEI if heading else SONG, size=font_size)
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment: p.alignment = alignment
    pf = p.paragraph_format; pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    return p

def add_label(doc, header, body, fs=10.5):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{header}："); set_cn(r1, cn=HEI, size=fs); r1.bold = True
    r1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    r2 = p.add_run(body); set_cn(r2, cn=SONG, size=fs)
    p.paragraph_format.space_after = Pt(7)

def add_src(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text); set_cn(run, cn=SONG, size=8)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75); run.italic = True
    p.paragraph_format.space_after = Pt(4)

def fmt_table(table, hdr="003366"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, hdr)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True; set_cn(r, cn=HEI, size=9)
    for i, row in enumerate(table.rows):
        if i == 0: continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: set_cn(r, cn=SONG, size=9)

# ── Build ────────────────────────────────────────────────────────────────────
doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"; doc.styles["Normal"].font.size = Pt(11)
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

add_p(doc, "晨会纪要 — 2026年6月18日", bold=True, font_size=15,
      color=(0x00, 0x33, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, heading=True)
add_p(doc, "航天 / 卫星互联 / AI算力基础设施", bold=True, font_size=11,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, heading=True)

# 首要观点
add_p(doc, "首要观点 — SPCX：勿追IPO逼空行情，维持「市场表现（Market Perform）」评级，目标价150美元",
      bold=True, font_size=12, color=(0x00, 0x33, 0x66), space_after=4, space_before=2, heading=True)
add_p(doc,
      f"SpaceX股价已升至{px('SPCX')}（市值约{cap('SPCX')}），较我们6月14日按{SPCX_INIT:.0f}美元给出的"
      f"首次评级上涨+{chg_init:.0f}%，较135美元的IPO发行价上涨+{chg_ipo:.0f}%。本轮上涨将估值推升至"
      f"约{ps_now:.0f}倍市销率（IPO定价时为94倍）——在S-1之后并无新的基本面数据出炉的情况下估值更贵了。"
      f"我们150美元的目标价目前隐含约{abs(dn_pt):.0f}%的下行空间。当前点位我们不建议买入，倾向于在"
      f"上涨中减持/逢高了结。",
      font_size=10.5, space_after=8)

# 行情快照
add_p(doc, "覆盖范围行情快照（实时报价）", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=2, heading=True)
rows = [
    ("代码", "最新价", "市值", "52周区间", "角色 / 关联解读"),
    ("SPCX", px("SPCX"), cap("SPCX"), rng("SPCX"), "重点标的——航天+星链+xAI"),
    ("RKLB", px("RKLB"), cap("RKLB"), rng("RKLB"), "上市的纯航天标的；最干净的可比"),
    ("NBIS", px("NBIS"), cap("NBIS"), rng("NBIS"), "xAI/Colossus的AI云可比标的"),
    ("NVDA", px("NVDA"), cap("NVDA"), rng("NVDA"), "Colossus建设的GPU供应商"),
]
t = doc.add_table(rows=len(rows), cols=5); t.style = "Table Grid"
for i, r in enumerate(rows):
    for j, v in enumerate(r):
        cell = t.cell(i, j); cell.text = ""
        run = cell.paragraphs[0].add_run(v); set_cn(run, cn=SONG, size=9)
fmt_table(t)
add_src(doc, "资料来源：yfinance 实时报价（盘中）。SPCX基本面数据来自S-1招股书 / 我们6月14日的首次评级报告。")
add_p(doc, "", space_after=2)

# 隔夜 / 盘前动态
add_p(doc, "隔夜 / 盘前动态", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=2, heading=True)
add_label(doc, "SPCX",
    "无新增披露——纯粹是IPO后的动量/稀缺性买盘（首只上市的航天+AI巨型市值标的；可能存在指数纳入预期炒作）。"
    "投资逻辑不变：星链是皇冠上的明珠（营收33亿美元、经营利润12亿美元、订户1030万），但xAI每季度烧钱"
    "(25)亿美元，且Q1合并自由现金流为(91)亿美元。价格更高，风险照旧。")
add_label(doc, "RKLB",
    f"（{px('RKLB')}，约{cap('RKLB')}）最接近的上市纯航天标的；已从52周高点回落，但年内仍是数倍涨幅。"
    "业务干净——发射+航天系统，无AI烧钱拖累。是在不承担SPCX「xAI税负」前提下获取航天敞口的最佳载体。")
add_label(doc, "NBIS",
    f"（{px('NBIS')}，约{cap('NBIS')}）AI云可比标的，逼近52周高点。关联解读是双向的——既印证了投资者"
    "赋予xAI/Colossus基础设施的价值，也暴露出这条赛道的拥挤与高资本开支（AWS/Azure/GCP/CoreWeave）。")
add_label(doc, "NVDA",
    f"（{px('NVDA')}，约{cap('NVDA')}）xAI的Colossus建设（22万张GPU、占SPCX一季度资本开支77亿美元）的"
    "直接受益方。SPCX年化约400亿美元的资本开支节奏，是对加速卡需求持续旺盛的边际印证。")

# 今日关注
add_p(doc, "今日关键事件 / 关注要点", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=2, heading=True)
for b in [
    "覆盖标的无业绩发布安排。关注SPCX上市首周的波动，以及卖方早期首次评级（可能进一步助推动量行情）。",
    "投资逻辑监控项：(1) 星链ARPU（已下滑33%至66美元/月）；(2) xAI烧钱速度 vs. 约750亿美元IPO募资"
    "所支撑的约两年现金跑道。",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(b); set_cn(r, cn=SONG, size=10.5)
    p.paragraph_format.space_after = Pt(5)

# 交易思路
add_p(doc, "交易思路", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=4, heading=True)
add_label(doc, "逢高减持 / 回避 SPCX",
    f"目标价150美元 = 约{abs(dn_pt):.0f}%下行空间。IPO跳涨已计入乐观情景；受xAI投资周期拖累，近端"
    "每股盈利可见度有限。风险：动量+潜在指数纳入+星链订户/ARPU再超预期，均可能延续逼空——这是估值层面的"
    "判断，而非基本面层面的做空。")
add_label(doc, "相对价值：做多RKLB / 做空SPCX（配对交易）",
    "在不承担xAI年化约(100)亿美元亏损与马斯克85%投票权治理隐忧的前提下，持有上市的航天成长性。"
    "风险：RKLB自身估值亦不便宜且体量小得多；SPCX的星链现金引擎无可比拟。")
add_label(doc, "NVDA关联解读（暂不操作）",
    "SPCX的资本开支披露强化了AI资本开支景气度；对已在覆盖范围内的NVDA构成边际利好。")

add_p(doc, "", space_after=2)
add_src(doc, "本纪要撰写于6月18日开盘前；SPCX为实时报价并在变动中——开盘后点位可能变化。"
             "仅供内部讨论，不构成投资建议。")

path = os.path.join(OUT, "SPCX_Morning_Note_2026-06-18_中文版.docx")
doc.save(path)
print(f"已保存：{path}")
