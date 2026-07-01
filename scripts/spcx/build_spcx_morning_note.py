#!/usr/bin/env python3
"""Generate SpaceX (SPCX) Morning Note — June 18, 2026 (English)."""

import os
import yfinance as yf
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "output", "SPCX")
os.makedirs(OUT, exist_ok=True)

# ── Market Data (dynamic via yfinance) ───────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        info = yf.Ticker(ticker).fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: Could not fetch market data for {ticker}: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

UNIVERSE = ["SPCX", "RKLB", "NBIS", "NVDA"]
mkt = {tk: get_market_data(tk) for tk in UNIVERSE}

def px(tk):  return f"${mkt[tk]['price']}" if mkt[tk]['price'] != 'N/A' else 'N/A'
def cap(tk):
    m = mkt[tk]['market_cap']
    if m == 'N/A': return 'N/A'
    return f"${m/1e12:.2f}T" if m >= 1e12 else f"${m/1e9:.0f}B"
def rng(tk):
    d = mkt[tk]
    if d['52w_low'] == 'N/A': return 'N/A'
    return f"${d['52w_low']}–${d['52w_high']}"

# SPCX move math vs. initiation
SPCX_INIT = 160.95   # price referenced in Jun 14 initiation
SPCX_IPO  = 135.00
spcx_now  = mkt["SPCX"]["price"]
PT        = 150.0
if spcx_now != 'N/A':
    chg_init = (spcx_now / SPCX_INIT - 1) * 100
    chg_ipo  = (spcx_now / SPCX_IPO - 1) * 100
    dn_pt    = (PT / spcx_now - 1) * 100
    ps_now   = 94 * (mkt["SPCX"]["market_cap"] / 2.11e12)  # scale 94x@$2.11T basis
else:
    chg_init = chg_ipo = dn_pt = ps_now = float('nan')

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def add_p(doc, text, bold=False, font_size=11, color=None,
          alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"; run.font.size = Pt(font_size)
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment: p.alignment = alignment
    pf = p.paragraph_format; pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    return p

def add_label(doc, header, body, fs=10.5):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{header}: "); r1.font.name = "Times New Roman"; r1.font.size = Pt(fs); r1.bold = True
    r1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    r2 = p.add_run(body); r2.font.name = "Times New Roman"; r2.font.size = Pt(fs)
    p.paragraph_format.space_after = Pt(7)
    return p

def add_src(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.name = "Times New Roman"; run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75); run.italic = True
    p.paragraph_format.space_after = Pt(4)

def fmt_table(table, hdr_color="003366"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, hdr_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True; r.font.size = Pt(9); r.font.name = "Times New Roman"
    for i, row in enumerate(table.rows):
        if i == 0: continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = "Times New Roman"

# ── Build ────────────────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(11)
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

add_p(doc, "MORNING NOTE — June 18, 2026", bold=True, font_size=15,
      color=(0x00, 0x33, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_p(doc, "Space / Satellite Connectivity / AI Infrastructure",
      bold=True, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

# Top Call
add_p(doc, "Top Call — SPCX: Don't chase the IPO melt-up; reiterate Market Perform, $150 PT",
      bold=True, font_size=12, color=(0x00, 0x33, 0x66), space_after=4, space_before=2)
add_p(doc,
      f"SpaceX has run to {px('SPCX')} (mkt cap ~{cap('SPCX')}), up +{chg_init:.0f}% since our "
      f"June 14 initiation at ${SPCX_INIT:.0f} and +{chg_ipo:.0f}% off the ${SPCX_IPO:.0f} IPO price. "
      f"The move has pushed the stock to ~{ps_now:.0f}x trailing sales (vs. 94x at the IPO print) — "
      f"richer, on no new fundamental data since the S-1. Our $150 PT now implies ~{abs(dn_pt):.0f}% "
      f"downside. We are not buyers here; we'd fade strength / trim into the momentum.",
      font_size=10.5, space_after=8)

# Market snapshot table
add_p(doc, "Coverage Snapshot (live quotes)", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=2)
rows = [
    ("Ticker", "Last", "Mkt Cap", "52-Wk Range", "Role / Read-Through"),
    ("SPCX", px("SPCX"), cap("SPCX"), rng("SPCX"), "Focus name — space + Starlink + xAI"),
    ("RKLB", px("RKLB"), cap("RKLB"), rng("RKLB"), "Listed space pure-play; cleanest peer"),
    ("NBIS", px("NBIS"), cap("NBIS"), rng("NBIS"), "AI-neocloud comp for xAI/Colossus"),
    ("NVDA", px("NVDA"), cap("NVDA"), rng("NVDA"), "GPU supplier to Colossus buildout"),
]
t = doc.add_table(rows=len(rows), cols=5); t.style = "Table Grid"
for i, r in enumerate(rows):
    for j, v in enumerate(r):
        cell = t.cell(i, j); cell.text = ""
        run = cell.paragraphs[0].add_run(v); run.font.name = "Times New Roman"; run.font.size = Pt(9)
fmt_table(t)
add_src(doc, "Source: yfinance live quotes (intraday). SPCX fundamentals per S-1 / our June 14 initiation.")
add_p(doc, "", space_after=2)

# Overnight / Pre-Market
add_p(doc, "Overnight / Pre-Market Developments", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=2)
add_label(doc, "SPCX",
    "No new disclosure — pure momentum / scarcity bid post-IPO (first listed space+AI mega-cap; "
    "likely index-inclusion speculation). Thesis unchanged: Starlink is the crown jewel ($3.3B rev, "
    "$1.2B op profit, 10.3M subs), but xAI is burning $(2.5)B/qtr and consolidated FCF was $(9.1)B in Q1. "
    "Higher price, same risk.")
add_label(doc, "RKLB",
    f"({px('RKLB')}, ~{cap('RKLB')}) Closest listed space pure-play; off its 52-wk high but a multi-bagger "
    "YoY. Trades clean — launch + Space Systems, no AI cash-burn drag. Best vehicle for space exposure "
    "without SPCX's xAI tax.")
add_label(doc, "NBIS",
    f"({px('NBIS')}, ~{cap('NBIS')}) AI-neocloud comp near 52-wk highs. Read-through is two-sided — "
    "validates the value investors ascribe to xAI/Colossus infra, but also flags how crowded and "
    "capital-intensive that race is (AWS/Azure/GCP/CoreWeave).")
add_label(doc, "NVDA",
    f"({px('NVDA')}, ~{cap('NVDA')}) Direct beneficiary of xAI's Colossus buildout (220K GPUs, $7.7B of "
    "SPCX's Q1 capex). SPCX's ~$40B annualized capex run-rate is incremental confirmation of sustained "
    "accelerator demand.")

# Key events
add_p(doc, "Key Events / Watch Items Today", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=2)
for b in [
    "No coverage earnings scheduled. Monitor SPCX for first-week-of-trading volatility and early "
    "sell-side initiations (could add to the momentum).",
    "Thesis monitors: (1) Starlink ARPU (already -33% to $66/mo) and (2) xAI burn vs. the ~2yr runway "
    "from ~$75B IPO proceeds.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(b); r.font.name = "Times New Roman"; r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(5)

# Trade ideas
add_p(doc, "Trade Ideas", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=4)
add_label(doc, "Trim / Avoid SPCX into strength",
    f"PT $150 = ~{abs(dn_pt):.0f}% downside. The IPO pop has priced optimistic scenarios; near-term EPS "
    "visibility is capped by the xAI investment cycle. Risk: momentum + potential index inclusion + "
    "another Starlink sub/ARPU beat can extend the squeeze — this is a valuation call, not a fundamental short.")
add_label(doc, "Relative value: Long RKLB / Short SPCX (pairs)",
    "Own listed space growth without the ~$(10)B/yr xAI annualized loss and 85% Musk voting overhang. "
    "Risk: RKLB is itself richly valued and far smaller scale; SPCX's Starlink cash engine has no peer.")
add_label(doc, "NVDA read-through (no action)",
    "SPCX capex disclosure reinforces the AI-capex tape; constructive at the margin for NVDA already "
    "under coverage.")

add_p(doc, "", space_after=2)
add_src(doc, "Written pre-open 6/18; SPCX quote is live and moving — levels may shift by the bell. "
             "For internal discussion only; not investment advice.")

path = os.path.join(OUT, "SPCX_Morning_Note_2026-06-18.docx")
doc.save(path)
print(f"Saved: {path}")
print(f"SPCX {px('SPCX')} | +{chg_init:.1f}% vs init | PT downside {dn_pt:.1f}% | ~{ps_now:.0f}x P/S")
